#!/usr/bin/env python3
"""
Generate Retail Control Architect™ — Complete VSL Script .docx
Follows the IEC Diamond VSL framework slide-by-slide, fully written out
for Aakash Savant's retail stock control consulting offer.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colour Palette ──────────────────────────────────────────────────
DEEP_BLUE    = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD  = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL  = RGBColor(0x17, 0xA2, 0xB8)
WARM_RED     = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS_GRN  = RGBColor(0x27, 0xAE, 0x60)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
MED_GRAY     = RGBColor(0x66, 0x66, 0x66)
BODY_BLACK   = RGBColor(0x2C, 0x2C, 0x2C)
ORANGE       = RGBColor(0xE6, 0x7E, 0x22)
PURPLE       = RGBColor(0x8E, 0x44, 0xAD)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BODY_BLACK

# ── Helpers ─────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)

def add_cover_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(60)
    p.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DEEP_BLUE

def add_cover_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT_GOLD

def add_section_header(text, emoji=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(4)
    full = f"{emoji}  {text}" if emoji else text
    r = p.add_run(full)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DEEP_BLUE
    line = doc.add_paragraph()
    line.space_after = Pt(8)
    rl = line.add_run("━" * 75)
    rl.font.size = Pt(6)
    rl.font.color.rgb = ACCENT_GOLD

def add_sub_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color or ACCENT_TEAL

def add_minor_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color or DEEP_BLUE

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_highlight_box(text, bg_hex="FFF8E1"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BODY_BLACK
    set_cell_shading(cell, bg_hex)
    cell.width = Inches(6.0)
    add_spacer(6)

def add_slide(slide_num, text, is_spoken=True):
    """Add a VSL slide with script text."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"SLIDE {slide_num}")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = ACCENT_GOLD
    
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    cp = cell.paragraphs[0]
    cr = cp.add_run(text)
    cr.font.size = Pt(12) if is_spoken else Pt(11)
    cr.font.color.rgb = DEEP_BLUE if is_spoken else MED_GRAY
    if is_spoken:
        cr.bold = True
    set_cell_shading(cell, "F7F9FC" if is_spoken else "FFFDE7")
    cell.width = Inches(6.0)
    add_spacer(4)

def add_direction(text):
    """Add a stage direction / instruction."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"📋 {text}")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = MED_GRAY

def add_bullet(text, prefix_emoji=None, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.clear()
    if prefix_emoji:
        re = p.add_run(f"{prefix_emoji} ")
        re.font.size = Pt(11)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(11)
        p.add_run(f" {text}").font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)

def add_numbered(text, num, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.color.rgb = ACCENT_TEAL
    r.font.size = Pt(11)
    if bold_label:
        rl = p.add_run(f"{bold_label}: ")
        rl.bold = True
        rl.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)

def styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        set_cell_shading(cell, "1B2A4A")
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, "F7F9FC")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    add_spacer(6)
    return table


# ════════════════════════════════════════════════════════════════════
#                       DOCUMENT CONTENT
# ════════════════════════════════════════════════════════════════════

# ── COVER PAGE ─────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

add_cover_title("RETAIL CONTROL ARCHITECT™")
add_cover_title("Complete VSL Script")
add_cover_subtitle(
    "Word-for-Word Video Sales Letter Script\n"
    "Following the IEC Diamond VSL Framework\n"
    "15-Minute VSL to Book Free Store Check-Up Calls"
)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("READY TO RECORD — EVERY WORD, EVERY SLIDE, EVERY DIRECTION")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = ACCENT_GOLD

doc.add_paragraph()

add_minor_header("VSL Structure Overview")
styled_table(
    ["Phase", "Slides", "Duration", "Purpose"],
    [
        ["1. Pattern Interrupt", "1-3", "0:00 — 0:45", "Grab attention with shocking opening"],
        ["2. Big Promise + Transition", "4-7", "0:45 — 2:00", "Hook them with the dream outcome"],
        ["3. Pain Questions (Yes Ladder)", "8-13", "2:00 — 3:30", "Make them nod along"],
        ["4. Exclusivity + Proof Blast", "14-20", "3:30 — 5:00", "Build desire + show results"],
        ["5. Credibility + Origin Story", "21-30", "5:00 — 7:30", "Build trust + authority"],
        ["6. Close the Loop + System Reveal", "31-38", "7:30 — 9:30", "Reveal the framework"],
        ["7. Results + Proof Stack", "39-44", "9:30 — 11:00", "Mount undeniable proof"],
        ["8. Objection Handling", "45-52", "11:00 — 12:30", "Kill every fear"],
        ["9. The Offer + Guarantee", "53-60", "12:30 — 14:00", "Present the opportunity"],
        ["10. CTA + Urgency Close", "61-70", "14:00 — 15:30", "Drive the booking"],
    ],
    col_widths=[1.8, 0.6, 1.0, 2.6]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 1: PATTERN INTERRUPT
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 1: PATTERN INTERRUPT", "⚡")
add_body("Goal: STOP them from scrolling. Use a shocking visual or statement that breaks their mental pattern.")

add_highlight_box(
    "FRAMEWORK: Use a pattern interrupt that you will RELATE BACK to your topic later.\n"
    "This opens a 'loop' in their mind that forces them to keep watching to close it.\n\n"
    "For RCA, we use the 'Invisible Thief' pattern interrupt — a concept every store owner can relate to.",
    "FFF8E1"
)

add_slide(1,
    "There is a thief inside your store right now.")
add_direction("VISUAL: Show a shadowy silhouette inside a clothing store. Bold white text on dark background.")

add_slide(2,
    "But here's the thing…\nYou hired him. You pay his salary.\nAnd he steals from you every single day.")
add_direction("VISUAL: Dramatic text reveal. Keep it mysterious. They should be confused — 'Wait, what?'")

add_slide(3,
    "In the next few minutes, I'm going to show you how this 'invisible thief' "
    "is silently draining ₹2 to 5 Lakhs from your retail store every year…\n\n"
    "And exactly how to catch him — in 30 days flat.")
add_direction("VISUAL: The number '₹2–5 Lakhs' should appear large and bold. This is the TRANSITION slide — connecting the pattern interrupt to the big promise.")

add_highlight_box(
    "WHY THIS WORKS:\n"
    "• The 'thief in your store' metaphor = instant attention grab\n"
    "• 'You hired him' = pattern break — they expect a real thief, not their own system\n"
    "• We open a LOOP — they MUST keep watching to find out who this 'thief' is\n"
    "• We'll close this loop later when we reveal that the 'thief' = their broken SYSTEM",
    "E8F5E9"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 2: BIG PROMISE + INTRO
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 2: BIG PROMISE + INTRODUCTION", "🎯")
add_body("Goal: State the specific result they'll get. Introduce yourself briefly. Then INTERRUPT the pattern again — don't give your full intro yet.")

add_slide(4,
    "How to stop losing ₹40,000 or more every single month from your clothing store…\n\n"
    "Without buying expensive software, without firing your staff, "
    "and without spending hours reconciling stock every night.\n\n"
    "Even if you've tried billing software before and it didn't work.")
add_direction("VISUAL: Headline text — use the How To / Without / Even If formula. This is your BIG PROMISE.")

add_highlight_box(
    "THE BIG PROMISE MUST BE:\n"
    "N = New (a new system, not just software)\n"
    "U = Unique (nobody else installs a stock control PROCESS)\n"
    "E = Exciting (₹40K/month savings — that's exciting!)\n"
    "E = Easy (without complex tech, without firing staff)\n"
    "P = Predictable (30 days, measurable reduction)\n"
    "H = Huge (₹2-5 Lakhs per year recovered)",
    "FFF8E1"
)

add_slide(5,
    "Hey, I'm Aakash Savant.")
add_direction("VISUAL: Your headshot or face on camera. Clean, professional. 1-2 seconds only.")

add_slide(6,
    "And let me ask you this…")
add_direction("VISUAL: Simple text. You are NOT giving your full introduction yet — you're breaking the pattern. They expect a bio. Instead, you ask questions.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 3: PAIN QUESTIONS (YES LADDER)
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 3: PAIN QUESTIONS — THE YES LADDER", "😤")
add_body("Goal: Ask 3-4 pain questions that make them unconsciously nod YES. This builds empathy and keeps attention.")

add_slide(7,
    "Are you sick and tired of your stock count NEVER matching "
    "what your billing software says?")
add_direction("VISUAL: Red frustrated emoji or stressed store owner visual.")

add_slide(8,
    "Do you wish you could stop spending 2 hours every night "
    "trying to find where items went missing?")

add_slide(9,
    "Are you looking for a way to finally run your store "
    "without depending on one 'trusted' staff member "
    "who may or may not be stealing from you?")

add_slide(10,
    "Or do you simply want to go home on time, knowing that "
    "every single item in your store is exactly where it should be?")

add_highlight_box(
    "WHY THIS WORKS:\n"
    "Each question puts them on a 'YES ladder' — they're unconsciously nodding.\n"
    "Every YES = higher engagement = higher chance they watch to the end.\n"
    "These are the EXACT pain points retail store owners experience daily.",
    "E8F5E9"
)

add_slide(11,
    "If this sounds like you… then you're in luck.")

add_slide(12,
    "Because in the next few minutes, I'm going to reveal "
    "the single most effective stock control breakthrough "
    "in the Indian retail industry.")

add_slide(13,
    "Sound good? Excellent.")
add_direction("VISUAL: Simple slide. Brief pause. Builds anticipation.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 4: EXCLUSIVITY + PROOF BLAST
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 4: EXCLUSIVITY + PROOF BLAST", "🏆")
add_body("Goal: Build desire by making it exclusive, then HIT them with undeniable proof.")

add_slide(14,
    "Now before I reveal how this works…")

add_slide(15,
    "Let me first say — this is NOT for everybody.")

add_slide(16,
    "I have to be extremely selective about who I work with.\n\n"
    "I only take on 3 new stores per month.\n\n"
    "For reasons you'll understand in just a moment…")
add_direction("VISUAL: '3 Stores Per Month' — big, bold. This opens another loop AND builds exclusivity.")

add_slide(17,
    "But first, let me SHOW you what I'm talking about…")

add_slide(18,
    "[INSERT PROOF: Screenshot of store owner WhatsApp message]\n\n"
    "'Aakash, our mismatch went from 17% to under 2% in just 45 days. "
    "We recovered ₹3.2 Lakhs that we didn't even know we were losing.'")
add_direction("VISUAL: Real WhatsApp screenshot or designed testimonial card. Show the RESULT clearly.")

add_slide(19,
    "[INSERT PROOF: Before/After stock report]\n\n"
    "Before RCA: 847 items mismatched across 2,100 SKUs.\n"
    "After RCA (Day 45): 12 items mismatched. Stock accuracy: 99.4%.")
add_direction("VISUAL: Side-by-side comparison. Red numbers on left, green on right.")

add_slide(20,
    "[INSERT PROOF: 3-4 more result screenshots]\n\n"
    "Store owner #2: Dead stock worth ₹4.2L identified and liquidated in 3 weeks.\n"
    "Store owner #3: Nightly reconciliation time reduced from 2 hours to 10 minutes.\n"
    "Store owner #4: Opened second branch within 6 months — operations were that clean.")
add_direction("VISUAL: Stack 3-4 proof points. Mix: WhatsApp screenshots, photos, data reports. Show different proof angles — recovery, time savings, expansion.")

add_highlight_box(
    "PROOF PRINCIPLE:\n"
    "Show proof from as many ANGLES as possible — different types of results, "
    "different kinds of store owners, different metrics.\n"
    "The more varied the proof, the more believable it becomes.\n"
    "If you don't have all of these yet → use the forensic audit numbers from your own analysis.",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 5: CREDIBILITY + ORIGIN STORY
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 5: CREDIBILITY + ORIGIN STORY", "👤")
add_body("Goal: Tell your story in a way that builds trust, shows empathy, and establishes authority. Be SPECIFIC and use DETAIL.")

add_slide(21,
    "Now before I tell you exactly how we achieve these results…\n\n"
    "Let me briefly explain who I am and why you should listen to me at all.")

add_slide(22,
    "My name is Aakash Savant and I'm the founder of Retail Control Architect™ — "
    "a specialized stock control system designed exclusively for owner-operated clothing stores in India.")

add_slide(23,
    "We help retail store owners like you eliminate stock mismatch, "
    "recover hidden losses, and build a store that runs on a system "
    "instead of depending on one person's memory.")

add_slide(24,
    "I'm passionate about this because I saw the exact same problem "
    "destroy the margins of store owners I knew — hardworking people "
    "who were bleeding money without realizing it.")

add_slide(25,
    "Here's what happened…")
add_direction("BEGIN ORIGIN STORY — make it personal, specific, and dramatic.")

add_slide(26,
    "I remember walking into a store and asking the owner one simple question: "
    "'How many units of your best-selling shirt do you have right now?'\n\n"
    "He looked at his screen. It said 34.\n\n"
    "We went to the rack. Counted. There were 21.\n\n"
    "13 shirts — gone. No record. No explanation. Just… gone.")
add_direction("VISUAL: Close-up of a store rack. Or show the numbers: 34 → 21 = 13 MISSING.")

add_slide(27,
    "And that was just ONE item out of 2,000+.\n\n"
    "When we did the full audit, we found the store was losing ₹4 Lakhs a year "
    "in stock that simply… disappeared.\n\n"
    "The owner thought his billing software was handling everything. It wasn't.\n"
    "The software tracks data. But who tracks the PROCESS?")
add_direction("VISUAL: The number ₹4,00,000 in large red text. Then 'Who tracks the process?' in gold.")

add_slide(28,
    "That's when I realized — this is NOT a software problem.\n"
    "It's NOT a staff problem.\n"
    "It is a SYSTEM problem.\n\n"
    "There was no step-by-step stock control process in the store. Period.")
add_direction("VISUAL: 3 lines with ❌ Software, ❌ Staff, ✅ SYSTEM. This is the BIG INSIGHT moment.")

add_slide(29,
    "So I built one.\n\n"
    "I spent months studying how the best-managed stores in the world "
    "track their inventory — from luxury retail in Europe to fast-fashion chains in Asia.\n\n"
    "And I distilled everything into a simple 5-step system "
    "that any store owner can implement in 30 days.")

add_slide(30,
    "Since then, I've helped store owners recover lakhs in hidden losses, "
    "reduce mismatch from double digits to under 2%, "
    "and build stores that run smoothly even when the owner is not physically there.\n\n"
    "All using the system I'm about to share with you.")
add_direction("VISUAL: Key stats — '₹3.2L recovered' / '17% → 2%' / 'Stores that run themselves'")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 6: CLOSE THE LOOP + SYSTEM REVEAL
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 6: CLOSE THE LOOP + SYSTEM REVEAL", "🔓")
add_body("Goal: Close the pattern interrupt loop, then reveal your system framework.")

add_slide(31,
    "Now, remember that invisible thief I mentioned at the start?")
add_direction("VISUAL: Callback to Slide 1. Same visual but now brighter — about to be revealed.")

add_slide(32,
    "That thief is not a person.\n\n"
    "It's not your staff member sitting at the counter.\n"
    "It's not the supplier short-shipping your orders.\n\n"
    "The invisible thief is your BROKEN SYSTEM — "
    "the absence of a step-by-step stock control process in your store.\n\n"
    "And you've been paying for it every single day you've been open.")
add_direction("VISUAL: The 'thief' silhouette morphs into a broken gear / broken system icon. Text: 'The Invisible Thief = Your Broken System'. LOOP CLOSED.")

add_highlight_box(
    "THIS IS THE MOST POWERFUL MOMENT IN THE VSL.\n"
    "The audience has been waiting to find out who the 'thief' is.\n"
    "When you reveal it's their own system (or lack thereof), "
    "it reframes everything they believed.\n"
    "Now they're primed for your solution.",
    "FDEDED"
)

add_slide(33,
    "And I guarantee this is NOT like anything you've heard before.")

add_slide(34,
    "So pay very close attention — because this could be "
    "the breakthrough your store has been waiting for.")

add_slide(35,
    "Here's the scoop…")

add_slide(36,
    "The Retail Control Architect™ system allows you to:\n\n"
    "• Find every single source of stock leakage in your store…\n"
    "• Build daily accountability checks that your staff actually follows…\n"
    "• Get a weekly report that takes 10 minutes to review…\n\n"
    "So you can achieve stock certainty, recover your hidden losses, "
    "and finally go home on time — all within 30 days.")
add_direction("VISUAL: The 5-step framework shown as a clean process diagram or icons.")

add_slide(37,
    "This is what we do day-in, day-out at Retail Control Architect™.")

add_slide(38,
    "And we've done it for store owners just like you…")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 7: RESULTS + PROOF STACK
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 7: RESULTS + PROOF STACK", "📸")
add_body("Goal: Mount undeniable proof that your system works. Stack it from every angle.")

add_slide(39,
    "[INSERT: 3-4 results screenshots]\n\n"
    "Rajesh, Pune: 'My stock matches exactly now. I sleep better at night.'\n"
    "Meera, Bengaluru: 'We found ₹1.8 Lakhs in dead stock we didn't know existed.'\n"
    "Vikram, Mumbai: 'My staff actually follows a process now. It's like a different store.'")
add_direction("VISUAL: Real testimonial screenshots — WhatsApp messages, video stills, photographs. Mix men and women, different cities, different results.")

add_slide(40,
    "[INSERT: More results from different angles]\n\n"
    "Revenue recovered. Time saved. Staff improved. Second store opened.")
add_direction("VISUAL: Stack more proof screenshots. Show emails, messages, before/after data.")

add_slide(41,
    "You see, at Retail Control Architect™, "
    "we find and eliminate hidden stock losses for retail store owners every single month.\n\n"
    "Using a system that is simple, repeatable, and guaranteed to work.")

add_slide(42,
    "We do all the detective work for you — "
    "the forensic audit, the SKU cleanup, the staff training — "
    "so you don't have to become a tech expert or an operations genius.")

add_slide(43,
    "Many of our store owners have recovered lakhs of rupees in hidden losses "
    "within the first 45 days.\n\n"
    "All without buying new software, firing their team, "
    "or spending hours on nightly reconciliation.")

add_slide(44,
    "Imagine: walking on the shop floor, knowing that every item "
    "on every rack is exactly where your system says it is.\n\n"
    "Imagine: leaving the store on time, with a clean daily report, "
    "and zero surprises waiting for you tomorrow.\n\n"
    "That's what we deliver. Guaranteed.")
add_direction("VISUAL: Paint the DREAM picture. Use aspirational imagery — a calm, organized store.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 8: OBJECTION HANDLING
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 8: OBJECTION HANDLING", "🛡️")
add_body("Goal: Raise and kill the top 3 objections BEFORE they stop watching. Use the 'And by the way, you may be thinking…' framework.")

add_slide(45,
    "And by the way… you might be thinking…")

add_slide(46,
    "'But Aakash, I already use billing software. "
    "Doesn't that handle my inventory?'")
add_direction("VISUAL: Show the objection in quotes — like it's coming from the viewer.")

add_slide(47,
    "I don't blame you. I used to think the same way.\n\n"
    "That was until I discovered that billing software only RECORDS data. "
    "It does NOT create the daily habits, checks, and processes "
    "that PREVENT stock from going missing.\n\n"
    "Your software is a register. My system is the discipline that makes "
    "the register trustworthy.")

add_slide(48,
    "Or maybe you're worried that your staff won't follow any new system?")

add_slide(49,
    "That's the exact reason I designed this system with "
    "built-in daily checks that physically CANNOT be skipped.\n\n"
    "The system doesn't depend on your staff's willingness. "
    "It's structured so that if a check is missed, you know about it "
    "before the store opens the next morning.\n\n"
    "We don't hope for compliance. We engineer it.")

add_slide(50,
    "Or are you thinking: 'My store is too small for this'?")

add_slide(51,
    "Here's something most people don't realize:\n\n"
    "Smaller stores actually lose a HIGHER percentage of their revenue "
    "to stock mismatch than large chains.\n\n"
    "At just ₹2,000 per day in hidden loss, that's over ₹7 Lakhs per year.\n\n"
    "This system was BUILT for stores like yours.")

add_slide(52,
    "Are you beginning to see how easy it can be for you to "
    "finally stop the bleeding in your store?")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 9: THE OFFER + GUARANTEE
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 9: THE OFFER + GUARANTEE", "💎")
add_body("Goal: Present the free check-up offer. Layer the guarantee. Make it irresistible.")

add_slide(53,
    "My 5-step Retail Control Architect™ system is…\n\n"
    "New — no one else installs a stock PROCESS in your store.\n"
    "Simple — it takes 10 minutes a day to maintain.\n"
    "Convenient — we do the heavy lifting for you.\n"
    "Predictable — measurable mismatch reduction in 30 days.\n"
    "And best of all…")

add_slide(54,
    "My team and I would love to find YOUR biggest stock leak — FOR FREE.")

add_slide(55,
    "And I'm so confident we can help you that I'll even guarantee it.")

add_slide(56,
    "Here's my promise to you:\n\n"
    "If after 30 days of working together, your stock mismatch does not reduce measurably, "
    "I will keep working with you for FREE until it does.\n\n"
    "We call it the Stock Certainty Guarantee™.\n\n"
    "We don't want your money if your stock doesn't match.")
add_direction("VISUAL: The guarantee in a green bordered box. Bold. Prominent. This is a HUGE moment.")

add_slide(57,
    "Want a completely risk-free way to find out how much your store is losing?\n\n"
    "Let me show you how…")

add_slide(58,
    "For a limited time, you can claim a FREE, no-obligation "
    "30-minute Store Check-Up with me personally.\n\n"
    "We'll look at your stock together, find your #1 leak source, "
    "and I'll tell you exactly how much it's costing you — every month.")
add_direction("VISUAL: 'FREE 30-Minute Store Check-Up' in large green text. Calendar icon. Button preview below.")

add_slide(59,
    "On this call, we'll also cover:\n\n"
    "• How you can reduce stock mismatch to under 2% in just 30 days.\n"
    "• The #1 thing you should NEVER do when trying to fix inventory problems.\n"
    "• And the 3 common mistakes that clothing store owners make "
    "that silently drain their profits.")

add_slide(60,
    "Remember — there is absolutely no cost and no obligation to move forward afterwards. "
    "If you feel it's not for you, that's completely fine.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PHASE 10: CTA + URGENCY CLOSE
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 10: CTA + URGENCY CLOSE", "🔥")
add_body("Goal: Create urgency. Drive the click. Close the VSL with conviction.")

add_slide(61,
    "But a word of warning…")

add_slide(62,
    "This is not some quick-fix magic solution.\n\n"
    "This is for legitimate, hard-working retail store owners "
    "who are ready to invest in fixing their operations "
    "and are willing to follow a proven system to make it happen.")

add_slide(63,
    "You must also be willing to commit to the process. "
    "I have a strict 'Single System Rule' — if you work with me, "
    "we eliminate all parallel diaries and manual shortcuts.\n\n"
    "If you can agree to that…")

add_slide(64,
    "Then this is truly the opportunity of a lifetime.\n\n"
    "All you need to do is click the button below and book your free call.")

add_slide(65,
    "Now this call is NOT a sales call.\n\n"
    "It's simply a free diagnostic session designed to find your biggest leak "
    "and see if we're a good fit.")

add_slide(66,
    "If we are — I'll share all the details, including the investment to get started.\n\n"
    "If not — I'll politely let you know. No pressure. No awkwardness.")

add_slide(67,
    "But please note — this video is being seen by hundreds of store owners.\n\n"
    "More people will apply for this than I have slots for.\n\n"
    "I take on only 3 new stores per month — first come, first served.")

add_slide(68,
    "So don't wait.\n\n"
    "Click the button below right NOW and book the next available slot.")
add_direction("VISUAL: Big green CTA button: 'BOOK MY FREE STORE CHECK-UP →'. Pulsing animation if possible.")

add_slide(69,
    "Lastly, I'd just like to say…\n\n"
    "You've got two options here.\n\n"
    "Option 1: Forget everything I just told you. "
    "Go back to doing what you've always done. "
    "And keep losing ₹40,000 or more every month without knowing where it goes.\n\n"
    "Let's face it — if you were happy with your stock situation, "
    "you wouldn't be watching this right now.")

add_slide(70,
    "Option 2: Accept that to fix this problem in the fastest way possible, "
    "it makes complete logical sense to get a proven roadmap "
    "from someone who has already solved it for stores just like yours.\n\n"
    "And avoid the frustration, stress, and financial loss "
    "of trying to figure it out alone.\n\n"
    "Click the button below this video right now.\n"
    "Book your free 30-minute Store Check-Up.\n\n"
    "And I'll speak to you soon.")
add_direction("VISUAL: Final slide. Green CTA button. Your headshot. 'Book Your Free Store Check-Up' in large text. Music fades.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# RECORDING GUIDE
# ════════════════════════════════════════════════════════════════════
add_section_header("RECORDING GUIDE — HOW TO PRODUCE THIS VSL", "🎬")

add_sub_header("Option A: PowerPoint/Keynote + Voiceover (Fastest)")
add_numbered("Create a slide deck with one slide per script section", 1)
add_numbered("Record your voiceover reading the script naturally (not robotic)", 2)
add_numbered("Use screen recording (OBS or Loom) to capture slides + voice", 3)
add_numbered("Add background music (subtle, professional — try Epidemic Sound)", 4)
add_numbered("Export as MP4. Upload to YouTube (unlisted) or Vimeo.", 5)

add_sub_header("Option B: Face-Cam + Slides (Better)")
add_numbered("Record yourself on camera for the intro and story sections", 1)
add_numbered("Cut to slides/screenshots for the proof, stats, and system sections", 2)
add_numbered("Use CapCut, Premiere, or DaVinci Resolve to edit", 3)
add_numbered("Mix face-cam and slides to keep it engaging", 4)

add_sub_header("Option C: Full Face-Cam (Best Conversion)")
add_numbered("Memorize key sections (or use a teleprompter app)", 1)
add_numbered("Record full face-cam with b-roll of store footage", 2)
add_numbered("Edit with text overlays for key stats and proof", 3)

add_sub_header("Technical Specs")
styled_table(
    ["Setting", "Recommendation"],
    [
        ["Video Length", "12-16 minutes (15 is the sweet spot)"],
        ["Resolution", "1080p minimum"],
        ["Lighting", "Ring light or window light — clean, professional"],
        ["Audio", "Use a lapel mic or USB condenser mic — audio is MORE important than video"],
        ["Background", "Clean. Solid color or blurred. No clutter."],
        ["Hosting", "YouTube (unlisted) or Vimeo — embed on landing page"],
        ["Captions", "Add subtitles — many watch on mute initially"],
    ],
    col_widths=[1.5, 4.5]
)

add_highlight_box(
    "CRITICAL RECORDING TIPS:\n\n"
    "1. Practice the script 3-5 times before recording.\n"
    "2. Speak naturally — like you're talking to ONE person, not a camera.\n"
    "3. Vary your tone — emphasize key phrases, pause before big reveals.\n"
    "4. The first 30 seconds determine if they keep watching. NAIL the pattern interrupt.\n"
    "5. Record in chunks — do Phase 1-3, then 4-6, then 7-10. Edit together.\n"
    "6. Audio quality > video quality. Always.",
    "FFF3E0"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NOW GO RECORD YOUR VSL. EVERY WORD IS WRITTEN. JUST HIT RECORD. 🎬🏆")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = ACCENT_GOLD


# ── Save ───────────────────────────────────────────────────────────
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "RCA_Complete_VSL_Script.docx"
)
doc.save(output_path)
print(f"✅ Successfully created: {output_path}")
