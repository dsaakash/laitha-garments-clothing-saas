#!/usr/bin/env python3
"""
Week 4: GAP Framework - 3 Months of Content Ready Using AI
IEC Diamond Week 4 - Marketing & Lead Generation
Beautiful DOCX guide for Aakash Savant / Retail Control Architect
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Colors
DARK_BG = RGBColor(0x0F, 0x0F, 0x1A)
ORANGE = RGBColor(0xFF, 0x6B, 0x35)
GOLD = RGBColor(0xD4, 0xA5, 0x00)
GREEN = RGBColor(0x00, 0xB3, 0x5F)
BLUE = RGBColor(0x3D, 0x8B, 0xE8)
RED = RGBColor(0xE8, 0x3D, 0x3D)
PURPLE = RGBColor(0x6C, 0x5C, 0xE7)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
BLACK = RGBColor(0x1A, 0x1A, 0x2E)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_para(doc, text, size=12, color=BLACK, bold=False, italic=False,
                    align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    r.font.name = 'Calibri'
    return p


def add_section_header(doc, text, color=ORANGE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"━━━  {text}")
    r.font.size = Pt(20)
    r.font.color.rgb = color
    r.bold = True
    r.font.name = 'Calibri'
    return p


def add_prompt_block(doc, step_num, title, prompt_text, notes=None, fill_fields=None):
    """Add a formatted AI prompt block."""
    # Step header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"🤖  AI PROMPT — STEP {step_num}: {title.upper()}")
    r.font.size = Pt(16)
    r.font.color.rgb = PURPLE
    r.bold = True
    r.font.name = 'Calibri'

    if fill_fields:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(f"⚠️  FILL IN YOUR INFO: {fill_fields}")
        r2.font.size = Pt(11)
        r2.font.color.rgb = RED
        r2.bold = True
        r2.italic = True
        r2.font.name = 'Calibri'

    # Prompt box as table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, 'F5F0FF')
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("📋 COPY THIS PROMPT INTO ChatGPT:")
    r.font.size = Pt(10)
    r.font.color.rgb = PURPLE
    r.bold = True
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)

    p2 = cell.add_paragraph()
    r2 = p2.add_run(prompt_text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = BLACK
    r2.font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(4)

    if notes:
        p3 = cell.add_paragraph()
        r3 = p3.add_run(f"💡 NOTE: {notes}")
        r3.font.size = Pt(10)
        r3.font.color.rgb = GREEN
        r3.italic = True
        r3.font.name = 'Calibri'

    doc.add_paragraph()  # spacer


def add_checklist_item(doc, text, is_done=False):
    icon = "✅" if is_done else "⬜"
    p = doc.add_paragraph(f"  {icon}  {text}")
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(12)
        r.font.name = 'Calibri'
    return p


def create_gap_framework_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    # ═══ COVER ═══
    for _ in range(4):
        doc.add_paragraph()

    add_styled_para(doc, "WEEK 4", size=16, color=ORANGE, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_para(doc, "GAP FRAMEWORK", size=36, color=BLACK, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_para(doc, "3 Months of Content Ready Using AI", size=20, color=ORANGE, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_styled_para(doc, "IEC DIAMOND — MARKETING & LEAD GENERATION", size=14, color=GRAY,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # Divider
    add_styled_para(doc, "━" * 40, size=12, color=ORANGE,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_styled_para(doc, "Content Ideas Through AI For Lead Generation", size=16,
                    color=BLACK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_para(doc, "A Step-by-Step Playbook for Aakash Savant", size=14,
                    color=GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_para(doc, "Retail Control Architect™", size=14,
                    color=ORANGE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    add_styled_para(doc, "How to Use This Guide:", size=14, color=BLACK, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    instructions = [
        "1️⃣  Read each section carefully",
        "2️⃣  Fill in YOUR information where marked with ⚠️",
        "3️⃣  Copy each prompt into ChatGPT one by one",
        "4️⃣  Film yourself answering the questions ChatGPT gives you",
        "5️⃣  Post the videos as Reels / Shorts / Stories",
    ]
    for inst in instructions:
        add_styled_para(doc, inst, size=12, color=GRAY,
                        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    doc.add_page_break()

    # ═══ SECTION 1: WHAT IS THE GAP FRAMEWORK? ═══
    add_section_header(doc, "SECTION 1: WHAT IS THE GAP FRAMEWORK?", ORANGE)

    add_styled_para(doc, "A personal brand that grows on social media has 3 core elements.",
                    size=14, color=BLACK, bold=True, space_after=10)

    # GAP explained
    gap_data = [
        ("🚀  G — GROWTH CONTENT", GREEN,
         "Talk about popular trends, figures, and viral topics in your industry.",
         "This drives traffic to your profile by using what people are already searching for.",
         [
            "How to use ChatGPT to manage your retail store",
            "Breakdown of how DMart or Reliance Retail keeps stock perfect",
            "React to a viral video about a store owner losing money",
            "Trending news in retail + your take on it",
         ]),
        ("🎓  A — AUTHORITY CONTENT", BLUE,
         "Show your skills and knowledge. Prove you know what you are talking about.",
         "This makes people follow you because you are helping them solve their problems.",
         [
            "How to do a stock count in 10 minutes",
            "3 mistakes every store owner makes with inventory",
            "Case study: How a store recovered Rs 3.2 Lakhs",
            "Tips, lessons, and how-to content",
         ]),
        ("❤️  P — PERSONAL CONTENT", PURPLE,
         "Share your stories, opinions, and what makes you unique as a person.",
         "This builds deep connection. People don't just follow you — they LOVE you.",
         [
            "Why I started helping store owners",
            "My biggest failure and what I learned",
            "A day in my life visiting stores",
            "My opinion on why most retail software fails",
         ]),
    ]

    for title, color, desc, why, examples in gap_data:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(title)
        r.font.size = Pt(16)
        r.font.color.rgb = color
        r.bold = True
        r.font.name = 'Calibri'

        add_styled_para(doc, desc, size=12, color=BLACK, space_after=2)
        add_styled_para(doc, f"Why it works: {why}", size=11, color=GRAY, italic=True, space_after=4)

        add_styled_para(doc, "Examples for your brand:", size=11, color=color, bold=True, space_after=2)
        for ex in examples:
            add_styled_para(doc, f"    •  {ex}", size=11, color=GRAY, space_after=2)

    # Visual diagram
    add_styled_para(doc, "", size=6)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, hex_color, emoji) in enumerate([
        ("GROWTH\nDrives Traffic", "E8F5E9", "🚀"),
        ("AUTHORITY\nBuilds Trust", "E3F2FD", "🎓"),
        ("PERSONAL\nCreates Fans", "F3E5F5", "❤️"),
    ]):
        cell = tbl.cell(0, i)
        set_cell_shading(cell, hex_color)
        cell.text = f"{emoji}\n{label}"
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.name = 'Calibri'
                r.bold = True

    add_styled_para(doc, "\nStranger → Follower → Fan → Customer", size=14,
                    color=ORANGE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_styled_para(doc, "When you create all 3 types, you bridge the GAP.", size=12,
                    color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ═══ SECTION 2: AI PROMPTS — STEP BY STEP ═══
    add_section_header(doc, "SECTION 2: AI PROMPTS — FOLLOW THESE STEPS", BLUE)

    add_styled_para(doc,
        "Follow these 5 prompts in order. Copy each one into ChatGPT. "
        "Wait for ChatGPT to reply before moving to the next prompt.",
        size=12, color=GRAY, italic=True, space_after=12)

    # ─── PROMPT 1 ───
    prompt1 = (
        "Please read this and once you understand, reply with the word yes:\n\n"
        "A personal brand that grows on social media has 3 core elements I call the GAP Framework.\n\n"
        "The GAP Framework consists of:\n"
        "Growth Content\n"
        "Authority Content\n"
        "Personal Content\n\n"
        "When you create these 3 types of content, it bridges the GAP from a stranger to die-hard fan.\n\n"
        "It's broken down like this:\n\n"
        "Growth content is when you talk about popular trends and figures. "
        "This leverages people's interest and drives traffic to your profile.\n\n"
        "For example, if I was a fitness coach when ChatGPT was trending, I could create content on "
        "how to use ChatGPT to create a nutrition plan. Or I could break down Arnold Schwarzenegger's "
        "workout routine.\n\n"
        "Since people in your niche have an interest in the growth content you're talking about, "
        "you're going to get more interest and reach.\n\n"
        "The A in the GAP framework is authority content.\n"
        "This is content that shows your competence at what your account is centered around.\n"
        "So things like:\n"
        "- Content that shows people how to solve their pain points or reach their goals\n"
        "- Lessons and tips\n"
        "- Case studies\n\n"
        "When you create this type of content, you incentivize people to follow you because you're "
        "actively helping the market by giving actionable information.\n\n"
        "The 3rd pillar of the GAP framework is personal content.\n"
        "Personal content is when you talk about the things that make you unique:\n"
        "- Your personal stories\n"
        "- Your personal worldviews\n"
        "- Your personal opinions on life or industry\n\n"
        "When you create personal content, you build a deep connection with your audience.\n"
        "And more importantly, you build an audience of true fans who adore you.\n\n"
        "Reply with \"Yes\" once you read and understand this."
    )
    add_prompt_block(doc, 1, "Teach ChatGPT the GAP Framework", prompt1,
                     notes="ChatGPT will reply 'Yes'. Then move to Step 2.")

    # ─── PROMPT 2 ───
    prompt2 = (
        "Here is some information about me:\n\n"
        "Authority:\n"
        "- I am the Founder of Retail Control Architect, a system that helps clothing & retail stores stop hidden stock losses\n"
        "- I have helped store owners recover Rs 3.2 Lakhs in hidden losses within 45 days\n"
        "- I reduced stock mismatch from 17% to 2% for a client store\n"
        "- I built a 5-step enforcement framework that any store can follow\n"
        "- I have worked with multiple retail stores facing inventory problems\n"
        "- I helped a store owner open their second branch confidently after fixing stock issues\n"
        "- Zero stock-shock guarantee — if mismatch does not reduce in 30 days, I work for free\n"
        "- I focus on fixing the PROCESS, not just installing software\n"
        "- I have deep knowledge of how Indian retail stores operate day-to-day\n\n"
        "Personal:\n"
        "- I believe retail store owners are the backbone of India's economy and deserve better systems\n"
        "- I once walked into a store where Rs 4 Lakhs had vanished and the owner thought it was 'normal'\n"
        "- I think expensive software is NOT the answer — simple processes are\n"
        "- I believe staff are not the enemy — lack of structure is\n"
        "- I am passionate about helping small business owners protect their hard-earned money\n"
        "- I believe every store owner deserves to sleep peacefully without worrying about stock\n"
        "- I come from a background of understanding ground-level retail operations\n"
        "- I think the biggest threat to retail is not online shopping — it is internal chaos\n"
        "- I want to build a movement where every Indian retail store runs with military-grade stock discipline\n\n"
        "Once you understand, reply with 'Yes'."
    )
    add_prompt_block(doc, 2, "Give ChatGPT Your Authority & Personal Info", prompt2,
                     notes="This is pre-filled with YOUR information. Edit/add more bullet points if needed.",
                     fill_fields="Review the bullet points above. Add more personal stories, achievements, or opinions if you have them.")

    # ─── PROMPT 3 ───
    prompt3 = (
        "My target audience is:\n\n"
        "- Gender: Mostly Male (some female store owners too)\n"
        "- Age: 28-55 years old\n"
        "- Occupation: Retail store owners, clothing shop owners, footwear/accessories store owners\n"
        "- Location: Tier 1, 2, and 3 cities in India\n"
        "- Goals: They want their store to run smoothly, stop losing money from stock mismatch, grow their business, maybe open a second branch\n"
        "- Pain Points: Stock never matches the computer, end-of-month surprises, staff not accountable, items go missing, stressed about profits, tried software but it did not help\n"
        "- Typical Day: Opens store at 10 AM, manages staff, handles customers, checks stock sometimes, deals with suppliers, closes at 9-10 PM, feels exhausted but unsure if the day was profitable\n"
        "- Current Situation: Running the store for 3-15+ years, has billing software but no real stock control process, loses Rs 2-5 Lakhs annually without knowing, blames staff but the real problem is lack of system\n"
        "- Language: Hindi-English mix, simple language, no technical jargon\n"
        "- Social Media: Uses WhatsApp daily, watches YouTube and Instagram Reels, follows business content in Hindi\n\n"
        "Once you understand, reply with \"Yes.\""
    )
    add_prompt_block(doc, 3, "Define Your Target Audience", prompt3,
                     notes="Pre-filled for your audience. Add more details about specific store owners you have met.",
                     fill_fields="Review the target audience. Add any specific details about real store owners you know.")

    # ─── PROMPT 4 ───
    prompt4 = (
        "My goal is to create short-form video content for social media that makes my target avatars know, like, and trust me.\n\n"
        "Using the GAP Framework, please ask me 25 questions for each category of the GAP Framework that does this.\n\n"
        "You can use the information I gave you about me to help, but feel free to ask questions outside that.\n\n"
        "The questions asked should set me up so I can answer them with a soundbite that's under 60 seconds.\n\n"
        "Make the questions in simple language that a store owner would understand. No English jargon. "
        "Think of questions a store owner would hear and immediately think 'I want to hear his answer to this.'"
    )
    add_prompt_block(doc, 4, "Get 75 Video Content Questions (25 per category)", prompt4,
                     notes="ChatGPT will give you 75 questions. Film yourself answering each one = 75 videos = 3 months of content!")

    add_styled_para(doc, "💡  POWER TIP: Want even MORE questions?", size=13, color=GREEN, bold=True, space_before=4)
    add_styled_para(doc, "Just tell ChatGPT:", size=12, color=GRAY)
    for extra in [
        '"Ask me 10 more personal questions"',
        '"Give me questions about common store owner mistakes"',
        '"Ask questions that will make store owners feel understood"',
        '"Give me controversial opinion questions about retail"',
    ]:
        add_styled_para(doc, f"    •  {extra}", size=11, color=PURPLE, italic=True, space_after=2)

    # ─── PROMPT 5 ───
    prompt5 = (
        "Write me a 60-second Instagram reel script that feels natural and conversational "
        "(not salesy or filled with buzzwords). The script must include:\n\n"
        "1. A strong hook that grabs attention in the first 3 seconds.\n"
        "2. Clear value that teaches or shares something useful.\n"
        "3. A short story or relatable example that makes it engaging.\n"
        "4. A simple and direct CTA at the end.\n\n"
        "This should be for content piece #1 in my Growth Content series. "
        "Use your years of copywriting experience to make it flow smoothly, "
        "like how a real person would speak.\n\n"
        "Write it in simple Hindi-English mix. No fancy English words. "
        "Think of how an Indian store owner talks to his friend."
    )
    add_prompt_block(doc, 5, "Generate Your First Reel Script", prompt5,
                     notes="Change '#1 in Growth Content' to '#1 in Authority Content' or '#1 in Personal Content' for other categories. Repeat for all 75 questions!")

    doc.add_page_break()

    # ═══ SECTION 3: CONTENT CALENDAR ═══
    add_section_header(doc, "SECTION 3: YOUR 3-MONTH CONTENT CALENDAR", GREEN)

    add_styled_para(doc, "Post 1 video per day. Rotate between the 3 content types:", size=13, color=BLACK, bold=True)

    # Weekly schedule table
    tbl = doc.add_table(rows=8, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["DAY", "CONTENT TYPE", "WHAT TO POST"]
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        set_cell_shading(cell, '1A1A2E')
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.size = Pt(11)
                r.bold = True
                r.font.name = 'Calibri'

    schedule = [
        ("Monday", "🚀 Growth", "Trending topic + your angle"),
        ("Tuesday", "🎓 Authority", "Tip / How-to / Lesson"),
        ("Wednesday", "❤️ Personal", "Your story / opinion"),
        ("Thursday", "🎓 Authority", "Case study / results"),
        ("Friday", "🚀 Growth", "React to industry news"),
        ("Saturday", "❤️ Personal", "Behind the scenes / day in life"),
        ("Sunday", "🎓 Authority", "Mistakes to avoid / myth-busting"),
    ]

    colors = {"🚀": "E8F5E9", "🎓": "E3F2FD", "❤️": "F3E5F5"}

    for i, (day, ctype, what) in enumerate(schedule):
        emoji = ctype.split()[0]
        color_hex = colors.get(emoji, "FFFFFF")
        for j, val in enumerate([day, ctype, what]):
            cell = tbl.cell(i + 1, j)
            set_cell_shading(cell, color_hex)
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Calibri'

    add_styled_para(doc, "\n📊  25 Growth + 25 Authority + 25 Personal = 75 Videos = ~11 Weeks of Daily Content!",
                    size=13, color=ORANGE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ═══ SECTION 4: CONTENT IDEAS BANK ═══
    add_section_header(doc, "SECTION 4: 30 READY-MADE CONTENT IDEAS", PURPLE)

    ideas = {
        "🚀 GROWTH CONTENT (10 Ideas)": (GREEN, [
            "How to use ChatGPT to manage your store stock",
            "What DMart does differently that small stores can copy",
            "React: Store owner loses Rs 10 Lakhs — here is why",
            "5 things I learned from visiting 50+ retail stores",
            "The future of retail stores in India (2025 and beyond)",
            "Why Reliance Retail never has stock mismatch",
            "How WhatsApp can help you manage your store better",
            "The Rs 4 Lakh mistake every clothing store makes",
            "What Amazon warehouses can teach small store owners",
            "Is retail dying? My honest opinion",
        ]),
        "🎓 AUTHORITY CONTENT (10 Ideas)": (BLUE, [
            "3 signs your store is silently losing money",
            "How to do a stock count in under 10 minutes",
            "The ONLY 3 rules your staff needs to follow",
            "Why billing software will NEVER fix your stock problem",
            "How one store owner recovered Rs 3.2 Lakhs in 45 days",
            "5 daily habits of profitable store owners",
            "The leaking bucket: why your profits disappear",
            "How to hold your staff accountable without fighting",
            "What to check every Saturday before closing your store",
            "My 5-step system that fixes stock mismatch permanently",
        ]),
        "❤️ PERSONAL CONTENT (10 Ideas)": (PURPLE, [
            "Why I started Retail Control Architect (my story)",
            "The day I saw Rs 4 Lakhs vanish from a store",
            "My biggest failure and what it taught me",
            "What my father taught me about business",
            "A day in my life helping store owners",
            "The store owner who made me cry (emotional story)",
            "Why I refuse to sell expensive software",
            "My opinion: Staff are NOT your enemy",
            "What keeps me going when things get tough",
            "The moment I knew this was my calling",
        ]),
    }

    for section_title, (color, idea_list) in ideas.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        r = p.add_run(section_title)
        r.font.size = Pt(14)
        r.font.color.rgb = color
        r.bold = True
        r.font.name = 'Calibri'

        for idx, idea in enumerate(idea_list, 1):
            add_styled_para(doc, f"    {idx}.  {idea}", size=11, color=GRAY, space_after=3)

    doc.add_page_break()

    # ═══ SECTION 5: FILMING CHECKLIST ═══
    add_section_header(doc, "SECTION 5: FILMING & POSTING CHECKLIST", GOLD)

    add_styled_para(doc, "Before You Film:", size=14, color=BLACK, bold=True, space_after=4)
    for item in [
        "Read the question from ChatGPT's list",
        "Think about your answer for 30 seconds",
        "Set up your phone (vertical, good lighting, clean background)",
        "Look into the camera lens, not the screen",
        "Keep your answer under 60 seconds",
    ]:
        add_checklist_item(doc, item)

    add_styled_para(doc, "\nWhile Filming:", size=14, color=BLACK, bold=True, space_after=4)
    for item in [
        "Start with a HOOK — the first 3 seconds must grab attention",
        "Speak naturally — like you are talking to a store owner friend",
        "Use simple Hindi-English mix — no English jargon",
        "Share ONE clear point per video — do not try to say everything",
        "End with a CTA — 'Follow for more' or 'Comment if this happens to you'",
    ]:
        add_checklist_item(doc, item)

    add_styled_para(doc, "\nAfter Filming:", size=14, color=BLACK, bold=True, space_after=4)
    for item in [
        "Edit using CapCut or InShot (add captions!)",
        "Write a caption with 3-5 relevant hashtags",
        "Post on Instagram Reels + YouTube Shorts + WhatsApp Status",
        "Reply to EVERY comment within 1 hour",
        "Save the video to your phone for WhatsApp sharing",
    ]:
        add_checklist_item(doc, item)

    doc.add_page_break()

    # ═══ SECTION 6: HASHTAG STRATEGY ═══
    add_section_header(doc, "SECTION 6: HASHTAGS & CAPTION TEMPLATES", BLUE)

    add_styled_para(doc, "Top Hashtags For Your Content:", size=14, color=BLACK, bold=True, space_after=4)
    hashtags = [
        "#RetailBusiness  #StoreOwner  #InventoryManagement",
        "#ClothingStore  #RetailTips  #BusinessGrowth",
        "#StockManagement  #SmallBusiness  #IndianRetail",
        "#ShopOwner  #RetailLife  #BusinessOwner",
        "#AakashSavant  #RetailControlArchitect",
    ]
    for h in hashtags:
        add_styled_para(doc, f"    {h}", size=11, color=BLUE, space_after=2)

    add_styled_para(doc, "\nCaption Templates:", size=14, color=BLACK, bold=True, space_after=4, space_before=8)

    captions = [
        ("Growth Post", "Did you know [trending fact]? Here is what it means for store owners like you... 👇\n#RetailTips #BusinessGrowth"),
        ("Authority Post", "Stop doing [common mistake]. Do this instead 👇\nI have seen this save store owners lakhs.\n#StoreOwner #RetailBusiness"),
        ("Personal Post", "[Personal story in one line].\nThis changed how I think about [topic].\nFull story in the video 👆\n#MyStory #RetailLife"),
    ]

    for label, template in captions:
        p = doc.add_paragraph()
        r = p.add_run(f"{label}:")
        r.font.size = Pt(12)
        r.font.color.rgb = ORANGE
        r.bold = True
        r.font.name = 'Calibri'
        add_styled_para(doc, template, size=11, color=GRAY, italic=True, space_after=8)

    doc.add_page_break()

    # ═══ SECTION 7: WEEKLY ACTION PLAN ═══
    add_section_header(doc, "SECTION 7: YOUR WEEK 4 ACTION PLAN", RED)

    days = [
        ("DAY 1 (Monday)", RED, [
            "Complete AI Prompt Steps 1-3 in ChatGPT",
            "Review and edit your Authority & Personal bullet points",
            "Save the ChatGPT conversation — you will need it all week",
        ]),
        ("DAY 2 (Tuesday)", ORANGE, [
            "Complete AI Prompt Step 4 — get your 75 questions",
            "Organize questions into 3 lists: Growth / Authority / Personal",
            "Pick your first 7 questions (one for each day this week)",
        ]),
        ("DAY 3 (Wednesday)", GOLD, [
            "Complete AI Prompt Step 5 — get your first reel script",
            "Practice reading the script out loud 3 times",
            "Film your first video today! Just press record and talk.",
        ]),
        ("DAY 4 (Thursday)", GREEN, [
            "Film 2 more videos (1 Authority + 1 Personal)",
            "Edit all 3 videos with captions in CapCut",
            "Post your first video on Instagram + YouTube Shorts",
        ]),
        ("DAY 5 (Friday)", BLUE, [
            "Film 2 more videos (1 Growth + 1 Authority)",
            "Post 1 video + share on WhatsApp Status",
            "Reply to all comments and DMs",
        ]),
        ("DAY 6 (Saturday)", PURPLE, [
            "Film remaining 2 videos for the week",
            "Schedule posts for next week using Meta Business Suite",
            "Review analytics: which video got the most views?",
        ]),
        ("DAY 7 (Sunday)", GRAY, [
            "Rest or batch-film next week's content",
            "Plan next week's 7 topics from your question list",
            "Celebrate — you created 7 pieces of content this week!",
        ]),
    ]

    for day_title, color, tasks in days:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(f"📅  {day_title}")
        r.font.size = Pt(14)
        r.font.color.rgb = color
        r.bold = True
        r.font.name = 'Calibri'
        for task in tasks:
            add_checklist_item(doc, task)

    # ═══ FINAL PAGE ═══
    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()

    add_styled_para(doc, "YOU NOW HAVE", size=18, color=GRAY,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_para(doc, "3 MONTHS OF CONTENT", size=32, color=ORANGE, bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_para(doc, "Ready To Film & Post", size=18, color=BLACK,
                    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_styled_para(doc, "75 Questions  •  30 Content Ideas  •  7-Day Action Plan",
                    size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_para(doc, "Hashtags  •  Caption Templates  •  Reel Scripts",
                    size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_styled_para(doc, "━" * 30, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_styled_para(doc, "Now stop reading and start filming. 🎬",
                    size=16, color=ORANGE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_styled_para(doc, "— Aakash Savant, Retail Control Architect™",
                    size=12, color=GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    save_path = "/Users/aakash/Desktop/Week_3/Week 4/Week4_GAP_Framework_Content_Playbook.docx"
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    return save_path


if __name__ == "__main__":
    path = create_gap_framework_doc()
    print(f"✅ Created: {path}")
