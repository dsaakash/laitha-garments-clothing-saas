#!/usr/bin/env python3
"""
Generate the Retail Control Architect High Ticket Closing (HTC) Script.
Adapted from Sid's 8-Figure Script.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ── Colors ────────────────────────────────────────────────────────
DEEP_BLUE   = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL = RGBColor(0x17, 0xA2, 0xB8)
WARM_RED    = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS_GRN = RGBColor(0x27, 0xAE, 0x60)
BODY        = RGBColor(0x2C, 0x2C, 0x2C)
DARK_GRAY   = RGBColor(0x55, 0x55, 0x55)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11); style.font.color.rgb = BODY

# ── Helpers ───────────────────────────────────────────────────────
def spacer(pts=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)

def cover_title(txt, sz=26):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(4); p.space_after = Pt(4)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = DEEP_BLUE

def cover_sub(txt, sz=13, c=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    r = p.add_run(txt); r.italic = True; r.font.size = Pt(sz); r.font.color.rgb = c or ACCENT_GOLD

def sec_hdr(txt, emoji=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.space_before = Pt(16); p.space_after = Pt(3)
    full = f"{emoji}  {txt}" if emoji else txt
    r = p.add_run(full); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = DEEP_BLUE
    ln = doc.add_paragraph(); ln.space_after = Pt(6)
    rl = ln.add_run("━" * 75); rl.font.size = Pt(5); rl.font.color.rgb = ACCENT_GOLD

def sub_hdr(txt, c=None):
    p = doc.add_paragraph()
    p.space_before = Pt(12); p.space_after = Pt(3)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = c or ACCENT_TEAL

def body(txt, bold_prefix=None, c=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = BODY
        txt = " " + txt
    rx = p.add_run(txt)
    rx.font.size = Pt(11)
    rx.font.color.rgb = c or BODY
    return p

def callout(txt, bg_color="FFF8E1"):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    c.text = ''
    p = c.paragraphs[0]
    r = p.add_run(txt)
    r.font.size = Pt(11)
    r.font.color.rgb = BODY
    r.bold = True
    spacer(4)


# ════════════════════════════════════════════════════════════════
#                             TITLE
# ════════════════════════════════════════════════════════════════

doc.add_paragraph()
cover_title("8-FIGURE HIGH TICKET CLOSING SCRIPT", 28)
cover_sub("Curated for Aakash Savant — Retail Control Architect™", 14, ACCENT_TEAL)
spacer(4)
callout(
    "🟥 PROCESS OVERVIEW YOU MUST FOLLOW:\n\n"
    "➡️ 1. Opening Phase + Compliment\n"
    "➡️ 2. Rapport Building + Agenda\n"
    "➡️ 3. Probing (CS = Current Situation)\n"
    "➡️ 4. Probing (DS = Desired Situation)\n"
    "➡️ 5. Gap & Pitch\n"
    "➡️ 6. Explain The Grand Slam Offer (Urgency & Scarcity)\n"
    "➡️ 7. Closing",
    "FCE4EC"
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
#                             STEP 1
# ════════════════════════════════════════════════════════════════
sec_hdr("➡️ Step 1: Opening Phase + Compliment (Break the Ice)")

body("Hi, this is Aakash (or your name from Team Aakash). Am I speaking to [Mr/Ms Prospect]?")
body("You recently booked a slot for our specific 'Stock Leakage/Forensic Audit' consultation. Out of the dozens of applications we received this week, we specifically selected only three dedicated retail owners, and you are one of them.")
body("Is this a good time to talk?")

sub_hdr("Compliment to Establish Authority:")
body("By the way, before we start, I noticed you've been running your store for [X] years. Building a retail brand that survives and grows in today's market is no easy task. You must be dealing with immense footfall and inventory scale.")
body(
    "Target Triggers: Age of Business, Resilience, Scale of Operations, Energy.",
    bold_prefix="Note:",
    c=DARK_GRAY
)


# ════════════════════════════════════════════════════════════════
#                             STEP 2
# ════════════════════════════════════════════════════════════════
sec_hdr("➡️ Step 2: Rapport Building + Agenda")

body("Try to find something common and smile: location, local market struggles, etc.")

sub_hdr("Setting the Agenda:")
body("So, over these calls, we use a specific Forensic Framework. First, I want to understand your exact Current Situation (CS) regarding stock control and hidden leakage.")
body("Then, we'll map out your Desired Situation (DS). If I see that our Authority System can actually plug your leak, we'll move forward. Does that align with what you're looking for?")


# ════════════════════════════════════════════════════════════════
#                             STEP 3
# ════════════════════════════════════════════════════════════════
sec_hdr("➡️ Step 3: Probing (CS = Current Situation)")
sub_hdr("Target: Expose Their Chaos Tax - The Mismatch")

qs_cs = [
    "What's your current situation? How big is your store and how many SKUs are we talking about?",
    "Why exactly did you book this Stock Leakage Audit call? What's the biggest pain you're facing?",
    "Does your physical stock accurately match your system stock right now?",
    "How much money or inventory are you losing every month? Are you currently paying the 'chaos tax'?",
    "Do your staff use a parallel 'Kachha' register or diary because they don't enter things on the software instantly?",
    "How long has this mismatch been going on?",
    "Why haven't you solved it yet? Have you bought software before that didn't fix the behavior?",
    "Why is it absolutely critical to plug this leak right now?",
    "How much time do you spend every night trying to reconcile missing stock? 1 hour? 2 hours?",
    "What have you tried to solve this so far?"
]

for i, q in enumerate(qs_cs, 1):
    body(q, bold_prefix=f"{i}.")


# ════════════════════════════════════════════════════════════════
#                             STEP 4
# ════════════════════════════════════════════════════════════════
sec_hdr("➡️ Step 4: Probing (DS = Desired Situation)")
sub_hdr("Target: 100% Accuracy & Scale")

qs_ds = [
    "If everything was perfect, what does your desired store operation look like?",
    "If system stock = physical stock 100% of the time, how would that change your peace of mind?",
    "Why do you want this exact control? (Is it to scale to Store #2? Or to go home to your family on time?)",
    "Why do you think you haven't been able to achieve this total control already?",
    "If you let this leak continue, what happens in the next 12 months?",
    "What do you think is stopping your staff from following the digital system correctly?"
]

for i, q in enumerate(qs_ds, 1):
    body(q, bold_prefix=f"{i}.")
body("(Keep digging deep: 'What else? What else? Why is that important to you?')", c=DARK_GRAY)


# ════════════════════════════════════════════════════════════════
#                             STEP 5
# ════════════════════════════════════════════════════════════════
sec_hdr("➡️ Step 5: Gap & Pitch")

body("So, [Prospect Name], I can clearly see your Current Situation (chaos, hidden leakage, parallel registers) and your Desired Situation (100% accuracy, the 10-minute close, opening Store #2).")
body("But right now, you have a massive gap. And relying on 'buying more software' or 'telling staff to be careful' isn't closing that gap. Do you agree?")

sub_hdr("Consult & Destroy:")
body("The reason you aren't getting results is because software doesn't create discipline. A lack of structure creates chaos. You have a 'Double Reality' running in your store.")

sub_hdr("Transition:")
body("What if I told you we don't just sell software... but we physically install the Retail Control Architecture™ in your store?")
body("Do you feel I am the best person to help you plug this leak once and for all? Where should we go from here?")
body("(STOP. Let them speak. Let them ask you 'How do you do it?').", c=SUCCESS_GRN)


# ════════════════════════════════════════════════════════════════
#                             STEP 6
# ════════════════════════════════════════════════════════════════
sec_hdr("➡️ Step 6: The Grand Slam Offer (Urgency & Scarcity)")

body("Let me explain exactly how we enforce this:")
body("We implement The 30-Day Stock Certainty System™.")

sub_hdr("The Offer Breakdown:")
body("1. Control Gap Audit (We measure your specific baseline leakage).")
body("2. Digital Gatekeeper & Sales Lock (No stock enters or leaves without strict system verification).")
body("3. 30-Day Discipline Installation (We remove parallel registers and enforce compliance until it becomes habit).")

sub_hdr("Adding Urgency & Scarcity:")
body("Because this requires forensic accuracy and my personal team's energy, we ONLY onboard 3 clothing stores per month.")
body("If we don't fix this today, your 'chaos tax' of ₹2,739 per day continues tomorrow.")

sub_hdr("The Guarantee:")
body("And because we know our system works, we offer the Stock Certainty Guarantee™: If after 30 days of compliance, your stock mismatch doesn't reduce measurably, we will work with you for FREE until it does.")


# ════════════════════════════════════════════════════════════════
#                             STEP 7
# ════════════════════════════════════════════════════════════════
sec_hdr("➡️ Step 7: Closing")

body("Listen, I talk to 3 types of store owners: Yes, Maybe, and No.")
body("The 'Yes' fixes their mismatch and scales. The 'No' stays the same. The 'Maybe' ends up going bankrupt from silent leakage over 5 years. Which one are you?")

sub_hdr("Objection Handling Check:")
body("1. Do you believe this structure will fix your mismatch?", bold_prefix="Q:")
body("2. How confident are you on a scale of 1-10?", bold_prefix="Q:")
body("3. What's stopping you from committing right now?", bold_prefix="Q:")

body("Remember the 3 Selling Truths:", c=DARK_GRAY)
body("• Confidence = Self", c=DARK_GRAY)
body("• Conviction = Product", c=DARK_GRAY)
body("• Clarity = Benefit of the Product", c=DARK_GRAY)

sub_hdr("The Final Ask:")
body("Every day you wait, you lose ₹3,000. For a one-time ₹85k installation, this pays for itself in just over 30 days. After that, everything recovered is pure profit.")
body("Do you have any final questions before we process your enrollment?")
body("How would you like to handle the investment today... Card, GPay, or Net Banking?")
body("BOOM. CLOSE THE SALE.", c=SUCCESS_GRN, bold_prefix="")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
#                       SECRETS & POINTERS
# ════════════════════════════════════════════════════════════════
sec_hdr("🔥 3 SECRETS OF SELLING OVER THE FIRST CALL")

body("When they explain their gap (lost stock, confused staff), tell them the situation is actually worse than they think. Extrapolate their daily loss to an annual 'Chaos Tax' to make the pain unignorable.", bold_prefix="1. PULL Current Situation:")
body("When they tell you they want to 'just organize things', push them further. Tell them why stop there? Why not achieve 100% precision and open Store #2?", bold_prefix="2. PUSH Desired Situation:")
body("If they think 'just telling staff to use the software better' is the solution, destroy that plan. Software without enforced structure equals chaos.", bold_prefix="3. DESTROY the Bridge:")


sec_hdr("💡 IMPORTANT CLOSING POINTERS")
pointers = [
    "Most store owners do the hard work AFTER crisis hits, not before. Make them fix it NOW.",
    "Maintain Straight Line Persuasion. Keep pivoting back to the Stock Mismatch pain.",
    "Your Tonality, Confidence, and Conviction are more important than feature explanations.",
    "NEVER talk bad about competitors (other software). Play with our positivity: 'They are just software; we are an Authority Installation.'",
    "People buy BENEFITS, not features. Sell them the transformation: The Peace of Mind, The 10-Minute Routine, The Freedom.",
    "Make them understand the Worst-Case Scenario: Massive internal theft or inventory rot.",
    "Money is useless until it's protecting profit. Money comes and goes, but lost inventory never returns.",
    "Ask them to use a pen and paper on the call. Do the ₹ leakage math with them live.",
    "Highlight Lalitha Garments Success Story—show them 98% accuracy is possible."
]

for p in pointers:
    body(p, bold_prefix="✔️", c=DARK_GRAY)

# ── Save ────────────────────────────────────────────────────────
out_file = "/Users/aakash/Desktop/Week_3/Aakash_Retail_Control_HTC_Script.docx"
doc.save(out_file)
print(f"✅ Success! Document created at: {out_file}")
