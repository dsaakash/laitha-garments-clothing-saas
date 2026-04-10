#!/usr/bin/env python3
"""
Generate Week 2 Training — Complete ChatGPT Prompt Structure & Grand Slam Offer Build
For Aakash Savant — Retail Control Architect
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colours ─────────────────────────────────────────────────────
DEEP_BLUE   = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL = RGBColor(0x17, 0xA2, 0xB8)
WARM_RED    = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS_GRN = RGBColor(0x27, 0xAE, 0x60)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
MED_GRAY    = RGBColor(0x66, 0x66, 0x66)
BODY        = RGBColor(0x2C, 0x2C, 0x2C)
ORANGE      = RGBColor(0xE6, 0x7E, 0x22)
PURPLE      = RGBColor(0x8E, 0x44, 0xAD)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.8); s.bottom_margin = Cm(1.8)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(10.5); style.font.color.rgb = BODY

# ── Helpers ─────────────────────────────────────────────────────
def shade(cell, hx):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hx}"/>'))

def spacer(pts=10):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(pts)

def cover_title(txt, sz=26):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_before = Pt(4); p.space_after = Pt(4)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = DEEP_BLUE

def cover_sub(txt, sz=13, c=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(6)
    r = p.add_run(txt); r.italic = True; r.font.size = Pt(sz); r.font.color.rgb = c or ACCENT_GOLD

def sec_hdr(txt, emoji=""):
    doc.add_paragraph()
    p = doc.add_paragraph(); p.space_before = Pt(16); p.space_after = Pt(3)
    r = p.add_run(f"{emoji}  {txt}" if emoji else txt); r.bold = True; r.font.size = Pt(17); r.font.color.rgb = DEEP_BLUE
    ln = doc.add_paragraph(); ln.space_after = Pt(6)
    rl = ln.add_run("━"*75); rl.font.size = Pt(5); rl.font.color.rgb = ACCENT_GOLD

def sub_hdr(txt, c=None):
    p = doc.add_paragraph(); p.space_before = Pt(12); p.space_after = Pt(3)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = c or ACCENT_TEAL

def minor_hdr(txt, c=None):
    p = doc.add_paragraph(); p.space_before = Pt(8); p.space_after = Pt(2)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = c or DEEP_BLUE

def body(txt):
    p = doc.add_paragraph(txt); p.paragraph_format.space_after = Pt(5); return p

def body_b(txt):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(10.5); return p

def quote(txt, c=None):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'"{txt}"'); r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = c or MED_GRAY

def prompt_box(txt):
    """Blue-tinted prompt box"""
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    c = t.rows[0].cells[0]; c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rb = p.add_run("💬 PROMPT:\n"); rb.bold = True; rb.font.size = Pt(9); rb.font.color.rgb = DEEP_BLUE
    r = p.add_run(txt); r.font.size = Pt(9.5); r.font.color.rgb = BODY
    shade(c, "EBF5FB"); c.width = Inches(6.0)
    spacer(6)

def output_box(txt, bg="F0FFF0"):
    """Green-tinted output box"""
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    c = t.rows[0].cells[0]; c.text = ''
    p = c.paragraphs[0]
    rb = p.add_run("📋 NICHE OUTPUT:\n"); rb.bold = True; rb.font.size = Pt(9); rb.font.color.rgb = SUCCESS_GRN
    r = p.add_run(txt); r.font.size = Pt(9.5); r.font.color.rgb = BODY
    shade(c, bg); c.width = Inches(6.0)
    spacer(6)

def highlight(txt, bg="FFF8E1"):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    c = t.rows[0].cells[0]; c.text = ''
    p = c.paragraphs[0]; r = p.add_run(txt); r.font.size = Pt(10); r.font.color.rgb = BODY
    shade(c, bg); c.width = Inches(6.0)
    spacer(6)

def warn_box(txt):
    highlight(txt, "FDEDED")

def bul(txt, bp=None, emoji=None):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(2); p.clear()
    if emoji: p.add_run(f"{emoji} ").font.size = Pt(10.5)
    if bp:
        rb = p.add_run(bp); rb.bold = True; rb.font.size = Pt(10.5)
        p.add_run(f" {txt}").font.size = Pt(10.5)
    else:
        p.add_run(txt).font.size = Pt(10.5)

def num(txt, n, bl=None):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{n}. "); r.bold = True; r.font.color.rgb = ACCENT_TEAL; r.font.size = Pt(10.5)
    if bl:
        rl = p.add_run(f"{bl}: "); rl.bold = True; rl.font.size = Pt(10.5)
        p.add_run(txt).font.size = Pt(10.5)
    else:
        p.add_run(txt).font.size = Pt(10.5)

def arrow(txt, bl=None):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("→ "); r.bold = True; r.font.color.rgb = ACCENT_GOLD; r.font.size = Pt(10.5)
    if bl:
        rl = p.add_run(f"{bl}: "); rl.bold = True; rl.font.size = Pt(10.5)
        p.add_run(txt).font.size = Pt(10.5)
    else:
        p.add_run(txt).font.size = Pt(10.5)

def stbl(headers, rows, cw=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE; shade(c, "1B2A4A")
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            r = p.add_run(str(v)); r.font.size = Pt(9)
            if ri % 2 == 0: shade(c, "F7F9FC")
    if cw:
        for i, w in enumerate(cw):
            for row in t.rows: row.cells[i].width = Inches(w)
    spacer(6)


# ════════════════════════════════════════════════════════════════
#                       COVER PAGE
# ════════════════════════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph()
cover_title("BUILD A GRAND SLAM OFFER", 28)
cover_title("COMPLETE PROMPT STRUCTURE", 22)
spacer(8)
cover_sub("Week 2 — ChatGPT Prompts & Niche Outputs", 14)
cover_sub("By Aakash Savant — Retail Control Architect", 12, MED_GRAY)
doc.add_paragraph()

highlight(
    "THE 6-STEP ROADMAP TO A PITCH-READY GRAND SLAM OFFER\n\n"
    "Step 1 → Set the Context (BizGPT Persona)\n"
    "Step 2 → Find Niche, Sub-Niche & Micro-Niche\n"
    "Step 3 → Build Unique Mechanism & Magnetic Positioning\n"
    "Step 4 → Validate the Market (4 Parameters)\n"
    "Step 5 → 7-Step Grand Slam Offer Framework\n"
    "Step 6 → Make it Pitch-Ready\n\n"
    "FORMAT: Each step shows the PROMPT (blue box) → then the NICHE OUTPUT (green box)",
    "F0F4FF"
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 1: SET THE CONTEXT
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 1: SET THE CONTEXT", "🔵")
body("Before anything, you set ChatGPT's persona. This ensures all outputs are business-grade, no fluff, action-oriented.")

prompt_box(
    "You are now BizGPT, my personal business coach and strategic growth advisor. "
    "You are an expert with over 20 years of experience in marketing, direct response copywriting, "
    "high-converting sales funnels, VSL funnels, grand slam offer creation, and business building. "
    "You have deeply studied and mastered the principles of Russell Brunson, Alex Hormozi, Sabri Suby, "
    "Dan Kennedy, and Sam Ovens.\n\n"
    "Your job is to give me clear, actionable strategies, roadmaps, scripts, and implementation plans "
    "to grow and scale my business fast. You will help me refine my offers, improve my funnels, boost "
    "conversions, and close more sales using proven frameworks.\n\n"
    "You are allowed to ask me smart, strategic questions to understand my vision, audience, niche, and goals. "
    "You always provide specific and practical next steps I can implement immediately.\n\n"
    "Only share what drives real results. No fluff. Keep it sharp, direct, and focused on generating "
    "more leads, clients, and revenue.\n\n"
    "Type (Yes) if you understood."
)
doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 2: FINDING NICHE
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 2: FINDING NICHE, SUB-NICHE & MICRO-NICHE", "🔵")

prompt_box(
    "I want help defining my business niche, sub-niche & micro-niche. Here's my current thinking:\n\n"
    "The problem I am solving (WHAT) = [Your answer]\n"
    "The people I am solving it for (WHO) = [Your answer]\n"
    "My solution (HOW) = [Your answer]\n\n"
    "Based on this, please suggest:\n"
    "1. The problems they are facing (all the problems)\n"
    "2. My niche\n"
    "3. Sub-niche\n"
    "4. Micro-niche"
)

sub_hdr("Niche Output for Retail Control Architect")

output_box(
    "WHAT: Clothing retailers lack structured operational control → stock mismatch, margin leakage, "
    "supplier confusion, staff dependency.\n\n"
    "WHO: Owner-operated, single-store clothing retailers doing ₹30L–₹1.5Cr annually, "
    "holding ₹20L+ inventory, struggling with stock accuracy.\n\n"
    "HOW: I install a 30-Day Retail Control Architecture™ system that structures inventory, "
    "supplier workflows, and sales integration to eliminate stock mismatch."
)

minor_hdr("All Problems They Face")
for cat, items in [
    ("Inventory Problems", [
        "10–20% stock mismatch", "No SKU discipline", "Duplicate entries",
        "Missing inward records", "Unrecorded shrinkage", "Dead inventory accumulation",
        "Slow-moving stock unnoticed", "No real inventory valuation clarity"
    ]),
    ("Financial Problems", [
        "Hidden margin leakage", "Supplier overbilling unnoticed", "No clear purchase tracking",
        "Inaccurate profit calculation", "Cash flow strain", "Capital stuck in unsold stock"
    ]),
    ("Operational Problems", [
        "1–2 hours daily reconciliation", "Parallel manual + digital systems",
        "No enforcement of process discipline", "Inconsistent data entry by staff"
    ]),
    ("Dependency Problems", [
        "One staff controls operational knowledge", "Owner cannot independently verify",
        "Business vulnerable if key employee leaves", "Fear of internal manipulation"
    ]),
    ("Growth Problems", [
        "Cannot confidently open second store", "Cannot delegate operations",
        "Scaling increases chaos", "Sales growth without operational stability"
    ]),
    ("Emotional Problems", [
        "Lack of trust in data", "Daily stress", "Constant checking",
        "Fear of hidden losses", "Feeling stuck despite revenue growth"
    ]),
]:
    body_b(cat)
    for item in items:
        bul(item)

minor_hdr("Niche Structure")
stbl(["Level", "Definition"], [
    ["Niche", "Retail Operational Control for Clothing Stores"],
    ["Sub-Niche", "Established, single-store clothing retailers\n₹30L–₹1.5Cr revenue, ₹20L+ inventory, Owner-operated"],
    ["Micro-Niche", "Tier 2 & Tier 3 Indian clothing retailers\n₹50L–₹1Cr revenue, ₹30L+ inventory\nUsing manual/basic billing, 10%+ mismatch\nOwner personally handling reconciliation"],
    ["NOT For", "Startups • E-commerce only • Multi-branch chains\nRetailers seeking cheap billing software"],
], cw=[1.2, 4.8])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 3: UNIQUE MECHANISM
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 3: UNIQUE MECHANISM & MAGNETIC POSITIONING", "🔵")

prompt_box(
    "Now you already have clarity on my niche, sub-niche, micro niche, my service/product, "
    "my audience & the problem I am solving.\n\n"
    "I don't want to sell my product/service as a commodity. I want to sell my solution as a system "
    "that is unique, new and different from the entire market.\n\n"
    "Goal: Create a magnetic positioning that makes my offer stand out completely.\n\n"
    "I want you to: Develop a Unique Mechanism — the proprietary name, framework, or process that "
    "explains why my solution works differently and better than anything else available.\n"
    "Make it sound new, exciting, and ownable, so competitors can't copy it."
)

sub_hdr("Market Truth (Simplified)")
body("Clothing retailers think: \"I need better billing software.\"")
body("What they actually suffer from: \"My stock in system never matches real stock.\"")
body_b("Your Unique Mechanism must revolve around that pain. Not ERP. Not AI. Stock match.")

sub_hdr("The Unique Mechanism")
highlight(
    "THE STOCK MATCH SYSTEM™\n"
    "Tagline: \"Make your system stock match your real stock.\"\n\n"
    "Full Version: The 30-Day Stock Match Formula™\n"
    "\"In 30 days, we make your billing system stock match your physical stock — and keep it that way.\"",
    "E8F8F5"
)

minor_hdr("The 4-Step Stock Match Formula™")
num("We check your current stock and find the mismatch.", 1, "Clean the Mess")
num("We organize your items properly — sizes, colors, categories.", 2, "Fix the Structure")
num("Every purchase and every sale must go through one system only.", 3, "Lock the Entries")
num("You follow one simple daily control routine.", 4, "Daily Control Rule")

minor_hdr("Why This Is Different")
stbl(["Others Say", "You Say"], [
    ["\"Here is software. Use it.\"", "\"We make sure your stock matches reality.\""],
    ["Sell the TOOL", "Sell the OUTCOME"],
    ["Compete on features", "Compete on transformation"],
], cw=[3.0, 3.0])

minor_hdr("Simple One-Line Positioning")
highlight("\"I fix stock mismatch in clothing stores in 30 days.\"", "FFF3E0")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 4: MARKET VALIDATION
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 4: MARKET VALIDATION (4 PARAMETERS)", "🔵")

prompt_box(
    "When we choose a market, we need to check on 4 parameters. If the answer is (yes) for all, "
    "the market is good:\n"
    "1. Pain (They have too much pain & want to solve it desperately)\n"
    "2. Growing (Market should be growing)\n"
    "3. Easy to Find (Customers should be easy to find)\n"
    "4. Money (They are ready to pay)"
)

stbl(["Parameter", "Verdict", "Details"], [
    ["1. PAIN", "✅ YES", "8–15% mismatch, hidden leakage, daily stress, staff dependency\n"
     "⚠ Chronic bleeding pain (not heart attack). Needs diagnostic selling."],
    ["2. GROWING", "✅ YES", "Indian clothing retail: massive, evergreen, expanding in Tier 2/3\n"
     "Highly fragmented, still largely unstructured"],
    ["3. EASY TO FIND", "✅ YES", "Google Maps, market clusters, local associations\n"
     "WhatsApp groups, Instagram, offline visits, referrals\nPhysically concentrated. Low CAC."],
    ["4. MONEY", "✅ YES", "Revenue ₹30L–1.5Cr, Inventory ₹20L–60L+\n"
     "5% leakage = ₹1–3L annually. ₹60K–1L setup is rational.\n"
     "⚠ Price sensitive — need ROI clarity"],
], cw=[1.0, 0.8, 4.2])

highlight(
    "CONCLUSION: This is a GOOD market.\n"
    "Not unicorn-SaaS explosive. But strong for authority-led, consulting-driven scaling.",
    "E8F8F5"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 5: 7-STEP GSO FRAMEWORK (Overview)
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 5: THE 7-STEP GSO FRAMEWORK", "🔵")

prompt_box(
    "Help me create a Grand Slam Offer using this 7-step framework (go step by step):\n\n"
    "Step 1: Identify the Dream Outcome — their \"heaven\" (20+ pointers)\n"
    "Step 2: List the Problems — pain points, beliefs, frustrations\n"
    "Step 3: Turn Problems Into Solutions — bridge Island #1 → Island #2\n"
    "Step 4: Design Delivery Vehicles — format + effort model\n"
    "Step 5: Supercharge the Offer — urgency, scarcity, guarantee, bonuses\n"
    "Step 6: Trim & Stack — simplify into clear, irresistible structure\n"
    "Step 7: Maximize Perceived Value — Hormozi's Value Equation"
)

body("Each step is detailed in the following sections (Steps 6–12 of this document).")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 6 (GSO Step 1): DREAM OUTCOME
# ════════════════════════════════════════════════════════════════
sec_hdr("GSO STEP 1: DREAM OUTCOME (25 Pointers)", "🟥")

for cat, items in [
    ("Operational Control Dreams", [
        "System stock always matches physical stock",
        "No daily stock confusion",
        "No surprise shortages",
        "Every purchase clearly tracked",
        "Every sale automatically reflected in inventory",
        "No duplicate or missing entries",
        "Clean, organized item structure",
        "One single system — no parallel registers",
        "Reconciliation takes 10–15 minutes, not 2 hours",
        "Clear supplier outstanding at any time",
    ]),
    ("Financial Control Dreams", [
        "Accurate profit visibility",
        "Reduced hidden margin leakage",
        "Dead stock identified early",
        "Better buying decisions based on real data",
        "Working capital not stuck in wrong inventory",
        "Fewer emergency cash shortages",
        "Confidence in numbers during tax time",
    ]),
    ("Independence & Stability Dreams", [
        "Owner can check stock without asking staff",
        "Business runs even if one key employee leaves",
        "No fear of internal manipulation",
        "Ability to leave store for a few days without anxiety",
        "Structured daily control routine",
    ]),
    ("Growth Dreams", [
        "Confidence to open a second store",
        "Stable system that scales without chaos",
        "Reputation as a disciplined, professional retailer",
    ]),
]:
    minor_hdr(cat)
    for item in items: bul(item, emoji="✅")

highlight(
    "EMOTIONAL CORE (The Real Dream):\n"
    "Control. Clarity. Certainty. Freedom from daily operational tension.\n\n"
    "They don't want software. They want:\n"
    "Peace of mind that their stock and money are under control.",
    "FFF8E1"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 7 (GSO Step 2): FULL PROBLEM BREAKDOWN
# ════════════════════════════════════════════════════════════════
sec_hdr("GSO STEP 2: FULL PROBLEM BREAKDOWN", "🟥")

for section_title, content_items in [
    ("1️⃣ Pain Points — What Are They Struggling With?", [
        ("Inventory & Money:", [
            "System stock doesn't match physical stock (5–20%)",
            "Missing pieces during rush hours", "Wrong sizes showing available",
            "Supplier outstanding unclear", "Cash stuck in dead inventory",
            "Profit calculation inaccurate", "Fear of internal shrinkage"
        ]),
        ("Operational:", [
            "Daily reconciliation stress", "Staff entering data late",
            "Parallel manual register + billing software",
            "No discipline in inward entries", "Owner personally checking everything"
        ]),
        ("Dependency:", [
            "One trusted staff controls everything",
            "Owner afraid of replacing key employee",
            "Business knowledge in people's heads"
        ]),
    ]),
    ("2️⃣ Current Situation — What Have They Tried?", [
        (None, [
            "Bought/switched billing software", "Hired more staff",
            "Increased supervision", "Maintained manual backup register",
            "Used Excel alongside system", "Performed random audits",
        ]),
    ]),
    ("3️⃣ Where Are They Stuck?", [
        (None, [
            "Don't know root cause of mismatch", "Think mismatch is 'normal in retail'",
            "Don't know how to enforce discipline", "Fear changing current system",
            "Believe structure will slow operations",
        ]),
    ]),
]:
    sub_hdr(section_title)
    for sub_label, items in content_items:
        if sub_label: body_b(sub_label)
        for item in items: bul(item)

sub_hdr("4️⃣ Internal Thoughts (Hindi Beliefs)")
for belief in [
    '"Retail mein mismatch hota hi hai."',
    '"Staff pe trust karna hi padta hai."',
    '"Software sab problem solve nahi karta."',
    '"Time nahi hai proper system banane ka."',
    '"Control chahiye, par complicated nahi."',
]:
    bul(belief)
body_b("They have normalized dysfunction.")

sub_hdr("5️⃣ External Frustrations")
for f in [
    "Competitors expanding to second branch",
    "Bigger stores using organized systems",
    "GST compliance pressure",
    "Customers expecting professional billing",
]:
    bul(f)

sub_hdr("6️⃣ False Beliefs")
stbl(["False Belief", "Truth"], [
    ["More staff = more control", "More unstructured staff = more chaos"],
    ["Manual backup = safer", "Two systems = two conflicting realities"],
    ["Billing software = automatic control", "Software records. It doesn't enforce."],
    ["Stock mismatch is unavoidable", "Avoidable with structure + enforcement"],
    ["Discipline slows sales", "Structured discipline speeds operations"],
    ["Structure is only for big stores", "Structure is what makes small stores big"],
], cw=[2.5, 3.5])

sub_hdr("7️⃣ Emotional Drivers")
body("Control • Certainty • Respect • Independence • Confidence • Less tension • Trust in numbers • Stability")
body_b("At core: They want peace of mind.")

sub_hdr("8️⃣ Desired Outcome")
highlight(
    'They don\'t want "better software."\n'
    "They want: A store that runs on system, not stress.",
    "E8F8F5"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 8 (GSO Step 3): THE BRIDGE
# ════════════════════════════════════════════════════════════════
sec_hdr("GSO STEP 3: PROBLEMS → SOLUTIONS (THE BRIDGE)", "🟥")

stbl(["🟥 ISLAND #1 (Chaos)", "🟩 ISLAND #2 (Control)"], [
    ["Stock confusion", "Stock matches reality"],
    ["Hidden leakage", "Leakage measured & stopped"],
    ["Parallel systems", "One system, one truth"],
    ["Staff dependency", "Owner independence"],
    ["Owner exhaustion", "10-min daily routine"],
    ["No trust in numbers", "Complete certainty"],
], cw=[3.0, 3.0])

for phase, color, problems, solution, outcome in [
    ("PHASE 1: Control Gap Diagnosis", WARM_RED,
     "Stock mismatch unknown • Hidden leakage • No trust in numbers",
     "Inventory Reality Audit™ → Count SKUs, compare system vs physical, calculate mismatch %, identify failure points",
     "Pain becomes measurable. No assumptions. Only numbers."),
    ("PHASE 2: Structure Reset", ORANGE,
     "Duplicate items • Wrong categorization • Disorganized SKU naming • Dead inventory invisible",
     "Inventory Structure Standardization™ → Clean SKU logic, uniform naming, category mapping, remove duplicates",
     "System becomes logically structured. Chaos removed at foundation level."),
    ("PHASE 3: Supplier Control Lock", ACCENT_TEAL,
     "Missing inward entries • Supplier outstanding unclear • Overbilling unnoticed",
     "Supplier Entry Protocol™ → No stock without system entry, real-time purchase logging, supplier dashboard",
     "Inventory inflow controlled. Leakage source #1 fixed."),
    ("PHASE 4: Sales-Stock Synchronization", ACCENT_TEAL,
     "Billing without stock deduction • Manual corrections • Batch entries",
     "Sales Deduction Lock™ → No sale without system billing, real-time stock deduction",
     "Inventory outflow traceable. Leakage source #2 fixed."),
    ("PHASE 5: Single Source of Truth", PURPLE,
     "Manual register duplication • Excel confusion • 'Later entry' culture",
     "One-System Enforcement™ → Remove manual records, disable duplicate processes, strict entry timing",
     "One reality. One number. One system. THE turning point."),
    ("PHASE 6: Owner Control Activation", SUCCESS_GRN,
     "Staff dependency • Owner can't verify independently • Fear of manipulation",
     "Owner Visibility Dashboard™ → Real-time stock access, 10-Min Daily Control Routine™",
     "Control shifts from staff to owner. Dependency risk reduced."),
    ("PHASE 7: Discipline Installation", SUCCESS_GRN,
     "Staff laziness • Inconsistent entries • Gradual decay",
     "30-Day Enforcement Period™ → Weekly review calls, compliance tracking, error correction",
     "Control becomes habit. Not temporary improvement."),
    ("PHASE 8: Validation & Scale Readiness", ACCENT_GOLD,
     "No measurable proof • Fear of expansion",
     "Before-After Validation Report™ → Recalculate mismatch %, reconciliation time, independence score",
     "Proof of transformation. Confidence to scale."),
]:
    sub_hdr(phase, color)
    body_b("Problems Solved: " + problems)
    body("Solution: " + solution)
    highlight("Outcome: " + outcome, "E8F8F5")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 9 (GSO Step 4): DELIVERY VEHICLE
# ════════════════════════════════════════════════════════════════
sec_hdr("GSO STEP 4: DELIVERY VEHICLE", "🟥")

stbl(["Decision", "Answer", "Rationale"], [
    ["Physical or Remote?", "Hybrid (Start Local)", "Retail discipline installs faster with physical visit\nTier 2/3 retailers trust more when you show up"],
    ["Ticket Level?", "Mid-High (₹80K–₹1.2L)", "Premium enough for authority\nAffordable via leakage reduction ROI"],
    ["Personal Involvement?", "First 10–15 installs\nthen systemize", "Need pattern recognition, objection refinement\ncase study strength, messaging validation"],
    ["Boutique or SaaS?", "Authority First →\nThen Scalable SaaS", "Authority first = category ownership\nSaaS first = feature competition"],
    ["Client Discipline?", "4/10 — Need strict\nenforcement", "DIY = failure. Group = diluted\nThey need structure + supervision + rules"],
], cw=[1.3, 1.5, 3.2])

highlight(
    "FINAL DELIVERY MODEL (LOCKED)\n\n"
    "Format: 1:1 Consulting-Led Implementation\n"
    "Effort: Done-With-You + Strict Enforcement\n"
    "Duration: 30-Day Installation → Monthly Control Maintenance\n"
    "Capacity: 3 stores/month maximum\n\n"
    "You design. You configure. You enforce. They comply.",
    "F0F4FF"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 10 (GSO Step 5): SUPERCHARGE
# ════════════════════════════════════════════════════════════════
sec_hdr("GSO STEP 5: SUPERCHARGE THE OFFER", "🟥")

sub_hdr("🔥 Strong Offer Name")
highlight("The 30-Day Stock Certainty System™\n\"Make your system stock match your shop.\"", "FFF3E0")

sub_hdr("⏳ Urgency (Real, Not Fake)")
body("Every month of delay: 8–15% mismatch continues, capital stuck, staff habits harden.")
body_b("Implementation cycles start on the 1st of each month only. Miss it → wait 30 days.")

sub_hdr("🚪 Scarcity (Capacity-Based)")
body("Maximum 3 clothing stores per month. Not marketing — enforcement requires oversight.")

sub_hdr("🛡️ Guarantee: The Stock Certainty Guarantee™")
highlight(
    "If after full compliance your stock mismatch does not reduce measurably, "
    "we continue working at no additional setup fee until control is achieved.\n\n"
    "\"I don't win unless you win.\"",
    "E8F8F5"
)

sub_hdr("🎁 Bonuses (Retail-Relevant)")
stbl(["#", "Bonus", "Purpose", "Value"], [
    ["1", "Hidden Leakage Exposure Report™", "Quantify annual loss in ₹", "₹5,000"],
    ["2", "Staff Control Rulebook™", "Discipline rules for staff", "₹3,000"],
    ["3", "10-Minute Daily Control Routine™", "Permanent control habit", "₹2,000"],
    ["4", "90-Day Stability Audit™", "Prevent system decay", "₹8,000"],
    ["5", "Expansion Readiness Scorecard™", "Scale-ready assessment", "₹4,000"],
], cw=[0.3, 2.2, 2.0, 0.7])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 11 (GSO Step 6): TRIM & STACK
# ════════════════════════════════════════════════════════════════
sec_hdr("GSO STEP 6: TRIM & STACK (FINAL OFFER)", "🟥")

highlight(
    "HEADLINE:\nFix Your Stock Mismatch in 30 Days — Or We Continue Free Until It's Fixed.\n\n"
    "CORE PROMISE:\nIn 30 days, your system stock will match your physical stock, "
    "daily reconciliation will drop under 15 minutes, and your store will run on one "
    "disciplined system — not staff memory.",
    "FFF3E0"
)

minor_hdr("The Clean Value Stack")
stbl(["#", "Component"], [
    ["1", "Stock Mismatch Audit — measure real mismatch & leakage points"],
    ["2", "Inventory Structure Reset — fix SKU logic, categories, sizes"],
    ["3", "Supplier Entry Lock — no stock without system entry"],
    ["4", "Sales Deduction Lock — every sale reduces stock automatically"],
    ["5", "Single System Rule — no manual register, no Excel"],
    ["6", "30-Day Enforcement — monitor, correct, enforce discipline"],
    ["7", "Before vs After Validation — measurable mismatch reduction"],
], cw=[0.3, 5.7])

minor_hdr("What Was REMOVED (Trim)")
bul("Payroll modules", emoji="❌")
bul("Fancy analytics", emoji="❌")
bul("Unnecessary ERP features", emoji="❌")
bul("Multi-store complexity", emoji="❌")
body_b("Stacked only what drives: Stock match • Leakage reduction • Owner control • Discipline.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 12 (GSO Step 7): VALUE EQUATION
# ════════════════════════════════════════════════════════════════
sec_hdr("GSO STEP 7: VALUE EQUATION (HORMOZI)", "🟥")

highlight(
    "Value = (Dream Outcome × Perceived Likelihood) ÷ (Time Delay × Effort & Sacrifice)\n\n"
    "Maximize numerator. Minimize denominator.",
    "F0F4FF"
)

stbl(["Variable", "Before", "After", "How You Optimize"], [
    ["Dream Outcome", "6/10", "9/10", "Identity upgrade + expansion anchor\n\"Run a store where system = reality\""],
    ["Perceived Likelihood", "5/10", "9/10", "Diagnosis first + process detail\nGuarantee + enforcement + case studies"],
    ["Time Delay", "6/10", "2/10", "\"Visible improvement in 14 days\"\n30-day total, weekly milestones"],
    ["Effort & Sacrifice", "5/10", "2/10", "3 rules + 10 min daily\nYou handle everything else"],
], cw=[1.3, 0.6, 0.6, 3.5])

body_b("Result: 20× higher perceived value than an average unoptimized offer.")

highlight(
    "KEY INSIGHT: Make primary claim BINARY.\n\n"
    "Not: \"Reduce mismatch by X%\"\n"
    "Instead: \"Make your stock match your shop.\"\n\n"
    "Binary claims convert better than percentage claims.",
    "FFF8E1"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 13: FINAL ASSEMBLED OFFER
# ════════════════════════════════════════════════════════════════
sec_hdr("FINAL ASSEMBLED GRAND SLAM OFFER", "🟥")

cover_title("The 30-Day Stock Certainty System™", 20)
cover_sub("By Aakash Savant — Retail Control Architect", 11, MED_GRAY)
spacer(6)

for step_num, title, content in [
    ("1", "DREAM OUTCOME", "Control + Certainty + Independence.\nSystem stock matches. 10-min routine. No staff dependency. Scale-ready."),
    ("2", "PROBLEMS (Island #1)", "8–15% mismatch. Manual duplication. Wrong entries. Staff dependency.\nNo trust in numbers. Growth fear. Normalized chaos."),
    ("3", "THE BRIDGE", "Retail Control Architecture™ powered by Stock Certainty Method™\n"
     "8 phases: Audit → Structure → Supplier Lock → Sales Lock → Single System → Dashboard → Enforcement → Validation"),
    ("4", "DELIVERY", "1:1 Consulting-Led, 30-Day Installation\nDone-With-You + Strict Enforcement\nMonthly Control Maintenance available"),
    ("5", "SUPERCHARGERS", "Name: The 30-Day Stock Certainty System™\nUrgency: Cycles start 1st only\nScarcity: 3 stores/month\n"
     "Guarantee: Stock Certainty Guarantee™\nBonuses: 5 aligned to control"),
    ("6", "TRIM & STACK", "Only: Stock match + Leakage reduction + Owner control + Discipline\nRemoved: Payroll, analytics, ERP features, multi-store"),
    ("7", "VALUE EQUATION", "Dream ↑ Likelihood ↑ Time ↓ Effort ↓\n= Premium, outcome-driven, authority-positioned offer"),
]:
    sub_hdr(f"STEP {step_num}: {title}")
    body(content)

sub_hdr("Revenue Model: Option B (Authority Path)")
stbl(["Metric", "Option A (Volume)", "Option B (Authority)"], [
    ["Clients/month", "8", "3"],
    ["Price", "₹40K", "₹1L"],
    ["Monthly Revenue", "₹3.2L", "₹3L"],
    ["Support Load", "2× higher", "Manageable"],
    ["Brand Strength", "Weak (commodity)", "Strong (authority)"],
    ["Case Study Quality", "Diluted", "Premium"],
    ["Scale Path", "More noise", "Clean growth"],
], cw=[1.5, 2.0, 2.5])

body_b("Option B wins on every metric except vanity volume.")

sub_hdr("Final Pitch")
highlight(
    "\"I don't sell billing software.\n"
    "I install Retail Control Architecture™ inside clothing stores.\n\n"
    "In 30 days, your system stock will match your shop, daily reconciliation drops under 15 minutes, "
    "and you stop depending on staff for control.\"\n\n"
    "— Aakash Savant, Retail Control Architect",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# POC STRATEGY
# ════════════════════════════════════════════════════════════════
sec_hdr("POC STRATEGY (PROOF OF CONCEPT)", "🟥")

highlight(
    "If you don't have POC, build it BEFORE charging money.\n"
    "Work for free — 3 people. Collect Video Testimonial from ALL 3.\n\n"
    "Client 1 → Convert to paid\nClient 2 → Testimonial → New client\nClient 3 → Referral → New client",
    "FDEDED"
)

for phase, title, items in [
    ("1", "Select Right 3 Free Clients", [
        "Clothing retailer, ₹30L+ revenue, ₹20L+ inventory",
        "Visible stock mismatch, owner actively involved",
        "Willing to follow single-source-of-truth rule",
        "If they resist discipline → REJECT them",
    ]),
    ("2", "Position It Properly", [
        "DON'T say: \"I'll do it free.\"",
        "SAY: \"I'm selecting 3 pilot stores for my framework in exchange for data + video testimonial.\"",
        "Position as: Pilot Program, Limited Selection, Structured Case Study",
        "Free = weak. Pilot = authority.",
    ]),
    ("3", "Execute for Each Client", [
        "Baseline: Mismatch %, reconciliation time, dependency risk",
        "30-Day: SKU discipline, supplier workflow, sales sync, enforcement",
        "Validate: Re-measure everything. Document improvement.",
        "Must collect HARD NUMBERS.",
    ]),
    ("4", "Video Testimonial (Non-Negotiable)", [
        "Script it: Problems before? Mismatch %? What changed? Measurable improvement? Recommend?",
        "2–3 minutes max. No vague praise. Only measurable statements.",
    ]),
    ("5", "Authority Packaging (After 3 POCs)", [
        "Create: 3 case studies + 3 video testimonials + before-after metrics",
        "Pitch changes from \"Trust me\" → \"We reduced mismatch by avg X% in 30 days\"",
        "After 3 successful POCs: NO MORE FREE WORK.",
    ]),
]:
    sub_hdr(f"Phase {phase}: {title}")
    for item in items: bul(item)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# HIDDEN LEAKAGE REPORT
# ════════════════════════════════════════════════════════════════
sec_hdr("HIDDEN LEAKAGE EXPOSURE REPORT™ (Sales Tool)", "🔍")
body("This is not a generic document. This is a CLOSING TOOL. Fill it live during audit.")

stbl(["Section", "Field", "Value"], [
    ["Store Snapshot", "Store Name / Monthly Revenue / Annual Revenue / Inventory Value", "[Fill Live]"],
    ["Stock Mismatch", "Sample SKUs Counted / System Stock / Physical Stock / Mismatch %", "[Fill Live]"],
    ["Margin Leakage", "Inventory at Risk × Avg Margin (25–35%)", "= ₹_____ annual"],
    ["Time Cost", "Daily reconciliation hours × Owner hourly value × 26 days × 12", "= ₹_____ annual"],
    ["Dead Stock", "Estimated slow/dead stock value", "= ₹_____ stuck capital"],
    ["Dependency Risk", "Can owner verify stock independently? Track suppliers? Do audit?", "High / Medium / Low"],
], cw=[1.2, 3.3, 1.5])

highlight(
    "CLOSING QUESTION:\n\n"
    "\"If ₹2–4 lakh is leaking silently every year,\n"
    "does investing ₹80,000–₹1,00,000 once to fix structure make sense?\"\n\n"
    "→ Pause. Let the numbers speak. That's your closing moment.",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# BUSINESS NOTES & GUIDELINES
# ════════════════════════════════════════════════════════════════
sec_hdr("IMPORTANT GUIDELINES & BUSINESS NOTES", "💡")

sub_hdr("Guidelines for Using ChatGPT Effectively")
num("Restart when needed — if responses go off-track, restart the process.", 1, "Restart")
num("Always double-check important details before taking action.", 2, "Verify")
num("Complete the AI training by IEC DIAMOND (Member Makeover section).", 3, "Training")
num("Use your own judgment to adapt outputs to your situation.", 4, "Apply Your Thinking")
num("Be as specific and detailed as possible. Better input = better output.", 5, "Input Quality")

sub_hdr("Business Notes", WARM_RED)
warn_box(
    "1. REVISE EVERYTHING — Each step should be revised as you learn more.\n\n"
    "2. THIS IS NOT THE FINAL OFFER — ChatGPT is for ideas & brainstorming.\n"
    "After dummy version → Use your own brain → Market feedback → Competitor analysis → Improve again.\n\n"
    "3. FOR SERVICE BUSINESSES:\n"
    "• On the call: Walk through pitch deck, sell value, show guarantee, handle objections.\n"
    "• If they're ready: Close directly (collect advance).\n"
    "• If they say \"Send details\": Send written proposal as RECAP, not first-time pricing."
)

spacer(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("With you every step of the way,\nSid Upadhyay\nAlways your biggest supporter ❤️✊")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = MED_GRAY

spacer(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— Aakash Savant, Retail Control Architect —")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = ACCENT_GOLD


# ── Save ───────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "Week_2_Complete_Prompt_Structure_Grand_Slam_Offer.docx")
doc.save(out)
print(f"✅ Successfully created: {out}")
