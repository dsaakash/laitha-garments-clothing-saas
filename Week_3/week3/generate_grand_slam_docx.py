#!/usr/bin/env python3
"""
Generate the Pitch-Ready Grand Slam Offer .docx
Includes: Grand Slam Offer, Forensic Audit, Service Agreement,
Daily Control Pulse, and Launch Recap.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colour Palette ──────────────────────────────────────────────────
DARK_NAVY    = RGBColor(0x0F, 0x1B, 0x2D)
DEEP_BLUE    = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD  = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL  = RGBColor(0x17, 0xA2, 0xB8)
WARM_RED     = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS_GRN  = RGBColor(0x27, 0xAE, 0x60)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG     = RGBColor(0xF7, 0xF9, 0xFC)
MED_GRAY     = RGBColor(0x66, 0x66, 0x66)
BODY_BLACK   = RGBColor(0x2C, 0x2C, 0x2C)
ORANGE       = RGBColor(0xE6, 0x7E, 0x22)

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BODY_BLACK

# ── Helper Functions ────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)

def add_cover_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(60)
    p.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DEEP_BLUE

def add_cover_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT_GOLD

def add_section_header(text, emoji=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(4)
    full = f"{emoji}  {text}" if emoji else text
    r = p.add_run(full)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DEEP_BLUE
    # gold line
    line = doc.add_paragraph()
    line.space_after = Pt(8)
    rl = line.add_run("━" * 75)
    rl.font.size = Pt(6)
    rl.font.color.rgb = ACCENT_GOLD

def add_sub_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color or ACCENT_TEAL

def add_minor_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color or DEEP_BLUE

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body_bold(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    return p

def add_quote_block(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'"{text}"')
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = color or MED_GRAY

def add_script_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY_BLACK

def add_highlight_box(text, bg_hex="FFF8E1"):
    """Simulated highlight box using a 1-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BODY_BLACK
    set_cell_shading(cell, bg_hex)
    cell.width = Inches(6.0)
    add_spacer(6)

def add_bullet(text, prefix_emoji=None, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.clear()
    if prefix_emoji:
        re = p.add_run(f"{prefix_emoji} ")
        re.font.size = Pt(11)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(11)
        p.add_run(f" {text}").font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_red_bullet(text):
    add_bullet(text, prefix_emoji="❌")

def add_green_bullet(text):
    add_bullet(text, prefix_emoji="✅")

def add_numbered(text, num, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.color.rgb = ACCENT_TEAL
    r.font.size = Pt(11)
    if bold_label:
        rl = p.add_run(f"{bold_label}: ")
        rl.bold = True
        rl.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)

def styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        set_cell_shading(cell, "1B2A4A")
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, "F7F9FC")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    add_spacer(6)
    return table

def add_signature_block():
    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    labels = [
        ("Authorized Signature (Architect):", "Authorized Signature (Partner):"),
        ("", ""),
        ("Name: ________________________", "Name: ________________________"),
        ("Date:  ________________________", "Date:  ________________________"),
    ]
    for ri, (left, right) in enumerate(labels):
        cl = tbl.rows[ri].cells[0]
        cr = tbl.rows[ri].cells[1]
        cl.text = ''
        cr.text = ''
        pl = cl.paragraphs[0]
        pr = cr.paragraphs[0]
        rl = pl.add_run(left)
        rr = pr.add_run(right)
        rl.font.size = Pt(10)
        rr.font.size = Pt(10)
        if ri == 0:
            rl.bold = True
            rr.bold = True
        cl.width = Inches(3.0)
        cr.width = Inches(3.0)
    add_spacer(8)


# ════════════════════════════════════════════════════════════════════
#                       DOCUMENT CONTENT
# ════════════════════════════════════════════════════════════════════

# ── COVER PAGE ─────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

add_cover_title("🌟 THE RETAIL CONTROL")
add_cover_title("AUTHORITY SYSTEM™")
add_cover_subtitle("The 30-Day Structural Installation to Fix Stock Mismatch\nand Restore Owner Freedom")

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PITCH-READY GRAND SLAM OFFER")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = ACCENT_GOLD

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("For Owner-Operated Clothing Stores  •  Tier 2 & 3 India")
r2.font.size = Pt(11)
r2.font.color.rgb = MED_GRAY

doc.add_paragraph()

# Key facts
info_tbl = doc.add_table(rows=1, cols=3)
info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
info_items = [
    ("INVESTMENT", "₹1,20,000 Setup\n+ ₹6,000/mo SaaS"),
    ("CAPACITY", "3 Stores / Month\nCycles Start on the 1st"),
    ("GUARANTEE", "Stock Certainty Guarantee™\nNo result = Free until fixed"),
]
for i, (label, val) in enumerate(info_items):
    cell = info_tbl.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + "\n")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = ACCENT_GOLD
    r2 = p.add_run(val)
    r2.font.size = Pt(9)
    r2.font.color.rgb = DEEP_BLUE
    set_cell_shading(cell, "F7F9FC")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PART 1: THE GRAND SLAM OFFER
# ════════════════════════════════════════════════════════════════════
add_section_header("THE GRAND SLAM OFFER", "🌟")

# ── Island 1: Current Reality ─────────────────────────────────────
add_sub_header("💔 The Current Reality (Island #1: Chaos)", WARM_RED)
add_body("Most clothing store owners are stuck in a cycle of silent leakage:")

add_red_bullet("8–15% Stock Mismatch: You've normalized losing inventory every day.")
add_red_bullet('The "Double Entry" Trap: Using a manual register because you don\'t trust the software.')
add_red_bullet("Staff Dependency: If your main manager leaves, the store's \"brain\" goes with them.")
add_red_bullet('Reconciliation Stress: Spending 2 hours every night just to "find" missing stock.')
add_red_bullet("Dead Stock Pile-up: Your capital is rotting on the racks because you don't have aging data.")

add_spacer(4)
add_highlight_box("You are working hard, but your store is leaking profit through the floor.", "FDEDED")


# ── Island 2: Dream Outcome ──────────────────────────────────────
add_sub_header("🚀 The Dream Outcome (Island #2: Certainty)", SUCCESS_GRN)
add_body("Imagine a store where:")

add_green_bullet("System Stock = Physical Stock: 100% accuracy, 100% of the time.")
add_green_bullet('The 10-Minute Close: Go home on time with a clean, verified daily report.')
add_green_bullet("Authority Restored: The system is the Boss; staff simply follows the protocol.")
add_green_bullet('Expansion Ready: Your operations are so clean you can open Store #2 tomorrow.')


# ── The Unique Mechanism ─────────────────────────────────────────
add_sub_header("⚙️ The Unique Mechanism: 3 Phases / 7 Steps")
add_body('We don\'t just "give you software." We install a Single Source of Truth through our proprietary process:')

add_minor_header("Phase 1: Measure the Leakage (Forensic Audit)", WARM_RED)
add_numbered("We physically verify system vs. real stock to find your baseline.", 1, "The Control Gap Audit™")
add_numbered("We show you the annual loss in Rupees (usually ₹2L–₹8L).", 2, "Hidden Leakage Exposure Report™")

add_minor_header("Phase 2: Fix the Structure (SaaS Customization)", ORANGE)
add_numbered("We clean SKU logic, fix categories, and size-color mapping.", 3, "Inventory DNA Mapping")
add_numbered('No stock enters the shop floor without a "Digital Gatekeeper."', 4, "Supplier Entry Lock™")
add_numbered("Every sale deducts inventory automatically with 0% manual error.", 5, "Sales-Stock Sync")

add_minor_header("Phase 3: Lock & Enforce (Discipline Installation)", SUCCESS_GRN)
add_numbered("We physically remove manual registers and Excel duplication.", 6, "Single-System Enforcement")
add_numbered("Daily accountability checks to ensure habits don't slip.", 7, "30-Day Compliance Monitoring")


# ── The Grand Slam Stack ─────────────────────────────────────────
add_sub_header("📦 The Grand Slam Stack (Value Recap)")

styled_table(
    ["#", "Component", "Value"],
    [
        ["1", "Retail Control SaaS Platform (12 Months) — Custom-configured dashboard", "₹1,20,000"],
        ["2", "Phase 1–3 Structural Installation — Hands-on audit and setup", "₹2,75,000"],
        ["🎁", "BONUS #1: Staff Control Rulebook™ — Pre-written discipline rules", "₹25,000"],
        ["🎁", "BONUS #2: 10-Min Daily Routine™ — Owner's reconciliation checklist", "₹20,000"],
        ["🎁", "BONUS #3: Dead Stock Recovery Plan — Turn old stock into cash", "₹35,000"],
        ["🎁", "BONUS #4: Expansion Readiness Scorecard — Scale-ready audit", "₹35,000"],
        ["", "TOTAL REAL-WORLD VALUE", "₹5,30,000+"],
    ],
    col_widths=[0.4, 4.2, 0.9]
)

add_highlight_box(
    "YOUR INVESTMENT TODAY:\n"
    "Option A: ₹1,20,000 Setup + ₹6,000/month SaaS\n"
    "Option B: ₹70,000/month for 3 Months All-Inclusive",
    "FFF8E1"
)


# ── Guarantee ────────────────────────────────────────────────────
add_sub_header("🔒 The Stock Certainty Guarantee™")
add_highlight_box(
    "\"If after 30 days of compliance, your stock mismatch does not reduce measurably, "
    "we will continue to work with you for FREE until it does.\n\n"
    "We don't want your money if your stock doesn't match your shop.\"",
    "E8F8F5"
)


# ── Scarcity & Urgency ──────────────────────────────────────────
add_sub_header("⚡ Scarcity & Urgency")
add_bullet("I only onboard 3 clothing stores per month to ensure forensic accuracy.", prefix_emoji="🏆", bold_prefix='The "Elite 3" Rule:')
add_bullet("Once I fix a store in your micro-market, I don't work with your direct neighbor for 6 months.", prefix_emoji="📍", bold_prefix="Geo-Exclusivity:")
add_bullet("Next implementation cycle starts on the 1st. Miss the slot, wait 30 days.", prefix_emoji="📅", bold_prefix="The 1st of the Month:")


# ── Closing Pitch ────────────────────────────────────────────────
add_sub_header("🎯 The Closing Pitch")
add_highlight_box(
    "\"Look, you can keep your current software and keep losing ₹40,000 every single month in hidden leakage. "
    "Or, you can make a one-time investment in the Retail Control Authority System™ to stop the bleeding.\n\n"
    "In 60 days, this system has paid for itself. Every month after that is pure profit.\n\n"
    "Shall we secure your slot for the 1st?\"",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PART 2: THE FORENSIC AUDIT CHECKLIST
# ════════════════════════════════════════════════════════════════════
add_section_header("PHASE 1: THE FORENSIC AUDIT CHECKLIST", "🔍")
add_body('The "Live Interrogation" during the sales call.')
add_body(
    'During the sales call, have them open their current software or registers and ask these specific questions. '
    'Each "No" or "I don\'t know" represents a specific amount of money they are losing.'
)

# Test 1
add_sub_header('1. The "Ghost Stock" Verification')
add_body_bold("Action: Ask the owner to pick any 3 high-moving items (e.g., a specific Levi's 511 Blue, Size 32).")
add_body('Question: "Look at your screen — how many does it say you have? Now, look at your shelf — how many are actually there?"')
add_highlight_box(
    'THE EXPOSURE: "If this happens with your best-seller, how much \'Ghost Stock\' '
    'is sitting in the rest of your 2,000 SKUs?"',
    "FDEDED"
)

# Test 2
add_sub_header('2. The "Inward Gatekeeper" Test')
add_body(
    'Question: "When a supplier sends 100 shirts, who verifies that exactly 100 entered the system? '
    'Do they check the size/color breakdown, or just the total count?"'
)
add_highlight_box(
    "THE EXPOSURE: If they just check total count, they are likely victims of Supplier Shrinkage "
    "(Short-shipping), which usually accounts for 2–3% of total leakage.",
    "FDEDED"
)

# Test 3
add_sub_header('3. The "Manual Parallel" Check')
add_body(
    'Question: "Do you or your staff write anything in a \'Kachha\' register or diary '
    'that isn\'t immediately put into the computer?"'
)
add_highlight_box(
    "THE EXPOSURE: Every manual entry is a Data Black Hole. It creates a 100% chance of human error "
    "and makes the software 0% reliable.",
    "FDEDED"
)

# Test 4
add_sub_header('4. The "Dead Stock" Aging Test')
add_body(
    'Question: "Can you show me, in 3 clicks, the total value of stock that hasn\'t moved for more than 120 days?"'
)
add_highlight_box(
    'THE EXPOSURE: If they can\'t, they have "Rotted Capital." They are paying interest on loans '
    "for stock that is essentially trash.",
    "FDEDED"
)


# ── The Live Leakage Calculation ─────────────────────────────────
add_sub_header("🧮 The \"Live Leakage\" Calculation (The Closer)")
add_body("After the audit, use this formula to show them their \"Annual Penalty\":")

add_script_block('"Based on what we just saw, here is your Forensic Math:"')

styled_table(
    ["Metric", "Value", "Note"],
    [
        ["Annual Revenue", "₹1 Crore", "(Enter their actual number)"],
        ["Est. Leakage (10%)", "₹10,00,000", "(Revenue × Mismatch %)"],
        ["Monthly Loss", "₹83,333", "(Annual ÷ 12)"],
        ["Daily Loss", "₹2,739", "(Monthly ÷ 30.4)"],
        ["Setup Fee", "₹1,20,000", "(Your investment)"],
        ["Payback Time", "44 Days", "(Setup ÷ Daily Loss)"],
    ],
    col_widths=[1.8, 1.5, 2.7]
)

add_highlight_box(
    "\"You are paying a 'Chaos Tax' of ₹2,739 every single day you wait. "
    "My setup fee of ₹1.2L is essentially covered by 44 days of your own recovered profit.\n\n"
    "Does it make sense to wait another 30 days and lose another ₹82,000?\"",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PART 3: THE SERVICE AGREEMENT
# ════════════════════════════════════════════════════════════════════
add_section_header("RETAIL CONTROL SERVICE AGREEMENT", "📜")

# Agreement header
tbl_hdr = doc.add_table(rows=2, cols=2)
tbl_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_hdr.style = 'Table Grid'

cell_a = tbl_hdr.rows[0].cells[0]
cell_a.text = ''
pa = cell_a.paragraphs[0]
ra = pa.add_run("BETWEEN:")
ra.bold = True
ra.font.size = Pt(10)
set_cell_shading(cell_a, "F7F9FC")

cell_b = tbl_hdr.rows[0].cells[1]
cell_b.text = ''
pb = cell_b.paragraphs[0]
pb.add_run('[Your Name / Agency Name]  ("The Architect")').font.size = Pt(10)
set_cell_shading(cell_b, "F7F9FC")

cell_c = tbl_hdr.rows[1].cells[0]
cell_c.text = ''
pc = cell_c.paragraphs[0]
rc = pc.add_run("AND:")
rc.bold = True
rc.font.size = Pt(10)
set_cell_shading(cell_c, "F7F9FC")

cell_d = tbl_hdr.rows[1].cells[1]
cell_d.text = ''
pd = cell_d.paragraphs[0]
pd.add_run('[Client Store Name]  ("The Partner")').font.size = Pt(10)
set_cell_shading(cell_d, "F7F9FC")

add_spacer(8)

# Section 1
add_sub_header("1. THE MISSION (The Objective)")
add_body(
    "To install the 30-Day Stock Certainty System™ within [Store Name] to achieve a measurable "
    "reduction in stock mismatch and establish a \"Single Source of Truth\" via the Retail Control SaaS platform."
)

# Section 2
add_sub_header("2. SCOPE OF INSTALLATION")
add_body("The Architect will execute the 3-Phase, 7-Step protocol including:")
add_bullet("Full calculation of current Mismatch % and Annual Leakage.", bold_prefix="Forensic Audit:")
add_bullet("Cleaning SKU logic, category mapping, and system locks.", bold_prefix="SaaS Customization:")
add_bullet('Removal of manual registers and installation of the "Single System Rule."', bold_prefix="Structural Enforcement:")
add_bullet("Staff training on Inward/Outward protocols.", bold_prefix="Discipline Training:")

# Section 3
add_sub_header("3. THE PARTNER'S COMMITMENT (Rules of Engagement)")
add_body("For the system to work, the Partner (Store Owner) agrees to:")

add_highlight_box(
    "⚠️ THE SINGLE SYSTEM RULE: 100% of transactions must go through the SaaS. "
    'No parallel "Kachha" diaries.\n\n'
    "⏱️ THE 10-MINUTE CLOSE: The owner must verify the Daily Control Checklist every evening.\n\n"
    "👨‍💼 STAFF AUTHORITY: Empower the Architect to enforce the \"Staff Control Rulebook™\" without interference.",
    "FFF8E1"
)

# Section 4
add_sub_header("4. INVESTMENT & TERMS")
styled_table(
    ["Item", "Amount"],
    [
        ["Setup & Installation Fee", "₹1,20,000 (One-time)"],
        ["SaaS Subscription", "₹6,000 / month (Billed monthly or annually)"],
        ["Payment Schedule", "50% Advance (to lock slot) / 50% on Day 15 (after Phase 2)"],
    ],
    col_widths=[3.0, 3.0]
)

# Section 5
add_sub_header("5. THE PERFORMANCE GUARANTEE")
add_highlight_box(
    "If, after 30 days of full compliance with the \"Rules of Engagement\" (Section 3), "
    "the physical stock does not measurably align with the system stock, The Architect will continue "
    "providing implementation support at ZERO additional cost until the target accuracy is achieved.",
    "E8F8F5"
)

# Section 6
add_sub_header("6. CONFIDENTIALITY")
add_body(
    "All financial data, supplier lists, and margin structures exposed during the Forensic Audit "
    "will remain strictly confidential and will not be shared with any third party."
)

add_spacer(20)
add_signature_block()

add_spacer(8)
add_minor_header("💡 Why This Agreement Closes the Deal:")
add_bullet('You aren\'t a vendor; you are a partner in their profit.', bold_prefix='The "Partner" Label:')
add_bullet("If they don't follow your rules, the guarantee is void — this forces actual software adoption.", bold_prefix="Section 3 Protects You:")
add_bullet('It removes the final "Will this work for me?" fear.', bold_prefix="The Guarantee is Highlighted:")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PART 4: THE DAILY CONTROL PULSE
# ════════════════════════════════════════════════════════════════════
add_section_header("THE DAILY CONTROL PULSE™", "📱")
add_body("WhatsApp Template — To be sent every evening at store closing time.")
add_body("This is what separates you from \"software support\" and makes you an Authority.")

add_spacer(8)
add_highlight_box(
    "Day [X] of Retail Control Reset™\n\n"
    "Hi [Owner Name], store closing time! 🕒\n\n"
    "To lock in today's discipline and keep our 30-Day Stock Match on track, "
    "please reply with 'DONE' once you've verified these 3 things:\n\n"
    "1️⃣ The Inward Lock: Did 100% of today's incoming stock get entered into the system "
    "before hitting the racks? [Yes/No]\n\n"
    "2️⃣ The Zero-Diary Rule: Are there any 'Kachha' slips or manual notes lying on the counter? "
    "(If yes, enter them NOW and shred the paper).\n\n"
    "3️⃣ The Daily Mismatch Check: Pick ONE random item from the rack. Does the system count "
    "match what you're holding?\n\n"
    "📸 Action Required: Please send a screenshot of your 'Daily Sales Summary' from the dashboard.\n\n"
    "Total control is built in 10 minutes a day. See you in the morning! 🚀",
    "E8F5E9"
)

add_spacer(8)
add_minor_header("💡 Why This Works for Your Business:")
add_bullet(
    "It makes the owner realize that THEY are the ones who let the leakage happen if they don't check.",
    bold_prefix="Psychological Ownership:"
)
add_bullet(
    "Clients who answer this message get 100% results. Clients who get results renew their SaaS subscription forever.",
    bold_prefix="Retention & Results:"
)
add_bullet(
    'You are collecting "Daily Wins" that you can use as screenshots (blurred) to sell your next 3 clients.',
    bold_prefix="Case Study Data:"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# PART 5: THE 9:00 AM LAUNCH PLAN
# ════════════════════════════════════════════════════════════════════
add_section_header("THE 9:00 AM LAUNCH PLAN", "🚀")
add_body("Your step-by-step execution plan for tomorrow morning:")

add_numbered("Open your laptop.", 1)
add_numbered('Send the "Forensic Hook" to your top 5 leads.', 2)
add_numbered(
    'If they ask "Price?" → Say: "It depends on the size of the leakage we find in the Audit. '
    'Let\'s do the 10-min Audit first."',
    3
)
add_numbered('On the Call: Use the Forensic Checklist to find their "Chaos Tax."', 4)
add_numbered("The Close: Send this Agreement + your Payment Link.", 5)

add_spacer(12)


# ════════════════════════════════════════════════════════════════════
# PART 6: YOUR FINAL ARSENAL (MISSION RECAP)
# ════════════════════════════════════════════════════════════════════
add_section_header("YOUR FINAL ARSENAL", "🏁")
add_body("You have officially built a high-ticket retail consulting agency. Here is your complete toolkit:")

styled_table(
    ["#", "Weapon", "Purpose"],
    [
        ["1", "The Grand Slam Offer", "₹1.2L Setup + ₹6k/mo — Value-based pricing"],
        ["2", "The Forensic Audit Checklist", "Live math that proves they are losing money"],
        ["3", "The Service Agreement", "Professional contract that protects your time"],
        ["4", "The Daily Control Pulse™", "System that ensures clients actually succeed"],
        ["5", "The Closing Pitch", "Script to convert diagnosis into deal"],
        ["6", "The Stock Certainty Guarantee™", "Risk reversal that removes final fear"],
        ["7", "Scarcity & Urgency Triggers", "Elite 3 Rule + Geo-Exclusivity + Cycle Dates"],
    ],
    col_widths=[0.4, 2.2, 3.4]
)

add_spacer(12)

# Final motivational box
add_highlight_box(
    "FINAL ADVICE:\n\n"
    "Tomorrow morning, don't overthink. Don't worry about the \"perfect\" font.\n"
    "Just send the first 5 messages.\n\n"
    "The first person who tells you \"Yes, I have a stock mismatch\" is the person "
    "who is going to pay for your next level of growth.\n\n"
    "You are not \"trying to start a business.\" You are running one.\n\n"
    "Go out there and turn those leaking stores into Retail Authorities. 🚀🏆",
    "FFF3E0"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("GO DOMINATE. 🚀🏆")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = ACCENT_GOLD


# ── Save ───────────────────────────────────────────────────────────
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Retail_Control_Authority_System_Grand_Slam_Offer.docx"
)
doc.save(output_path)
print(f"✅ Successfully created: {output_path}")
