#!/usr/bin/env python3
"""
Generate Retail Control Architect™ — High Ticket Funnel Blueprint .docx
Covers: Low vs High Ticket, Webinar Call Booking Funnel, High Ticket Closing Script,
Revenue Calculations, and Complete Landing Page System.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colour Palette ──────────────────────────────────────────────────
DARK_NAVY    = RGBColor(0x0F, 0x1B, 0x2D)
DEEP_BLUE    = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD  = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL  = RGBColor(0x17, 0xA2, 0xB8)
WARM_RED     = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS_GRN  = RGBColor(0x27, 0xAE, 0x60)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG     = RGBColor(0xF7, 0xF9, 0xFC)
MED_GRAY     = RGBColor(0x66, 0x66, 0x66)
BODY_BLACK   = RGBColor(0x2C, 0x2C, 0x2C)
ORANGE       = RGBColor(0xE6, 0x7E, 0x22)
PURPLE       = RGBColor(0x8E, 0x44, 0xAD)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BODY_BLACK

# ── Helper Functions ────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)

def add_cover_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(60)
    p.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DEEP_BLUE

def add_cover_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT_GOLD

def add_section_header(text, emoji=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(4)
    full = f"{emoji}  {text}" if emoji else text
    r = p.add_run(full)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DEEP_BLUE
    line = doc.add_paragraph()
    line.space_after = Pt(8)
    rl = line.add_run("━" * 75)
    rl.font.size = Pt(6)
    rl.font.color.rgb = ACCENT_GOLD

def add_sub_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color or ACCENT_TEAL

def add_minor_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color or DEEP_BLUE

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body_bold(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    return p

def add_highlight_box(text, bg_hex="FFF8E1"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BODY_BLACK
    set_cell_shading(cell, bg_hex)
    cell.width = Inches(6.0)
    add_spacer(6)

def add_bullet(text, prefix_emoji=None, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.clear()
    if prefix_emoji:
        re = p.add_run(f"{prefix_emoji} ")
        re.font.size = Pt(11)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(11)
        p.add_run(f" {text}").font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_numbered(text, num, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.color.rgb = ACCENT_TEAL
    r.font.size = Pt(11)
    if bold_label:
        rl = p.add_run(f"{bold_label}: ")
        rl.bold = True
        rl.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)

def styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        set_cell_shading(cell, "1B2A4A")
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, "F7F9FC")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    add_spacer(6)
    return table

def add_script_line(speaker, text, color=None):
    """Add a sales script line with speaker label."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{speaker}: ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = color or ACCENT_TEAL
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BODY_BLACK


# ════════════════════════════════════════════════════════════════════
#                       DOCUMENT CONTENT
# ════════════════════════════════════════════════════════════════════

# ── COVER PAGE ─────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

add_cover_title("RETAIL CONTROL ARCHITECT™")
add_cover_title("High-Ticket Funnel Blueprint")
add_cover_subtitle(
    "The Complete Webinar Call-Booking Funnel\n"
    "High-Ticket Closing Script • Revenue Calculations\n"
    "& Landing Page System for ₹1.2L+ Offer"
)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SELL YOUR ₹1,20,000 RETAIL CONTROL SYSTEM — PREDICTABLY")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = ACCENT_GOLD

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Applied to Aakash Savant's Retail Consulting Offer  •  High Ticket Closing")
r2.font.size = Pt(11)
r2.font.color.rgb = MED_GRAY

doc.add_paragraph()

# Table of contents
add_minor_header("What This Blueprint Covers")
styled_table(
    ["Section", "Topic"],
    [
        ["1", "Low Ticket vs High Ticket — Why RCA MUST Be High Ticket"],
        ["2", "The High-Ticket Funnel Architecture (Webinar Call Booking)"],
        ["3", "Step-by-Step Funnel Breakdown — All 8 Pages"],
        ["4", "Revenue Calculations — RCA Goal Math"],
        ["5", "High-Ticket Closing Script — The 1-on-1 Sales Call"],
        ["6", "The Pre-Call & Post-Call System"],
        ["7", "Landing Page Blueprints (All Pages)"],
        ["8", "Objection Handling Arsenal for ₹1.2L Close"],
        ["9", "Funnel Metrics & Optimization Playbook"],
        ["10", "Tasks & Implementation Checklist"],
    ],
    col_widths=[0.6, 5.4]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 1: LOW TICKET VS HIGH TICKET
# ════════════════════════════════════════════════════════════════════
add_section_header("LOW TICKET VS HIGH TICKET — WHY RCA MUST BE HIGH TICKET", "⚡")

add_body(
    "Before we build the funnel, you must deeply understand WHY a ₹1,20,000 offer "
    "is 10x easier to build a business on than a ₹5,000 product."
)

styled_table(
    ["Metric", "LOW Ticket (₹5,000)", "HIGH Ticket (₹1,20,000)"],
    [
        ["Leads Needed", "5,000 leads (massive scale)", "500 leads (very achievable)"],
        ["Conversion Rate", "1% = 50 sales", "1% = 5 sales"],
        ["Revenue", "50 × ₹5,000 = ₹2,50,000", "5 × ₹1,20,000 = ₹6,00,000"],
        ["Support Queries", "40+ queries to handle", "1-2 queries (personalised)"],
        ["Staff Needed", "Hire team for fulfilment", "No staff needed (maybe 1)"],
        ["Profit Margins", "Low margins, high refunds", "70%+ margins, happy clients"],
        ["Customer Experience", "Dissatisfied, no attention", "Personalised, high-touch results"],
        ["Your Time", "Spent on operations", "Spent on delivering results"],
    ],
    col_widths=[1.5, 2.3, 2.3]
)

add_highlight_box(
    "THE MATH IS CLEAR:\n\n"
    "LOW TICKET: 5,000 leads → 50 sales → ₹2,50,000 → High stress, low margins\n"
    "HIGH TICKET: 500 leads → 5 sales → ₹6,00,000 → Low stress, 70%+ margins\n\n"
    "Same effort. 2.4x more revenue. 10x less headache.\n\n"
    "For Retail Control Architect™, the math is even better because your service "
    "RECOVERS ₹2-5 Lakhs for the store owner — making your ₹1.2L fee look small.",
    "E8F5E9"
)

add_sub_header("🎯 What Counts as High Ticket in India?", ACCENT_GOLD)
add_highlight_box(
    "₹50,000 is NOT truly high ticket.\n\n"
    "When you charge ₹1 Lakh — ₹2 Lakhs in India, THEN it is genuinely high ticket.\n\n"
    "RCA PRICING:\n"
    "• Setup + Installation: ₹1,20,000 (one-time)\n"
    "• SaaS Subscription: ₹6,000/month\n"
    "• Full Overhaul (Multi-Branch): ₹75,000 — ₹2,00,000+\n\n"
    "This puts RCA firmly in the HIGH TICKET category.\n"
    "You do NOT need thousands of clients. You need 3-5 per month.",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 2: HIGH TICKET FUNNEL ARCHITECTURE
# ════════════════════════════════════════════════════════════════════
add_section_header("THE HIGH-TICKET FUNNEL ARCHITECTURE", "🏗️")

add_body(
    "This is the Webinar Call-Booking Funnel — the #1 funnel for selling "
    "high-ticket offers in the ₹25K — ₹2 Lakh range. It combines the trust-building power "
    "of a webinar with the closing power of a 1-on-1 discovery call."
)

add_sub_header("⚠️ Important Prerequisites", WARM_RED)
add_highlight_box(
    "DO NOT use this funnel unless:\n\n"
    "1. You have completed at least 30-50 one-to-one sales calls already\n"
    "2. You are selling a high-ticket service/product (₹25K — ₹2L)\n"
    "3. You understand your customer's objections from real conversations\n\n"
    "This is the EASIEST way to sell high ticket — but only when you have the "
    "experience from real calls.\n\n"
    "DO NOT use this funnel for products under ₹10,000. Use the direct webinar funnel instead.",
    "FDEDED"
)

add_sub_header("The Complete Funnel Flow Map")
add_body("Here is the exact flow for RCA's high-ticket funnel:")

styled_table(
    ["Step", "Page / Action", "RCA Application", "Goal"],
    [
        ["1", "Traffic (Ads)", "Meta/Google ads targeting retail store owners", "Drive targeted clicks"],
        ["2", "Webinar Registration Page", "'Free Masterclass: Stop Silent Stock Losses'", "Capture lead + register"],
        ["3", "Thank You Page", "Confirm registration + pre-frame the webinar", "Build anticipation"],
        ["4", "Checkout Page (if paid)", "₹500 paid workshop (optional)", "Filter serious attendees"],
        ["5", "Live Webinar", "60-90 min value + sell 1-on-1 call", "Create desire for call"],
        ["6", "Replay Page", "Recording for those who missed it", "Catch non-attendees"],
        ["7", "Calendly Link", "Book 1-on-1 Store Check-Up call", "Get call booked"],
        ["8", "1-on-1 Discovery Call", "Forensic Audit + close ₹1.2L deal", "CLOSE THE SALE"],
        ["9", "Sales Page (optional)", "Full RCA offer details for review", "Support the decision"],
        ["10", "Checkout / Payment", "Razorpay payment link for ₹1.2L", "Collect payment"],
    ],
    col_widths=[0.4, 1.5, 2.3, 1.3]
)

add_highlight_box(
    "KEY DIFFERENCE FROM DIRECT WEBINAR FUNNEL:\n\n"
    "Direct Webinar: Webinar → Sell product directly on webinar (₹3K-₹10K)\n"
    "High-Ticket Funnel: Webinar → Book 1-on-1 call → Sell on the call (₹25K-₹2L)\n\n"
    "The 1-on-1 call is where the ₹1.2L close happens. The webinar's job is to create "
    "enough trust and desire that they WANT the call.",
    "FFF8E1"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 3: STEP-BY-STEP FUNNEL BREAKDOWN
# ════════════════════════════════════════════════════════════════════
add_section_header("STEP-BY-STEP FUNNEL BREAKDOWN — ALL 8 PAGES", "📄")

# Page 1 — Registration
add_sub_header("PAGE 1: Webinar Registration Page (Lead Generation)", SUCCESS_GRN)
add_body("This is the gateway. Every visitor either registers or leaves. Optimize ruthlessly.")

add_minor_header("Headline for RCA Webinar Registration:")
add_highlight_box(
    "FREE LIVE MASTERCLASS\n\n"
    "How to Find & Stop the ₹2-5 Lakhs Your Retail Store is Silently Losing Every Year\n\n"
    "Without buying expensive software, firing your staff, or hiring an accountant\n"
    "Even if you've tried billing software before and it didn't work\n\n"
    "📅 [Date] | ⏰ [Time] | 🎥 Live on Zoom\n\n"
    "👉 Register FREE — Only 100 Seats Available",
    "E8F5E9"
)

add_minor_header("What to Include on Registration Page:")
add_numbered("Compelling headline (use the formula above)", 1)
add_numbered("3-4 bullet points of what they'll learn", 2)
add_numbered("Aakash's photo + short credibility line", 3)
add_numbered("Timer (countdown to webinar date)", 4)
add_numbered("Simple form: Name, Email, WhatsApp Number", 5)
add_numbered("Social proof: 'X store owners already registered'", 6)
add_numbered("Bonus for attending: 'Stay till the end → Free Staff Control Rulebook™ (₹25K value)'", 7)

# Page 2 — Thank You
add_sub_header("PAGE 2: Thank You Page (Post-Registration)", ACCENT_TEAL)
add_body("The TY page is NOT a dead end. It is your first opportunity to build trust and ensure show-up.")

add_minor_header("What to Include on Thank You Page:")
add_numbered("Confirmation: 'You're IN! 🎉 Your seat is reserved for [Date + Time]'", 1)
add_numbered("Pre-frame video (2-3 min): Aakash explains what to expect + why this matters", 2)
add_numbered("Zoom/meeting link with 'Add to Calendar' button", 3)
add_numbered("Pre-webinar homework: 'Before the masterclass, write down your top 3 stock frustrations'", 4)
add_numbered("ALL social media links — Instagram, YouTube, LinkedIn", 5)
add_numbered("Short case study teaser: 'Meet Rajesh — he recovered ₹3.2L after applying these principles'", 6)

# Page 3 — Webinar
add_sub_header("PAGE 3: The Live Webinar (60-90 Minutes)", ACCENT_GOLD)
add_body(
    "The webinar's ONLY purpose in a high-ticket funnel is to create enough trust & desire "
    "that they book a 1-on-1 call. You are NOT closing on the webinar."
)

add_minor_header("RCA Webinar Structure (90 Minutes):")
styled_table(
    ["Time", "Section", "What to Cover", "Goal"],
    [
        ["0-10 min", "INTRO", "Aakash's credentials, stores helped, ₹3.2L case", "Build authority"],
        ["10-15 min", "SET CONTEXT", "What they'll learn, why stay till end, bonus reveal", "Lock attention"],
        ["15-30 min", "PROBLEM", "5 pain points all retail stores face, ₹28L vs ₹24L case", "Create urgency"],
        ["30-35 min", "INSIGHT", "'It's NOT software. It's a SYSTEM problem.'", "Break false beliefs"],
        ["35-55 min", "SOLUTION", "The 5-Step Retail Control Framework overview", "Show the path"],
        ["55-65 min", "PROOF", "3-4 real testimonials with specific numbers", "Social proof"],
        ["65-75 min", "THE OFFER", "Introduce the free 1-on-1 Store Check-Up call", "Create desire"],
        ["75-90 min", "Q&A + CTA", "'Book your free call NOW — only 10 slots this week'", "Drive bookings"],
    ],
    col_widths=[0.8, 1.0, 2.7, 1.0]
)

add_highlight_box(
    "⚠️ CRITICAL RULE FOR HIGH-TICKET WEBINAR:\n\n"
    "You are NOT selling the ₹1.2L package on the webinar.\n"
    "You are selling the FREE 1-on-1 Store Check-Up CALL.\n\n"
    "The webinar CTA is: 'Book Your Free 30-Minute Store Check-Up'\n"
    "The ₹1.2L close happens on the 1-on-1 call, NOT on the webinar.\n\n"
    "Optional: Sell the ₹500 paid 1-on-1 call on the webinar to filter serious buyers.",
    "FFF3E0"
)

# Page 4 — Replay
add_sub_header("PAGE 4: Replay Page (For Non-Attendees)", PURPLE)
add_body("Send the replay within 2 hours of the webinar ending. Include:")
add_numbered("Full webinar recording (hosted on YouTube/Vimeo — unlisted)", 1)
add_numbered("Key timestamps for easy navigation", 2)
add_numbered("The same CTA: 'Book Your Free Store Check-Up' prominently below the video", 3)
add_numbered("Urgency: 'Replay available for 48 hours only'", 4)

# Page 5 — Sales Page
add_sub_header("PAGE 5: Sales Page — Full RCA Offer (Optional)", ORANGE)
add_body(
    "This is an optional page that lives separately. Send it to prospects who want "
    "to review the full offer before the call. It contains:"
)
add_numbered("Full headline + subheadline", 1)
add_numbered("Complete value stack (all components + bonuses)", 2)
add_numbered("The 5-step system explained", 3)
add_numbered("4-6 testimonials with specific results", 4)
add_numbered("The Stock Certainty Guarantee™", 5)
add_numbered("Price reveal + payment options", 6)
add_numbered("FAQ section addressing top objections", 7)
add_numbered("CTA: 'Book Your Free Store Check-Up' OR direct payment link", 8)

# Page 6 — Calendly
add_sub_header("PAGE 6: Calendly Booking Page", SUCCESS_GRN)
add_body("Set up your calendar for 1-on-1 discovery calls:")
add_numbered("Use Calendly or Cal.com", 1)
add_numbered("Offer 30-minute slots", 2)
add_numbered("Limit to 5-8 slots per week initially (scarcity + quality)", 3)
add_numbered("Include a pre-call questionnaire (embedded in Calendly)", 4)

add_minor_header("Pre-Call Questionnaire Fields:")
add_numbered("Store name and location", 1)
add_numbered("Number of staff members", 2)
add_numbered("Current billing software used", 3)
add_numbered("Estimated monthly stock mismatch (if known)", 4)
add_numbered("'What's your single biggest stock frustration right now?'", 5)
add_numbered("Annual revenue range (dropdown)", 6)

# Page 7 — Checkout
add_sub_header("PAGE 7: Checkout / Payment Page", WARM_RED)
add_body("Once closed on the call, send the payment link immediately.")
add_numbered("Use Razorpay / Instamojo / PayU for payment collection", 1)
add_numbered("Offer payment options: Full ₹1,20,000 OR 50% advance (₹60,000) + 50% on Day 15", 2)
add_numbered("Include GST (18%) in the displayed price or note it separately", 3)
add_numbered("Add testimonial + guarantee box on the checkout page", 4)
add_numbered("Confirmation email + WhatsApp message auto-triggered on payment", 5)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 4: REVENUE CALCULATIONS
# ════════════════════════════════════════════════════════════════════
add_section_header("REVENUE CALCULATIONS — RCA GOAL MATH", "🧮")

add_body(
    "This is the exact math for the Webinar Call-Booking Funnel applied to RCA's "
    "₹1,20,000 offer. Understand these numbers — they dictate your decisions."
)

add_sub_header("Scenario: Free Webinar → Paid 1-on-1 Call → High Ticket Close")

styled_table(
    ["Metric", "Value", "Notes"],
    [
        ["Ad Spend", "₹20,000", "Starting budget for testing"],
        ["Lead Cost (CPL)", "₹50", "Retail store owner leads"],
        ["Leads Generated", "400", "Ad spend ÷ CPL"],
        ["Webinar Show-up (25%)", "100 people", "With pre-framing & reminders"],
        ["1-on-1 Call Price", "₹500 (optional)", "Filters serious prospects"],
        ["Live Call Bookings", "15 people", "15% of show-ups book a call"],
        ["1-on-1 Calls Completed", "12 people", "80% show-up to booked calls"],
        ["Conversion Rate (20%)", "2-3 closures", "High-ticket close rate"],
        ["Product Price", "₹1,20,000", "RCA Setup + Installation"],
        ["Revenue from High Ticket", "₹2,40,000 — ₹3,60,000", "2-3 clients × ₹1.2L"],
    ],
    col_widths=[1.8, 1.5, 2.7]
)

add_sub_header("Additional Revenue Streams from Same Webinar")
styled_table(
    ["Revenue Source", "Calculation", "Amount"],
    [
        ["Paid Call Bookings", "15 calls × ₹500", "₹7,500"],
        ["High Ticket Closes (live)", "2-3 × ₹1,20,000", "₹2,40,000 — ₹3,60,000"],
        ["Follow-up Closes (5-day)", "1-2 more × ₹1,20,000", "₹1,20,000 — ₹2,40,000"],
        ["TOTAL REVENUE", "From ₹20K ad spend", "₹3,67,500 — ₹6,07,500"],
        ["ROAS", "Revenue ÷ Ad Spend", "18x — 30x 🔥"],
    ],
    col_widths=[1.8, 1.8, 2.4]
)

add_highlight_box(
    "THE POWER OF HIGH TICKET:\n\n"
    "₹20,000 ad spend → ₹3.6L — ₹6L revenue = 18x — 30x ROAS\n\n"
    "Compare to low ticket:\n"
    "₹20,000 ad spend → ₹2.5L revenue = 12.5x ROAS (but with 50 clients to manage)\n\n"
    "High ticket: 3 clients, ₹3.6L+ revenue, personalised delivery\n"
    "Low ticket: 50 clients, ₹2.5L revenue, overwhelmed with support",
    "E8F5E9"
)

add_sub_header("Monthly & Annual Revenue Targets")
styled_table(
    ["Period", "Clients Needed", "Revenue", "Ad Spend", "Net Profit (70%)"],
    [
        ["Per Month (Conservative)", "3 clients", "₹3,60,000", "₹20,000-₹40,000", "₹2,24,000"],
        ["Per Month (Growth)", "5 clients", "₹6,00,000", "₹40,000-₹60,000", "₹3,78,000"],
        ["Per Quarter (with Upsells)", "12 core + 3 retainers", "₹14,40,000 + ₹90,000", "₹1,20,000", "₹9,50,000+"],
        ["Annual Target", "36-60 clients", "₹43L — ₹72L", "₹5-7L", "₹30L — ₹50L"],
    ],
    col_widths=[1.5, 1.3, 1.5, 1.3, 1.2]
)

add_highlight_box(
    "THE GOAL: 3 New RCA Clients Per Month = ₹43+ Lakhs Per Year\n\n"
    "Add SaaS subscription revenue (₹6K/month per client × 36 clients = ₹25.9L/year)\n"
    "Add Rung 4 upsells (Quarterly Audit ₹30K+ per client)\n"
    "Add Rung 5 deals (Multi-branch rollout ₹1-2L per deal)\n\n"
    "REALISTIC YEAR 1 TARGET: ₹50-70 Lakhs total revenue.\n"
    "REALISTIC YEAR 2 TARGET: ₹1-1.5 Crore (with repeat + referral clients).",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 5: HIGH-TICKET CLOSING SCRIPT
# ════════════════════════════════════════════════════════════════════
add_section_header("HIGH-TICKET CLOSING SCRIPT — THE 1-ON-1 DISCOVERY CALL", "🎯")

add_body(
    "This is the most important section of this entire document. The 1-on-1 call is where "
    "the ₹1,20,000 deal is won or lost. Follow this script framework precisely."
)

add_highlight_box(
    "REMEMBER: You will learn advanced high-ticket sales in Week 5.\n"
    "This script gets you started with a proven framework.",
    "FFF8E1"
)

# Phase 1
add_sub_header("PHASE 1: RAPPORT & DISCOVERY (Minutes 0-8)", SUCCESS_GRN)
add_body("Goal: Understand their situation. Make them feel heard. Identify the real pain.")

add_minor_header("Opening Script:")
add_script_line("AAKASH", '"Hi [Name], thanks for taking the time today. I\'m really looking forward to '
    'understanding your store. Before we dive in, tell me a little bit about your store — how long '
    'have you been running it, how many staff, and what\'s it like on a typical day?"')

add_spacer(6)
add_body("LET THEM TALK. Listen for pain signals. Then go deeper:")

add_script_line("AAKASH", '"That\'s really interesting. Now I\'m curious — when you think about your stock, '
    'your inventory, the day-to-day numbers... what\'s the one thing that bothers you the most?"')

add_minor_header("5 Diagnostic Questions to Ask:")
add_numbered('"How are you currently tracking your daily stock movement?"', 1,
    "Listen for")
add_body("   → Red flags: 'We check at month-end' or 'Our billing software handles it'")

add_numbered('"When was the last time your physical stock count matched your system exactly?"', 2,
    "The Aha Moment")
add_body("   → Most owners will pause. They've NEVER had a perfect match.")

add_numbered('"How do you record when items move between racks, storage, or branches?"', 3,
    "The Hidden Leak")
add_body("   → Untracked internal movement = most common hidden leak.")

add_numbered('"When a staff member discovers a missing item — who is accountable?"', 4,
    "The Blame Gap")
add_body("   → Exposes the zero-accountability culture.")

add_numbered('"If I asked you right now how many units of your top 5 SKUs you have — could you tell me in 10 minutes?"', 5,
    "The Closer")
add_body("   → If they can't answer, the problem is proven. The solution sells itself.")

# Phase 2
add_sub_header("PHASE 2: AGITATE THE PROBLEM (Minutes 8-15)", WARM_RED)
add_body("Goal: Make them FEEL the cost of their current chaos. Quantify the pain.")

add_script_line("AAKASH", '"Based on what you\'ve told me, let me show you something that might shock you. '
    'Let\'s do some quick math together..."')

add_spacer(4)
add_minor_header("Live Leakage Calculation (Do This On Screen):")
styled_table(
    ["Metric", "Their Numbers", "Calculation"],
    [
        ["Annual Revenue", "₹___ (ask them)", "Whatever they share"],
        ["Estimated Mismatch %", "8-15%", "Based on diagnostic answers"],
        ["Annual Leakage", "Revenue × Mismatch %", "Usually ₹2L — ₹8L"],
        ["Monthly Loss", "Annual ÷ 12", "Usually ₹16K — ₹66K"],
        ["Daily Loss", "Monthly ÷ 30", "Usually ₹500 — ₹2,200"],
        ["3-Year Total Loss", "Annual × 3", "Usually ₹6L — ₹24L 😱"],
    ],
    col_widths=[1.5, 1.5, 3.0]
)

add_script_line("AAKASH", '"So you\'re paying a \'Chaos Tax\' of ₹___ every single day you wait. '
    'Over the next 3 years, that\'s ₹___ walking out the door. '
    'And this isn\'t theoretical — we just saw it in your own numbers."', WARM_RED)

# Phase 3
add_sub_header("PHASE 3: REFRAME & SOLUTION (Minutes 15-22)", ACCENT_TEAL)
add_body("Goal: Break their false beliefs. Position yourself as the answer.")

add_script_line("AAKASH", '"Here\'s the thing that most store owners don\'t realize — and I didn\'t either until '
    'I started doing forensic audits. This is NOT a software problem. It is NOT a staff problem. '
    'It is a SYSTEM problem. You don\'t have a step-by-step stock control PROCESS in your store."')

add_spacer(4)
add_body("Then reveal the 5-Step System briefly:")
add_numbered("Find the Leaks — Full forensic audit of real vs system stock", 1)
add_numbered("Organize Every Item — Unique tagging for every SKU", 2)
add_numbered("Staff Accountability — Daily checklists per staff member", 3)
add_numbered("Weekly Tracking — One-page report, 10 minutes to review", 4)
add_numbered("Lock the Process — Written store rules that run without the owner", 5)

add_script_line("AAKASH", '"I\'ve installed this exact system in [X] stores. One owner in Pune recovered '
    '₹3.2 Lakhs in 45 days. His mismatch went from 17% to under 2%. And the best part? '
    'His store now runs even when he\'s not there."')

# Phase 4
add_sub_header("PHASE 4: THE PITCH (Minutes 22-26)", ACCENT_GOLD)
add_body("Goal: Present the offer naturally. The transition should feel like help, not selling.")

add_script_line("AAKASH", '"So based on everything we\'ve discussed today, I genuinely believe I can help you '
    'fix this. Here\'s what working together looks like..."')

add_spacer(4)
add_body("Walk through the complete offer:")
add_numbered("Phase 1: Forensic Audit — We find every leak source in your store (Day 1-3)", 1)
add_numbered("Phase 2: System Setup — We configure your SaaS, clean SKUs, build locks (Day 4-15)", 2)
add_numbered("Phase 3: Discipline Installation — Staff training, daily routines, monitoring (Day 16-30)", 3)
add_numbered("BONUS: Staff Control Rulebook™ + Daily Control Pulse™ + Dead Stock Recovery Plan", 4)
add_numbered("GUARANTEE: If mismatch doesn't reduce in 30 days — I keep working for FREE", 5)

add_script_line("AAKASH", '"The total investment for everything is ₹1,20,000 for setup and installation, '
    'plus ₹6,000 per month for the SaaS platform. And if you compare that to the ₹___ you\'re '
    'losing every single month... the system pays for itself in about 44 days."')

# Phase 5
add_sub_header("PHASE 5: CLOSE (Minutes 26-30)", WARM_RED)
add_body("Goal: Get a YES or create urgency for a 48-hour decision.")

add_minor_header("If They're Ready:")
add_script_line("AAKASH", '"Great! I\'ll send you the payment link right now. We can lock in your slot '
    'for the next implementation cycle starting on the 1st. I only take 3 stores per month, '
    'and 1 slot is already taken."')

add_minor_header("If They Need Time:")
add_script_line("AAKASH", '"I completely understand. Here\'s what I\'ll do — I\'ll send you a WhatsApp summary '
    'of everything we discussed, including the leakage calculation we just did. '
    'I can hold your slot at this price for 48 hours. After that, pricing goes back to standard. '
    'What specific concern can I help you with right now?"')

add_minor_header("If They Say It's Expensive:")
add_script_line("AAKASH", '"I hear you. Let me put it this way — ₹1,20,000 divided by 365 days is ₹328 per day. '
    'You\'re currently LOSING ₹___ per day in stock leakage. So the question isn\'t whether you can afford it — '
    'it\'s whether you can afford NOT to fix it for another year."')

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 6: PRE-CALL & POST-CALL SYSTEM
# ════════════════════════════════════════════════════════════════════
add_section_header("PRE-CALL & POST-CALL SYSTEM", "📲")

add_sub_header("🟩 Pre-Call Sequence (Between Booking & Call)", SUCCESS_GRN)
styled_table(
    ["When", "Channel", "Message", "Purpose"],
    [
        ["Immediately", "Email + WA", "'You're confirmed! Here's your pre-call homework: Write down your top 3 stock frustrations'",
         "Small commitment ↑ show-up 30-40%"],
        ["24 hrs before", "WhatsApp", "'Raj from Pune had the same issues. Here's what we found... [brief story]'",
         "Prime them with proof"],
        ["2 hrs before", "WhatsApp", "'Your call is in 2 hours. Quick note: write down your top 3 frustrations. Zoom link: [link]'",
         "Reduce no-shows"],
        ["15 min before", "WhatsApp", "'See you in 15! I've looked at your pre-call form — already spotted one area to explore 🔍'",
         "Personal touch + excitement"],
    ],
    col_widths=[0.8, 0.8, 2.8, 1.6]
)

add_sub_header("🟩 Post-Call Follow-Up — 5-Day Sequence (Non-Buyers)", ACCENT_TEAL)
styled_table(
    ["Day", "Channel", "Hook", "Content"],
    [
        ["Day 1", "WhatsApp", "Call summary", "Summary of their specific leaks + payment link"],
        ["Day 2", "Email + WA", "'The math that surprised you'", "Show annual cost: ₹___/day × 365 = ₹___ vs ₹1.2L program fee"],
        ["Day 3", "WhatsApp", "'45 days later...'", "Full ₹3.2L recovery case story with details"],
        ["Day 4", "Email", "'3 questions everyone asks'", "Address top 3 objections + guarantee prominently"],
        ["Day 5", "WA + Call", "'Last chance at this price'", "Price urgency + offer 10-min clarity call"],
    ],
    col_widths=[0.5, 0.8, 1.5, 3.2]
)

add_highlight_box(
    "FOLLOW-UP TRUTH:\n\n"
    "Most high-ticket sales happen on Days 3-5 of follow-up, NOT on the call itself.\n"
    "Consistent follow-up is the difference between 5% and 25% close rate.\n\n"
    "NEVER give up after the call. The fortune is in the follow-up.",
    "E8F5E9"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 7: LANDING PAGE BLUEPRINTS
# ════════════════════════════════════════════════════════════════════
add_section_header("LANDING PAGE BLUEPRINTS — ALL PAGES", "🖥️")

add_body("Build these pages on Flexi-Funnel, Systeme.io, or WordPress + Elementor.")

add_sub_header("Page 1: Webinar Registration — Section Flow")
styled_table(
    ["#", "Section", "Content", "Tactic"],
    [
        ["1", "Gold Header Bar", "'FREE Masterclass — Only 100 Seats — Closing Soon'", "Scarcity"],
        ["2", "Hero Headline", "Main headline + date/time + subheadline", "Headline formula"],
        ["3", "What You'll Learn", "4 bullet points of webinar value", "Curiosity"],
        ["4", "Registration Form", "Name + Email + WhatsApp", "Lead capture"],
        ["5", "About Aakash", "Photo + 3-line credibility", "Trust"],
        ["6", "Social Proof", "'347 store owners attended last session'", "FOMO"],
        ["7", "Bonus Reveal", "'Stay till end → Free Staff Rulebook™ (₹25K)'", "Incentive"],
        ["8", "Final CTA", "Green button: 'Reserve My FREE Seat Now →'", "Conversion"],
    ],
    col_widths=[0.3, 1.2, 2.8, 0.9]
)

add_sub_header("Page 2: Thank You Page — Section Flow")
styled_table(
    ["#", "Section", "Content"],
    [
        ["1", "Confirmation", "'🎉 You're In! Check WhatsApp for your Zoom link'"],
        ["2", "Pre-Frame Video", "2-min video: What to expect + why this matters"],
        ["3", "Calendar Add", "'Add to Google Calendar' button"],
        ["4", "Pre-Homework", "'Write down your top 3 stock frustrations before the call'"],
        ["5", "Social Links", "Instagram / YouTube / LinkedIn icons"],
        ["6", "Teaser", "'Meet Rajesh — recovered ₹3.2L' (mini case study)"],
    ],
    col_widths=[0.3, 1.2, 4.5]
)

add_sub_header("Page 3: Replay Page — Section Flow")
styled_table(
    ["#", "Section", "Content"],
    [
        ["1", "Headline", "'You Missed the Live Session — Watch the Replay Now'"],
        ["2", "Video Embed", "Full webinar recording (YouTube/Vimeo unlisted)"],
        ["3", "Key Timestamps", "Clickable timestamps for major sections"],
        ["4", "CTA Below Video", "'Book Your Free 1-on-1 Store Check-Up →'"],
        ["5", "Urgency", "'Replay expires in 48 hours'"],
    ],
    col_widths=[0.3, 1.2, 4.5]
)

add_sub_header("Essential Technical Setup")
styled_table(
    ["Tool", "Purpose", "Setup Action"],
    [
        ["Flexi-Funnel / Systeme.io", "Build all landing pages", "Create account + set up domain"],
        ["Calendly / Cal.com", "Booking 1-on-1 calls", "Set up 30-min slots, max 5-8/week"],
        ["Razorpay / Instamojo", "Payment collection", "Activate account, set up ₹1.2L payment link"],
        ["Zoom / Google Meet", "Webinar + 1-on-1 calls", "Pro plan for webinars (or LVRG/Our-Area)"],
        ["LetConvert", "Social proof pop-ups", "Set up booking notifications"],
        ["TinyPNG", "Image compression", "Compress ALL images below 100KB"],
        ["Google PageSpeed", "Speed testing", "Test every page before launching ads"],
    ],
    col_widths=[1.5, 1.8, 2.7]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 8: OBJECTION HANDLING ARSENAL
# ════════════════════════════════════════════════════════════════════
add_section_header("OBJECTION HANDLING ARSENAL FOR ₹1.2L CLOSE", "🛡️")

add_body(
    "At ₹1,20,000 the objections are bigger. You MUST have a pre-loaded response "
    "for every objection. Write these down. Practice them."
)

styled_table(
    ["Objection", "Their Real Fear", "Your Kill-Shot Response"],
    [
        ['"Too expensive"',
         "Fear of losing money",
         "'₹1.2L ÷ 365 = ₹328/day. You're already losing ₹___/day. The system pays for itself in 44 days. After that, it's pure profit recovery.'"],
        ['"My staff won\'t follow it"',
         "Past failures with systems",
         "'The system has built-in daily checks that PHYSICALLY cannot be skipped. Plus, the Staff Rulebook makes accountability automatic.'"],
        ['"I already use Tally/billing software"',
         "Belief that tools = solutions",
         "'Software records data. It doesn't create the PROCESS. Does Tally tell your staff what to do every morning? That's what the system does.'"],
        ['"Let me think about it"',
         "Decision paralysis",
         "'Absolutely. I'll send you the summary. But remember — every day you think costs ₹___. I can hold your slot for 48 hours at this price.'"],
        ['"Can you do it cheaper?"',
         "Budget concern",
         "'I could — but then I'd have to cut the forensic audit or the 30-day monitoring. Those are the parts that GUARANTEE results. Without them, you'd be paying for hope, not certainty.'"],
        ['"I need to talk to my partner"',
         "Not the sole decision maker",
         "'Of course. Would it help if I joined a 10-minute call with both of you? I can walk through the leakage numbers so they see the same picture you do.'"],
        ['"What if it doesn\'t work?"',
         "Risk aversion",
         "'That's exactly why we have the Stock Certainty Guarantee™. If mismatch doesn't reduce in 30 days of compliance, I keep working for FREE. You literally cannot lose.'"],
        ['"My store is too small"',
         "Doesn't feel qualified",
         "'Smaller stores actually lose a HIGHER percentage. ₹2K/day leakage in a small store is more painful than ₹5K/day in a large one. This is designed FOR stores like yours.'"],
    ],
    col_widths=[1.2, 1.0, 3.8]
)

add_highlight_box(
    "GOLDEN OBJECTION-HANDLING RULE:\n\n"
    "Never argue. Always AGREE first, then REFRAME.\n\n"
    "'I understand. And that's exactly why...' → Then deliver your response.\n\n"
    "The goal is not to defeat their objection. It's to make them see the situation differently.",
    "FFF8E1"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 9: FUNNEL METRICS & OPTIMIZATION
# ════════════════════════════════════════════════════════════════════
add_section_header("FUNNEL METRICS & OPTIMIZATION PLAYBOOK", "📊")

add_sub_header("Key Metrics to Track Weekly")
styled_table(
    ["Metric", "Formula", "Target", "If Below Target"],
    [
        ["CPL (Cost Per Lead)", "Ad Spend ÷ Registrations", "Below ₹80", "Change ad creative or targeting"],
        ["Registration Rate", "Registrations ÷ Page Visitors", "25-40%", "Improve headline or offer"],
        ["Show-up Rate", "Attendees ÷ Registrations", "25-35%", "Improve reminder sequence"],
        ["Call Booking Rate", "Calls Booked ÷ Attendees", "15-20%", "Improve webinar CTA section"],
        ["Call Show-up Rate", "Calls Completed ÷ Booked", "75-85%", "Improve pre-call sequence"],
        ["Close Rate (on call)", "Sales ÷ Calls Completed", "20-30%", "Improve sales script"],
        ["Follow-up Close Rate", "Follow-up Sales ÷ Non-buyers", "15-25%", "Improve follow-up sequence"],
        ["Overall Funnel Conv.", "Total Sales ÷ Total Leads", "3-5%", "Fix the weakest metric above"],
        ["ROAS", "Revenue ÷ Ad Spend", "15x+", "Scale if above, pause if below 5x"],
    ],
    col_widths=[1.3, 1.5, 0.8, 2.4]
)

add_sub_header("Optimization Decision Tree")
add_highlight_box(
    "IF registration rate is low → Test new headlines (4 variations)\n"
    "IF show-up rate is low → Add more WhatsApp reminders + bonuses for attending\n"
    "IF call booking rate is low → Strengthen the webinar's last 15 minutes\n"
    "IF call show-up rate is low → Add pre-call questionnaire + personal touch messages\n"
    "IF close rate is low → Rewrite the discovery call script + practice more\n"
    "IF follow-up close is low → Add more case studies + urgency to Day 3-5 messages\n\n"
    "⚠️ SCALE RULE: Do NOT increase ad spend until close rate is consistently above 20%.\n"
    "Scaling a broken funnel = burning money.",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 10: TASKS & IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════
add_section_header("TASKS & IMPLEMENTATION CHECKLIST", "✅")

add_sub_header("Week 1: Foundation Setup", SUCCESS_GRN)
add_numbered("Rewatch the High-Ticket Funnel training and internalize every concept.", 1)
add_numbered("Set up payment gateway (Razorpay/Instamojo) — make it ready to receive ₹1.2L.", 2)
add_numbered("Set up Calendly with 30-min discovery call slots (5-8 per week).", 3)
add_numbered("Build webinar registration page on Flexi-Funnel.", 4)
add_numbered("Build thank you page with pre-frame video + social links.", 5)
add_numbered("Write the pre-call WhatsApp sequence (4 messages).", 6)

add_sub_header("Week 2: Content & Script Prep", ORANGE)
add_numbered("Write the 90-minute webinar script using the 7-section structure.", 1)
add_numbered("Practice the high-ticket closing script at least 5 times (record yourself).", 2)
add_numbered("Prepare the live leakage calculation spreadsheet for calls.", 3)
add_numbered("Collect and prepare 3-4 testimonials/case studies for the webinar.", 4)
add_numbered("Build the replay page.", 5)

add_sub_header("Week 3: First Webinar + Calls", ACCENT_TEAL)
add_numbered("Plan and announce your FIRST free webinar.", 1)
add_numbered("Launch Meta ads at ₹500/day to the registration page.", 2)
add_numbered("Activate the full pre-webinar reminder sequence.", 3)
add_numbered("Conduct the live webinar.", 4)
add_numbered("Book and complete 1-on-1 discovery calls from webinar attendees.", 5)
add_numbered("Activate the 5-day follow-up sequence for non-buyers.", 6)

add_sub_header("Week 4+: Optimize & Scale", PURPLE)
add_numbered("Review all metrics — find the weakest point and fix it.", 1)
add_numbered("A/B test registration page headlines (4 variations).", 2)
add_numbered("Add LetConvert social proof pop-ups.", 3)
add_numbered("Increase ad budget on winning creatives.", 4)
add_numbered("Create the webinar registration + thank you pages for the next webinar.", 5)
add_numbered("Begin monthly webinar rhythm (1-2 webinars per month).", 6)

add_spacer(12)

add_highlight_box(
    "MOST IMPORTANT TASK:\n\n"
    "If you haven't done 30-50 one-to-one calls yet — do those FIRST.\n"
    "Use the VSL → Call Booking funnel from Part 1.\n\n"
    "Once you've done 50+ calls and know the objections inside out → "
    "THEN launch this high-ticket webinar funnel.\n\n"
    "The calls teach you what the market is really thinking.\n"
    "The webinar amplifies what the calls have already proven.\n\n"
    "Don't skip steps. Master the call → Then master the funnel.",
    "FDEDED"
)

add_spacer(12)

# Final motivational box
add_highlight_box(
    "THE HIGH-TICKET MINDSET:\n\n"
    "You are not 'charging ₹1.2 Lakhs for a service.'\n"
    "You are saving a store owner ₹5-10 Lakhs per year in hidden leakage.\n\n"
    "Your fee is 44 days of their current loss.\n"
    "After Day 45, every rupee recovered is THEIR pure profit.\n\n"
    "You are not expensive. You are the cheapest insurance policy\n"
    "against ₹2,739 disappearing from their store EVERY SINGLE DAY.\n\n"
    "Price with confidence. Close with conviction.\n"
    "The numbers are on your side. 🚀🏆",
    "FFF3E0"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NOW GO CLOSE YOUR FIRST ₹1.2L DEAL. 🚀🏆")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = ACCENT_GOLD


# ── Save ───────────────────────────────────────────────────────────
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "RCA_High_Ticket_Funnel_Blueprint.docx"
)
doc.save(output_path)
print(f"✅ Successfully created: {output_path}")
