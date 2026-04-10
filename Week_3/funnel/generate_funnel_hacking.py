#!/usr/bin/env python3
"""
Generate RCA Funnel Hacking / Competitor Study Blueprint .docx
Covers: What is funnel hacking, How to hack competitors, Tracking sheet,
and actionable implementation — all applied to Retail Control Architect™
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colour Palette ──────────────────────────────────────────────────
DARK_NAVY    = RGBColor(0x0F, 0x1B, 0x2D)
DEEP_BLUE    = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD  = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL  = RGBColor(0x17, 0xA2, 0xB8)
WARM_RED     = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS_GRN  = RGBColor(0x27, 0xAE, 0x60)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
MED_GRAY     = RGBColor(0x66, 0x66, 0x66)
BODY_BLACK   = RGBColor(0x2C, 0x2C, 0x2C)
ORANGE       = RGBColor(0xE6, 0x7E, 0x22)
PURPLE       = RGBColor(0x8E, 0x44, 0xAD)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BODY_BLACK

# ── Helper Functions ────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)

def add_cover_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(60)
    p.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DEEP_BLUE

def add_cover_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT_GOLD

def add_section_header(text, emoji=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(4)
    full = f"{emoji}  {text}" if emoji else text
    r = p.add_run(full)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DEEP_BLUE
    line = doc.add_paragraph()
    line.space_after = Pt(8)
    rl = line.add_run("━" * 75)
    rl.font.size = Pt(6)
    rl.font.color.rgb = ACCENT_GOLD

def add_sub_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color or ACCENT_TEAL

def add_minor_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color or DEEP_BLUE

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body_bold(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    return p

def add_highlight_box(text, bg_hex="FFF8E1"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BODY_BLACK
    set_cell_shading(cell, bg_hex)
    cell.width = Inches(6.0)
    add_spacer(6)

def add_bullet(text, prefix_emoji=None, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.clear()
    if prefix_emoji:
        re = p.add_run(f"{prefix_emoji} ")
        re.font.size = Pt(11)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(11)
        p.add_run(f" {text}").font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_numbered(text, num, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.color.rgb = ACCENT_TEAL
    r.font.size = Pt(11)
    if bold_label:
        rl = p.add_run(f"{bold_label}: ")
        rl.bold = True
        rl.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)

def styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        set_cell_shading(cell, "1B2A4A")
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, "F7F9FC")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    add_spacer(6)
    return table


# ════════════════════════════════════════════════════════════════════
#                       DOCUMENT CONTENT
# ════════════════════════════════════════════════════════════════════

# ── COVER PAGE ─────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

add_cover_title("RETAIL CONTROL ARCHITECT™")
add_cover_title("Funnel Hacking Playbook")
add_cover_subtitle(
    "How to Study, Hack & Outperform Your Competitors\n"
    "in the Retail Inventory Consulting Space\n"
    "Complete Tracking System & Implementation Guide"
)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("STUDY THE MARKET → STEAL THE STRATEGIES → DOMINATE THE NICHE")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = ACCENT_GOLD

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Applied to Aakash Savant's Retail Consulting Offer")
r2.font.size = Pt(11)
r2.font.color.rgb = MED_GRAY

doc.add_paragraph()

add_minor_header("What This Playbook Covers")
styled_table(
    ["Section", "Topic"],
    [
        ["1", "What is Funnel Hacking & Why It Matters for RCA"],
        ["2", "What Exactly to Study in Your Competitors"],
        ["3", "Your Direct & Indirect Competitors — Who to Hack"],
        ["4", "7 Methods to Funnel Hack (Step-by-Step)"],
        ["5", "The RCA Competitor Tracking Sheet — Ready to Use"],
        ["6", "Deep-Dive: Hacking Each Funnel Element for RCA"],
        ["7", "What to Do After Funnel Hacking — Implementation"],
        ["8", "10 Competitor Funnel Hack Assignments"],
    ],
    col_widths=[0.6, 5.4]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 1: WHAT IS FUNNEL HACKING
# ════════════════════════════════════════════════════════════════════
add_section_header("WHAT IS FUNNEL HACKING & WHY IT MATTERS FOR RCA", "🔍")

add_body(
    "Whatever you are selling right now — retail store consulting, inventory systems, "
    "stock control services — there are people in the market already selling similar solutions."
)

add_highlight_box(
    "THE GOLDEN RULE OF FUNNEL HACKING:\n\n"
    "Do NOT try to re-invent the wheel.\n"
    "IMPLEMENT & IMPROVE whatever is already working in the market.\n\n"
    "Your competitors have already spent lakhs testing ads, landing pages, "
    "offers, and pricing. Learn from their investment FOR FREE.\n\n"
    "Funnel Hacking = Studying a competitor's ENTIRE customer journey — "
    "from the first ad they see to the final payment — and finding ways to do it BETTER.",
    "E8F5E9"
)

add_sub_header("Why This is CRITICAL for Retail Control Architect™")
add_body("The retail consulting / inventory management space in India has:")
add_bullet("SaaS companies selling inventory software (Zoho, Vyapar, mArgERP, Busy)", prefix_emoji="🏢")
add_bullet("Individual consultants offering store operations advice", prefix_emoji="👤")
add_bullet("Business coaches teaching retail owners how to scale", prefix_emoji="🎓")
add_bullet("Billing/POS companies marketing to the same audience", prefix_emoji="💻")
add_bullet("Franchise consultants targeting store owners", prefix_emoji="🏪")

add_body(
    "Each competitor is running ads, building funnels, writing sales pages, "
    "and solving objections. By studying ALL of them, you can build the ultimate "
    "Retail Control Architect™ funnel — one that beats everyone."
)

add_highlight_box(
    "YOUR COMPETITIVE EDGE:\n\n"
    "Most competitors sell SOFTWARE (a tool).\n"
    "You sell a SYSTEM (process + tool + accountability + guarantee).\n\n"
    "By funnel hacking software companies, you learn their messaging.\n"
    "Then you position RCA as the answer to why software ALONE doesn't work.\n\n"
    "This is the most powerful positioning move you can make.",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 2: WHAT EXACTLY TO STUDY
# ════════════════════════════════════════════════════════════════════
add_section_header("WHAT EXACTLY TO STUDY IN YOUR COMPETITORS", "📋")

add_body(
    "When you funnel hack a competitor, you don't just glance at their website. "
    "You study EVERY element of their customer journey systematically."
)

add_sub_header("The 12 Elements to Study for Every Competitor")

styled_table(
    ["#", "Element to Study", "What to Look For", "RCA Application"],
    [
        ["1", "Their Funnel Structure",
         "What pages do they have? Landing page → TY page → Email → Webinar → Sales?",
         "Map their entire flow and compare to your VSL → Call Booking funnel"],
        ["2", "Their Offer & Pricing",
         "What exactly do they sell? At what price? One-time or recurring?",
         "Compare to RCA's ₹1.2L setup + ₹6K/mo SaaS pricing"],
        ["3", "Their Delivery Method",
         "DFY / DWY / DIY? Software-only? Consulting? Course?",
         "Most sell software-only. RCA sells system installation. Highlight the gap."],
        ["4", "Their Marketing Strategy",
         "Where do they advertise? What platforms? What messaging?",
         "Find gaps they're not covering (WhatsApp, local retail events)"],
        ["5", "Their Customer Acquisition",
         "How do they get customers? Ads? SEO? Referrals? Cold calling?",
         "Adopt what works. Improve what doesn't."],
        ["6", "Their OTO (One-Time Offers)",
         "Do they upsell? What's their value ladder look like?",
         "Model your Rung 2 → 3 → 4 → 5 ladder against theirs"],
        ["7", "Their Emails",
         "Sign up for their list. Study the entire email sequence.",
         "Note subject lines, frequency, tone, CTAs — steal the best ideas"],
        ["8", "Their Sales Process",
         "Do they sell via call? Webinar? Direct checkout?",
         "Study their scripts if they do calls. Apply to your discovery call."],
        ["9", "Their Webinars",
         "Do they run webinars? What structure? What do they pitch?",
         "Attend and take detailed notes. Improve for your RCA webinar."],
        ["10", "Their Ads",
         "What ad creatives are they using? Video? Image? Carousel?",
         "Screenshot their top ads. Model (don't copy) for your campaigns."],
        ["11", "Their Lead Magnets & LPs",
         "What free offers do they use? E-books? Templates? Free trials?",
         "Compare to your '7 Hidden Reasons' PDF and free store check-up."],
        ["12", "Their Bonuses, Guarantee & Urgency",
         "What bonuses? What guarantees? How do they create urgency?",
         "Compare to your Stock Certainty Guarantee™ and Elite 3 Rule"],
    ],
    col_widths=[0.3, 1.3, 2.0, 2.4]
)

add_highlight_box(
    "KEY PRINCIPLE:\n\n"
    "You are NOT copying. You are STUDYING what the market has validated.\n"
    "Then you IMPROVE and make it your own.\n\n"
    "The best funnel is one that combines the winning elements of 10 competitors.",
    "FFF8E1"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 3: YOUR COMPETITORS — WHO TO HACK
# ════════════════════════════════════════════════════════════════════
add_section_header("YOUR DIRECT & INDIRECT COMPETITORS — WHO TO HACK", "🎯")

add_body(
    "For Retail Control Architect™, your competitors fall into 3 categories. "
    "You must study ALL three because your prospects are comparing you to all of them."
)

add_sub_header("Category 1: Direct Competitors — Retail Consulting Services", WARM_RED)
add_body("These are people offering similar consulting for retail stores.")
styled_table(
    ["Competitor Type", "Examples to Find", "What to Study", "Your Advantage"],
    [
        ["Retail operations consultants",
         "Search: 'retail store consultant India'",
         "Their pricing, process, and case studies",
         "You have a SYSTEM, not just advice"],
        ["Inventory management consultants",
         "Search: 'inventory consultant clothing store'",
         "Their offer structure and delivery",
         "You include SaaS + process + guarantee"],
        ["Business coaches for retail",
         "Search: 'retail business coaching India'",
         "Their funnels, webinars, email sequences",
         "You deliver hands-on installation, not just teaching"],
        ["Local store optimization freelancers",
         "Check Upwork, Fiverr, LinkedIn",
         "Their pricing and service descriptions",
         "You're a specialist, not a generalist"],
    ],
    col_widths=[1.3, 1.5, 1.5, 1.7]
)

add_sub_header("Category 2: Indirect Competitors — SaaS / Software Companies", ORANGE)
add_body("These are the tools your clients THINK will solve their problem.")
styled_table(
    ["Company", "What They Sell", "Study Their...", "Your Counter-Position"],
    [
        ["Zoho Inventory", "Cloud inventory software", "Ads, landing pages, pricing tiers",
         "'Software records data. It doesn't fix staff behaviour.'"],
        ["Vyapar App", "Mobile billing + inventory", "Facebook ads, user testimonials",
         "'Vyapar is great for billing. But does it tell your staff what to do each morning?'"],
        ["mArgERP", "Retail ERP software", "Sales pages, demo funnels",
         "'ERP installs software. RCA installs DISCIPLINE.'"],
        ["Busy Accounting", "Billing + inventory", "Google ads, pricing structure",
         "'You've had billing software for years. Are your numbers matching yet?'"],
        ["Tally", "Accounting software", "Their positioning, market share",
         "'Tally is the register. RCA is the system that makes the register trustworthy.'"],
        ["GoFrugal / RetailEasy", "Retail POS solutions", "Feature pages, case studies",
         "'POS tracks sales. RCA tracks the gap between sales and reality.'"],
    ],
    col_widths=[1.0, 1.2, 1.5, 2.3]
)

add_sub_header("Category 3: Aspirational Competitors — Big Players to Model", PURPLE)
add_body("These are successful consultants/agencies in adjacent niches. Study their FUNNELS, not their niche.")
styled_table(
    ["Who", "Their Niche", "What to Study for RCA"],
    [
        ["Sid Upadhyay / IEC", "Internet business coaching",
         "Funnel structure, webinar format, email sequences, landing page design"],
        ["Ankur Warikoo", "Low-ticket courses", "Ad creatives, social media strategy, content approach"],
        ["Raj Shamani / Similar coaches", "Business coaching",
         "High-ticket funnel flow, YouTube → Funnel pipeline"],
        ["International: Sam Ovens", "Consulting accelerator",
         "VSL structure, call booking funnel, objection handling"],
        ["International: Alex Hormozi", "Business scaling",
         "Grand slam offer framework, lead magnet strategy, value stacking"],
        ["EOS / Traction consultants", "Business operations",
         "How they sell systems installation (similar to RCA model)"],
    ],
    col_widths=[1.5, 1.5, 3.0]
)

add_highlight_box(
    "YOUR HIT LIST:\n\n"
    "You need to funnel hack at least 10 competitors total:\n"
    "• 3-4 Direct competitors (retail consultants)\n"
    "• 3-4 Indirect competitors (SaaS companies)\n"
    "• 2-3 Aspirational competitors (big players to model)\n\n"
    "Document EVERYTHING in the tracking sheet (next section).",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 4: 7 METHODS TO FUNNEL HACK
# ════════════════════════════════════════════════════════════════════
add_section_header("7 METHODS TO FUNNEL HACK — STEP BY STEP", "🛠️")

add_body(
    "Here are the 7 exact methods to study your competitors' funnels, "
    "with specific instructions for the retail inventory consulting niche."
)

# Method 1
add_sub_header("METHOD 1: Facebook Ads Library 📱", SUCCESS_GRN)
add_highlight_box(
    "HOW TO USE:\n\n"
    "1. Go to facebook.com/ads/library\n"
    "2. Select country: India\n"
    "3. Select category: All Ads\n"
    "4. Search for competitor names or keywords\n\n"
    "SEARCH TERMS FOR RCA NICHE:\n"
    "• 'inventory management'\n"
    "• 'retail store'\n"
    "• 'stock management'\n"
    "• 'billing software retail'\n"
    "• 'Zoho Inventory'\n"
    "• 'Vyapar'\n"
    "• 'retail business'\n"
    "• 'store management'\n"
    "• 'POS software'\n"
    "• 'clothing store business'\n\n"
    "WHAT TO NOTE:\n"
    "• Ad format (video / image / carousel)\n"
    "• Ad copy — what hooks do they use?\n"
    "• CTA — what action do they ask for?\n"
    "• How long has the ad been running? (Longer = it's working)\n"
    "• Landing page URL from the ad",
    "E8F5E9"
)

# Method 2
add_sub_header("METHOD 2: Google Search 🔎", ACCENT_TEAL)
add_highlight_box(
    "SEARCH THESE EXACT PHRASES ON GOOGLE:\n\n"
    "• 'retail store consultant India'\n"
    "• 'inventory management consultant'\n"
    "• 'stock mismatch solution retail'\n"
    "• 'clothing store operations consultant'\n"
    "• 'retail inventory audit India'\n"
    "• 'how to reduce stock loss in retail store'\n"
    "• 'retail store billing software India'\n"
    "• 'POS software for clothing store'\n\n"
    "WHAT TO DO:\n"
    "• Click on the top 5 organic results AND the top 3 ads\n"
    "• Study their landing pages — headline, layout, CTA, pricing\n"
    "• Screenshot everything\n"
    "• Note what's good and what you can do BETTER",
    "E8F5E9"
)

# Method 3
add_sub_header("METHOD 3: Google Site Search (Advanced) 🧠", ACCENT_GOLD)
add_highlight_box(
    "USE THE 'site:' OPERATOR TO DEEP-DIVE INTO COMPETITOR SITES:\n\n"
    "Example searches:\n"
    "• site:zoho.com inventory retail\n"
    "• site:vyaparapp.in retail clothing\n"
    "• site:gofrugal.com clothing store\n"
    "• site:margcompusoft.com retail\n\n"
    "This shows you EVERY page on their site related to your keywords.\n"
    "Study their blog posts, case studies, pricing pages, and feature pages.\n\n"
    "ALSO USE:\n"
    "• site:youtube.com 'retail inventory management India'\n"
    "  → Find their YouTube content and study the topics they cover\n"
    "• site:linkedin.com 'retail consultant India'\n"
    "  → Find individual consultants in your space",
    "FFF8E1"
)

# Method 4
add_sub_header("METHOD 4: Paid Tools (AdvertSuite & Others) 💰", ORANGE)
add_body("Paid tools give you deeper insight into competitor ad strategies:")
styled_table(
    ["Tool", "What It Does", "Cost", "RCA Use Case"],
    [
        ["AdvertSuite", "Search ALL Facebook/Instagram ads by keyword",
         "Paid (one-time)", "Find every retail/inventory ad running in India"],
        ["SimilarWeb", "See competitor website traffic & sources",
         "Free version available", "See where Zoho/Vyapar get their traffic from"],
        ["SEMrush", "Competitor keyword & ad analysis",
         "Paid", "Find what keywords retail software companies bid on"],
        ["SpyFu", "See competitor Google Ads history",
         "Paid", "See exact Google ads Vyapar/Zoho run in India"],
        ["BuiltWith", "See what tech a website uses",
         "Free version", "Check what funnel tools competitors use"],
    ],
    col_widths=[1.0, 2.0, 1.0, 2.0]
)

# Method 5
add_sub_header("METHOD 5: Enroll in Their Products 🎯", WARM_RED)
add_highlight_box(
    "THE MOST POWERFUL HACK — BECOME THEIR CUSTOMER:\n\n"
    "FOR SaaS COMPETITORS:\n"
    "• Sign up for Zoho Inventory free trial → Study their onboarding emails\n"
    "• Sign up for Vyapar App → See their upsell flow\n"
    "• Request a GoFrugal demo → Study their sales call script\n\n"
    "FOR CONSULTING COMPETITORS:\n"
    "• Book a free consultation call → Study how they diagnose, present, and pitch\n"
    "• Download their lead magnets → Study the content and follow-up sequence\n"
    "• Attend their free webinars → Take notes on structure, content, and sales pitch\n\n"
    "FOR COURSES/COACHES:\n"
    "• Buy their cheapest product (₹500-₹2000) → Study the entire buyer journey\n"
    "• Join their WhatsApp groups → See how they nurture and sell\n\n"
    "NOTE: You learn MORE from enrolling in 3 competitor products than from "
    "100 hours of guessing. This is the single best investment of ₹5K-₹10K you'll make.",
    "FDEDED"
)

# Method 6
add_sub_header("METHOD 6: Engage with Ads on Social Media 📲", PURPLE)
add_body("Start actively engaging with retail/inventory-related ads to trigger the algorithm:")
add_numbered("Follow retail software company pages on Instagram/Facebook", 1)
add_numbered("Click on every inventory/billing software ad you see", 2)
add_numbered("Comment on their posts (even generic comments)", 3)
add_numbered("Watch their video ads completely (the algorithm shows you MORE)", 4)
add_numbered("Visit their websites from the ads (you'll get retargeted)", 5)
add_numbered("Sign up for their lead magnets and newsletters", 6)

add_highlight_box(
    "WITHIN 48 HOURS:\n\n"
    "Your Instagram and Facebook feeds will be FLOODED with competitor ads.\n"
    "This is exactly what you want — free, continuous competitor research.\n\n"
    "Screenshot EVERY ad. Save the landing page URLs.\n"
    "Build a swipe file of the best ads, headlines, and offers.",
    "E8F5E9"
)

# Method 7
add_sub_header("METHOD 7: Social Media Deep Dive 🌐", ACCENT_TEAL)
add_body("Study competitor social media presence systematically:")
styled_table(
    ["Platform", "What to Search", "What to Note"],
    [
        ["Instagram", "'inventory management', 'retail software', competitor handles",
         "Content style, reels format, engagement, bio (funnel link?), story highlights"],
        ["YouTube", "'retail store management India', 'inventory software demo'",
         "Video topics, view counts, comment sentiments, how they link to funnels"],
        ["LinkedIn", "'retail consultant', 'inventory consultant', competitor names",
         "How they position themselves, content strategy, lead gen posts"],
        ["Twitter/X", "Competitor handles, '#retailtech', '#inventorymanagement'",
         "Thought leadership, industry conversations, customer complaints"],
        ["Quora", "'stock mismatch in retail store', 'inventory loss India'",
         "What questions store owners ask — use these in your ad copy"],
        ["Reddit", "r/retail, r/smallbusiness, r/IndianBusinesses",
         "Real pain points, complaints about existing solutions"],
    ],
    col_widths=[1.0, 2.2, 2.8]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 5: COMPETITOR TRACKING SHEET
# ════════════════════════════════════════════════════════════════════
add_section_header("THE RCA COMPETITOR TRACKING SHEET", "📊")

add_body(
    "Use this tracking sheet for EVERY competitor you study. "
    "Create a Google Sheet or Notion database with these exact columns."
)

add_sub_header("Master Tracking Sheet — Column Structure")
styled_table(
    ["Column", "What to Record", "Example Entry"],
    [
        ["Competitor Name", "Company or person name", "Zoho Inventory"],
        ["Category", "Direct / Indirect / Aspirational", "Indirect (SaaS)"],
        ["Website URL", "Main website link", "zoho.com/inventory"],
        ["Offer", "What they sell", "Cloud inventory software — ₹1,499/mo"],
        ["Price Point", "Exact pricing", "Free plan → ₹1,499/mo → ₹2,999/mo"],
        ["Delivery Model", "DFY / DWY / DIY / Software", "DIY — self-serve SaaS"],
        ["Target Customer", "Who are they selling to?", "SMBs, retail & wholesale"],
        ["Funnel Type", "VSL / Webinar / Direct / Free trial", "Free trial → Onboarding emails → Upgrade"],
        ["Lead Magnet", "Free offer to get leads", "14-day free trial"],
        ["Landing Page URL", "Main conversion page", "[URL]"],
        ["Headline", "Their main headline text", "'Manage inventory effortlessly'"],
        ["Ad Platforms", "Where they advertise", "Google Ads, Facebook, Instagram"],
        ["Ad Creative Type", "Video / Image / Carousel", "Mix of product demos + testimonials"],
        ["Email Sequence", "How many emails? What content?", "7-email onboarding, then weekly tips"],
        ["OTO / Upsell", "What they upsell", "Annual plan discount, premium features"],
        ["Guarantee", "What guarantee do they offer?", "30-day money-back guarantee"],
        ["Bonuses", "Extra offers included", "'Free migration support'"],
        ["Urgency Tactic", "How they create urgency", "'Limited time: 20% off annual plans'"],
        ["Strengths", "What they do WELL", "Clean UI, strong brand, large user base"],
        ["Weaknesses", "Where they FAIL", "No process installation, no hands-on support"],
        ["RCA Opportunity", "How you can beat them", "'Software alone doesn't fix behaviour. RCA installs the SYSTEM.'"],
    ],
    col_widths=[1.2, 2.0, 2.8]
)

add_highlight_box(
    "TRACKING SHEET TEMPLATE:\n\n"
    "Create a Google Sheet with 10+ rows (one per competitor).\n"
    "Use the columns above as your header row.\n"
    "Fill in every cell for every competitor.\n\n"
    "This becomes your COMPETITIVE INTELLIGENCE DATABASE.\n"
    "Update it monthly as competitors change their strategies.",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 6: DEEP DIVE — HACKING EACH ELEMENT
# ════════════════════════════════════════════════════════════════════
add_section_header("DEEP-DIVE: HACKING EACH FUNNEL ELEMENT FOR RCA", "🔬")

add_body(
    "Now let's go deep into each funnel element and show exactly how to hack it "
    "for the Retail Control Architect™ niche."
)

# Element 1 — Ads
add_sub_header("🎨 Hacking Their Ads", WARM_RED)
styled_table(
    ["What to Steal", "How to Improve for RCA", "Example"],
    [
        ["Hook / opening line",
         "Replace generic hooks with store-specific pain",
         "Their: 'Manage inventory easily'\nYours: 'Your store lost ₹40K last month. You just don't know it yet.'"],
        ["Ad format (video vs image)",
         "Use what's working + add your POV",
         "If their video demos work, create a 'live store audit' video showing real leaks"],
        ["Social proof in ads",
         "Use YOUR specific results",
         "Their: '10,000 businesses trust us'\nYours: '₹3.2 Lakhs recovered for one Pune store in 45 days'"],
        ["CTA in the ad",
         "Make the CTA lower-friction",
         "Their: 'Start free trial'\nYours: 'Get a free 30-min store check-up — zero obligation'"],
    ],
    col_widths=[1.2, 2.0, 2.8]
)

# Element 2 — Landing Pages
add_sub_header("🖥️ Hacking Their Landing Pages", ACCENT_TEAL)
styled_table(
    ["What to Study", "Questions to Ask", "RCA Action"],
    [
        ["Headline",
         "What benefit do they promise? Is it specific or vague?",
         "Make yours MORE specific: '₹2-5 Lakhs' > 'Save money on inventory'"],
        ["Page layout",
         "What sections appear? In what order?",
         "Map their section flow. Apply the 14-strategy system from Part 2."],
        ["Trust elements",
         "Logos? Testimonials? Case studies? Numbers?",
         "Add more SPECIFIC proof: exact rupee amounts, store names, timelines"],
        ["Form fields",
         "How many fields? Too many kills conversions.",
         "Keep yours to 3 fields max: Name, Email, WhatsApp"],
        ["Page speed",
         "How fast does their page load?",
         "Beat them on speed. Under 2.5s on mobile."],
        ["Mobile design",
         "How does it look on phone?",
         "Your page MUST look perfect on mobile — 80%+ traffic is mobile."],
    ],
    col_widths=[1.0, 2.2, 2.8]
)

# Element 3 — Offer
add_sub_header("💰 Hacking Their Offer & Pricing", ACCENT_GOLD)
styled_table(
    ["Competitor Type", "Their Typical Offer", "Your Superior RCA Position"],
    [
        ["SaaS (Zoho, Vyapar)", "Software subscription ₹500-₹3K/mo",
         "Software + Process + Training + Guarantee — at ₹1.2L one-time.\n'Their tool sits there. My system gets USED.'"],
        ["Generic consultants", "Hourly consulting ₹2K-₹5K/hr",
         "Fixed outcome: Mismatch reduces or you work free.\n'I don't charge for hours. I charge for results.'"],
        ["Business coaches", "Group course ₹5K-₹50K",
         "1-on-1 implementation in YOUR store, not generic advice.\n'I don't teach theory. I install the system.'"],
        ["POS companies", "Hardware + software package ₹50K-₹2L",
         "You don't replace their POS. You fix the PROCESS around it.\n'I work WITH your existing software.'"],
    ],
    col_widths=[1.2, 1.8, 3.0]
)

# Element 4 — Email
add_sub_header("📧 Hacking Their Email Sequences", PURPLE)
add_body("Sign up for every competitor's email list and track their entire sequence:")
styled_table(
    ["Day", "Their Typical Email", "What to Note", "RCA Improvement"],
    [
        ["Day 0", "Welcome + delivery", "Subject line, tone, what they deliver",
         "Open with a shocking stat: 'Your store lost money while you read this.'"],
        ["Day 1-2", "Feature education", "What features do they highlight?",
         "Share a case study instead: 'Rajesh was losing ₹4L/year...'"],
        ["Day 3-4", "Social proof", "Testimonials, logos, numbers",
         "Share specific recovery numbers: '₹3.2L in 45 days'"],
        ["Day 5-7", "Urgency / sale", "Discounts, deadlines, CTAs",
         "Use genuine scarcity: 'Only 3 slots/month' + guarantee reminder"],
    ],
    col_widths=[0.5, 1.3, 1.5, 2.7]
)

# Element 5 — Guarantee & Urgency
add_sub_header("🛡️ Hacking Their Guarantees & Urgency Tactics", SUCCESS_GRN)
styled_table(
    ["What Competitors Usually Offer", "The RCA Superior Version"],
    [
        ["'30-day money-back guarantee'",
         "Stock Certainty Guarantee™: 'If mismatch doesn't reduce in 30 days, I work for FREE until it does.' (Performance guarantee > refund guarantee)"],
        ["'Cancel anytime'",
         "'We only take 3 stores/month to ensure forensic accuracy.' (Real capacity constraint, not fake scarcity)"],
        ["'Limited time 20% off'",
         "'Implementation cycles start on the 1st. Miss the slot, wait 30 days.' (Process-based urgency)"],
        ["'X users already signed up'",
         "'Once I fix a store in your area, I don't work with your direct competitor for 6 months.' (Geo-exclusivity = real scarcity)"],
    ],
    col_widths=[2.5, 3.5]
)

add_highlight_box(
    "YOUR GUARANTEE IS YOUR COMPETITIVE WEAPON:\n\n"
    "Most competitors offer money-back guarantees (risk transfer).\n"
    "RCA offers a PERFORMANCE guarantee (risk elimination).\n\n"
    "'We don't want your money if your stock doesn't match.'\n\n"
    "This single line will beat EVERY competitor's guarantee.\n"
    "Make sure it's prominent on every landing page, in every email, and on every call.",
    "E8F5E9"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 7: WHAT TO DO AFTER FUNNEL HACKING
# ════════════════════════════════════════════════════════════════════
add_section_header("WHAT TO DO AFTER FUNNEL HACKING — IMPLEMENTATION", "🚀")

add_body(
    "Funnel hacking is USELESS unless you take what you've learned and "
    "implement it. Here's the exact process:"
)

add_sub_header("Step 1: Compile Your Insights", SUCCESS_GRN)
add_body("After studying 10+ competitors, organize your findings into 3 categories:")
styled_table(
    ["Category", "Description", "Action"],
    [
        ["🟢 STEAL & IMPROVE", "Things competitors do well that you can do BETTER",
         "Directly adopt and improve these elements in your funnel"],
        ["🟡 FILL THE GAP", "Things NO competitor is doing that you should",
         "These become your unique competitive advantages"],
        ["🔴 AVOID", "Things competitors do badly or that don't work",
         "Learn from their mistakes without repeating them"],
    ],
    col_widths=[1.2, 2.5, 2.3]
)

add_sub_header("Step 2: Update Your Funnel Elements", ACCENT_TEAL)
add_body("Go through each element of your RCA funnel and improve based on hacking insights:")

add_numbered("Rewrite your ad copy using the best hooks you found", 1, "Ads")
add_numbered("Update your page headline, layout, and CTA based on what converts", 2, "Landing Page")
add_numbered("Reposition your offer to clearly beat competitors' weaknesses", 3, "Offer Positioning")
add_numbered("Rewrite your follow-up emails using the best subject lines and stories", 4, "Email Sequence")
add_numbered("Strengthen your guarantee and urgency based on what competitors lack", 5, "Guarantee")
add_numbered("Add bonuses that competitors DON'T offer", 6, "Value Stack")
add_numbered("Update your webinar content to address gaps competitors miss", 7, "Webinar")
add_numbered("Refine your sales script based on objections you learned from competitor calls", 8, "Sales Script")

add_sub_header("Step 3: Create Your 'Unfair Advantage' Positioning Statement", ACCENT_GOLD)
add_highlight_box(
    "AFTER HACKING 10 COMPETITORS, FILL IN THIS TEMPLATE:\n\n"
    "'Unlike [COMPETITOR TYPE] who only [WHAT THEY DO],\n"
    "Retail Control Architect™ [WHAT YOU DO DIFFERENTLY]\n"
    "so that [RESULT THE CLIENT GETS].'\n\n"
    "EXAMPLE:\n"
    "'Unlike billing software companies who only give you a tool and leave you alone,\n"
    "Retail Control Architect™ installs a complete 5-step stock control SYSTEM in your store\n"
    "so that your physical stock matches your system within 30 days — GUARANTEED.'\n\n"
    "Use this positioning statement in your ads, landing pages, webinars, and sales calls.",
    "FFF3E0"
)

add_sub_header("Step 4: Build Your Competitive Comparison Chart", WARM_RED)
add_body("Create a comparison chart for your landing page & sales calls:")
styled_table(
    ["Feature", "Billing Software", "Generic Consultant", "RCA™"],
    [
        ["Actual store audit", "❌", "Sometimes", "✅ Full forensic audit"],
        ["Daily staff accountability", "❌", "❌", "✅ Built-in checklists"],
        ["Process installation", "❌", "Advice only", "✅ Hands-on implementation"],
        ["Ongoing monitoring", "❌", "❌", "✅ 30-day compliance tracking"],
        ["Performance guarantee", "Money-back", "None", "✅ Free until it works"],
        ["Dead stock recovery", "❌", "❌", "✅ Included"],
        ["Single system enforcement", "❌", "❌", "✅ No more parallel diaries"],
        ["Staff training", "PDF/video", "Generic", "✅ In-store, customized"],
    ],
    col_widths=[1.5, 1.2, 1.2, 2.1]
)

add_highlight_box(
    "PUT THIS COMPARISON CHART:\n"
    "• On your landing page (Section 10 — Comparison)\n"
    "• In your webinar (Proof section)\n"
    "• In your sales call (Reframe section)\n"
    "• In your follow-up emails (Day 4)\n\n"
    "This chart does more selling than 10 paragraphs of text.",
    "FFF8E1"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 8: 10 COMPETITOR HACK ASSIGNMENTS
# ════════════════════════════════════════════════════════════════════
add_section_header("YOUR 10 COMPETITOR FUNNEL HACK ASSIGNMENTS", "✅")

add_body(
    "Complete ALL 10 of these assignments. Each one will directly improve your "
    "Retail Control Architect™ funnel."
)

# Assignment 1
add_sub_header("Assignment 1: Facebook Ads Library — 5 Competitors", SUCCESS_GRN)
add_highlight_box(
    "ACTION:\n"
    "1. Go to facebook.com/ads/library\n"
    "2. Search for: 'inventory management', 'retail software', 'stock management'\n"
    "3. Find 5 competitors actively running ads\n"
    "4. Screenshot their best ad creatives\n"
    "5. Note the landing page URL from each ad\n"
    "6. Record in your tracking sheet\n\n"
    "DELIVERABLE: 5 competitor ad screenshots + landing page URLs in your tracking sheet",
    "E8F5E9"
)

# Assignment 2
add_sub_header("Assignment 2: Google Search — Map 5 Competing Funnels", ACCENT_TEAL)
add_highlight_box(
    "ACTION:\n"
    "1. Google: 'retail store consultant India' and 'inventory management consultant'\n"
    "2. Click on top 5 results (both organic AND ads)\n"
    "3. Map their entire funnel: Landing page → Thank You → Email → Webinar → Sales\n"
    "4. Record their headline, CTA, pricing, and offer\n\n"
    "DELIVERABLE: 5 competitor funnel maps in your tracking sheet",
    "E8F5E9"
)

# Assignment 3
add_sub_header("Assignment 3: Sign Up for 3 Competitor Email Lists", ACCENT_GOLD)
add_highlight_box(
    "ACTION:\n"
    "1. Sign up for Zoho Inventory, Vyapar, and one retail consultant's email list\n"
    "2. For 7 days, screenshot EVERY email they send\n"
    "3. Note: Subject line, content type, CTA, frequency\n"
    "4. Identify the best 3 emails — model these for your RCA sequence\n\n"
    "DELIVERABLE: 21+ email screenshots + notes on the 3 best emails",
    "E8F5E9"
)

# Assignment 4
add_sub_header("Assignment 4: Request 2 Demo/Sales Calls", ORANGE)
add_highlight_box(
    "ACTION:\n"
    "1. Request a demo call from GoFrugal or any retail POS company\n"
    "2. Book a free consultation with a retail business coach\n"
    "3. On the call, take notes on: How they diagnose, how they present, how they pitch\n"
    "4. Note every objection-handling technique they use\n\n"
    "DELIVERABLE: Call notes for 2 competitor sales calls with key takeaways",
    "E8F5E9"
)

# Assignment 5
add_sub_header("Assignment 5: Attend 1 Competitor Webinar", PURPLE)
add_highlight_box(
    "ACTION:\n"
    "1. Find a retail business coach or SaaS company running webinars\n"
    "2. Register and attend the entire webinar\n"
    "3. Take detailed notes on: Intro (how long?), Content (what 3 chunks?), Pitch (how did they transition?)\n"
    "4. Note what worked and what felt weak\n\n"
    "DELIVERABLE: Full webinar breakdown with improvement notes for your RCA webinar",
    "E8F5E9"
)

# Assignment 6-10
add_sub_header("Assignments 6-10: Quick Wins", WARM_RED)
styled_table(
    ["#", "Assignment", "Time", "Deliverable"],
    [
        ["6", "Create a swipe file of the 10 best competitor ads (screenshots)", "1 hour",
         "Folder with 10 ad screenshots + notes"],
        ["7", "List 5 lead magnets competitors offer (PDFs, trials, templates)", "30 min",
         "List + URLs + ideas for better RCA lead magnets"],
        ["8", "Compare 5 competitors' guarantees to your Stock Certainty Guarantee™", "20 min",
         "Comparison table showing your guarantee is superior"],
        ["9", "Review 3 competitors' social media profiles (Instagram/LinkedIn)", "45 min",
         "Notes on their content strategy + what to model for RCA"],
        ["10", "Complete your full tracking sheet with ALL competitors", "2 hours",
         "Filled-in Google Sheet with 10+ competitors"],
    ],
    col_widths=[0.3, 3.0, 0.7, 2.0]
)

add_spacer(12)

add_highlight_box(
    "TOTAL TIME INVESTMENT: 8-10 hours over 1 week\n"
    "TOTAL MONEY INVESTMENT: ₹0 — ₹5,000 (if you buy a competitor's cheapest product)\n\n"
    "RETURN: A funnel that's built on PROVEN market data instead of guesswork.\n"
    "This is the difference between a funnel that converts at 2% and one that converts at 8%.\n\n"
    "Do NOT skip this. Your competitors have already paid for the research.\n"
    "All you have to do is study it for free.",
    "FFF3E0"
)

add_spacer(12)

# Final box
add_highlight_box(
    "THE FUNNEL HACKING MINDSET:\n\n"
    "1. Start managing your tracking sheet — update it weekly.\n"
    "2. Funnel hack at least 10 competitors — document EVERYTHING.\n"
    "3. Get insights and takeaways — then IMPLEMENT in your business.\n\n"
    "The best funnel in the world is a mashup of the best ideas from 10 good funnels.\n\n"
    "Your Retail Control Architect™ funnel should be the final boss — "
    "built from the strengths of every competitor, with none of their weaknesses.\n\n"
    "Go hack. Go build. Go dominate. 🚀🏆",
    "E8F5E9"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("START HACKING YOUR COMPETITORS TODAY. 🔍🏆")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = ACCENT_GOLD


# ── Save ───────────────────────────────────────────────────────────
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "RCA_Funnel_Hacking_Playbook.docx"
)
doc.save(output_path)
print(f"✅ Successfully created: {output_path}")
