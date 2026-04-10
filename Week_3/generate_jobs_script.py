#!/usr/bin/env python3
"""
Generate the Steve Jobs Style Keynote Script .docx
For Aakash Savant — Retail Control Architect
Layman terms, high impact, demonstration-led.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ── Colours ─────────────────────────────────────────────────────
BLACK       = RGBColor(0x00, 0x00, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY   = RGBColor(0x33, 0x33, 0x33)
BLUE_SF     = RGBColor(0x00, 0x7A, 0xFF) # Apple Blue

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

# ── Helpers ─────────────────────────────────────────────────────
def add_slide_box(slide_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ SLIDE: {slide_text} ]")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BLUE_SF
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_spoken_line(person, text):
    p = doc.add_paragraph()
    r_name = p.add_run(f"{person}: ")
    r_name.bold = True
    r_name.font.size = Pt(11)
    
    r_text = p.add_run(text)
    r_text.font.size = Pt(11)

def add_stage_direction(text):
    p = doc.add_paragraph()
    r = p.add_run(f"({text})")
    r.italic = True
    r.font.color.rgb = DARK_GRAY
    r.font.size = Pt(10)

# ── Content ─────────────────────────────────────────────────────

# Header
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("THE KEYNOTE SCRIPT\n")
r.bold = True; r.font.size = Pt(22)
r2 = title.add_run("The 'Steve Jobs' Launch of Retail Control")
r2.font.size = Pt(14); r2.font.color.rgb = DARK_GRAY

doc.add_paragraph("\n")

# -- ACT 1: THE VILLAIN --
add_stage_direction("Lights dim. Aakash walks onto stage. Plain black shirt. He looks at the audience for 5 seconds of silence.")

add_slide_box("A Single Image of a Leaking Metal Bucket")

add_spoken_line("AAKASH", "Every day, you wake up and go to your store. You work hard. You buy the best collection. You treat your customers like family.")
add_spoken_line("AAKASH", "But there is a hole in your floor. A hole you’ve tried to ignore.")

add_slide_box("₹2,739 / Day")

add_spoken_line("AAKASH", "This is the 'Chaos Tax.' It’s the money that goes missing because your system says you have a Blue Shirt in Size 32... but your shelf is empty.")
add_spoken_line("AAKASH", "You’ve accepted this. You’ve been told that retail is messy. You’ve been told that 'thoda mismatch toh hota hi hai.'")
add_spoken_line("AAKASH", "Well... we don't believe that.")

# -- ACT 2: THE BREAKTHROUGH --
add_slide_box("Images of 3 complex ERP softwares with 100 buttons each")

add_spoken_line("AAKASH", "The world tried to fix this with software. Complex, ugly, expensive software that requires a PhD to use. Your staff hates it. You don't trust it. So what do you do?")

add_stage_direction("He mocks a writing motion")

add_spoken_line("AAKASH", "You go back to the 'Kachha' register. The diary. You end up doing double the work... for half the accuracy.")
add_spoken_line("AAKASH", "Today, we are going to fix that. We are going to do something that has never been done in retail before.")

add_slide_box("RETAIL CONTROL AUTHORITY™ (Big, Bold Gold Text)")

add_spoken_line("AAKASH", "It’s not just software. It’s an installation. It’s a way of life for your store.")

# -- ACT 3: THE DEMONSTRATION --
add_stage_direction("He reaches into his pocket and pulls out a smartphone")

add_slide_box("A Simple Screenshot of a 'DONE' WhatsApp Message")

add_spoken_line("AAKASH", "This is it. This is the entire 'Control System' in the palm of your hand.")
add_spoken_line("AAKASH", "Most owners spend 2 hours at night 'finding' stock. Our owners? They spend 10 minutes.")
add_spoken_line("AAKASH", "They check three things. Inward. Outward. Random count. Done.")

add_slide_box("Video: A manual register being thrown into a trash bin")

add_spoken_line("AAKASH", "We don't 'add' things to your store. We REMOVE things. We remove the confusion. We remove the parallel diaries. We remove the dependency on that one manager who holds all the keys.")

# -- ACT 4: THE VALUE --
add_slide_box("Comparison Table: OLD WAY vs. AUTHORITY WAY")

add_spoken_line("AAKASH", "On the left? Chaos. ₹40,000 lost every month. Stress. No freedom.")
add_spoken_line("AAKASH", "On the right? Certainty. System = Shop. 100% accuracy. Freedom to open Store #2.")

add_spoken_line("AAKASH", "People ask me... 'Aakash, how much does it cost to fix my store forever?'")
add_spoken_line("AAKASH", "They expect it to cost ₹5 Lakh. Why? Because the leakage is costing them ₹10 Lakh a year.")

add_slide_box("₹1.2 Lakh (Setup) + ₹6k/mo (SaaS)")

add_spoken_line("AAKASH", "We’ve made it so affordable that it pays for itself in less than 45 days. After that? Everything you recover is pure, 100% profit.")

# -- ACT 5: THE ONE MORE THING --
add_stage_direction("He starts to walk off stage... then stops and turns back.")

add_spoken_line("AAKASH", "Oh... one more thing.")

add_slide_box("THE STOCK CERTAINTY GUARANTEE™")

add_spoken_line("AAKASH", "We are so confident in our 30-day process... that if your stock doesn't match your shop after we finish... we will keep working for free. Permanently.")
add_spoken_line("AAKASH", "We don't want your money if your stock isn't fixed.")

add_slide_box("3 Slots Remaining for March")

add_spoken_line("AAKASH", "We only do this for 3 stores a month. Because forensic accuracy requires focus. If you want to stop the bleeding, come talk to us.")
add_spoken_line("AAKASH", "Let's build some authority. Thank you.")

add_stage_direction("Fade to black.")

# Save
out_path = "/Users/aakash/Desktop/Week_3/The_Keynote_Script_Steve_Jobs_Style.docx"
doc.save(out_path)
print(f"DONE: {out_path}")
