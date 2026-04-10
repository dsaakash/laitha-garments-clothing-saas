#!/usr/bin/env python3
"""
Generate RCA Funnel Part 2 — Advanced Landing Pages, Pricing Economics,
Funnel Types & Webinar Selling System applied to Retail Control Architect™
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colour Palette ──────────────────────────────────────────────────
DARK_NAVY    = RGBColor(0x0F, 0x1B, 0x2D)
DEEP_BLUE    = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD  = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL  = RGBColor(0x17, 0xA2, 0xB8)
WARM_RED     = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS_GRN  = RGBColor(0x27, 0xAE, 0x60)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG     = RGBColor(0xF7, 0xF9, 0xFC)
MED_GRAY     = RGBColor(0x66, 0x66, 0x66)
BODY_BLACK   = RGBColor(0x2C, 0x2C, 0x2C)
ORANGE       = RGBColor(0xE6, 0x7E, 0x22)
PURPLE       = RGBColor(0x8E, 0x44, 0xAD)

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BODY_BLACK

# ── Helper Functions ────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)

def add_cover_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(60)
    p.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = DEEP_BLUE

def add_cover_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT_GOLD

def add_section_header(text, emoji=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(4)
    full = f"{emoji}  {text}" if emoji else text
    r = p.add_run(full)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DEEP_BLUE
    line = doc.add_paragraph()
    line.space_after = Pt(8)
    rl = line.add_run("━" * 75)
    rl.font.size = Pt(6)
    rl.font.color.rgb = ACCENT_GOLD

def add_sub_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color or ACCENT_TEAL

def add_minor_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color or DEEP_BLUE

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body_bold(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    return p

def add_highlight_box(text, bg_hex="FFF8E1"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BODY_BLACK
    set_cell_shading(cell, bg_hex)
    cell.width = Inches(6.0)
    add_spacer(6)

def add_bullet(text, prefix_emoji=None, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.clear()
    if prefix_emoji:
        re = p.add_run(f"{prefix_emoji} ")
        re.font.size = Pt(11)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(11)
        p.add_run(f" {text}").font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_numbered(text, num, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.color.rgb = ACCENT_TEAL
    r.font.size = Pt(11)
    if bold_label:
        rl = p.add_run(f"{bold_label}: ")
        rl.bold = True
        rl.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)

def styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        set_cell_shading(cell, "1B2A4A")
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, "F7F9FC")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    add_spacer(6)
    return table


# ════════════════════════════════════════════════════════════════════
#                       DOCUMENT CONTENT
# ════════════════════════════════════════════════════════════════════

# ── COVER PAGE ─────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

add_cover_title("RETAIL CONTROL ARCHITECT™")
add_cover_title("Funnel Mastery — Part 2")
add_cover_subtitle(
    "Advanced Landing Pages • Pricing Economics • Funnel Types\n"
    "& The Complete Webinar Selling System\n"
    "Applied to Aakash Savant's Retail Consulting Offer"
)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("WEEK 3 — FUNNELS PART 2 BLUEPRINT")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = ACCENT_GOLD

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("IEC Diamond Membership 💎 • Applied Framework for Retail Consulting")
r2.font.size = Pt(11)
r2.font.color.rgb = MED_GRAY

doc.add_paragraph()

# Table of contents
add_minor_header("What This Document Covers")
styled_table(
    ["Section", "Topic"],
    [
        ["1", "Landing Page Tools & Setup for RCA"],
        ["2", "14 Advanced Landing Page Strategies — Applied to RCA"],
        ["3", "Pricing Economics — Does Low-Ticket Make Sense for RCA?"],
        ["4", "The 3 Delivery Models (DFY / DWY / DIY) for RCA"],
        ["5", "5 Funnel Types You Must Deploy for RCA"],
        ["6", "The Complete Webinar Selling System for RCA"],
        ["7", "Before & After Webinar Execution Playbook"],
        ["8", "Webinar Selling Formula — Master Pointers"],
        ["9", "Implementation Timeline & Action Steps"],
    ],
    col_widths=[0.6, 5.4]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 1: LANDING PAGE TOOLS & SETUP
# ════════════════════════════════════════════════════════════════════
add_section_header("LANDING PAGE TOOLS & SETUP FOR RCA", "🛠️")

add_body(
    "You do NOT need to code landing pages. Multiple no-code platforms exist "
    "that let you build professional, high-converting pages in hours."
)

add_sub_header("Recommended Tools for RCA Landing Pages")
styled_table(
    ["Tool", "Best For", "RCA Recommendation"],
    [
        ["Flexi-Funnel", "Full funnel building, Indian market", "⭐ PRIMARY — Use this for all RCA funnels"],
        ["Systeme.io", "Free plan, email + funnel combo", "BACKUP — Good for email sequences"],
        ["Elementor + WordPress", "Full website + landing pages", "For retailcontrolarchitect.com main site"],
        ["ClickFunnels", "Premium funnel builder", "Optional — higher cost, more features"],
        ["Unbounce", "A/B testing focus", "For split-testing headlines later"],
        ["WordPress", "Blog + SEO content", "For content marketing & case studies"],
    ],
    col_widths=[1.2, 1.8, 3.0]
)

add_sub_header("🏠 2 Essential Things You Need First")
add_highlight_box(
    "DOMAIN = retailcontrolarchitect.com (Your digital address — like your shop's signboard)\n\n"
    "HOSTING = Hostinger or SiteGround (The plot of land where your website lives)\n\n"
    "💡 Think of it this way:\n"
    "• Domain = Your store's name board on the street\n"
    "• Hosting = The actual building your store operates from",
    "E8F5E9"
)

add_body(
    "A landing page for RCA consists of many elements working together — "
    "headlines, trust signals, CTAs, video, testimonials, and urgency triggers. "
    "Let's break down each one with advanced strategies."
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 2: 14 ADVANCED LANDING PAGE STRATEGIES
# ════════════════════════════════════════════════════════════════════
add_section_header("14 ADVANCED LANDING PAGE STRATEGIES — APPLIED TO RCA", "🎯")

add_body(
    "These are the strategies that the top 0.5% use to create "
    "high-converting landing pages. Each one is applied directly to the "
    "Retail Control Architect™ offer."
)

# Strategy 1 — Headline
add_sub_header("1️⃣ HEADLINE — The Most Important Element", WARM_RED)
add_body(
    "The headline is the MOST IMPORTANT thing on your landing page. "
    "It must depict the BIG BENEFIT the reader gets. Split test 4 variations."
)

add_minor_header("The Headline Formula (Applied to RCA):")
add_highlight_box(
    "How to — (Desire) — Stop Losing ₹2-5 Lakhs Silently Every Year From Your Retail Store\n\n"
    "Without — Buying expensive software, firing your staff, or hiring a full-time accountant\n\n"
    "Even if — (Subheading) — You've tried billing software before and it didn't work\n\n"
    "While — (Bracket) — Running your store yourself with limited time and technical knowledge",
    "FFF8E1"
)

add_minor_header("4 Split-Test Headline Variations for RCA:")
add_numbered(
    '"How to Stop Silent Stock Losses in Your Retail Store — Without Buying Expensive Software — '
    'Even if You Have Tried Everything Before"', 1, "Variation A"
)
add_numbered(
    '"Discover Exactly Where Your Clothing Store is Losing ₹40,000+ Every Month — '
    'And Stop It in 30 Days"', 2, "Variation B"
)
add_numbered(
    '"The Simple 5-Step System That Helped One Store Owner Recover ₹3.2 Lakhs in 45 Days"',
    3, "Variation C"
)
add_numbered(
    '"Why 87% of Retail Store Owners Are Losing Money They Don\'t Even Know About — '
    'Free 30-Min Store Check-Up Reveals Your Biggest Leak"', 4, "Variation D"
)

# Strategy 2 — Page Speed
add_sub_header("2️⃣ PAGE SPEED — Don't Kill Conversions with Slow Loading")
add_body("Retail store owners browse primarily on MOBILE. Page must load in under 2.5 seconds.")
add_bullet("Compress ALL images below 100KB using TinyPNG.com", prefix_emoji="📸")
add_bullet("Host VSL video on YouTube/Vimeo — never self-host", prefix_emoji="🎥")
add_bullet("Test on Google PageSpeed Insights before launching any ads", prefix_emoji="⚡")
add_bullet("Mobile-first design — 80%+ of your traffic will be mobile", prefix_emoji="📱")

# Strategy 3 — Build Trust
add_sub_header("3️⃣ BUILD TRUST — Borrow Credibility Before the Ask")
add_body("For RCA, trust is built through:")
add_bullet("Photos with recognized retail industry leaders or business coaches", prefix_emoji="📷")
add_bullet("Display numbers: '₹3.2 Lakhs recovered' / '17% → 2% mismatch reduction'", prefix_emoji="📊")
add_bullet("Real store owner testimonials with names, locations, and specific results", prefix_emoji="💬")
add_bullet("'Featured in' badges — industry publications, podcasts, local business awards", prefix_emoji="🏅")

# Strategy 4 — Timer
add_sub_header("4️⃣ TIMER ON YOUR LANDING PAGE")
add_body("Add a countdown timer showing limited availability for the free store check-up.")
add_highlight_box(
    "RCA APPLICATION:\n"
    "'⏰ Free Store Check-Up Slots Close in: [COUNTDOWN TIMER]'\n"
    "'Only 5 slots available this week — 3 already booked'\n\n"
    "Use Deadline Funnel or built-in Flexi-Funnel timer to create evergreen urgency.",
    "FFF3E0"
)

# Strategy 5 — Social Proof Popups
add_sub_header("5️⃣ SOCIAL PROOF POP-UPS")
add_body("Use LetConvert or similar tools to show real-time notifications:")
add_highlight_box(
    "RCA POP-UP EXAMPLES:\n"
    "🔔 'Rajesh from Pune just booked a Free Store Check-Up — 2 minutes ago'\n"
    "🔔 'Meera from Bengaluru recovered ₹1.8 Lakhs — via Retail Control Architect'\n"
    "🔔 '12 store owners booked a check-up this week'",
    "E8F5E9"
)

# Strategy 6 — Value Stacking
add_sub_header("6️⃣ VALUE STACKING")
add_body("Stack all the value the client gets to make the price feel tiny by comparison.")
styled_table(
    ["What They Get", "Real-World Value"],
    [
        ["Full Forensic Store Audit (Phase 1)", "₹75,000"],
        ["SaaS Platform Setup & Customization (12 months)", "₹1,20,000"],
        ["Staff Control Rulebook™", "₹25,000"],
        ["10-Min Daily Routine™ Checklist", "₹20,000"],
        ["Dead Stock Recovery Plan", "₹35,000"],
        ["Expansion Readiness Scorecard", "₹35,000"],
        ["30-Day Compliance Monitoring", "₹65,000"],
        ["BONUS: WhatsApp Daily Control Pulse™", "₹30,000"],
        ["TOTAL REAL-WORLD VALUE", "₹4,05,000+"],
        ["YOUR INVESTMENT TODAY", "₹1,20,000"],
    ],
    col_widths=[3.5, 1.5]
)

# Strategy 7 — Price Cut & Reveal
add_sub_header("7️⃣ PRICE CUT & PRICE REVEAL")
add_body("Never reveal the price immediately. Build value FIRST, then reveal strategically.")
add_highlight_box(
    "RCA PRICE REVEAL SCRIPT:\n\n"
    "'The total real-world value of everything you're getting is ₹4,05,000+.\n\n"
    "But you're NOT paying ₹4 Lakhs. Not even ₹2 Lakhs.\n\n"
    "Your total investment for the complete Retail Control Authority System™ is just ₹1,20,000 setup "
    "+ ₹6,000/month.\n\n"
    "That's less than the ₹40,000 you're ALREADY LOSING every single month in hidden stock leakage.'",
    "FFF3E0"
)

# Strategy 8 — Scarcity & Urgency
add_sub_header("8️⃣ SCARCITY & URGENCY")
add_body("Real, legitimate scarcity for RCA:")
add_bullet("Only 3 stores onboarded per month (forensic accuracy requires this)", prefix_emoji="🏆")
add_bullet("Geo-exclusivity: Won't work with competitor store in same area for 6 months", prefix_emoji="📍")
add_bullet("Implementation cycles start on the 1st — miss it, wait 30 days", prefix_emoji="📅")

# Strategy 9 — CTA Colors
add_sub_header("9️⃣ CTA BUTTON COLORS")
styled_table(
    ["CTA Position", "Button Text", "Color"],
    [
        ["Below VSL video", "Book My FREE Store Check-Up →", "Green (#2DC653)"],
        ["After social proof", "Claim My Free Check-Up Now →", "Orange (#E8A020)"],
        ["Final section", "YES — Find My Biggest Stock Leak Today", "Orange (#E8A020)"],
    ],
    col_widths=[1.5, 2.8, 1.7]
)

# Strategy 10 — Comparison
add_sub_header("🔟 COMPARISON")
add_body("Show what life looks like WITH vs WITHOUT the Retail Control Architect™ system.")
styled_table(
    ["WITHOUT RCA (Current Reality)", "WITH RCA (After 30 Days)"],
    [
        ["8-15% stock mismatch monthly", "Under 2% mismatch consistently"],
        ["₹2-5 Lakhs silent annual loss", "₹0 unaccounted leakage"],
        ["2-hour stressful nightly reconciliation", "10-minute Daily Control Pulse™"],
        ["Staff says 'I don't know'", "Staff follows written accountability rules"],
        ["Can't open 2nd store — too chaotic", "Expansion-ready with documented systems"],
        ["Paying interest on dead stock", "Dead stock identified & liquidated in Week 2"],
    ],
    col_widths=[3.0, 3.0]
)

# Strategy 11 — Kill Objections
add_sub_header("1️⃣1️⃣ MAXIMIZE BENEFITS & KILL OBJECTIONS")
add_body("Address every objection directly on the landing page with before/after proof.")
styled_table(
    ["Objection", "Kill Shot Response"],
    [
        ["'My staff won't follow any system'", "System has built-in daily checks that physically cannot be skipped"],
        ["'I already use billing software'", "Software records transactions. It doesn't create the daily habits that prevent loss"],
        ["'My store is too small'", "Smaller stores lose a HIGHER % of revenue. ₹2K/day loss = ₹7.3L/year"],
        ["'Sounds expensive'", "Compare to ₹40K/month you're already losing. System pays for itself in 44 days"],
        ["'Let me think about it'", "Every day you wait costs ₹2,739. We hold your slot for 48 hours only"],
    ],
    col_widths=[2.0, 4.0]
)

# Strategy 12 — Bonus
add_sub_header("1️⃣2️⃣ BONUS OFFERS")
add_body("Bonuses increase perceived value without increasing your cost.")
add_numbered("Staff Control Rulebook™ — Pre-written discipline rules (Value: ₹25,000)", 1, "BONUS")
add_numbered("10-Min Daily Routine™ — Owner's evening reconciliation checklist (Value: ₹20,000)", 2, "BONUS")
add_numbered("Dead Stock Recovery Plan — Turn aging inventory into cash (Value: ₹35,000)", 3, "BONUS")
add_numbered("Expansion Readiness Scorecard — Know when you're ready for Store #2 (Value: ₹35,000)", 4, "BONUS")

# Strategy 13 — Hero of Landing Page
add_sub_header("1️⃣3️⃣ HERO OF THE LANDING PAGE — ONE PURPOSE ONLY")
add_highlight_box(
    "YOUR LANDING PAGE HAS ONE HERO — ONE PURPOSE:\n\n"
    "🎯 'Book Your FREE 30-Minute Store Check-Up'\n\n"
    "Everything on the page exists to convince them to do this ONE thing.\n"
    "Do NOT add multiple CTAs (don't ask them to follow you on social media, "
    "download a PDF, AND book a call on the same page).\n\n"
    "RCA FUNNEL FLOW:\n"
    "Ad → VSL Landing Page (HERO: Book Check-Up) → Thank You Page (HERO: Attend the Call) "
    "→ Email Sequence (HERO: Show Up) → WhatsApp Reminders (HERO: Show Up) → Discovery Call → Sale",
    "E8F5E9"
)

# Strategy 14 — CTA After Each Section
add_sub_header("1️⃣4️⃣ CALL TO ACTION AFTER EACH SECTION")
add_body(
    "After every section of convincing content, add a CTA button. "
    "One page = One CTA action (booking the check-up), but the button appears after EVERY convincing section."
)
add_highlight_box(
    "RCA PAGE FLOW:\n"
    "Section 1: Hero Headline + VSL → [CTA: Book Check-Up]\n"
    "Section 2: Pain Points → [CTA: Book Check-Up]\n"
    "Section 3: Case Study + Proof → [CTA: Book Check-Up]\n"
    "Section 4: The 5-Step System → [CTA: Book Check-Up]\n"
    "Section 5: Value Stack + Guarantee → [CTA: Book Check-Up]\n"
    "Section 6: Final Urgency + Timer → [CTA: Book Check-Up]\n\n"
    "Same action, repeated after every convincing argument.",
    "FFF8E1"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 3: PRICING ECONOMICS
# ════════════════════════════════════════════════════════════════════
add_section_header("PRICING ECONOMICS — WHAT MAKES SENSE FOR RCA?", "💰")

add_sub_header("❌ Does Selling a ₹2,000 Product Make Sense for RCA?", WARM_RED)
add_body("Short answer: NO. Here's the math:")
styled_table(
    ["Metric", "Value", "Problem"],
    [
        ["Product Price", "₹2,000", "—"],
        ["Ad Spend (CAC)", "₹1,000", "50% of revenue gone to acquire the customer"],
        ["GST (18%)", "₹360", "Can't avoid this"],
        ["Platform/Tools", "₹200", "Payment gateway + hosting"],
        ["Total Expense", "₹1,560", "—"],
        ["NET PROFIT", "₹440", "Not sustainable at scale"],
    ],
    col_widths=[1.5, 1.0, 3.5]
)
add_highlight_box(
    "VERDICT: A ₹2,000 retail audit checklist PDF is NOT worth building. "
    "The margins are too thin to sustain ad spend and your time.\n\n"
    "Exception: ONLY as a tripwire (Rung 2 of your value ladder) to get them into your ecosystem, "
    "where the real money is made on the ₹1.2L setup.",
    "FDEDED"
)

add_sub_header("⚠️ Does Selling ₹5K-₹10K Through a Webinar Make Sense?", ORANGE)
add_body("Let's look at the math for a ₹5,000 product sold via webinar:")
styled_table(
    ["Metric", "Small Scale", "Large Scale (3-Day Event)"],
    [
        ["Ad Spend", "₹20,000", "₹5,00,000"],
        ["Lead Cost", "₹50", "₹50"],
        ["Total Leads", "400", "10,000"],
        ["Show-up (25%)", "100", "2,500"],
        ["Conversion (8%)", "8 sales", "200 sales"],
        ["Revenue", "₹40,000", "₹10,00,000"],
        ["ROAS", "2x", "2x"],
    ],
    col_widths=[1.5, 1.5, 3.0]
)

add_body("Now let's see the REAL profit per sale at ₹5,000:")
styled_table(
    ["Expense", "Amount"],
    [
        ["Product Price", "₹5,000"],
        ["GST (18%)", "₹760"],
        ["CAC (Ad cost per sale)", "₹2,500"],
        ["Revenue Share %", "₹500"],
        ["Tools & Platform", "₹500"],
        ["Total Expense", "₹4,260"],
        ["NET PROFIT PER SALE", "₹740 😬"],
    ],
    col_widths=[3.0, 3.0]
)

add_highlight_box(
    "🔑 THE BIG INSIGHT FOR RCA:\n\n"
    "This is EXACTLY why Retail Control Architect™ is priced at ₹1,20,000 setup.\n\n"
    "You need HIGH-TICKET products and services to make real money.\n"
    "₹740 profit per sale = burnout. ₹80,000+ profit per client = freedom.\n\n"
    "The free check-up → ₹1.2L close model is the ONLY path to sustainable revenue.",
    "E8F5E9"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 4: THE 3 DELIVERY MODELS
# ════════════════════════════════════════════════════════════════════
add_section_header("THE 3 DELIVERY MODELS — APPLIED TO RCA", "🔄")

styled_table(
    ["Model", "DFY (Done For You)", "DWY (Done With You)", "DIY (Do It Yourself)"],
    [
        ["Work Split", "80% us, 20% client", "50% us, 50% client", "20% us, 80% client"],
        ["Pricing", "Super High Ticket", "Mid-High Ticket", "Normal High Ticket"],
        ["Scalability", "Less scalable", "Medium scalable", "Highly scalable"],
        ["Format", "Agency services", "Live group coaching", "Course / templates"],
        ["RCA Product", "Full Systems Overhaul\n(Rung 5: ₹75K-₹2L)", "30-Day Stock Control\n(Rung 3: ₹5K-₹15K)", "Self-Audit Toolkit\n(Rung 2: ₹499-₹999)"],
        ["RCA Scope", "Aakash does the full\naudit, setup, training", "Aakash guides; owner\nimplements with support", "Owner follows video\ncourse & templates"],
    ],
    col_widths=[1.0, 1.5, 1.5, 1.5]
)

add_highlight_box(
    "RCA STRATEGY:\n\n"
    "Start with DWY (Done With You) — the 30-Day Stock Control System at ₹5K-₹15K.\n"
    "This is your CORE offer. It builds trust and proves results.\n\n"
    "Then UPGRADE proven clients to DFY (Full Overhaul) at ₹75K-₹2L.\n"
    "This is where the real revenue lives.\n\n"
    "DIY (Self-Audit Toolkit) is ONLY used as a lead magnet / tripwire to enter the ecosystem.",
    "FFF8E1"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 5: 5 FUNNEL TYPES
# ════════════════════════════════════════════════════════════════════
add_section_header("5 FUNNEL TYPES YOU MUST DEPLOY FOR RCA", "🔀")

# Funnel 1
add_sub_header("1️⃣ Direct Webinar Funnel → Sell Core Offer", SUCCESS_GRN)
add_body("Best for: Selling RCA products in the ₹3,000-₹10,000 range.")
add_highlight_box(
    "RCA APPLICATION:\n"
    "Ad → Webinar Registration Page → Reminder Sequence → Live Webinar → Sell ₹5K-₹10K "
    "30-Day Stock Control Package\n\n"
    "⚠️ IMPORTANT: Only use this if your offer is ₹10K or below.\n"
    "If selling the ₹1.2L setup, use Funnel #2 (VSL → Call Booking) instead.",
    "E8F5E9"
)

# Funnel 2
add_sub_header("2️⃣ VSL Funnel → Appointment/Call Booking", ACCENT_GOLD)
add_body("Best for: Selling RCA high-ticket ₹1.2L+ services through discovery calls.")
add_highlight_box(
    "RCA APPLICATION (PRIMARY FUNNEL):\n"
    "Ad → VSL Landing Page (15-min video) → Book Free 30-Min Store Check-Up → "
    "Pre-Call Sequence → Discovery Call → Close ₹1.2L Deal\n\n"
    "This is your MAIN funnel from Part 1. It converts cold traffic into high-ticket clients.",
    "FFF3E0"
)

# Funnel 3
add_sub_header("3️⃣ E-Book / Template Funnel → Appointment Book", PURPLE)
add_body("Best for: Building an email list of retail store owners at low CPL.")
add_highlight_box(
    "RCA APPLICATION:\n"
    "Ad → Landing Page: 'Download Free: 7 Hidden Reasons Your Retail Store is Losing Stock' → "
    "Email Opt-in → Deliver PDF → 3-Email Nurture Sequence → Invite to Book Free Check-Up\n\n"
    "Lead Magnet Ideas for RCA:\n"
    "• '7 Hidden Reasons Your Retail Store is Losing Stock (and How to Spot Them This Week)'\n"
    "• 'The Retail Store Owner's Daily Stock Checklist Template'\n"
    "• 'How One Store Owner Recovered ₹3.2 Lakhs — The Complete Case Study'",
    "E8F5E9"
)

# Funnel 4
add_sub_header("4️⃣ Webinar → Call Booking Funnel (Hybrid)", ACCENT_TEAL)
add_body("Best for: Using webinars as a qualification step before the sales call.")
add_highlight_box(
    "RCA APPLICATION:\n"
    "Ad → Webinar Registration → Live 60-Min Webinar (teach stock control concepts) → "
    "CTA: Book Your Personalized Store Check-Up → Discovery Call → Close\n\n"
    "This COMBINES the trust-building of a webinar with the closing power of a 1-on-1 call.\n"
    "Best to use once you have done 30-40 one-to-one calls and are ready to scale.",
    "FFF8E1"
)

# Funnel 5
add_sub_header("5️⃣ 1-Day / 2-Day / 3-Day Event Funnel (Free or Paid)", WARM_RED)
add_body("Best for: Establishing authority and selling at scale once RCA is proven.")
add_highlight_box(
    "RCA APPLICATION:\n"
    "Ad → Register for 'Retail Control Masterclass — 2-Day Free Event' → "
    "Day 1: Diagnose + Educate → Day 2: Case Studies + Sell ₹1.2L Package → Close\n\n"
    "MATH: Spend ₹5L ads → 10K leads → 2,500 show up → 10% convert to L2 (₹50K event) → "
    "150 attend event → 10% close at ₹1.2L = ₹1.8 Crore revenue\n\n"
    "⚠️ ADVANCED: Only attempt after you have a proven offer and 50+ sales calls under your belt.",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 6: THE WEBINAR SELLING SYSTEM
# ════════════════════════════════════════════════════════════════════
add_section_header("THE COMPLETE WEBINAR SELLING SYSTEM FOR RCA", "🎤")

add_body(
    "The webinar is divided into 3 parts. The golden rule: BE TRUTHFUL ALWAYS. "
    "People buy YOU first, before your product."
)

add_sub_header("⏱️ Webinar Time Structure for RCA")
styled_table(
    ["Session Type", "Total Duration", "Content Time", "Sales Time", "Sweet Spot"],
    [
        ["Free Webinar", "2 Hours", "First 60 min", "Last 60 min", "2 hours"],
        ["Paid Workshop", "3-4 Hours", "First 90-120 min", "Last 60 min", "3-4 hours"],
    ],
    col_widths=[1.2, 1.0, 1.0, 1.0, 1.0]
)

# Part 1 — Intro
add_sub_header("PART 1: WEBINAR INTRO (10-12 Minutes)", ACCENT_GOLD)
add_minor_header("Building Authority & Setting Context")

add_numbered("Build Authority: Show Aakash's retail consulting credentials and store results.", 1)
add_numbered("Show photos with recognized retail industry leaders or business coaches.", 2)
add_numbered("Show achievements: ₹3.2L recovered, 17%→2% mismatch, stores helped across India.", 3)
add_numbered("Do NOT tell a rags-to-riches story — it doesn't work anymore. Show PROOF instead.", 4)
add_numbered("Keep the intro between 10-12 minutes maximum.", 5)

add_minor_header("Setting Expectations & Hook")
add_numbered("Preview the content: 'In the next 2 hours, you'll learn the exact 5-step system I use to find ₹2-5 Lakhs of hidden stock loss in ANY retail store.'", 6)
add_numbered("Share industry stats: '87% of clothing store owners in India are losing money they don't even know about.'", 7)
add_numbered("Why stay till the end: 'At the end, I'm giving away my Staff Control Rulebook™ worth ₹25,000 — but only to those who stay.'", 8)
add_numbered("Show them what they'll get as a bonus for staying.", 9)
add_numbered("The tagline 'Retail Control Architect™ — Stop the Silent Bleed' should be on the first page.", 10)

add_minor_header("Engagement & Belief Building")
add_numbered("Set the context: 'This webinar is 90 minutes. If you have an urgent meeting, you can leave — but I promise the insights in the next hour will save you lakhs.'", 11)
add_numbered("Why this topic matters NOW: 'Every day you wait, you lose ₹2,739. By the end of this webinar, you'll know EXACTLY where that money is going.'", 12)
add_numbered("'In about 30-40 minutes, there will be an Aha Moment where everything clicks.'", 13)
add_numbered("Keep them engaged: 'Type YES in the chat if this has ever happened in your store.'", 14)
add_numbered("Make them believe it's achievable: 'One store owner I worked with had the SAME problems you do. Today, his stock matches 100%.'", 15)

doc.add_page_break()

# Part 2 — Content
add_sub_header("PART 2: CONTENT DELIVERY (45-60 Minutes)", SUCCESS_GRN)
add_minor_header("The 3-Chunk Content Structure")

add_highlight_box(
    "CRITICAL: Over-deliver on content. Show them the COMPLETE path.\n\n"
    "Chunk 1 (15 min): THE PROBLEM — Why Stock Mismatch Happens\n"
    "• Walk through the 5 pain points all retail stores face\n"
    "• Show the ₹28L vs ₹24L real case story\n"
    "• 'The owner thought this was normal. It is NOT.'\n\n"
    "Chunk 2 (15 min): THE SYSTEM — The 5-Step Retail Control Framework\n"
    "• Introduce each of the 5 steps (Find Leaks, Organize Items, Staff Accountability, Weekly Tracking, Lock Process)\n"
    "• Show the process diagram — make it visual\n"
    "• 'It is NOT software. NOT your staff. It is a SYSTEM problem.'\n\n"
    "Chunk 3 (15 min): THE PROOF — Real Store Transformations\n"
    "• Walk through 2-3 detailed case studies with actual numbers\n"
    "• Before/after screenshots of stock reports\n"
    "• Store owner video testimonials (if available)",
    "E8F5E9"
)

add_minor_header("Content Engagement Techniques")
add_bullet("Use stories for better connection — walk through YOUR journey of discovering these problems", prefix_emoji="📖")
add_bullet("Show them where they ARE and where they SHOULD BE", prefix_emoji="🗺️")
add_bullet("People buy RESULTS: 'It took me 2 years to figure this out because I didn't have a mentor'", prefix_emoji="🎯")
add_bullet("Add humor — keep the energy high", prefix_emoji="😄")
add_bullet("'Type NEW in the comment box whenever you learn something new'", prefix_emoji="💬")
add_bullet("Show 6-10 testimonials WITHIN the content (not just at the end)", prefix_emoji="⭐")
add_bullet("Attention is the new currency — HOLD their attention with real examples", prefix_emoji="👁️")

add_highlight_box(
    "KEY MINDSET:\n"
    "'You have already lost so many opportunities because of missing stock control skills.\n"
    "The amount you'll invest in this system is NOTHING compared to what you're already losing every month.'\n\n"
    "Make them FEEL the cost of inaction BEFORE you pitch.",
    "FFF3E0"
)

# Part 3 — Sales Pitch
add_sub_header("PART 3: SALES PITCH (Final 60 Minutes)", WARM_RED)
add_minor_header("The Natural Transition")
add_highlight_box(
    "CRITICAL RULE: The transition from content to promotion MATTERS A LOT.\n"
    "People should NOT feel like you suddenly started selling.\n"
    "Be in a NATURAL state. The pitch should feel like a continuation of the help.\n\n"
    "TRANSITION SCRIPT:\n"
    "'So now you understand the 5-step system. You've seen the results. You know this works.\n"
    "The question is — who is going to implement this IN YOUR STORE?\n"
    "That's exactly what I want to talk to you about next...'",
    "FFF8E1"
)

add_minor_header("Sales Sequence")
add_numbered("Introduce the product: Show ALL deliverables of the Retail Control System.", 1)
add_numbered("Show testimonials immediately after deliverables.", 2)
add_numbered("Ask: 'Based on everything I've shown you, what do you think this should cost?'", 3)
add_numbered("Strategic price reveal: Total value ₹4L+ → Your investment: ₹1.2L → That's ₹109/day.", 4)
add_numbered("Content & offer must MATCH — don't sell something unrelated to what you taught.", 5)

add_minor_header("Creating Urgency & Breaking Objections")
add_bullet("'People don't buy products. People buy a BETTER VERSION of themselves.'", prefix_emoji="💡")
add_bullet("'Not even God can teach you everything in 2 hours. That's why the system exists.'", prefix_emoji="🙏")
add_bullet("Tell them the problems they're facing — make them NOD along.", prefix_emoji="🎯")

add_minor_header("Objection Handling in the Webinar")
add_body("Write down and address ALL objections they're thinking:")
add_numbered("'I'm not a good fit' → 'If you own a retail store with stock, this system works for you.'", 1)
add_numbered("'I don't have time' → 'The system takes 10 minutes a day. How many hours do you spend worrying now?'", 2)
add_numbered("'I already tried software' → 'Software ≠ System. We install the PROCESS, not just the tool.'", 3)

add_minor_header("Self-Close Questions")
add_body("Ask these questions to let them close THEMSELVES:")
add_bullet("'Why do you want to fix your stock mismatch problem?'", prefix_emoji="❓")
add_bullet("'How much are you losing every month because of this?'", prefix_emoji="❓")
add_bullet("'When was the last time you spent ₹1.2L on something for your business?'", prefix_emoji="❓")
add_bullet("'Don't you think your store DESERVES to run without chaos?'", prefix_emoji="❓")
add_bullet("'You want your money sitting as dead stock, or generating growth?'", prefix_emoji="❓")
add_bullet("Break the pricing barrier: '₹1,20,000 ÷ 365 = ₹328/day. That's less than your daily chai budget.'", prefix_emoji="💰")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 7: BEFORE & AFTER WEBINAR PLAYBOOK
# ════════════════════════════════════════════════════════════════════
add_section_header("BEFORE & AFTER WEBINAR — EXECUTION PLAYBOOK", "📋")

add_sub_header("🟩 BEFORE THE WEBINAR — Maximize Show-Up Rate", SUCCESS_GRN)
styled_table(
    ["#", "Action", "RCA Application"],
    [
        ["1", "Create hype around the masterclass", "Teaser posts: 'I helped a store recover ₹3.2L — want to know how?'"],
        ["2", "Topic determines show-up rate", "Use: 'Stop Silent Stock Loss' (not generic 'retail tips')"],
        ["3", "Multiple reminders", "Calls + SMS + Email + WhatsApp + Automated calls"],
        ["4", "Add everyone to WhatsApp group", "Even if manually — this is the highest engagement channel"],
        ["5", "Pre-framing (Soap Opera Sequence)", "3-day email story sequence building anticipation"],
        ["6", "Share success stories before webinar", "Send 1 case study per day leading up to the webinar"],
        ["7", "Keep messaging during webinar", "WhatsApp group messages for first 1 hour of webinar"],
        ["8", "Show-up & stay bonuses", "'Staff Control Rulebook™ free to anyone who stays till the end'"],
        ["9", "Educate about YOU before webinar", "Short bio video: 'Who is Aakash Savant?'"],
        ["10", "Social media on Thank You page", "All links to Instagram, YouTube, LinkedIn on TY page"],
        ["11", "Small commitment before webinar", "Pre-webinar form: 'What's your biggest stock frustration?'"],
        ["12", "Soap opera email sequence", "5-email story-driven sequence before the webinar date"],
    ],
    col_widths=[0.3, 2.0, 3.7]
)

add_sub_header("🟩 AFTER THE WEBINAR — Close the Remaining Sales", ACCENT_TEAL)
styled_table(
    ["#", "Action", "RCA Application"],
    [
        ["1", "Provide recording (if webinar is >1 week away)", "Send replay link within 2 hours"],
        ["2", "Feedback form", "Google Form: 'Rate today's session + biggest takeaway'"],
        ["3", "Call non-buyers", "Personal call within 24 hours to top prospects"],
        ["4", "5-day follow-up with closing angle", "Day 1: Summary | Day 2: Math | Day 3: Case study | Day 4: FAQ | Day 5: Last chance"],
        ["5", "Re-invite to next webinar", "Those who didn't buy → invite again next week"],
        ["6", "Connect on social media", "Follow-up DM with value-add content"],
        ["7", "Keep providing value", "Weekly WhatsApp broadcast with retail tips"],
    ],
    col_widths=[0.3, 2.5, 3.2]
)

add_highlight_box(
    "⚠️ LEAKAGE CHECK:\n"
    "If you get fewer sales than expected, check where the funnel is leaking:\n\n"
    "Intro = 100% attendance = 100 people\n"
    "Content section = 20% drop = 80 remaining\n"
    "Pitch starts = 10% more drop = 72 remaining\n"
    "Link drop = Final buyers = 5-8% of original\n\n"
    "Find where the DROP happens and fix THAT section first.",
    "FDEDED"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 8: WEBINAR SELLING FORMULA
# ════════════════════════════════════════════════════════════════════
add_section_header("WEBINAR SELLING FORMULA — MASTER POINTERS", "🟥")

add_body("These 11 pointers separate successful webinars from failed ones:")

add_numbered("Teaching HURTS sales in a live webinar. Don't over-teach — create DESIRE to learn more.", 1)
add_numbered("Your main purpose is to BREAK THEIR FALSE BELIEFS (not to teach them everything).", 2)
add_numbered("Value in a webinar = Teaching them HOW TO THINK, not what to do step-by-step.", 3)
add_numbered("Value stacking is SO important — show the mountain of value before revealing price.", 4)
add_numbered("Solve their queries one by one — every unanswered question is a lost sale.", 5)
add_numbered("They should BELIEVE IN YOU as a coach/architect who can deliver results.", 6)
add_numbered("They should BELIEVE IN THEMSELVES that they can implement the system.", 7)
add_numbered("They should BELIEVE IN YOUR SYSTEM & PRODUCT — that it actually works.", 8)
add_numbered("People forget INFORMATION but never forget EXPERIENCE. Make the webinar an experience.", 9)
add_numbered("In the live webinar, show them BONUSES live — open the Staff Rulebook, show the checklist.", 10)
add_numbered("Show at least 6-10 testimonials in the FIRST 60 minutes woven into the lessons.", 11)

add_highlight_box(
    "THE 3 BELIEFS YOU MUST BUILD IN EVERY WEBINAR:\n\n"
    "1. BELIEF IN YOU: 'Aakash has done this for real stores and gotten real results.'\n"
    "2. BELIEF IN THEMSELVES: 'I can fix MY store with this system.'\n"
    "3. BELIEF IN THE SYSTEM: 'The 5-step Retail Control system actually works.'\n\n"
    "If ALL 3 beliefs are established → the sale is almost automatic.",
    "E8F5E9"
)

add_sub_header("Zoom Alternatives for Budget-Conscious Start")
styled_table(
    ["Tool", "Feature", "Cost"],
    [
        ["LVRG", "Zoom-like webinar platform", "Affordable alternative"],
        ["Our-area.com (oa1.in)", "Indian webinar platform", "Budget-friendly"],
        ["Google Meet", "Free for small groups", "Free (up to 100 participants)"],
        ["Zoom Free", "40-minute limit", "Free for testing"],
    ],
    col_widths=[1.5, 2.5, 2.0]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════
# SECTION 9: IMPLEMENTATION TIMELINE
# ════════════════════════════════════════════════════════════════════
add_section_header("IMPLEMENTATION TIMELINE & ACTION STEPS", "🚀")

add_sub_header("⚠️ Important Note Before You Start", WARM_RED)
add_highlight_box(
    "1. If you are a COMPLETE FRESHER and have never conducted any webinars → "
    "Do NOT start with webinars yet.\n\n"
    "2. First complete 30-40 one-to-one sales calls using the VSL → Call Booking funnel from Part 1.\n\n"
    "3. If you are already good at giving presentations → You can start webinars now.\n\n"
    "4. In Week 6, you'll learn in-depth Platform Closing / One-to-Many selling.",
    "FFF3E0"
)

add_sub_header("Week-by-Week Implementation Plan")

add_minor_header("Week 1: Landing Page Foundation", SUCCESS_GRN)
add_numbered("Set up domain (retailcontrolarchitect.com) and hosting (Hostinger/SiteGround).", 1)
add_numbered("Build VSL landing page on Flexi-Funnel with all 14 strategies applied.", 2)
add_numbered("Write and split-test 4 headline variations.", 3)
add_numbered("Compress all images via TinyPNG. Test page speed on mobile.", 4)
add_numbered("Set up Calendly with 5 discovery call slots per week.", 5)

add_minor_header("Week 2: Traffic & Testing", ORANGE)
add_numbered("Launch Meta ads at ₹500/day targeting retail store owners.", 1)
add_numbered("Test 2 ad creatives: VSL teaser video + ₹3.2L case study image.", 2)
add_numbered("Set up LetConvert social proof pop-ups.", 3)
add_numbered("Complete all discovery calls. Note every objection raised.", 4)

add_minor_header("Week 3: Optimize & First Webinar Prep", ACCENT_TEAL)
add_numbered("Review call close rate — if below 15%, rewrite diagnostic questions.", 1)
add_numbered("Build the e-book lead magnet funnel (Funnel #3).", 2)
add_numbered("Write the webinar script using the 3-part structure.", 3)
add_numbered("Set up webinar registration page + reminder sequences.", 4)

add_minor_header("Week 4+: Scale & Webinar Launch", PURPLE)
add_numbered("Conduct first webinar (ONLY if you have 30+ sales calls completed).", 1)
add_numbered("Activate referral program for completed Rung 3 clients.", 2)
add_numbered("Set up retargeting ads (14-day window) for re-engagement.", 3)
add_numbered("Scale winning ad creatives to ₹1,000/day.", 4)
add_numbered("Begin upselling Rung 3 clients to Rung 4 (Quarterly Audit retainer).", 5)

add_spacer(12)

# ── Task Section ──────────────────────────────────────────────────
add_section_header("YOUR TASKS — DO THIS NOW", "✅")

add_highlight_box(
    "TASK 1: Rewatch the Funnels Part 2 training and internalize every concept.\n\n"
    "TASK 2: Make your payment gateway ready (Razorpay / Instamojo / PayU).\n\n"
    "TASK 3: If you're already running webinars, upgrade them with today's 11 pointers + "
    "3-part structure.\n\n"
    "TASK 4: Plan your first RCA webinar — but ONLY if you've completed 50+ sales calls.\n\n"
    "TASK 5: Build your landing page this week using Flexi-Funnel with ALL 14 strategies applied.\n\n"
    "TASK 6: Write 4 headline variations and prepare to A/B test them.",
    "E8F5E9"
)

add_spacer(12)

# Final box
add_highlight_box(
    "REMEMBER:\n\n"
    "The free store check-up funnel (Part 1) is your FOUNDATION.\n"
    "The advanced strategies in Part 2 are your AMPLIFIER.\n\n"
    "Master the 1-on-1 calls first. Then scale with webinars.\n"
    "Every store you fix becomes a case study that sells the next 10 stores.\n\n"
    "You are not selling software. You are selling CERTAINTY.\n"
    "And certainty, in the retail world, is worth every rupee. 🚀🏆",
    "FFF3E0"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NOW GO BUILD THAT FUNNEL. 🚀🏆")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = ACCENT_GOLD


# ── Save ───────────────────────────────────────────────────────────
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "rca_funnel_part2_advanced.docx"
)
doc.save(output_path)
print(f"✅ Successfully created: {output_path}")
