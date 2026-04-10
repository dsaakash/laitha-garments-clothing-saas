"""
Generate a beautifully formatted 'Old You vs New You' DOCX for Aakash Savant
- Two-column comparison table style
- Professional formatting with colors
- Personalized to Aakash's Retail Control Architect journey
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── PAGE MARGINS ─────────────────────────────────────────────
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

# ─── STYLES ─────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(50, 50, 50)

# ─── HELPER FUNCTIONS ──────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_para(cell, text, bold=False, size=11, color=None, alignment=None):
    """Add a paragraph to a cell with styling."""
    p = cell.paragraphs[0] if cell.paragraphs[0].text == '' else cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    return p

def add_bullet_to_cell(cell, text, color=None):
    """Add a bullet point paragraph to a cell."""
    p = cell.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    return p

# ═══════════════════════════════════════════════════════════
# COVER / TITLE
# ═══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('💎')
run.font.size = Pt(48)
title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2.add_run('OLD YOU vs NEW YOU')
run.bold = True
run.font.size = Pt(30)
run.font.color.rgb = RGBColor(30, 30, 30)
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('IEC Diamond — Week 1 Transformation Blueprint')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_paragraph()
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name_para.add_run('Aakash Savant')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(255, 107, 53)  # orange

role_para = doc.add_paragraph()
role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = role_para.add_run('Retail Control Architect™')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# INTRO PARAGRAPH
# ═══════════════════════════════════════════════════════════
intro = doc.add_paragraph()
run = intro.add_run('About This Document')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(30, 30, 30)

doc.add_paragraph(
    'This document maps the transformation from who you are today to who you are becoming. '
    'The "Old You" is not someone to hate — it\'s someone who got you here. '
    'The "New You" is the version you are deliberately building through the IEC Diamond framework. '
    'Use this as a mirror, a compass, and a daily reminder.\n\n'
    'Every great transformation starts with honest self-awareness.'
)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# COMPARISON SECTIONS
# ═══════════════════════════════════════════════════════════

OLD_COLOR = 'F8D7DA'  # light red
NEW_COLOR = 'D4EDDA'  # light green
HEADER_OLD = 'C0392B'  # dark red
HEADER_NEW = '27AE60'  # dark green
SECTION_BG = 'F0F0F0'  # light gray

comparisons = [
    {
        'section': '1. Identity & Purpose',
        'question': 'Who are you and what do you do?',
        'old': [
            'Software engineer and product technologist with 5 years of backend experience',
            'Expert in Python, FastAPI, Django, SQLAlchemy, PostgreSQL',
            'Founder of Nirvriksh — helped startups with tech build',
            'Skilled builder, but stuck in misaligned environments',
            'Knew how to build software, but not how to build a business around it',
        ],
        'new': [
            'Retail Control Architect — helping clothing store owners stop losing money',
            'Built a complete system: VSL funnel, admin panel, lead capture, CRM, pitch deck',
            'Uses tech skills to CREATE a business, not just work in someone else\'s',
            'Founder of a consulting practice with a clear niche and offer',
            'Identity = problem solver for retail store owners, not a generic developer',
        ],
    },
    {
        'section': '2. Character Traits',
        'question': 'How would you describe your character?',
        'old': [
            'Direct, analytical, systems-thinker — but sometimes too internal',
            'Valued "uncomfortable truth" but didn\'t always share it at the right time',
            'Builder who preferred depth over speed — sometimes at the cost of action',
            'High standards that occasionally led to overthinking and delay',
            'Wary of hype — which sometimes meant being slow to adopt what works',
        ],
        'new': [
            'Still direct and analytical — but now action-biased',
            'Uses truth to build trust with store owners in simple language',
            'Builder who ships: landing pages, pitch decks, CRM — all done',
            'High standards applied to EXECUTION speed, not just quality of thought',
            'Wary of hype, but embraces proven frameworks (GAP, Grand Slam Offer)',
        ],
    },
    {
        'section': '3. Reason For Existence',
        'question': 'Why do you do what you do?',
        'old': [
            'Wanted to build meaningful technology that solved real problems',
            'Cared about clarity, discipline, and ownership',
            'But direction was unclear — no specific niche or dream outcome',
            'Energy spent compensating for broken systems in misaligned jobs',
            'Felt capable but underutilized',
        ],
        'new': [
            'Exists to help store owners stop losing lakhs in stock leakage',
            'Purpose is crystal clear: bring big-retail discipline to small stores',
            'Uses tech + business frameworks to create real financial impact',
            'Every action is aligned: content, outreach, sales, delivery — all connected',
            'Feels alive because the work MATTERS to real people',
        ],
    },
    {
        'section': '4. Skills & Talents',
        'question': 'What skills do you possess?',
        'old': [
            'Backend engineering: Python, Django, FastAPI, PostgreSQL',
            'AI/ML integration and clean architecture design',
            'Problem framing, documentation, technical communication',
            'Could explain complex systems in simple terms',
            'All skills pointed INWARD (building for others)',
        ],
        'new': [
            'ALL old skills retained + NEW business skills added',
            'Sales: cold calling scripts, WhatsApp outreach, closing',
            'Marketing: GAP Framework content, VSL funnels, lead magnets',
            'Offer design: Grand Slam Offer with 30-day guarantee',
            'Skills now point OUTWARD (building YOUR OWN business)',
        ],
    },
    {
        'section': '5. Environment',
        'question': 'Where do you work? How does it look?',
        'old': [
            'Quiet, minimal home workspace — good for coding',
            'Worked alone most of the time',
            'Environment designed for focus, but sometimes for isolation',
            'Controlled inputs, but didn\'t have structured business routines',
        ],
        'new': [
            'Same minimal workspace — but now with SYSTEMS on the wall',
            'Daily tracking sheet, content calendar, outreach CRM on screen',
            'Work includes going OUT — visiting stores, meeting owners, presenting',
            'Environment designed for both deep work AND business execution',
        ],
    },
    {
        'section': '6. Biggest Problem',
        'question': 'What\'s the biggest challenge you face?',
        'old': [
            'Misalignment between capability and environment',
            'Surrounded by roles that underutilized his skills',
            'Spent energy fixing broken systems instead of building his own',
            'Knew he could do more but didn\'t know the VEHICLE',
            'Dealt with this for 2-3 years',
        ],
        'new': [
            'Challenge is now GROWTH, not survival',
            'How to get 5 paying clients in 90 days',
            'How to build personal brand so leads come to you',
            'How to deliver results consistently at scale',
            'Problems are exciting, not draining',
        ],
    },
    {
        'section': '7. Daily Routine',
        'question': 'How does your day look?',
        'old': [
            'Morning: thinking, learning, deep technical work',
            'Midday: meetings, collaboration, reviews',
            'Evening: reflection, lighter work, content consumption',
            'Structured but not REVENUE-producing',
            'Missing: outreach, marketing, sales activities',
        ],
        'new': [
            'Morning: 📖 Reading + 🧘 Meditation + 📝 Planning',
            'Deep Work Block: 💻 3 hours on business-building activities',
            'Midday: 📞 Outreach calls + WhatsApp DMs + store visits',
            'Afternoon: 🎬 Content creation (GAP Framework)',
            'Evening: 📊 Review tracking sheet + 📝 Daily reflection',
            'EVERY hour is tied to a revenue-producing activity',
        ],
    },
    {
        'section': '8. Habits',
        'question': 'What are your good and bad habits?',
        'old_title': '🟥 Old Bad Habits',
        'new_title': '🟩 New Good Habits',
        'old': [
            'Overthinking before executing',
            'Staying too long in misaligned situations',
            'Impatient with slow thinkers (hurt relationships)',
            'Consuming too much content without acting',
            'Comparing timelines with others',
        ],
        'new': [
            'Execute first, refine later — shipping beats perfecting',
            'Exit misalignment within 1 week, not 1 year',
            'Patient with store owners — explains in layman\'s terms',
            'Content CREATION > content consumption (3:1 ratio)',
            'Track OWN metrics only: calls made, meetings booked, revenue',
        ],
    },
    {
        'section': '9. Money & Spending',
        'question': 'Where does your money go?',
        'old': [
            'Essentials, learning tools, books, hardware',
            'Selective health/travel spending',
            'Waste: unused subscriptions, convenience spending, tool hoarding',
            'Spending to compensate for lack of direction',
        ],
        'new': [
            'Investing in BUSINESS tools: CRM, hosting, domains',
            'Spending on marketing: ad creatives, content tools',
            'Revenue reinvested: every ₹1 spent must generate ₹5+',
            'Eliminated: aimless tool shopping, unused subscriptions',
            'New rule: no spend without tracking ROI',
        ],
    },
    {
        'section': '10. Myers-Briggs & Strengths',
        'question': 'What is your personality type?',
        'old': [
            'INTJ-A: Analytical, internally driven, future-oriented',
            'Comfortable working independently',
            'Decisive once logic is clear — but sometimes too rigid',
            'Low tolerance for ambiguity and inefficiency',
        ],
        'new': [
            'Still INTJ-A — but now channeled into SALES & BUSINESS',
            'Uses analytical thinking to design irresistible offers',
            'Independence = perfect for running own consulting practice',
            'Decisiveness applied to: "Is this store owner a fit? Yes/No → Act"',
            'Low tolerance for inefficiency → helps stores fix their chaos',
        ],
    },
    {
        'section': '11. Revenue Goals (Chunked Down)',
        'question': 'What\'s your big audacious goal broken into steps?',
        'old': [
            'Goal: ₹20L/month in profit (vague SaaS dream)',
            'Need 20 customers at ₹1L/month each',
            'Need 800 inbound leads from 16,000 monthly viewers',
            'No clear niche or vehicle to get there',
        ],
        'new': [
            'Goal: ₹3L/month in 6 months (specific & achievable)',
            'Price: Audit is FREE → Implementation package = ₹50K-1L',
            'Need: 3-5 paying clients',
            'Pipeline: 25 calls/month at 20% close rate = 5 clients',
            'Lead source: WhatsApp outreach + content + store visits',
            'Vehicle: Retail Control Architect consulting',
        ],
    },
    {
        'section': '12. Ownership vs Blame',
        'question': 'Do you take ownership or blame circumstances?',
        'old': [
            'Took ownership intellectually — but stayed in bad situations too long',
            'Understood the concept but didn\'t always ACT on it',
            'Sometimes blamed "environment" for slow progress',
        ],
        'new': [
            'Takes ownership by DEFAULT — no exceptions',
            'If a store visit fails → "What did I say wrong? Fix the script."',
            'If no leads → "Am I doing enough outreach? Increase volume."',
            'No complaints. Only adjustments.',
        ],
    },
]

for i, comp in enumerate(comparisons):
    # Section Header
    header = doc.add_paragraph()
    run = header.add_run(comp['section'])
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(30, 30, 30)

    q_para = doc.add_paragraph()
    run = q_para.add_run(f'❓ {comp["question"]}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True

    # Two-column table
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for cell in table.columns[0].cells:
        cell.width = Cm(8.5)
    for cell in table.columns[1].cells:
        cell.width = Cm(8.5)

    # Header row
    old_title = comp.get('old_title', '🟥 OLD YOU')
    new_title = comp.get('new_title', '🟩 NEW YOU')
    
    hdr_old = table.rows[0].cells[0]
    set_cell_shading(hdr_old, HEADER_OLD)
    p = hdr_old.paragraphs[0]
    run = p.add_run(old_title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr_new = table.rows[0].cells[1]
    set_cell_shading(hdr_new, HEADER_NEW)
    p = hdr_new.paragraphs[0]
    run = p.add_run(new_title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Content row
    content_row = table.add_row()
    old_cell = content_row.cells[0]
    new_cell = content_row.cells[1]
    set_cell_shading(old_cell, OLD_COLOR)
    set_cell_shading(new_cell, NEW_COLOR)

    for item in comp['old']:
        add_bullet_to_cell(old_cell, item, color=(180, 50, 50))

    for item in comp['new']:
        add_bullet_to_cell(new_cell, item, color=(30, 130, 60))

    # Remove first empty paragraph
    if old_cell.paragraphs[0].text == '':
        old_cell.paragraphs[0]._element.getparent().remove(old_cell.paragraphs[0]._element)
    if new_cell.paragraphs[0].text == '':
        new_cell.paragraphs[0]._element.getparent().remove(new_cell.paragraphs[0]._element)

    # Table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="6" w:color="333333"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════
# CLOSING SECTION
# ═══════════════════════════════════════════════════════════
doc.add_page_break()

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run('💎 THE TRANSFORMATION SUMMARY')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(255, 107, 53)

doc.add_paragraph()

# Summary table
summary = doc.add_table(rows=1, cols=3)
summary.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
cats = ['AREA', 'OLD AAKASH', 'NEW AAKASH']
colors = ['333333', HEADER_OLD, HEADER_NEW]
for j, (cat, clr) in enumerate(zip(cats, colors)):
    cell = summary.rows[0].cells[j]
    set_cell_shading(cell, clr)
    p = cell.paragraphs[0]
    run = p.add_run(cat)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

summary_data = [
    ('Identity', 'Software Engineer', 'Retail Control Architect'),
    ('Niche', 'None (generic tech)', 'Clothing stores with stock mismatch'),
    ('Offer', 'No offer', 'Free audit → ₹50K-1L implementation'),
    ('Revenue Goal', '₹20L/mo (vague)', '₹3L/mo in 6 months (clear)'),
    ('Marketing', 'None', 'GAP framework + daily outreach'),
    ('Sales', 'None', '15-slide pitch + closing scripts'),
    ('Content', 'Consumer', 'Creator (24 ideas ready)'),
    ('Tracking', 'Mental notes', 'Daily tracking sheet (365 days data)'),
    ('CRM', 'None', 'Built-in outreach CRM'),
    ('Mindset', 'Overthink → delay', 'Execute → refine → repeat'),
]

for area, old, new in summary_data:
    row = summary.add_row()
    cells = row.cells
    p = cells[0].paragraphs[0]
    run = p.add_run(area)
    run.bold = True
    run.font.size = Pt(10)
    set_cell_shading(cells[0], 'F5F5F5')

    p = cells[1].paragraphs[0]
    run = p.add_run(old)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 50, 50)
    set_cell_shading(cells[1], 'FFF5F5')

    p = cells[2].paragraphs[0]
    run = p.add_run(new)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(30, 130, 60)
    set_cell_shading(cells[2], 'F0FFF0')

# Table borders for summary
tbl = summary._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="4" w:color="CCCCCC"/>'
    '  <w:left w:val="single" w:sz="4" w:color="CCCCCC"/>'
    '  <w:bottom w:val="single" w:sz="4" w:color="CCCCCC"/>'
    '  <w:right w:val="single" w:sz="4" w:color="CCCCCC"/>'
    '  <w:insideH w:val="single" w:sz="4" w:color="CCCCCC"/>'
    '  <w:insideV w:val="single" w:sz="4" w:color="CCCCCC"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

doc.add_paragraph()

# Final quote
quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = quote.add_run('"You don\'t rise to the level of your goals.\nYou fall to the level of your systems."')
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('— James Clear, Atomic Habits')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('IEC Diamond 💎 Week 1 — Transformation Complete\nAakash Savant | Retail Control Architect™')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(180, 180, 180)

# ─── SAVE ─────────────────────────────────────────────────
output_path = '/Users/aakash/Desktop/Week_3/Week1 /Old_You_vs_New_You_Aakash_Formatted.docx'
doc.save(output_path)
print(f'✅ Saved: {output_path}')
