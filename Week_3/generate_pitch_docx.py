#!/usr/bin/env python3
"""
Generate a professionally formatted .docx for
The 30-Day Stock Certainty System™ — Pitch-Ready Sales Toolkit
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour palette ──────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL = RGBColor(0x17, 0xA2, 0xB8)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY  = RGBColor(0xF2, 0xF2, 0xF2)
MED_GRAY    = RGBColor(0x66, 0x66, 0x66)
BLACK       = RGBColor(0x33, 0x33, 0x33)
RED_ACCENT  = RGBColor(0xC0, 0x39, 0x2B)
GREEN_OK    = RGBColor(0x27, 0xAE, 0x60)

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BLACK

# ── Helper functions ────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Apply background colour to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_title_block(text, subtitle=None):
    """Large centred title with optional subtitle."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(36)
    p.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.space_after = Pt(18)
        run2 = p2.add_run(subtitle)
        run2.italic = True
        run2.font.size = Pt(13)
        run2.font.color.rgb = ACCENT_GOLD

def add_section_heading(text):
    """Major section heading with gold underline feel."""
    doc.add_paragraph()  # spacer
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_BLUE

    # thin gold line
    line = doc.add_paragraph()
    line.space_after = Pt(10)
    r = line.add_run("━" * 72)
    r.font.size = Pt(6)
    r.font.color.rgb = ACCENT_GOLD

def add_sub_heading(text):
    """Sub-section heading."""
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = ACCENT_TEAL

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_quote(text):
    """Indented italic quote block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = MED_GRAY

def add_script(text):
    """Indented script / verbatim block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    run.font.name = 'Calibri'

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        # clear default and re-add
        p.clear()
        p.add_run(text).font.size = Pt(11)
    return p

def add_numbered(text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.color.rgb = ACCENT_TEAL
    p.add_run(text).font.size = Pt(11)

def styled_table(headers, rows, col_widths=None):
    """Create a styled table with dark header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1B2A4A")

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, "F7F9FC")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacer
    return table


# ════════════════════════════════════════════════════════════════════
#                        DOCUMENT CONTENT
# ════════════════════════════════════════════════════════════════════

# ── COVER / TITLE ──────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
add_title_block(
    "THE 30-DAY STOCK CERTAINTY SYSTEM™",
    "Pitch-Ready Sales Toolkit"
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"I fix stock mismatch in clothing stores in 30 days — permanently."')
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = MED_GRAY

doc.add_paragraph()

# Key facts box as a single-row table
info_table = doc.add_table(rows=1, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_items = [
    ("TARGET", "Owner-Operated\nClothing Stores\nTier 2 & 3 India"),
    ("REVENUE", "₹30L – 1.5Cr\nAnnual Revenue\n₹20L+ Inventory"),
    ("INVESTMENT", "₹85,000\nOne-Time Setup"),
    ("CAPACITY", "3 Stores/Month\nCycles Start\non the 1st Only"),
]
for i, (label, val) in enumerate(info_items):
    cell = info_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + "\n")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = ACCENT_GOLD
    r2 = p.add_run(val)
    r2.font.size = Pt(9)
    r2.font.color.rgb = DARK_BLUE
    set_cell_shading(cell, "F7F9FC")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 1: THREE PITCH VERSIONS
# ══════════════════════════════════════════════════════════════════
add_section_heading("SECTION 1: THE THREE PITCH VERSIONS")
add_body("Use the right version for the right situation. Always.")

add_sub_heading("VERSION 1: One Sentence — Bio, Introductions, Casual")
add_quote("I fix stock mismatch in clothing stores in 30 days — permanently.")

add_sub_heading("VERSION 2: 30 Seconds — Networking, Cold Outreach, Referrals")
add_quote(
    "I run The Stock Certainty System™ for clothing stores. We come in, measure your exact "
    "stock mismatch in rupees, fix your inventory structure, lock your supplier and sales entries, "
    "remove every parallel system, and enforce discipline for 30 days straight. At the end, your "
    "system stock matches your shop, and you can verify everything in 10 minutes flat — without "
    "asking staff. Only 3 stores per month. Next cycle starts on the 1st."
)

add_sub_heading("VERSION 3: Full Positioning Statement — Sales, Website, Proposals")
add_script(
    "\"I don't sell billing software. I install The Stock Certainty System™ inside clothing stores.\n\n"
    "In 30 days: your system stock matches your physical stock, reconciliation drops under "
    "15 minutes, and you stop depending on staff memory for control.\n\n"
    "Only 3 stores per month. Next cycle starts on the 1st.\n\n"
    "If your stock already matches perfectly — don't work with me.\""
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 2: SALES MEETING SCRIPT
# ══════════════════════════════════════════════════════════════════
add_section_heading("SECTION 2: THE SALES MEETING SCRIPT (15–20 MIN)")
add_body("Goal: Get the \"YES\" for the 30-Day Installation.")
add_body("Rule: Never pitch before diagnosis. Never quote price before showing leakage.")

# Step 1
add_sub_heading("STEP 1: THE OPENING — Set the \"Doctor\" Frame (2 min)")
add_script(
    "\"Look, [Owner Name], I work with clothing store owners who are doing the revenue but "
    "missing the control. They're losing ₹2–6 Lakhs a year in silent stock leakage — money "
    "that walks out the door every day without making a sound.\n\n"
    "Most people think it's theft. It's not. It's a Structure Gap.\n\n"
    "You're working hard, but your store is running on 'adjustments' instead of 'control.' "
    "I'm not here to demo software — I'm here to install a system that makes your system stock "
    "match your shop in 30 days. Permanently.\""
)

# Step 2
add_sub_heading("STEP 2: BUILD THE DREAM + AGITATE CHAOS (5 min)")
add_body("Start with their dream:")
add_script(
    "\"Before we get into anything — tell me what a perfect store day looks like for you. "
    "What would need to be true for you to walk in, feel completely in control, and focus on "
    "growing rather than firefighting?\""
)
add_body("→ Let them describe their own dream. Reflect it back. Trust is instant.")

add_body("Then diagnose:")
add_script(
    "\"I need to see if your store is even fixable for our next 30-day cycle. Let me ask "
    "you three things:\"\n\n"
    "  1. \"If I pick a random Blue Shirt from that rack right now, will your computer show the exact count?\"\n\n"
    "  2. \"Do you still use a manual register or diary because you don't fully trust the software?\"\n\n"
    "  3. \"If your head staff member leaves tomorrow, do you lose the 'brain' of your store?\""
)
add_body("→ Every \"yes\" is a confirmation of chaos. Let them hear it from their own mouth.")

# Step 3
add_sub_heading("STEP 3: RUN THE DIAGNOSIS — Hidden Leakage Exposure Report™ (3 min)")
add_script(
    "\"Before I tell you anything about what I do — let me show you what is actually "
    "happening in your store right now.\""
)
add_body("→ Run the Control Gap Audit live. Build the Hidden Leakage Report™ in front of them.")

add_body("Quick Calculation (use during the meeting):")
styled_table(
    ["Input Field", "Value", "Formula / Note"],
    [
        ["Annual Revenue", "₹ __________", "(Enter their turnover)"],
        ["Est. Mismatch %", "____%", "(From 3-item spot check)"],
        ["Annual Leakage (₹)", "₹ __________", "(Revenue × Mismatch %)"],
        ["Monthly Leakage (₹)", "₹ __________", "(Annual ÷ 12)"],
        ["Daily \"Chaos Tax\" (₹)", "₹ __________", "(Monthly ÷ 30.4)"],
        ["Cost of System", "₹85,000", "(Your Setup Fee)"],
        ["System Payback Time", "____ Days", "(₹85,000 ÷ Daily Tax)"],
    ],
    col_widths=[2.0, 1.5, 2.5]
)
add_body("Example: ₹1Cr revenue × 10% mismatch = ₹10L annual leakage. Monthly = ₹83,333. Daily = ₹2,739. System pays for itself in ~31 days.")

# Step 4
add_sub_heading("STEP 4: PRESENT THE PROCESS — The 30-Day Installation (3 min)")
add_script("\"Here is exactly what we do over 30 days — week by week.\"")

styled_table(
    ["Week", "Phase", "Key Activities"],
    [
        ["WEEK 1", "THE FORENSIC\nFOUNDATION",
         "Control Gap Audit • Hidden Leakage Report™ delivered\n"
         "Inventory Foundation Reset begins: Clean SKU logic, categories, sizes\n"
         "Remove all duplicate codes and ghost inventory"],
        ["WEEK 2", "THE STRUCTURAL\nLOCK",
         "Inventory Reset completed\n"
         "Supplier Entry Lock™ activated: No stock enters without system entry\n"
         "Sales Deduction Lock™ activated: Every sale deducts in real time\n"
         "All parallel manual systems physically removed"],
        ["WEEK 3", "THE DISCIPLINE\nLOCK",
         "Single System Rule fully enforced\n"
         "Owner Dashboard activated with real-time views\n"
         "10-Minute Daily Control Routine™ trained and installed\n"
         "Compliance monitoring begins with daily tracking"],
        ["WEEK 4", "THE VALIDATION",
         "Full compliance review\n"
         "Mismatch remeasured against original baseline\n"
         "Before & After Validation Report™ delivered\n"
         "Expansion Readiness Scorecard™ completed"],
    ],
    col_widths=[0.8, 1.5, 3.7]
)

add_script("\"Most stores feel the difference within the first 14 days.\"")
add_body("→ 30 days becomes 14 psychologically. Same timeline. Higher perceived speed.")

# Step 5
add_sub_heading("STEP 5: MINIMIZE EFFORT — The Owner's Simple Role (1 min)")
add_script("\"Your job is three rules. We handle everything else.\"")
add_numbered("Follow the entry protocols we define.", 1)
add_numbered("Stop all parallel systems permanently.", 2)
add_numbered("Enforce the rules with your staff.", 3)
add_script("\"Your daily ongoing effort after installation? 10 minutes. That's it.\"")

# Step 6
add_sub_heading("STEP 6: THE GRAND SLAM OFFER — The Value Stack (3 min)")
add_script("\"You aren't paying for a tool. You're paying to stop the bleeding.\"")

styled_table(
    ["#", "Component", "Value"],
    [
        ["1", "Control Gap Audit — exact mismatch % + annual rupee cost", "₹5,000"],
        ["2", "Inventory Foundation Reset — SKU cleanup, categories, sizes", "₹15,000"],
        ["3", "Supplier Entry Lock™ — no stock without system entry", "₹10,000"],
        ["4", "Sales Deduction Lock™ — real-time automatic inventory", "₹10,000"],
        ["5", "Single System Enforcement — remove all parallel systems", "₹8,000"],
        ["6", "30-Day Compliance Monitoring — weekly calls + tracking", "₹20,000"],
        ["7", "Before & After Validation Report™ — measurable proof", "₹7,000"],
        ["🎁", "BONUS: Hidden Leakage Exposure Report™", "₹5,000"],
        ["🎁", "BONUS: Staff Control Rulebook™", "₹3,000"],
        ["🎁", "BONUS: 10-Minute Daily Control Routine™", "₹2,000"],
        ["🎁", "BONUS: 90-Day Stability Audit™", "₹8,000"],
        ["🎁", "BONUS: Expansion Readiness Scorecard™", "₹4,000"],
        ["", "TOTAL VALUE", "₹97,000"],
        ["", "YOUR INVESTMENT (one-time setup)", "₹85,000"],
    ],
    col_widths=[0.4, 4.0, 0.9]
)

add_body("Investment Logic (share with prospect):")
add_bullet("10% mismatch on ₹40L inventory = ₹4L in uncontrolled stock annually")
add_bullet("Even a 50% correction in Year 1 recovers far more than ₹85,000")
add_bullet("If the Hidden Leakage Report shows less than ₹85,000 in annual impact — we tell you honestly and don't ask for your investment")

# Step 7
add_sub_heading("STEP 7: THE GUARANTEE — The Stock Certainty Guarantee™ (1 min)")
add_script(
    "\"Here is my promise:\n\n"
    "If after full 30-day compliance your stock mismatch does not reduce measurably — "
    "I continue working with you at zero additional fee until control is achieved.\n\n"
    "I don't win unless you win.\""
)
add_body("→ Removes final resistance. Signals complete confidence in your method.")

# Step 8
add_sub_heading("STEP 8: THE CLOSE (1 min)")
add_script(
    "\"Based on your leakage numbers — does investing ₹85,000 once to stop ₹[their number] "
    "in annual leakage make financial sense to you?\""
)
add_body("→ Then STOP TALKING. Let them answer. Silence closes.")
add_body("If they hesitate:")
add_script(
    "\"I only have 3 slots for [Next Month] starting on the 1st. Should we lock your store in?\""
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 3: SOCIAL MEDIA TEMPLATES
# ══════════════════════════════════════════════════════════════════
add_section_heading("SECTION 3: SOCIAL MEDIA ANNOUNCEMENT TEMPLATES")
add_body("Your announcement should NOT look like an ad. It should look like a Public Service Announcement for retailers.")

add_sub_heading('OPTION 1: The "Profit Leak" Post — LinkedIn / Facebook')
add_body("HEADLINE:")
p = add_body("Most clothing store owners are paying a \"Chaos Tax\" of ₹30,000 to ₹50,000 every single month.")
p.runs[0].bold = True

add_body("BODY:")
add_script(
    "They see the sales. They see the customers. But when they look at the racks, the numbers don't add up.\n\n"
    "❌ The Symptom: System says 10 units, rack has 7.\n"
    "❌ The Fix (Wrong): Buying another \"easy\" billing software.\n"
    "❌ The Result: The same 8–15% mismatch, just on a prettier screen.\n\n"
    "Software doesn't fix a store. Discipline and Structure do.\n\n"
    "I am opening 3 slots for [Month] for The 30-Day Stock Certainty System™.\n"
    "We don't just give you a dashboard — we come in and:\n\n"
    "→ Audit your forensic leakage in rupees.\n"
    "→ Lock your supplier entry protocols.\n"
    "→ Enforce a \"Single System\" rule for your staff.\n"
    "→ Validate that your stock matches your shop.\n\n"
    "The Guarantee: If your mismatch doesn't reduce in 30 days, I work for free until it does.\n\n"
    "⚠️ I only work with 3 stores per month. Implementation starts on the 1st only.\n\n"
    "👇 Comment \"AUDIT\" or DM me to see if your store is eligible for a Forensic Leakage Report."
)

add_sub_heading('OPTION 2: The "Urgency" Post — Instagram / WhatsApp Status')
add_body("Image: Photo of a messy stockroom OR a clean dashboard showing mismatch.")
add_script(
    "🛑 STOP LEAKING PROFIT. 🛑\n\n"
    "If your system stock doesn't match your physical stock, you are losing ₹2,000–₹5,000 every single day.\n\n"
    "It's not a software problem. It's a Control Problem.\n\n"
    "I'm looking for 3 clothing store owners who want to fix their stock mismatch permanently in the next 30 days.\n\n"
    "✅ Goal: 100% Stock Accuracy + 15-Min Daily Close\n"
    "✅ Method: The 30-Day Stock Certainty System™\n"
    "✅ Guarantee: No measurable result = I work for free\n\n"
    "Current Status for [Month]:\n"
    "  Slot 1: [AVAILABLE / FILLED]\n"
    "  Slot 2: [AVAILABLE / FILLED]\n"
    "  Slot 3: [AVAILABLE / FILLED]\n\n"
    "Reply \"CONTROL\" to book your 10-minute Forensic Audit."
)

add_sub_heading("Strategy: The Multi-Channel Blast")
add_numbered('Send WhatsApp DM "Hook" to 10 store owners', 1)
add_numbered("Post Long-Form on LinkedIn/Facebook", 2)
add_numbered("Post Short-Form on WhatsApp Status + Instagram Story", 3)

add_body('WHEN SOMEONE REPLIES "What is the price?":')
add_script(
    "→ DO NOT give the price.\n"
    "→ Say: \"It depends on your store's leakage profile. Let me do a 10-minute Forensic Audit first. Are you at the store at 2 PM?\""
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 4: RE-ENGAGEMENT SCRIPTS
# ══════════════════════════════════════════════════════════════════
add_section_heading("SECTION 4: RE-ENGAGEMENT & FOLLOW-UP SCRIPTS")
add_body("Rule: Never follow up with \"Did you decide yet?\" Follow up with VALUE + FOMO.")

add_sub_heading('SCRIPT 1: The "Case Study" Follow-Up (7–10 days after going cold)')
add_script(
    "\"Hi [Owner Name], hope the week is going well.\n\n"
    "I just finished the Week 2 review for one of the stores we onboarded this month. "
    "We found that their Supplier Entry Leakage was actually 4% higher than they estimated "
    "— basically found them an extra ₹15,000 in 'lost' profit in just 14 days.\n\n"
    "It reminded me of your [Specific Item/Issue] we discussed. I'd hate for you to keep "
    "losing that margin while you 'wait for the right time.'\n\n"
    "I have one slot left for the next cycle starting on the 1st. Should we lock it in for you?\""
)

add_sub_heading('SCRIPT 2: The "Financial Penalty" Follow-Up (For the logical price-staller)')
add_script(
    "\"Hey [Owner Name], was just looking over your Audit notes.\n\n"
    "Based on your [X]% mismatch, your store has likely leaked another ₹12,000–₹15,000 "
    "since we spoke last week.\n\n"
    "My goal is to turn that 'Chaos Tax' into profit for you. If the setup fee was the only "
    "thing holding you back, let me know — I'm happy to discuss a 2-part payment plan for "
    "this final slot so we can stop the leakage today.\n\n"
    "What do you think?\""
)

add_sub_heading('SCRIPT 3: The "Expansion" Follow-Up (For the second-store planner)')
add_script(
    "\"Hi [Name], saw your latest collection post — looks great!\n\n"
    "Quick thought: I know you're planning for Store #2. I've seen 3 retailers this year "
    "try to scale before fixing their Stock Structure, and it usually led to a 200% increase "
    "in stress, not profit.\n\n"
    "Let's fix the foundation of Store #1 this month so Store #2 is a 'Copy-Paste' success. "
    "I'm finalizing the [Month] calendar tomorrow — do you want that last spot?\""
)

add_sub_heading('THE "GHOSTING" PROTOCOL — Break-Up Text (48 hrs after no reply)')
add_script(
    "\"Hi [Name], I'm assuming Retail Control isn't a priority right now, so I'm going to "
    "release your reserved slot to the next store on the waitlist.\n\n"
    "If your mismatch becomes a bigger headache down the road, feel free to reach out and "
    "we can see if a future slot is open.\n\n"
    "Wish you the best with the season! — [Your Name]\""
)
add_body("→ Why it works: Flips the script. They are losing the opportunity to work with YOU.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 5: CASE STUDY FRAMEWORK
# ══════════════════════════════════════════════════════════════════
add_section_heading("SECTION 5: CASE STUDY FRAMEWORK")
add_body("One good case study is worth 100 sales pitches. Use this after every completed 30-day installation.")

add_sub_heading("📈 CASE STUDY: The Stock Certainty System™ — [Store Name], [City]")

styled_table(
    ["Field", "Details"],
    [
        ["Store Type", "[e.g., Multi-Brand Clothing Outlet]"],
        ["Location", "[City, Tier]"],
        ["Timeline", "30 Days"],
        ["Investment", "₹85,000"],
    ],
    col_widths=[1.5, 4.5]
)

add_sub_heading('1. THE "BEFORE" (The Chaos)')
add_bullet("The Problem: Owner was losing approximately ₹____/month in \"Ghost Stock\"")
add_bullet("The Baseline: Physical audit showed ___% mismatch between software and racks")
add_bullet("The Daily Pain: Owner spent ___ minutes every night manually reconciling")
add_bullet("The Quote: \"[Pain point in their own words]\"")

add_sub_heading("2. THE FORENSIC DIAGNOSIS")
add_body("The Control Gap Audit revealed:")
add_bullet("[Finding 1: e.g., Supplier short-shipping — 3% missing at entry]")
add_bullet("[Finding 2: e.g., SKU duplication — same item entered 4 different ways]")
add_bullet("[Finding 3: e.g., Staff keeping manual notes, bypassing system]")
add_body("Total Annual Leakage Calculated: ₹____________")

add_sub_heading("3. THE STRUCTURAL INSTALLATION (What We Did)")
add_numbered("Reset entire SKU logic — every item now has a \"Digital DNA\"", 1)
add_numbered("Installed Supplier Entry Lock™ — no box opened without 100% invoice entry", 2)
add_numbered("Enforced Single System Rule — physically removed all manual diaries", 3)
add_numbered("Trained staff on 10-Minute Daily Control Routine™", 4)
add_numbered("30-day compliance monitoring with weekly reviews", 5)

add_sub_heading('4. THE "AFTER" (The Result)')
add_bullet("✅ Mismatch:       From ___% → ___% in 30 days")
add_bullet("✅ Time Saved:     Daily reconciliation from ___ min → ___ min")
add_bullet("✅ Found Money:    ₹_____ of dead stock identified and cleared")
add_bullet("✅ Owner Freedom:  Owner now checks store via phone once a day")
add_bullet("✅ Payback Period: System paid for itself in ___ days")

add_sub_heading("5. THE CLIENT VERDICT")
add_quote("[Testimonial in their own words — focus on emotional transformation]")
add_body("— [Client Name], Owner, [Store Name]")

add_sub_heading("How to Use This Case Study:")
add_numbered('THE "TEASER" POST: Share Before vs After stats on WhatsApp Status', 1)
add_numbered('THE "PROOF" PDF: Send full case study to prospects who ask "Will this work for my store?"', 2)
add_numbered("THE PITCH DECK: Insert as Slide #5 to prove the mechanism works", 3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 6: CRM TRACKER
# ══════════════════════════════════════════════════════════════════
add_section_heading("SECTION 6: CRM TRACKER — THE \"RETAIL CONTROL\" SYSTEM")
add_body("Build this in Google Sheets today. Migrates to any CRM later.")

add_sub_heading("Sheet 1: Prospect Pipeline")
styled_table(
    ["Lead / Store", "Contact", "Status", "Audit Date", "Est. Annual Leak", "Objection", "Next Action"],
    [
        ["Zara Fashion", "9988XXXXXX", "Hook Sent", "Mar 2", "₹4,50,000", "Price", "Send Case Study"],
        ["Royal Ethnic", "9876XXXXXX", "Audit Done", "Feb 28", "₹6,20,000", "Staff Fear", "Send Rulebook"],
    ],
    col_widths=[1.0, 0.9, 0.7, 0.6, 0.8, 0.6, 0.9]
)
add_body("Status Options: New Lead → Hook Sent → Audit Booked → Audit Done → Agreement Sent → Closed Won / Closed Lost / Follow-up Required")

add_sub_heading("Sheet 2: Forensic Audit Calculator")
add_body("Use during the sales call to calculate their \"Chaos Tax\" live.")

add_sub_heading("Sheet 3: Delivery & Installation Tracker")
styled_table(
    ["Client", "Start Date", "Day 1 Mismatch", "Day 30 Mismatch", "SKU Reset", "Staff Trained", "Guarantee Met?"],
    [
        ["[Store Name]", "[Date]", "____%", "____%", "YES / NO", "YES / NO", "YES / NO"],
    ],
    col_widths=[1.0, 0.7, 0.9, 0.9, 0.7, 0.7, 0.7]
)

# ══════════════════════════════════════════════════════════════════
# SECTION 7: WEEKLY REVIEW
# ══════════════════════════════════════════════════════════════════
add_section_heading("SECTION 7: WEEKLY AGENCY SUCCESS REVIEW")
add_body("Complete every Sunday. 20 minutes. Non-negotiable.")

add_sub_heading("1. The Revenue Scorecard")
add_bullet("Total Outreach Sent This Week: _____ (Target: 50+)")
add_bullet("Forensic Audits Booked: _____ (Target: 5–8)")
add_bullet("Agreements Sent: _____")
add_bullet("Cash Collected: ₹_____ ← The only number that truly matters")

add_sub_heading("2. The Pipeline Audit")
add_bullet("Who Ghosted Me? → Move to \"Break-up\" protocol")
add_bullet("Who is \"Thinking\"? → Send Case Study or Value-Add tonight")
add_bullet("Who is a \"Hard Yes\"? → Send Onboarding Message Monday morning")

add_sub_heading("3. Operations & Delivery")
add_bullet("Active Installations: _____ / 3 max")
add_bullet("Client Win of the Week: ______________________________")
add_bullet("The \"Leak\" in My Service: ______________________________")
add_bullet("Solution: ______________________________")

add_sub_heading("4. Strategic Questions")
add_bullet("#1 Objection This Week: ______________________________")
add_bullet("If Price → Improve the Leakage Math demo")
add_bullet("If Staff → Show the Staff Rulebook™ earlier")
add_bullet("If Trust → Share Case Study proof")
add_bullet("Did I post an Authority Hook this week?  Y / N")
add_bullet("ONE BIG GOAL for Next Week: ______________________________")

# ══════════════════════════════════════════════════════════════════
# SECTION 8: DAILY SCHEDULE
# ══════════════════════════════════════════════════════════════════
add_section_heading("SECTION 8: DAILY LAUNCH SCHEDULE")
add_body("Use this schedule every working day until all 3 monthly slots are filled.")

styled_table(
    ["Time", "Activity"],
    [
        ["9:00 AM", "Send \"Forensic Hook\" message to 10 NEW leads"],
        ["9:15 AM", "Post Long-Form on LinkedIn/Facebook"],
        ["9:15 AM", "Post Short-Form on WhatsApp Status + Instagram Story"],
        ["11:00 AM", "Send Re-engagement Script to any OLD leads"],
        ["2:00–5:00 PM", "Conduct Forensic Audits (in-person or video meetings)"],
        ["6:00 PM", "Send Agreements + Payment links to \"Yes\" leads"],
        ["8:30 PM", "Send \"Break-up\" text to anyone ghosting from last week"],
    ],
    col_widths=[1.5, 4.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# SECTION 9: ONE-PAGE CHEAT SHEET
# ══════════════════════════════════════════════════════════════════
add_section_heading("SECTION 9: THE ONE-PAGE CHEAT SHEET")
add_body("Print this. Keep it in front of you during every sales conversation.")

cheat_table = doc.add_table(rows=12, cols=2)
cheat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cheat_table.style = 'Table Grid'

cheat_data = [
    ("OFFER NAME", "The 30-Day Stock Certainty System™"),
    ("TAGLINE", "Make your system stock match your shop — permanently."),
    ("CORE PROMISE", "In 30 days → system stock = physical stock, reconciliation under 15 min, one system"),
    ("TARGET", "Owner-operated clothing stores, ₹30L–1.5Cr revenue, ₹20L+ inventory"),
    ("FORMAT", "1:1 Consulting-Led, Done-With-You + Strict Enforcement"),
    ("DURATION", "30 Days + Optional Monthly Maintenance"),
    ("CAPACITY", "3 stores/month max. Cycles start 1st only."),
    ("INVESTMENT", "₹85,000 one-time setup"),
    ("TOTAL VALUE", "₹97,000 (7 components + 5 bonuses)"),
    ("GUARANTEE", "Stock Certainty Guarantee™ — continue free if mismatch doesn't reduce"),
    ("URGENCY", "Leakage costs ₹8K–12K/month. Delay has a rupee cost."),
    ("ONE SENTENCE", "\"I fix stock mismatch in clothing stores in 30 days — permanently.\""),
]

for i, (label, value) in enumerate(cheat_data):
    cell_l = cheat_table.rows[i].cells[0]
    cell_r = cheat_table.rows[i].cells[1]
    cell_l.text = ''
    cell_r.text = ''

    p_l = cell_l.paragraphs[0]
    r_l = p_l.add_run(label)
    r_l.bold = True
    r_l.font.size = Pt(10)
    r_l.font.color.rgb = WHITE
    set_cell_shading(cell_l, "1B2A4A")

    p_r = cell_r.paragraphs[0]
    r_r = p_r.add_run(value)
    r_r.font.size = Pt(10)

    cell_l.width = Inches(1.8)
    cell_r.width = Inches(4.2)

doc.add_paragraph()

# ── Final Reminders ────────────────────────────────────────────────
add_section_heading("REMEMBER")

reminders = [
    "This is not billing software. This is a Stock Control Installation.",
    "Software is a tool. Behavior creates mismatch. Tools don't fix behavior.",
    "Structure is what makes small stores big. Chaos keeps them small.",
    "Three perfect installs beat six mediocre ones every single time.",
    "Quality of installation → quality of case study → ease of next sale.",
    "Your only job is to start conversations. Be helpful to a stressed store owner.",
    "You are not \"trying to start a business.\" You are running one.",
]
for r in reminders:
    add_bullet(r)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Go dominate. 🚀")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = ACCENT_GOLD

# ── Save ───────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "The_30_Day_Stock_Certainty_System_Pitch_Ready_Toolkit.docx")
doc.save(output_path)
print(f"✅ Successfully created: {output_path}")
