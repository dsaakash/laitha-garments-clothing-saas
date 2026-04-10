#!/usr/bin/env python3
"""
GAP Framework TEMPLATE — Fill-in-the-Blank Step-by-Step Guide
IEC Diamond Week 4: 3 Months of Content Ready Using AI
Blank template with writing spaces, prompts, and instructions.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors
ORANGE = RGBColor(0xFF, 0x6B, 0x35)
GOLD = RGBColor(0xD4, 0xA5, 0x00)
GREEN = RGBColor(0x00, 0xA8, 0x55)
BLUE = RGBColor(0x3D, 0x8B, 0xE8)
RED = RGBColor(0xE8, 0x3D, 0x3D)
PURPLE = RGBColor(0x6C, 0x5C, 0xE7)
BLACK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, hex_color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), hex_color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)


def styled(doc, text, sz=12, c=BLACK, b=False, i=False, a=WD_ALIGN_PARAGRAPH.LEFT, sa=6, sb=0):
    p = doc.add_paragraph()
    p.alignment = a
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.color.rgb = c
    r.bold = b
    r.italic = i
    r.font.name = 'Calibri'
    return p


def header(doc, text, color=ORANGE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"━━━  {text}")
    r.font.size = Pt(20)
    r.font.color.rgb = color
    r.bold = True
    r.font.name = 'Calibri'


def fill_box(doc, label, lines=3, hint=""):
    """Create a fill-in box with label and blank lines."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"✏️  {label}")
    r.font.size = Pt(12)
    r.font.color.rgb = ORANGE
    r.bold = True
    r.font.name = 'Calibri'
    if hint:
        r2 = p.add_run(f"  ({hint})")
        r2.font.size = Pt(10)
        r2.font.color.rgb = LIGHT
        r2.italic = True
        r2.font.name = 'Calibri'

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, 'FFF8F0')
    cell.text = ""
    for _ in range(lines):
        p2 = cell.add_paragraph("_" * 80)
        p2.paragraph_format.space_after = Pt(8)
        for r in p2.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
            r.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def bullet_fill(doc, label, count=5, examples=None):
    """Bullet point fill-in section."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"✏️  {label}")
    r.font.size = Pt(12)
    r.font.color.rgb = ORANGE
    r.bold = True
    r.font.name = 'Calibri'

    if examples:
        styled(doc, f"Examples to inspire you:", sz=10, c=LIGHT, i=True, sa=2)
        for ex in examples:
            styled(doc, f"    💡 {ex}", sz=10, c=GREEN, i=True, sa=1)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, 'FFF8F0')
    cell.text = ""
    for n in range(count):
        p2 = cell.add_paragraph(f"  {n+1}.  " + "_" * 70)
        p2.paragraph_format.space_after = Pt(6)
        for r in p2.runs:
            r.font.size = Pt(11)
            r.font.color.rgb = GRAY
            r.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def prompt_box(doc, step, title, prompt_text, instructions, fill_note=None):
    """Formatted AI prompt step."""
    # Step badge
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"STEP {step}")
    r.font.size = Pt(12)
    r.font.color.rgb = WHITE
    r.bold = True
    r.font.name = 'Calibri'
    # Fake badge with background via table
    # Title
    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"🤖  STEP {step}: {title}")
    r2.font.size = Pt(18)
    r2.font.color.rgb = PURPLE
    r2.bold = True
    r2.font.name = 'Calibri'
    p2.paragraph_format.space_after = Pt(4)

    # Instructions
    for inst in instructions:
        styled(doc, f"    → {inst}", sz=11, c=GRAY, sa=2)

    if fill_note:
        p3 = doc.add_paragraph()
        r3 = p3.add_run(f"⚠️  {fill_note}")
        r3.font.size = Pt(11)
        r3.font.color.rgb = RED
        r3.bold = True
        r3.italic = True
        r3.font.name = 'Calibri'
        p3.paragraph_format.space_after = Pt(6)

    # Prompt in styled box
    styled(doc, "📋 COPY THIS PROMPT INTO ChatGPT:", sz=11, c=PURPLE, b=True, sa=4, sb=6)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, 'F5F0FF')
    cell.text = ""
    p_prompt = cell.add_paragraph(prompt_text)
    p_prompt.paragraph_format.space_after = Pt(4)
    for r in p_prompt.runs:
        r.font.size = Pt(11)
        r.font.color.rgb = BLACK
        r.font.name = 'Calibri'

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def create_template():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(12)

    # ═══════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    styled(doc, "IEC DIAMOND — WEEK 4", sz=14, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER)
    styled(doc, "GAP FRAMEWORK", sz=40, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    styled(doc, "FILL-IN TEMPLATE", sz=24, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    styled(doc, "3 Months of Content Ready Using AI", sz=16, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=20)
    styled(doc, "━" * 35, sz=12, c=ORANGE, a=WD_ALIGN_PARAGRAPH.CENTER, sa=16)

    styled(doc, "Content Ideas Through AI for Lead Generation", sz=14, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=8)

    # Instructions box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, 'FFF3E0')
    cell.text = ""
    lines = [
        "📌 HOW TO USE THIS TEMPLATE:",
        "",
        "1️⃣  Fill in ALL the ✏️ sections with YOUR information",
        "2️⃣  Copy each 📋 prompt into ChatGPT (one at a time, in order)",
        "3️⃣  Wait for ChatGPT to reply 'Yes' before moving to the next step",
        "4️⃣  After Step 4, ChatGPT gives you 75 questions — film yourself answering them",
        "5️⃣  After Step 5, ChatGPT writes reel scripts — film those too",
        "",
        "⚠️  Parts marked with ✏️ MUST be filled in with YOUR info",
        "💡  The more detail you give ChatGPT, the better your content will be",
        "",
        "Need more questions? Just tell ChatGPT:",
        '     "Ask me more questions"',
        '     "Ask more personal questions"',
        '     "More questions on [whatever topic]"',
    ]
    for line in lines:
        p = cell.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = 'Calibri'
            if line.startswith("📌"):
                r.font.color.rgb = ORANGE
                r.bold = True
                r.font.size = Pt(13)
            elif line.startswith("⚠️"):
                r.font.color.rgb = RED
                r.bold = True
            elif line.startswith("💡"):
                r.font.color.rgb = GREEN
            else:
                r.font.color.rgb = GRAY

    doc.add_page_break()

    # ═══════════════════════════════════════
    # SECTION 1: UNDERSTAND THE GAP FRAMEWORK
    # ═══════════════════════════════════════
    header(doc, "SECTION 1: UNDERSTAND THE GAP FRAMEWORK", ORANGE)

    styled(doc, "Before you start, understand what the 3 types of content are:", sz=13, c=BLACK, b=True, sa=8)

    # GAP Table
    tbl = doc.add_table(rows=4, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_data = ["", "TYPE", "WHAT IT DOES", "EXAMPLE"]
    for i, h in enumerate(headers_data):
        cell = tbl.cell(0, i)
        shade(cell, '1A1A2E')
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.size = Pt(10)
                r.bold = True
                r.font.name = 'Calibri'

    gap_rows = [
        ("🚀", "GROWTH", "Drives traffic using trends & popular topics", "How ChatGPT can help your business", "E8F5E9"),
        ("🎓", "AUTHORITY", "Shows your skills, builds trust", "3 mistakes store owners make", "E3F2FD"),
        ("❤️", "PERSONAL", "Builds deep connection, creates fans", "Why I started my business", "F3E5F5"),
    ]

    for i, (emoji, type_name, what, example, bg) in enumerate(gap_rows):
        for j, val in enumerate([emoji, type_name, what, example]):
            cell = tbl.cell(i + 1, j)
            shade(cell, bg)
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Calibri'

    styled(doc, "", sz=6, sa=4)
    styled(doc, "Stranger  →  Sees your Growth content  →  Follows for Authority  →  Stays for Personal  →  Becomes a FAN",
           sz=12, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=10)

    styled(doc, "That is the GAP. Your content bridges this gap.", sz=12, c=GRAY, i=True, a=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ═══════════════════════════════════════
    # SECTION 2: FILL IN YOUR INFORMATION
    # ═══════════════════════════════════════
    header(doc, "SECTION 2: FILL IN YOUR INFORMATION (Before Using Prompts)", RED)

    styled(doc, "Complete ALL of these BEFORE you go to ChatGPT. The more you write, the better your content will be.",
           sz=12, c=RED, b=True, i=True, sa=10)

    # --- Your Name & Brand ---
    styled(doc, "📝  ABOUT YOU & YOUR BRAND", sz=16, c=BLACK, b=True, sa=6, sb=10)
    fill_box(doc, "Your Full Name:", lines=1, hint="e.g., Aakash Savant")
    fill_box(doc, "Your Brand / Business Name:", lines=1, hint="e.g., Retail Control Architect")
    fill_box(doc, "What do you do? (In one simple sentence):", lines=2,
             hint="e.g., I help retail store owners stop losing money from stock problems")
    fill_box(doc, "Your Social Media Handles:", lines=2,
             hint="Instagram, YouTube, etc.")

    # --- Authority ---
    styled(doc, "🎓  YOUR AUTHORITY (Proof that you know your stuff)", sz=16, c=BLUE, b=True, sa=6, sb=14)
    styled(doc, "List everything that shows you are good at what you do. Think: results, experience, achievements.",
           sz=11, c=GRAY, i=True, sa=6)

    bullet_fill(doc, "Your Accomplishments & Results:", count=8, examples=[
        "I helped a client recover Rs 3.2 Lakhs in 45 days",
        "I have worked with 50+ stores",
        "I reduced stock mismatch from 17% to 2%",
        "I have 10 years of experience in retail",
    ])

    bullet_fill(doc, "Results You Have Gotten for Clients:", count=5, examples=[
        "Client opened their second branch after working with me",
        "One client went from Rs 0 profit to Rs 1.5 Lakh/month profit",
    ])

    bullet_fill(doc, "Certifications, Awards, or Recognition:", count=3, examples=[
        "Featured in a local newspaper",
        "Certified business consultant",
    ])

    doc.add_page_break()

    # --- Personal ---
    styled(doc, "❤️  YOUR PERSONAL SIDE (What makes you unique)", sz=16, c=PURPLE, b=True, sa=6, sb=10)
    styled(doc, "List your opinions, stories, beliefs, and experiences. This is what makes people LOVE you.",
           sz=11, c=GRAY, i=True, sa=6)

    bullet_fill(doc, "Your Personal Worldviews & Beliefs:", count=6, examples=[
        "I believe small business owners are the backbone of India",
        "I think expensive software is NOT the answer",
        "Staff are not the enemy — lack of structure is",
    ])

    bullet_fill(doc, "Interesting Personal Stories:", count=5, examples=[
        "I once walked into a store where Rs 4 Lakhs had disappeared",
        "My biggest failure taught me the most important lesson",
        "The moment I knew this was my calling",
    ])

    bullet_fill(doc, "Your Opinions that Most People Might Disagree With:", count=4, examples=[
        "I think billing software alone is useless for stock control",
        "I believe most retail problems are process problems, not people problems",
    ])

    doc.add_page_break()

    # --- Target Audience ---
    styled(doc, "🎯  YOUR TARGET AUDIENCE (Who are you making content for?)", sz=16, c=GREEN, b=True, sa=6, sb=10)
    styled(doc, "Describe your ideal viewer / customer in detail. Be as specific as possible.",
           sz=11, c=GRAY, i=True, sa=6)

    fill_box(doc, "Gender:", lines=1, hint="e.g., Mostly male, some female")
    fill_box(doc, "Age Range:", lines=1, hint="e.g., 28-55 years old")
    fill_box(doc, "Occupation / Job:", lines=1, hint="e.g., Retail store owners, clothing shop owners")
    fill_box(doc, "Location:", lines=1, hint="e.g., Tier 1, 2, 3 cities in India")
    fill_box(doc, "Their Goals (What do they want?):", lines=3,
             hint="e.g., Run store smoothly, stop losing money, grow business, open second branch")
    fill_box(doc, "Their Pain Points (What keeps them up at night?):", lines=4,
             hint="e.g., Stock mismatch, staff not accountable, month-end surprises, unsure about real profit")
    fill_box(doc, "Their Typical Day:", lines=3,
             hint="e.g., Opens store at 10 AM, manages staff, handles customers, closes at 9 PM, feels exhausted")
    fill_box(doc, "Their Current Situation:", lines=3,
             hint="e.g., Running store for 5+ years, uses billing software but no stock process, loses money without knowing")
    fill_box(doc, "How They Talk / Language:", lines=2,
             hint="e.g., Hindi-English mix, simple language, no technical words")
    fill_box(doc, "Social Media They Use:", lines=2,
             hint="e.g., WhatsApp daily, watches YouTube and Instagram Reels")

    doc.add_page_break()

    # ═══════════════════════════════════════
    # SECTION 3: AI PROMPTS — STEP BY STEP
    # ═══════════════════════════════════════
    header(doc, "SECTION 3: AI PROMPTS — COPY INTO CHATGPT (In Order)", PURPLE)

    styled(doc, "Now take everything you filled in above and follow these 5 steps.", sz=13, c=BLACK, b=True)
    styled(doc, "Go to ChatGPT → Start a new chat → Copy each prompt one by one → Wait for reply before next step.",
           sz=12, c=GRAY, i=True, sa=12)

    # ─── STEP 1 ───
    prompt_box(doc, 1, "TEACH CHATGPT THE GAP FRAMEWORK",
        "Please read this and once you understand, reply with the word yes:\n\n"
        "A personal brand that grows on social media has 3 core elements I call the GAP Framework.\n\n"
        "The GAP Framework consists of:\n"
        "Growth Content\n"
        "Authority Content\n"
        "Personal Content\n\n"
        "When you create these 3 types of content, it bridges the GAP from a stranger to die-hard fan.\n\n"
        "It's broken down like this:\n\n"
        "Growth content is when you talk about popular trends and figures. "
        "This leverages people's interest and drives traffic to your profile.\n\n"
        "For example, if I was a fitness coach when ChatGPT was trending, I could create content on "
        "how to use ChatGPT to create a nutrition plan. Or I could break down Arnold Schwarzenegger's workout routine.\n\n"
        "Since people in your niche have an interest in the growth content you're talking about, "
        "you're going to get more interest and reach.\n\n"
        "The A in the GAP framework is authority content.\n"
        "This is content that shows your competence at what your account is centered around.\n"
        "So things like:\n"
        "- Content that shows people how to solve their pain points or reach their goals\n"
        "- Lessons and tips\n"
        "- Case studies\n\n"
        "When you create this type of content, you incentivize people to follow you because you're "
        "actively helping the market by giving actionable information.\n\n"
        "The 3rd pillar of the GAP framework is personal content.\n"
        "Personal content is when you talk about the things that make you unique:\n"
        "- Your personal stories\n"
        "- Your personal worldviews\n"
        "- Your personal opinions on life or industry\n\n"
        "When you create personal content, you build a deep connection with your audience.\n"
        "And more importantly, you build an audience of true fans who adore you.\n\n"
        'Reply with "Yes" once you read and understand this.',
        instructions=[
            "Copy this prompt EXACTLY as-is into ChatGPT",
            "ChatGPT will reply with 'Yes'",
            "Then move to Step 2",
        ])

    # ─── STEP 2 ───
    prompt_box(doc, 2, "GIVE CHATGPT YOUR AUTHORITY & PERSONAL INFO",
        "Here is some information about me:\n\n"
        "Authority:\n"
        "[PASTE YOUR AUTHORITY BULLET POINTS FROM SECTION 2 HERE]\n\n"
        "Personal:\n"
        "[PASTE YOUR PERSONAL BULLET POINTS FROM SECTION 2 HERE]\n\n"
        "Once you understand, reply with 'Yes'.",
        instructions=[
            "Go back to Section 2 and copy your Authority bullet points",
            "Copy your Personal bullet points",
            "Paste them into this prompt replacing the [...] sections",
            "The more bullet points you give, the better ChatGPT's questions will be",
        ],
        fill_note="REPLACE the [...] sections with your actual bullet points from Section 2!")

    # Where to paste - visual guide
    styled(doc, "📍 YOUR AUTHORITY POINTS TO PASTE:", sz=12, c=BLUE, b=True, sa=2, sb=4)
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade(cell, 'E3F2FD')
    cell.text = "Copy your Authority bullet points here before pasting into ChatGPT:\n\n- \n- \n- \n- \n- "
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = 'Calibri'
            r.font.color.rgb = GRAY

    styled(doc, "📍 YOUR PERSONAL POINTS TO PASTE:", sz=12, c=PURPLE, b=True, sa=6, sb=8)
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade(cell, 'F3E5F5')
    cell.text = "Copy your Personal bullet points here before pasting into ChatGPT:\n\n- \n- \n- \n- \n- "
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = 'Calibri'
            r.font.color.rgb = GRAY

    doc.add_page_break()

    # ─── STEP 3 ───
    prompt_box(doc, 3, "TELL CHATGPT YOUR TARGET AUDIENCE",
        "My target audience is:\n\n"
        "[PASTE YOUR TARGET AUDIENCE DETAILS FROM SECTION 2 HERE]\n\n"
        "(Include: gender, age, occupation, goals, pain points, typical day in the life, and current situation)\n\n"
        'Once you understand, reply with "Yes."',
        instructions=[
            "Go back to Section 2 and copy your Target Audience details",
            "Paste them into this prompt replacing the [...] section",
            "Include ALL the details: gender, age, job, goals, pains, daily life, situation",
        ],
        fill_note="REPLACE the [...] with your Target Audience from Section 2!")

    styled(doc, "📍 YOUR TARGET AUDIENCE TO PASTE:", sz=12, c=GREEN, b=True, sa=2, sb=4)
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade(cell, 'E8F5E9')
    cell.text = ("Copy your Target Audience details here before pasting into ChatGPT:\n\n"
                 "Gender: \nAge: \nOccupation: \nGoals: \nPain Points: \n"
                 "Typical Day: \nCurrent Situation: \nLanguage: \nSocial Media: ")
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = 'Calibri'
            r.font.color.rgb = GRAY

    # ─── STEP 4 ───
    prompt_box(doc, 4, "GET 75 VIDEO CONTENT QUESTIONS",
        "My goal is to create short-form video content for social media that makes my target avatars know, like, and trust me.\n\n"
        "Using the GAP Framework, please ask me 25 questions for each category of the GAP Framework that does this.\n\n"
        "You can use the information I gave you about me to help, but feel free to ask questions outside that.\n\n"
        "The questions asked should set me up so I can answer them with a soundbite that's under 60 seconds.",
        instructions=[
            "Copy this prompt EXACTLY as-is — no changes needed!",
            "ChatGPT will give you 75 questions (25 Growth + 25 Authority + 25 Personal)",
            "SAVE these questions — each one becomes a video!",
            "75 questions = 75 videos = almost 3 months of daily content!",
        ])

    # Results tracking
    styled(doc, "✅  SAVE YOUR 75 QUESTIONS HERE:", sz=14, c=ORANGE, b=True, sa=4, sb=10)

    for category, color, emoji, bg in [
        ("GROWTH QUESTIONS (25)", GREEN, "🚀", "E8F5E9"),
        ("AUTHORITY QUESTIONS (25)", BLUE, "🎓", "E3F2FD"),
        ("PERSONAL QUESTIONS (25)", PURPLE, "❤️", "F3E5F5"),
    ]:
        styled(doc, f"{emoji}  {category}", sz=13, c=color, b=True, sa=2, sb=8)
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.cell(0, 0)
        shade(cell, bg)
        cell.text = "Paste ChatGPT's questions here:\n\n1.\n2.\n3.\n4.\n5.\n...(continue to 25)"
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = 'Calibri'
                r.font.color.rgb = GRAY

    styled(doc, "\n💡 Want MORE questions? Tell ChatGPT:", sz=12, c=GREEN, b=True, sa=4)
    for extra in [
        '"Ask me 10 more questions"',
        '"Ask more personal questions"',
        '"More questions on [any topic you want]"',
        '"Give me controversial questions my audience would love"',
    ]:
        styled(doc, f"    → {extra}", sz=11, c=PURPLE, i=True, sa=2)

    doc.add_page_break()

    # ─── STEP 5 ───
    prompt_box(doc, 5, "GENERATE YOUR FIRST REEL SCRIPT",
        "Write me a 60-second Instagram reel script that feels natural and conversational "
        "(not salesy or filled with buzzwords). The script must include:\n\n"
        "A strong hook that grabs attention in the first 3 seconds.\n\n"
        "Clear value that teaches or shares something useful.\n\n"
        "A short story or relatable example that makes it engaging.\n\n"
        "A simple and direct CTA at the end.\n\n"
        "This should be for content piece #1 in my Growth Content series. "
        "Use your years of copywriting experience to make it flow smoothly, "
        "like how a real person would speak.",
        instructions=[
            "Copy this prompt as-is for your FIRST reel script",
            "To get scripts for other questions, change '#1 in Growth' to the question number and category",
            "Example: 'content piece #5 in my Authority Content series'",
            "Repeat for ALL 75 questions to get 75 ready-made scripts!",
        ])

    styled(doc, "🔄  HOW TO GET ALL 75 SCRIPTS:", sz=14, c=ORANGE, b=True, sa=4, sb=8)

    tbl = doc.add_table(rows=4, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = [("CHANGE THIS PART", "1A1A2E"), ("TO GET THIS", "1A1A2E")]
    for i, (h, bg) in enumerate(header_cells):
        cell = tbl.cell(0, i)
        shade(cell, bg)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.size = Pt(11)
                r.font.name = 'Calibri'
                r.bold = True

    swap_data = [
        ("#1 in my Growth Content series", "Script for Growth question 1"),
        ("#3 in my Authority Content series", "Script for Authority question 3"),
        ("#7 in my Personal Content series", "Script for Personal question 7"),
    ]
    for i, (change, result) in enumerate(swap_data):
        tbl.cell(i+1, 0).text = change
        tbl.cell(i+1, 1).text = result
        for j in range(2):
            for p in tbl.cell(i+1, j).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Calibri'

    doc.add_page_break()

    # ═══════════════════════════════════════
    # SECTION 4: CONTENT CALENDAR
    # ═══════════════════════════════════════
    header(doc, "SECTION 4: YOUR WEEKLY CONTENT CALENDAR", GREEN)

    styled(doc, "Post 1 video per day. Rotate between the 3 types like this:", sz=13, c=BLACK, b=True, sa=6)

    tbl = doc.add_table(rows=8, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["DAY", "CONTENT TYPE", "WHAT TO POST"]):
        cell = tbl.cell(0, i)
        shade(cell, '1A1A2E')
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.size = Pt(11)
                r.bold = True
                r.font.name = 'Calibri'

    schedule = [
        ("Monday", "🚀 Growth", "Trending topic + your angle", "E8F5E9"),
        ("Tuesday", "🎓 Authority", "Tip / How-to / Lesson", "E3F2FD"),
        ("Wednesday", "❤️ Personal", "Your story / opinion", "F3E5F5"),
        ("Thursday", "🎓 Authority", "Case study / results", "E3F2FD"),
        ("Friday", "🚀 Growth", "React to industry news", "E8F5E9"),
        ("Saturday", "❤️ Personal", "Behind the scenes / day in life", "F3E5F5"),
        ("Sunday", "🎓 Authority", "Mistakes to avoid", "E3F2FD"),
    ]
    for i, (day, ctype, what, bg) in enumerate(schedule):
        for j, val in enumerate([day, ctype, what]):
            cell = tbl.cell(i+1, j)
            shade(cell, bg)
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Calibri'

    styled(doc, "\n25 Growth + 25 Authority + 25 Personal = 75 Videos", sz=14, c=ORANGE, b=True,
           a=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    styled(doc, "= About 11 Weeks = Almost 3 MONTHS of Daily Content! 🎉", sz=14, c=GREEN, b=True,
           a=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ═══════════════════════════════════════
    # SECTION 5: FILMING CHECKLIST
    # ═══════════════════════════════════════
    header(doc, "SECTION 5: FILMING & POSTING CHECKLIST", GOLD)

    for phase, color, items in [
        ("BEFORE Filming", BLUE, [
            "Pick today's question from your list",
            "Read the ChatGPT script for that question",
            "Practice once out loud (just once!)",
            "Set up phone vertically, good light on your face",
            "Clean background behind you",
        ]),
        ("WHILE Filming", GREEN, [
            "Start with your HOOK — first 3 seconds matter most!",
            "Look at the camera lens, not the screen",
            "Speak naturally — like talking to a friend",
            "Keep it under 60 seconds",
            "Share ONE clear point — do not try to cover everything",
        ]),
        ("AFTER Filming", PURPLE, [
            "Edit in CapCut or InShot — add captions/subtitles!",
            "Write a short caption + 3-5 hashtags",
            "Post on: Instagram Reels + YouTube Shorts + WhatsApp Status",
            "Reply to every comment within 1 hour",
            "Share the video to your WhatsApp contacts & groups",
        ]),
    ]:
        styled(doc, f"📋  {phase}:", sz=14, c=color, b=True, sa=4, sb=8)
        for item in items:
            p = doc.add_paragraph(f"  ⬜  {item}")
            p.paragraph_format.space_after = Pt(4)
            for r in p.runs:
                r.font.size = Pt(12)
                r.font.name = 'Calibri'

    # ═══════════════════════════════════════
    # FINAL PAGE
    # ═══════════════════════════════════════
    doc.add_page_break()
    for _ in range(5):
        doc.add_paragraph()

    styled(doc, "YOU ARE NOW READY TO CREATE", sz=18, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    styled(doc, "3 MONTHS", sz=40, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    styled(doc, "OF CONTENT", sz=40, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=12)

    styled(doc, "Step 1: Fill in Section 2  ✏️", sz=14, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    styled(doc, "Step 2: Copy prompts into ChatGPT  🤖", sz=14, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    styled(doc, "Step 3: Film yourself answering  🎬", sz=14, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    styled(doc, "Step 4: Post daily  📱", sz=14, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    styled(doc, "Step 5: Watch your audience grow  🚀", sz=14, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=16)

    styled(doc, "━" * 30, c=ORANGE, a=WD_ALIGN_PARAGRAPH.CENTER, sa=10)
    styled(doc, "Stop planning. Start creating. 🎬", sz=16, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    import os
    save_path = "/Users/aakash/Desktop/Week_3/Week 4/Week4_GAP_Framework_TEMPLATE.docx"
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    return save_path


if __name__ == "__main__":
    path = create_template()
    print(f"Created: {path}")
