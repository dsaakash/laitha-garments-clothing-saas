#!/usr/bin/env python3
"""
Generate TWO Warm Reachout DOCX files:
1. Generic Template (fill-in-the-blank)
2. Niche-specific for Retail Control Architect
IEC Diamond Week 4
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ORANGE = RGBColor(0xFF, 0x6B, 0x35)
GOLD = RGBColor(0xD4, 0xA5, 0x00)
GREEN = RGBColor(0x00, 0xA8, 0x55)
BLUE = RGBColor(0x3D, 0x8B, 0xE8)
RED = RGBColor(0xE8, 0x3D, 0x3D)
PURPLE = RGBColor(0x6C, 0x5C, 0xE7)
BLACK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEAL = RGBColor(0x00, 0x96, 0x88)


def shade(cell, hex_color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), hex_color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)


def styled(doc, text, sz=12, c=BLACK, b=False, i=False, a=WD_ALIGN_PARAGRAPH.LEFT, sa=6, sb=0):
    p = doc.add_paragraph()
    p.alignment = a
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.color.rgb = c
    r.bold = b
    r.italic = i
    r.font.name = 'Calibri'
    return p


def section_header(doc, text, color=ORANGE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"━━━  {text}")
    r.font.size = Pt(20)
    r.font.color.rgb = color
    r.bold = True
    r.font.name = 'Calibri'


def script_box(doc, text, bg_hex='FFF8F0', text_color=BLACK, label=None, label_color=ORANGE):
    """Create a highlighted script box."""
    if label:
        styled(doc, label, sz=12, c=label_color, b=True, sa=4, sb=6)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, bg_hex)
    cell.text = ""
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(12)
        r.font.color.rgb = text_color
        r.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def fill_line(doc, label, hint=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"✏️  {label}: ")
    r.font.size = Pt(12)
    r.font.color.rgb = ORANGE
    r.bold = True
    r.font.name = 'Calibri'
    r2 = p.add_run("_" * 50)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    if hint:
        r3 = p.add_run(f"  ({hint})")
        r3.font.size = Pt(10)
        r3.font.color.rgb = LIGHT
        r3.italic = True
        r3.font.name = 'Calibri'


def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(12)
    return doc


# ═══════════════════════════════════════════════════════════════
# DOCUMENT 1: GENERIC WARM REACHOUT TEMPLATE
# ═══════════════════════════════════════════════════════════════
def create_generic_template():
    doc = setup_doc()

    # Cover
    for _ in range(3):
        doc.add_paragraph()

    styled(doc, "IEC DIAMOND — WEEK 4", sz=14, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER)
    styled(doc, "WARM REACHOUT", sz=40, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    styled(doc, "BASIC TEMPLATE", sz=24, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    styled(doc, "Fill-In Guide for Any Niche", sz=16, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=20)
    styled(doc, "💎", sz=30, a=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
    styled(doc, "━" * 35, c=ORANGE, a=WD_ALIGN_PARAGRAPH.CENTER, sa=16)

    # Instructions box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, 'FFF3E0')
    cell.text = ""
    for line in [
        "📌  HOW THIS WORKS:",
        "",
        "Warm reachout = You reach out to people you already KNOW.",
        "Friends, family, old contacts, WhatsApp contacts, social media connections.",
        "",
        "You are NOT selling to them. You are asking if they KNOW someone.",
        "This removes all pressure and awkwardness.",
        "",
        "1️⃣  Fill in the blanks in Section 1 with YOUR offer details",
        "2️⃣  Use the scripts in Section 2 to reach out",
        "3️⃣  Follow the tips in Section 3 for best results",
    ]:
        p = cell.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = 'Calibri'
            if line.startswith("📌"):
                r.font.color.rgb = ORANGE
                r.bold = True
                r.font.size = Pt(13)
            elif "NOT selling" in line:
                r.font.color.rgb = GREEN
                r.bold = True
            else:
                r.font.color.rgb = GRAY

    doc.add_page_break()

    # ─── SECTION 1: FILL IN YOUR DETAILS ───
    section_header(doc, "SECTION 1: FILL IN YOUR OFFER DETAILS", RED)

    styled(doc, "Complete these BEFORE writing your warm reachout messages.", sz=12, c=RED, b=True, i=True, sa=10)

    fill_line(doc, "Your Name", "e.g., Aakash Savant")
    fill_line(doc, "Your Business/Service Name", "e.g., Retail Control Architect")
    fill_line(doc, "Who Do You Help?", "e.g., Clothing store owners")
    fill_line(doc, "Their Biggest Struggle", "e.g., Losing money from stock mismatch")
    fill_line(doc, "Dream Outcome You Deliver", "e.g., Stop stock losses and save lakhs")
    fill_line(doc, "Time Frame", "e.g., In 30 days")
    fill_line(doc, "How Many Free Case Studies?", "e.g., 5")
    fill_line(doc, "Effort/Sacrifice They Avoid", "e.g., Without buying expensive software")
    fill_line(doc, "Your Guarantee", "e.g., If mismatch does not reduce, I work free")
    fill_line(doc, "Success Story #1 — Name", "e.g., Lalitha Textiles")
    fill_line(doc, "Success Story #1 — Result", "e.g., Recovered Rs 3.2 Lakhs in 45 days")
    fill_line(doc, "Success Story #2 — Name", "e.g., Ramesh Clothing Store")
    fill_line(doc, "Success Story #2 — Result", "e.g., Opened second branch confidently")

    doc.add_page_break()

    # ─── SECTION 2: THE SCRIPTS ───
    section_header(doc, "SECTION 2: WARM REACHOUT SCRIPTS", BLUE)

    # --- Script A: The Core Template ---
    styled(doc, "🔵  SCRIPT A: THE CORE WARM REACHOUT (Master Template)", sz=16, c=BLUE, b=True, sa=4, sb=10)
    styled(doc, "This is the foundation. Learn this one first.", sz=11, c=GRAY, i=True, sa=6)

    script_box(doc,
        '"...By the way, do you know anybody who is [DESCRIBE THEIR STRUGGLES] '
        'looking to [DREAM OUTCOME] in [TIME DELAY]?\n\n'
        'I am taking on [NUMBER] case studies for free, because that is all I can handle. '
        'I just want to get some testimonials for my service.\n\n'
        'I help them [DREAM OUTCOME] without [EFFORT AND SACRIFICE]. '
        'It works. I even guarantee people get [DREAM OUTCOME] or I work with them until they do.\n\n'
        'I just had [NAME] work with me and they [DREAM OUTCOME] even though they [SAME STRUGGLE]. '
        'I also had another person who [DREAM OUTCOME #2] and it was their first time.\n\n'
        'I would just like more testimonials to show it works across different scenarios. '
        'Does anyone you like come to mind?\n\n'
        '...and if they say no...\n\n'
        'Haha, well... does anyone you HATE come to mind? 😄"',
        bg_hex='E3F2FD', label="📋 THE SCRIPT (Fill in the [...] parts):", label_color=BLUE)

    styled(doc, "💡  KEY PRINCIPLES:", sz=14, c=GREEN, b=True, sa=6, sb=10)
    for tip in [
        "You are asking if they KNOW someone — not pitching THEM directly",
        "This removes pressure. They do not feel sold to.",
        "Often, THEY will say 'Actually... I might be interested myself!'",
        "The joke at the end breaks any awkwardness",
        "Always start with 'By the way...' — makes it feel casual, not planned",
        "Keep the number small (3-5 case studies) — creates urgency",
        "The guarantee removes all risk from their mind",
    ]:
        styled(doc, f"    ✅  {tip}", sz=11, c=GRAY, sa=2)

    doc.add_page_break()

    # --- Script B: WhatsApp Version ---
    styled(doc, "💬  SCRIPT B: WHATSAPP MESSAGE VERSION", sz=16, c=GREEN, b=True, sa=4, sb=10)
    styled(doc, "Copy-paste this into WhatsApp. Personalize the name.", sz=11, c=GRAY, i=True, sa=6)

    script_box(doc,
        'Hey [FRIEND NAME]! Hope you are doing well 😊\n\n'
        'Quick question — do you know anyone who is [STRUGGLING WITH X] '
        'and wants to [DREAM OUTCOME] in [TIME]?\n\n'
        'I am taking on [NUMBER] people for free as case studies. '
        'Just want to get some good testimonials.\n\n'
        'I help them [DREAM OUTCOME] without [EFFORT/SACRIFICE]. '
        'Already helped [NAME] who [RESULT].\n\n'
        'Does anyone come to mind? Would really appreciate if you could connect us! 🙏',
        bg_hex='E8F5E9', label="📋 WHATSAPP SCRIPT:", label_color=GREEN)

    # --- Script C: Phone Call Version ---
    styled(doc, "📞  SCRIPT C: PHONE CALL VERSION", sz=16, c=ORANGE, b=True, sa=4, sb=10)
    styled(doc, "Use this when calling friends, family, or past contacts.", sz=11, c=GRAY, i=True, sa=6)

    script_box(doc,
        '"Hey [NAME], how are you doing? Long time!\n\n'
        '(Small talk for 1-2 minutes...)\n\n'
        'By the way, I wanted to ask you something. '
        'Do you know anyone who is [STRUGGLING] and wants to [DREAM OUTCOME]?\n\n'
        'I recently started helping people with this, and I am looking for [NUMBER] case studies. '
        'Completely free — I just want to build up my testimonials.\n\n'
        'I help them [DREAM OUTCOME] without [SACRIFICE]. '
        'I had [NAME] work with me and they [RESULT]. Pretty cool, right?\n\n'
        'So... does anyone come to mind?\n\n'
        '(If no) Haha no worries! Does anyone you hate come to mind? 😄\n\n'
        '(If they show interest themselves) Oh wait — is this something YOU would want? '
        'Because I would love to help you out!"',
        bg_hex='FFF3E0', label="📋 PHONE SCRIPT:", label_color=ORANGE)

    doc.add_page_break()

    # --- Script D: Instagram/Social DM ---
    styled(doc, "📱  SCRIPT D: INSTAGRAM / SOCIAL MEDIA DM", sz=16, c=PURPLE, b=True, sa=4, sb=10)

    script_box(doc,
        'Hey [NAME]! Saw your recent post — loved it! 😊\n\n'
        'Random question — do you happen to know anyone who is [STRUGGLING] '
        'and wants to [DREAM OUTCOME]?\n\n'
        'I am doing [NUMBER] free case studies and just need a few more people. '
        'Already got amazing results for [NAME] who [RESULT].\n\n'
        'Would love a referral if anyone comes to mind! No pressure at all 🙏',
        bg_hex='F3E5F5', label="📋 SOCIAL DM SCRIPT:", label_color=PURPLE)

    # --- Script E: In-Person ---
    styled(doc, "🤝  SCRIPT E: IN-PERSON / FACE-TO-FACE", sz=16, c=TEAL, b=True, sa=4, sb=10)

    script_box(doc,
        '(During natural conversation...)\n\n'
        '"By the way, I have been working on something interesting lately. '
        'I have been helping people who [STRUGGLE] to [DREAM OUTCOME].\n\n'
        'I am looking for a few more case studies — totally free. '
        'Just want testimonials. Do you know anyone who might be interested?\n\n'
        'I helped [NAME] and they [RESULT] in just [TIME]. Pretty amazing.\n\n'
        'Anyone come to mind?"',
        bg_hex='E0F2F1', label="📋 IN-PERSON SCRIPT:", label_color=TEAL)

    doc.add_page_break()

    # ─── SECTION 3: TIPS & STRATEGY ───
    section_header(doc, "SECTION 3: WARM REACHOUT STRATEGY & TIPS", GOLD)

    styled(doc, "📊  WHO TO REACH OUT TO (Make a list of 100 people):", sz=14, c=BLACK, b=True, sa=6, sb=8)

    tbl = doc.add_table(rows=9, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["#", "CATEGORY", "HOW MANY?"]):
        cell = tbl.cell(0, i)
        shade(cell, '1A1A2E')
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.size = Pt(10)
                r.bold = True
                r.font.name = 'Calibri'

    categories = [
        ("1", "Close friends", "10-15"),
        ("2", "Family members", "10-15"),
        ("3", "WhatsApp contacts you talk to", "15-20"),
        ("4", "Old school/college friends", "10-15"),
        ("5", "People in your industry", "10-15"),
        ("6", "Social media connections", "10-15"),
        ("7", "Neighbours & local contacts", "5-10"),
        ("8", "Past clients / customers", "5-10"),
    ]
    for i, (num, cat, count) in enumerate(categories):
        for j, val in enumerate([num, cat, count]):
            tbl.cell(i + 1, j).text = val
            for p in tbl.cell(i + 1, j).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Calibri'

    styled(doc, "", sz=6)
    styled(doc, "🎯  DAILY TARGET:", sz=14, c=ORANGE, b=True, sa=4, sb=8)
    for item in [
        "Send 10-15 warm reachout messages per day",
        "Make 5-8 phone calls per day",
        "Follow up with everyone who did not reply after 3 days",
        "Track every conversation in a notebook or spreadsheet",
    ]:
        styled(doc, f"    ➡️  {item}", sz=12, c=GRAY, sa=2)

    styled(doc, "\n✅  DO's:", sz=14, c=GREEN, b=True, sa=4, sb=8)
    for item in [
        "Make it about THEM, not you",
        "Always ask 'do you know anyone' — not 'do YOU want this'",
        "Keep it casual and natural",
        "Share real results and names",
        "Use the 'anyone you hate' joke to break ice",
        "Follow up if they say 'let me think'",
    ]:
        styled(doc, f"    ✅  {item}", sz=11, c=GRAY, sa=2)

    styled(doc, "\n❌  DON'Ts:", sz=14, c=RED, b=True, sa=4, sb=4)
    for item in [
        "Do NOT start with a pitch — start with small talk",
        "Do NOT send the same copy-paste to everyone",
        "Do NOT be pushy if they say no",
        "Do NOT talk about price in the first message",
        "Do NOT send voice notes on first contact",
    ]:
        styled(doc, f"    ❌  {item}", sz=11, c=GRAY, sa=2)

    # Tracking
    doc.add_page_break()
    section_header(doc, "SECTION 4: TRACKING SHEET", PURPLE)
    styled(doc, "Write down everyone you contact and their response:", sz=12, c=GRAY, i=True, sa=8)

    tbl = doc.add_table(rows=11, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["#", "NAME", "METHOD", "RESPONSE", "FOLLOW-UP DATE"]):
        cell = tbl.cell(0, i)
        shade(cell, '1A1A2E')
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.size = Pt(9)
                r.bold = True
                r.font.name = 'Calibri'
    for row in range(1, 11):
        tbl.cell(row, 0).text = str(row)
        for col in range(5):
            for p in tbl.cell(row, col).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Calibri'

    save_path = "/Users/aakash/Desktop/Week_3/Week 4/warm_reachout/Guide_Warm_Reachout_Template.docx"
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    return save_path


# ═══════════════════════════════════════════════════════════════
# DOCUMENT 2: NICHE WARM REACHOUT (Retail Control Architect)
# ═══════════════════════════════════════════════════════════════
def create_niche_reachout():
    doc = setup_doc()

    # Cover
    for _ in range(3):
        doc.add_paragraph()

    styled(doc, "IEC DIAMOND — WEEK 4", sz=14, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER)
    styled(doc, "WARM REACHOUT", sz=40, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    styled(doc, "FOR RETAIL CONTROL ARCHITECT", sz=20, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    styled(doc, "Ready-To-Send Scripts for Aakash Savant", sz=16, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=20)
    styled(doc, "💎", sz=30, a=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
    styled(doc, "━" * 35, c=ORANGE, a=WD_ALIGN_PARAGRAPH.CENTER, sa=16)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade(cell, 'FFF3E0')
    cell.text = ""
    for line in [
        "📌  These scripts are READY TO USE.",
        "Just replace [FRIEND NAME] with the actual name and SEND.",
        "",
        "Your Details (Pre-filled):",
        "  Name: Aakash Savant",
        "  Service: Retail Control Architect",
        "  Who You Help: Clothing & retail store owners",
        "  Struggle: Losing money from stock mismatch without knowing",
        "  Dream Outcome: Stop stock losses and save lakhs in 30 days",
        "  Case Studies: 5 free spots",
        "  Guarantee: If mismatch does not reduce in 30 days, I work free",
    ]:
        p = cell.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(11)
            r.font.name = 'Calibri'
            if "📌" in line:
                r.font.color.rgb = ORANGE
                r.bold = True
            elif "Pre-filled" in line:
                r.font.color.rgb = GREEN
                r.bold = True
            else:
                r.font.color.rgb = GRAY

    doc.add_page_break()

    # ─── SCRIPT 1: Core Warm Reachout ───
    section_header(doc, "SCRIPT 1: THE CORE WARM REACHOUT", BLUE)
    styled(doc, "Use with: Friends, family, anyone you already know", sz=11, c=GRAY, i=True, sa=6)

    script_box(doc,
        '"...By the way, do you know any store owner or shopkeeper '
        'who keeps finding their stock count does not match the computer — '
        'and wants to fix it and save lakhs in just 30 days?\n\n'
        'I am taking on 5 case studies for free, because that is all I can handle right now. '
        'I just want to get some testimonials for my service.\n\n'
        'I help them stop losing money from stock mismatch without buying any expensive software. '
        'It works. I even guarantee that if the stock mismatch does not reduce in 30 days, '
        'I work with them for free until it does.\n\n'
        'I just had a store owner work with me — they recovered Rs 3.2 Lakhs in just 45 days, '
        'even though they thought stock mismatch was normal and nothing could be done. '
        'I also had another store owner whose mismatch went from 17% to just 2%, '
        'and he confidently opened his second branch.\n\n'
        'I just want more testimonials to show it works for different types of stores. '
        'Does anyone you know come to mind?\n\n'
        '...and if they say no...\n\n'
        'Haha, well... does anyone you HATE come to mind? 😄"',
        bg_hex='E3F2FD', label="📋  COPY THIS SCRIPT:", label_color=BLUE)

    doc.add_page_break()

    # ─── SCRIPT 2: WhatsApp DM ───
    section_header(doc, "SCRIPT 2: WHATSAPP MESSAGE (Ready to Send)", GREEN)
    styled(doc, "Copy-paste directly into WhatsApp. Just change [FRIEND NAME].", sz=11, c=GRAY, i=True, sa=6)

    script_box(doc,
        'Hey [FRIEND NAME]! Hope you are doing great 😊\n\n'
        'Quick question — do you know any store owner or shopkeeper '
        'whose stock count never matches the computer? Like, items keep going missing '
        'and they do not know where the money is going?\n\n'
        'I am helping 5 store owners for FREE right now as case studies. '
        'I just want to collect testimonials.\n\n'
        'I have a simple system that stops stock losses in 30 days — '
        'without any expensive software. One store owner I helped recovered Rs 3.2 Lakhs.\n\n'
        'Does anyone come to mind? Would really appreciate a connection! 🙏\n\n'
        'And if nobody comes to mind... '
        'do you know any shop owner you do NOT like? 😂 (kidding!)',
        bg_hex='E8F5E9', label="💬  WHATSAPP MESSAGE:", label_color=GREEN)

    # Shorter version
    styled(doc, "📝  SHORTER VERSION (for busy contacts):", sz=13, c=GREEN, b=True, sa=4, sb=10)

    script_box(doc,
        'Hey [FRIEND NAME]! 😊\n\n'
        'Do you know any shop owner who keeps losing stock / money '
        'and wants to fix it in 30 days? I am doing 5 free case studies '
        'for testimonials. Already helped a store save Rs 3.2 Lakhs.\n\n'
        'Anyone come to mind? 🙏',
        bg_hex='F1F8E9')

    doc.add_page_break()

    # ─── SCRIPT 3: Phone Call ───
    section_header(doc, "SCRIPT 3: PHONE CALL SCRIPT", ORANGE)

    script_box(doc,
        '"Hey [NAME]! How are you? It has been a while!\n\n'
        '(Chat for 1-2 minutes about life, family, etc.)\n\n'
        'Listen, I wanted to ask you something. Do you know any shop owner — '
        'maybe clothing store, footwear, any retail shop — '
        'who keeps finding that their stock does not match what the computer shows?\n\n'
        'I have recently started helping store owners fix this exact problem. '
        'I have a simple 5-step system — no expensive software, just a clear process.\n\n'
        'I am looking for 5 case studies. Completely free. '
        'I just want to build testimonials.\n\n'
        'One store owner I helped — their Rs 4 Lakhs had disappeared and they thought it was normal! '
        'We fixed it in 30 days. Another one recovered Rs 3.2 Lakhs in 45 days.\n\n'
        'Does anyone come to mind?\n\n'
        '(If NO) Haha, okay... anyone you do NOT like who owns a shop? 😄\n\n'
        '(If they say "wait, I might need this...") '
        'Oh really? Tell me more — what is happening in your store? '
        'I would love to help you out!"',
        bg_hex='FFF3E0', label="📞  PHONE SCRIPT:", label_color=ORANGE)

    doc.add_page_break()

    # ─── SCRIPT 4: Instagram DM ───
    section_header(doc, "SCRIPT 4: INSTAGRAM / SOCIAL MEDIA DM", PURPLE)

    script_box(doc,
        'Hey [NAME]! Love your content / your store looks amazing! 👏\n\n'
        'Random question — do you happen to know any retail store owner '
        'who struggles with stock mismatch? Like their computer says 100 items '
        'but only 85 are on the shelf?\n\n'
        'I help fix this in 30 days with a simple process (no expensive software). '
        'Looking for 5 free case studies right now.\n\n'
        'Already helped a store recover Rs 3.2 Lakhs! '
        'Would love a referral if anyone comes to mind 🙏',
        bg_hex='F3E5F5', label="📱  SOCIAL DM:", label_color=PURPLE)

    # ─── SCRIPT 5: For people who ARE store owners ───
    section_header(doc, "SCRIPT 5: WHEN YOU KNOW THEY ARE A STORE OWNER", RED)
    styled(doc, "Use this when reaching out to someone who actually owns a store.", sz=11, c=RED, i=True, sa=6)

    script_box(doc,
        'Hey [NAME] ji! Hope business is going well 😊\n\n'
        'I have been working on something interesting — I help clothing store owners '
        'who face stock mismatch problems. You know, when the computer shows different '
        'numbers than what is actually on the shelves?\n\n'
        'I have a simple system that fixes this in 30 days. No expensive software. '
        'Just a clear process your staff can follow.\n\n'
        'I am doing 5 FREE case studies right now to build testimonials. '
        'One owner I worked with recovered Rs 3.2 Lakhs in 45 days!\n\n'
        'Would you be interested? Or do you know someone who might be? '
        'Either way, I would really appreciate it! 🙏\n\n'
        'I guarantee — if mismatch does not reduce in 30 days, I work free until it does. '
        'Zero risk.',
        bg_hex='FFEBEE', label="📋  DIRECT STORE OWNER SCRIPT:", label_color=RED)

    doc.add_page_break()

    # ─── SCRIPT 6: Follow Up ───
    section_header(doc, "SCRIPT 6: FOLLOW-UP MESSAGES", GOLD)
    styled(doc, "Send these 3-4 days after your first message if no reply.", sz=11, c=GRAY, i=True, sa=6)

    script_box(doc,
        'Hey [FRIEND NAME]! Just following up on my message from the other day. '
        'No worries if you do not know anyone right now!\n\n'
        'But if any shop owner friend of yours ever complains about '
        '"stock count not matching" or "items going missing", '
        'just remember me 😊 I would love to help them.\n\n'
        'Thanks [NAME]! 🙏',
        bg_hex='FFF8E1', label="💬  FOLLOW-UP MESSAGE:", label_color=GOLD)

    styled(doc, "📱  IF THEY REFERRED SOMEONE:", sz=13, c=GREEN, b=True, sa=4, sb=10)

    script_box(doc,
        'Thank you so much [FRIEND NAME]! Really appreciate the connection! 🙏\n\n'
        'I will reach out to them today and take great care of them. '
        'If this works out, chai-party on me! 😄\n\n'
        'By the way, if anyone else comes to mind later, do send them my way. '
        'Always happy to help!',
        bg_hex='E8F5E9')

    doc.add_page_break()

    # ─── DAILY PLAN ───
    section_header(doc, "YOUR DAILY WARM REACHOUT PLAN", ORANGE)

    plan = [
        ("9:30 AM", "Make a list of 15 people to contact today", BLUE),
        ("10:00 AM", "Send 10 WhatsApp messages (Script 2)", GREEN),
        ("11:00 AM", "Make 5 phone calls (Script 3)", ORANGE),
        ("2:00 PM", "Send 3-5 Instagram DMs (Script 4)", PURPLE),
        ("4:00 PM", "Follow up on yesterday's messages (Script 6)", GOLD),
        ("6:00 PM", "Update your tracking sheet", GRAY),
    ]

    for time, task, color in plan:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell1 = tbl.cell(0, 0)
        cell2 = tbl.cell(0, 1)
        shade(cell1, '1A1A2E')
        cell1.text = time
        cell2.text = task
        for p in cell1.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.size = Pt(11)
                r.bold = True
                r.font.name = 'Calibri'
        for p in cell2.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.name = 'Calibri'

    styled(doc, "\n🎯  WEEKLY TARGET: Contact 100 people in 7 days", sz=14, c=ORANGE, b=True,
           a=WD_ALIGN_PARAGRAPH.CENTER, sa=4, sb=10)
    styled(doc, "Even if only 5% say yes = 5 FREE case studies = 5 testimonials!", sz=13, c=GREEN, b=True,
           a=WD_ALIGN_PARAGRAPH.CENTER)

    # Final
    doc.add_page_break()
    for _ in range(5):
        doc.add_paragraph()

    styled(doc, "YOUR SCRIPTS ARE READY ✅", sz=28, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=8)
    styled(doc, "Just copy, personalize the name, and SEND.", sz=16, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
    styled(doc, "Target: 15 messages per day", sz=14, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    styled(doc, "= 100 people per week", sz=14, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    styled(doc, "= 5 case studies minimum", sz=14, c=GREEN, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
    styled(doc, "= Testimonials + paying clients! 💰", sz=14, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER)

    save_path = "/Users/aakash/Desktop/Week_3/Week 4/warm_reachout/Retail_Control_Warm_Reachout.docx"
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    return save_path


if __name__ == "__main__":
    p1 = create_generic_template()
    print(f"Created: {p1}")
    p2 = create_niche_reachout()
    print(f"Created: {p2}")
