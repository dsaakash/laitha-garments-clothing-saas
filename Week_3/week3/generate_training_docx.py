#!/usr/bin/env python3
"""
Generate a professionally formatted .docx for the Week 2 Training Document:
Build a Grand Slam Offer — Retail Control Architecture™
By Aakash Savant
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colour Palette ──────────────────────────────────────────────
DARK_NAVY    = RGBColor(0x0F, 0x1B, 0x2D)
DEEP_BLUE    = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD  = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL  = RGBColor(0x17, 0xA2, 0xB8)
WARM_RED     = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS_GRN  = RGBColor(0x27, 0xAE, 0x60)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
MED_GRAY     = RGBColor(0x66, 0x66, 0x66)
BODY_BLACK   = RGBColor(0x2C, 0x2C, 0x2C)
ORANGE       = RGBColor(0xE6, 0x7E, 0x22)
PURPLE       = RGBColor(0x8E, 0x44, 0xAD)

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BODY_BLACK

# ── Helpers ─────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)

def add_cover_title(text, size=26):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = DEEP_BLUE

def add_cover_sub(text, size=14, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(size)
    r.font.color.rgb = color or ACCENT_GOLD

def add_section_header(text, emoji=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(4)
    full = f"{emoji}  {text}" if emoji else text
    r = p.add_run(full)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DEEP_BLUE
    line = doc.add_paragraph()
    line.space_after = Pt(8)
    rl = line.add_run("━" * 75)
    rl.font.size = Pt(6)
    rl.font.color.rgb = ACCENT_GOLD

def add_sub_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color or ACCENT_TEAL

def add_minor_header(text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = color or DEEP_BLUE

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body_bold(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    return p

def add_quote(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'"{text}"')
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = color or MED_GRAY

def add_highlight(text, bg="FFF8E1"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = BODY_BLACK
    set_cell_shading(cell, bg)
    cell.width = Inches(6.0)
    add_spacer(6)

def add_bullet(text, bold_prefix=None, emoji=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.clear()
    if emoji:
        p.add_run(f"{emoji} ").font.size = Pt(11)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(11)
        p.add_run(f" {text}").font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_numbered(text, num, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.color.rgb = ACCENT_TEAL
    r.font.size = Pt(11)
    if bold_label:
        rl = p.add_run(f"{bold_label}: ")
        rl.bold = True
        rl.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)

def add_arrow_item(text, bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("→ ")
    r.bold = True
    r.font.color.rgb = ACCENT_GOLD
    r.font.size = Pt(11)
    if bold_label:
        rl = p.add_run(f"{bold_label}: ")
        rl.bold = True
        rl.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)

def styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        set_cell_shading(cell, "1B2A4A")
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, "F7F9FC")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    add_spacer(6)
    return table


# ════════════════════════════════════════════════════════════════
#                     DOCUMENT CONTENT
# ════════════════════════════════════════════════════════════════

# ── COVER PAGE ──────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_cover_title("BUILD A GRAND SLAM OFFER", 28)
add_cover_title("RETAIL CONTROL ARCHITECTURE™", 24)

doc.add_paragraph()
add_cover_sub("Week 2 — Training Document", 14, ACCENT_GOLD)
add_cover_sub("By Aakash Savant — Retail Control Architect", 12, MED_GRAY)

doc.add_paragraph()
doc.add_paragraph()

# Roadmap box
add_highlight(
    "THE 6-STEP ROADMAP\n\n"
    "Step 1: What problem are you solving?\n"
    "Step 2: For whom are you solving it?\n"
    "Step 3: Magnetic Positioning (Niche & Unique Mechanism)\n"
    "Step 4: 7 Steps to make the GSO ready\n"
    "Step 5: Assemble & Improve\n"
    "Step 6: Make it Pitch-Ready",
    "F0F4FF"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# PART A: GRAND SLAM OFFER OVERVIEW
# ════════════════════════════════════════════════════════════════
add_section_header("GRAND SLAM OFFER OVERVIEW", "🟥")

# ── STEP 1 ────────────────────────────────────────────────────
add_sub_header("STEP 1: What Problem Are You Solving?")

add_minor_header("Core Problem")
add_body("Established clothing retailers lack structured operational control. This leads to:")
add_bullet("Stock mismatch")
add_bullet("Hidden margin leakage")
add_bullet("Supplier confusion")
add_bullet("Staff dependency")
add_bullet("No single source of truth")
add_bullet("Daily reconciliation stress")

add_minor_header("Deeper Truth")
add_body("Retailers think they have a billing problem. They actually have a Control Architecture failure.")
add_highlight(
    "Software records transactions.\n"
    "It does NOT enforce discipline.\n\n"
    "I solve the Control Gap between physical inventory and system-recorded inventory.",
    "FDEDED"
)


# ── STEP 2 ────────────────────────────────────────────────────
add_sub_header("STEP 2: For Whom Are You Solving It?")

styled_table(
    ["Category", "Details"],
    [
        ["Primary Niche", "Established single-store clothing retailers"],
        ["Sub-Niche", "Owner-operated clothing stores\n₹30L–₹1.5Cr annual revenue\n₹20L+ inventory\n5–20 staff"],
        ["Micro-Niche", "Tier 2 & Tier 3 Indian city clothing retailers\nUsing manual or basic billing software\nFacing 10%+ stock mismatch\nOwner actively involved in daily operations"],
        ["NOT For", "Startups • Ecommerce-only brands\nMulti-branch chains\nRetailers seeking cheap billing software"],
    ],
    col_widths=[1.5, 4.5]
)


# ── STEP 3 ────────────────────────────────────────────────────
add_sub_header("STEP 3: Magnetic Positioning")

add_minor_header("🔹 Identity")
styled_table(
    ["Element", "Details"],
    [
        ["Title", "Retail Control Architect™"],
        ["NOT", "Software vendor. NOT ERP provider."],
        ["IS", "Architect of operational control."],
    ],
    col_widths=[1.2, 4.8]
)

add_minor_header("🔹 Category")
add_body("Retail Control Architecture™")
add_body("You are not competing in software. You are defining structural retail control.")

add_minor_header("🔹 Unique Mechanism: The Control Gap Method™")
add_body("A structured 3-layer architecture that closes the gap between physical stock and recorded stock inside clothing retail businesses.")

add_spacer(4)
add_minor_header("3-Layer Retail Control Architecture™")

styled_table(
    ["Layer", "Focus", "Components"],
    [
        ["Layer 1", "Structural Discipline", "SKU standardization\nSupplier workflow control\nSingle source of truth"],
        ["Layer 2", "Transaction Integrity", "Every inward tracked\nEvery sale deducts stock\nNo invisible movement"],
        ["Layer 3", "Owner Visibility", "Real-time dashboard\nIndependent access\nDependency elimination"],
    ],
    col_widths=[0.8, 1.5, 3.7]
)


# ── STEP 4 ────────────────────────────────────────────────────
add_sub_header("STEP 4: 7 Steps to Make the GSO Ready")

styled_table(
    ["#", "Step", "Details"],
    [
        ["1", "Define Dream Outcome", "Full operational control without staff dependency"],
        ["2", "Identify All Problems", "Stock mismatch, leakage, reconciliation chaos"],
        ["3", "Bridge Island #1 → #2", "Audit → Structure → Integrate → Enforce → Validate"],
        ["4", "Choose Delivery Vehicle", "1:1 Consulting-led SaaS\nDone-with-you enforcement\n30-day implementation"],
        ["5", "Supercharge Offer", "Scarcity (3 stores/month)\nUrgency (monthly cycles)\nGuarantee (measurable reduction or extended support)\nBonuses (Audit Report™, SOP™, Stability Review™)"],
        ["6", "Trim & Stack", "Remove non-core modules\nFocus only on inventory control"],
        ["7", "Maximize Perceived Value", "High dream outcome\nLow time delay (30 days)\nStructured process\nMeasurable proof"],
    ],
    col_widths=[0.3, 1.7, 4.0]
)


# ── STEP 5 ────────────────────────────────────────────────────
add_sub_header("STEP 5: Assemble & Improve")

add_minor_header("Offer Name")
add_body_bold("Retail Control Architecture™ Installation")

add_minor_header("Core Promise")
add_highlight(
    "In 30 days, we install a structured control system inside your clothing store that "
    "measurably reduces stock mismatch and eliminates staff dependency.",
    "E8F8F5"
)

add_minor_header("Delivery Components")
add_numbered("Measure stock mismatch and reconciliation time.", 1, "Inventory Reality Audit")
add_numbered("Clean categorization and pricing discipline.", 2, "SKU Standardization")
add_numbered("Structured inward and outstanding tracking.", 3, "Supplier Workflow Structuring")
add_numbered("Every sale deducts stock automatically.", 4, "Sales-Stock Integration")
add_numbered("No parallel manual systems.", 5, "Single Source of Truth Enforcement")
add_numbered("Real-time control access.", 6, "Owner Dashboard Activation")
add_numbered("Before vs After measurable comparison.", 7, "Control Validation Review")

add_minor_header("Guarantee")
add_highlight(
    "If measurable stock mismatch reduction is not achieved with full compliance, "
    "implementation support continues at no additional setup fee.\n\n"
    "No empty promises. Only measurable architecture.",
    "E8F8F5"
)

add_minor_header("Pricing Position")
add_bullet("Premium consulting-led SaaS")
add_bullet("Not price competition")
add_bullet("Outcome-based pricing")


# ── STEP 6 ────────────────────────────────────────────────────
add_sub_header("STEP 6: Pitch Version (Short)")
add_highlight(
    "\"We don't sell billing software.\n"
    "We install Retail Control Architecture inside established clothing stores.\n"
    "In 30 days, we close your Control Gap and restore measurable operational control.\"",
    "FFF3E0"
)

add_minor_header("Final Positioning Line")
add_body_bold("From stock confusion to structural control in 30 days.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# PART B: POC STRATEGY
# ════════════════════════════════════════════════════════════════
add_section_header("POC STRATEGY", "🟥")
add_body_bold("Retail Control Architecture™ — Proof of Concept Plan")

add_highlight(
    "If you don't have 3 strong proofs, you don't have authority.\n"
    "You build proof first. Then scale.\n\n"
    "NOTE: If you don't have POC, you need to build it BEFORE charging money.\n"
    "You work for free — 3 people.",
    "FDEDED"
)

# POC Monetization path
add_minor_header("The POC Monetization Ladder")
styled_table(
    ["Client", "Strategy", "Outcome"],
    [
        ["Client 1 (Free)", "Prove value inside their store", "→ Convert to paid subscription"],
        ["Client 2 (Free)", "Collect video testimonial", "→ Use testimonial to close next paid client"],
        ["Client 3 (Free)", "Ask for direct referral", "→ Referral introduction to new client"],
    ],
    col_widths=[1.2, 2.5, 2.8]
)

add_body_bold("Most Important: Collect a Video Testimonial from ALL 3.")

# Phase 1
add_sub_header("Phase 1: Select the Right 3 Free Clients", SUCCESS_GRN)
add_body("Not random stores. Qualification criteria:")
add_bullet("Clothing retailer", emoji="✅")
add_bullet("₹30L+ revenue", emoji="✅")
add_bullet("₹20L+ inventory", emoji="✅")
add_bullet("Visible stock mismatch", emoji="✅")
add_bullet("Owner actively involved", emoji="✅")
add_bullet("Willing to follow single-source-of-truth rule", emoji="✅")

add_highlight(
    "If they resist discipline, REJECT them.\n"
    "Free client must be your ideal client.",
    "FFF8E1"
)

# Phase 2
add_sub_header('Phase 2: Position It Properly (NOT "Free Work")', ORANGE)
add_body("You do NOT say:")
add_highlight('"I\'ll do it free."   ← NEVER say this.', "FDEDED")

add_body("You say:")
add_highlight(
    '"I\'m selecting 3 pilot stores to implement my Retail Control Architecture™ framework '
    'in exchange for measurable data and a video testimonial."',
    "E8F8F5"
)

add_body("Position it as:")
add_bullet("Pilot Program")
add_bullet("Limited Selection")
add_bullet("Structured Case Study")
add_body_bold("Free = weak positioning.   Pilot = authority positioning.")

# Phase 3
add_sub_header("Phase 3: Execution Plan for Each Client", ACCENT_TEAL)
add_body("For each of the 3 pilot stores:")

add_minor_header("A. Baseline Measurement")
add_bullet("Stock mismatch %")
add_bullet("Reconciliation time")
add_bullet("Dependency risk")
add_bullet("Supplier clarity level")

add_minor_header("B. 30-Day Implementation")
add_bullet("SKU discipline")
add_bullet("Supplier workflow")
add_bullet("Sales-stock sync")
add_bullet("Enforcement")

add_minor_header("C. Validation")
add_bullet("Re-measure mismatch")
add_bullet("Re-measure reconciliation time")
add_bullet("Document improvement")

add_highlight(
    "You must collect HARD NUMBERS.\n"
    "If you cannot quantify, the POC is weak.",
    "FFF8E1"
)

# Phase 4
add_sub_header("Phase 4: Video Testimonial Collection (Non-Negotiable)", WARM_RED)
add_body("You script it. Ask them to answer:")
add_numbered("What problems were you facing before?", 1)
add_numbered("What was your stock mismatch?", 2)
add_numbered("What changed after implementation?", 3)
add_numbered("What measurable improvement did you see?", 4)
add_numbered("Would you recommend this to other clothing retailers?", 5)

add_body_bold("Keep it structured. 2–3 minutes max. No vague praise. Only measurable statements.")

# Phase 5
add_sub_header("Phase 5: Monetization Ladder", ACCENT_GOLD)
add_arrow_item("Convert to paid subscription → You've proven value inside their store.", "Client 1")
add_arrow_item("Use testimonial to close next paid client.", "Client 2")
add_arrow_item("Ask for direct referral introduction → Retail clusters talk.", "Client 3")

# Phase 6
add_sub_header("Phase 6: Authority Packaging", PURPLE)
add_body("After 3 POCs, you create:")
add_bullet("3 written case studies")
add_bullet("3 video testimonials")
add_bullet("1 before-after metrics summary")
add_bullet("Average % mismatch reduction")
add_bullet("Average time saved")

add_body("Now your pitch changes from:")
add_highlight(
    'FROM: "Trust me."\n\n'
    'TO: "In 3 established clothing stores, we reduced mismatch by an average of X% within 30 days."\n\n'
    "That's authority.",
    "E8F8F5"
)

# Critical Rule
add_sub_header("⚠️ Critical Discipline Rule", WARM_RED)
add_highlight(
    "Even during FREE POC:\n"
    "→ No parallel manual systems allowed.\n"
    "→ If they resist discipline: TERMINATE the pilot.\n"
    "→ Weak POC damages future positioning.\n\n"
    "After 3 successful POCs: NO MORE FREE WORK.",
    "FDEDED"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# PART C: FULLY ASSEMBLED GRAND SLAM OFFER
# ════════════════════════════════════════════════════════════════
add_section_header("THE FULLY ASSEMBLED GRAND SLAM OFFER", "🟥")
add_body_bold("Retail Control Architecture™ — By Aakash Savant")
add_body("Clean. Document-ready. No placeholders. Use as sales page, proposal PDF, VSL script, or discovery call spine.")

# Step 1
add_sub_header("STEP 1: WHAT PROBLEM ARE YOU SOLVING?")

add_minor_header("Core Problem")
add_body(
    "Established clothing retailers lose control because inventory, supplier tracking, "
    "and sales systems are unstructured and staff-dependent."
)
add_body_bold("The store runs on people, not on systems.")

add_minor_header("Deeper Truth")
add_bullet("Most retailers buy billing software before designing control architecture.")
add_bullet("They hire more staff instead of fixing workflows.")
add_bullet("They automate transactions without enforcing structure.")

add_highlight(
    "Software records entries. It does NOT enforce discipline.\n\n"
    "Stock mismatch continues because ownership and operational architecture "
    "were never designed properly.",
    "FDEDED"
)

add_minor_header("Cost of the Problem (Time, Money, Risk)")
styled_table(
    ["Impact Area", "Annual Cost"],
    [
        ["Stock Mismatch", "10–20% of inventory value"],
        ["Margin Leakage", "₹1–3L annually"],
        ["Daily Reconciliation", "1–2 hours wasted every day"],
        ["Supplier Overbilling", "Unnoticed for months"],
        ["Staff Dependency Risk", "Business collapses if key staff leaves"],
        ["Owner Independence", "Cannot step away confidently"],
        ["Scale Readiness", "Second store feels too risky"],
    ],
    col_widths=[2.5, 3.5]
)
add_body_bold("The business grows. Control does not.")


# Step 2
add_sub_header("STEP 2: FOR WHOM ARE YOU SOLVING IT?")

add_minor_header("Target Market")
add_body("Established, owner-operated clothing retail stores.")

add_minor_header("Ideal Buyer Profile")
styled_table(
    ["Attribute", "Criteria"],
    [
        ["Revenue", "₹30L – ₹1.5Cr annually"],
        ["Inventory Value", "₹20L+"],
        ["Team Size", "5–20 staff"],
        ["Current System", "Manual or basic billing software"],
        ["Pain Signal", "Visible stock mismatch"],
        ["Owner Role", "Involved in daily operations"],
    ],
    col_widths=[2.0, 4.0]
)

add_minor_header("This Offer Is NOT For")
add_bullet("Early-stage retail startups", emoji="❌")
add_bullet("Ecommerce-only brands", emoji="❌")
add_bullet("Multi-branch chains", emoji="❌")
add_bullet("Retailers looking for cheap billing software", emoji="❌")
add_bullet("Stores below ₹20L annual revenue", emoji="❌")
add_body_bold("This is for retailers serious about control.")


# Step 3
add_sub_header("STEP 3: MAGNETIC POSITIONING")

add_minor_header("🔹 Niche (What + Whom)")
add_body(
    "I help established clothing retailers eliminate stock mismatch and operational leakage "
    "by installing structured Retail Control Architecture™ inside their stores."
)

add_minor_header("🔹 Unique Mechanism: Retail Control Architecture™")
add_body("Powered by the Control Gap Method™")

add_highlight(
    "A proprietary 30-day implementation system that closes the gap between physical inventory "
    "and system-recorded inventory through structured discipline, transaction integrity, "
    "and owner-level visibility.",
    "F0F4FF"
)

add_minor_header("Why It Works (Sequence Competitors Miss)")
add_numbered("Structure before software", 1)
add_numbered("SKU discipline before automation", 2)
add_numbered("Supplier workflow before reporting", 3)
add_numbered("Single Source of Truth enforcement", 4)
add_numbered("Measurable validation before scaling", 5)
add_body_bold("Most vendors sell tools. I install control architecture.")


# Step 4
add_sub_header("STEP 4: 7 STEPS TO MAKE THE GRAND SLAM OFFER READY")

styled_table(
    ["#", "Step", "Details"],
    [
        ["1", "Dream Outcome Defined", "A clothing store that runs on structured control, not guesswork"],
        ["2", "Problems Identified", "Stock mismatch, leakage, staff dependency, reconciliation stress"],
        ["3", "Solution Path", "Audit → Structure → Integrate → Enforce → Validate"],
        ["4", "Delivery Vehicle", "1:1 Consulting-Led Implementation\n30 Days • Done-With-You Enforcement"],
        ["5", "Superchargers Added", "Urgency • Scarcity • Guarantee • Bonuses"],
        ["6", "Trim & Stack", "Only components that directly reduce mismatch and restore control"],
        ["7", "Perceived Value ↑", "Less effort for owner • Faster implementation\nHigher certainty • Bigger vision (scale readiness)"],
    ],
    col_widths=[0.3, 1.7, 4.0]
)


# Step 5
add_sub_header("STEP 5: ASSEMBLE THE OFFER (FINAL STACK)")

add_minor_header("Offer Name")
add_body_bold("Retail Control Architecture™ Installation")

add_minor_header("Core Promise (1 Sentence)")
add_highlight(
    "In 30 days, your clothing store will operate on measurable inventory control — "
    "eliminating stock mismatch, reducing leakage, and removing staff dependency.",
    "E8F8F5"
)

add_minor_header("Main Offer (Core Delivery)")
styled_table(
    ["#", "Component", "Outcome"],
    [
        ["1", "Inventory Reality Audit", "Measure stock mismatch and reconciliation time"],
        ["2", "SKU Structure Standardization", "Clean categorization and pricing discipline"],
        ["3", "Supplier Workflow Architecture", "Structured inward and outstanding tracking"],
        ["4", "Sales-Stock Integration", "Every sale deducts stock automatically"],
        ["5", "Single Source of Truth Enforcement", "No parallel manual systems"],
        ["6", "Owner Visibility Dashboard", "Real-time control access"],
        ["7", "Control Validation Review", "Before vs After measurable comparison"],
    ],
    col_widths=[0.3, 2.2, 3.5]
)

add_minor_header("Bonuses (Risk Reducers)")
styled_table(
    ["#", "Bonus", "Purpose"],
    [
        ["🎁 1", "Inventory Leakage Audit Report™", "Quantifies hidden loss"],
        ["🎁 2", "Staff Operational Discipline SOP™", "Ready-to-use training framework"],
        ["🎁 3", "Supplier Optimization Checklist™", "Streamline supplier processes"],
        ["🎁 4", "90-Day Stability Review", "Ensure results are permanent"],
        ["🎁 5", "Expansion Readiness Blueprint™", "Scale-ready audit"],
    ],
    col_widths=[0.5, 2.5, 3.0]
)

add_minor_header("Guarantee (Ethical & Believable)")
add_highlight(
    "If after full compliance with the 30-day framework, measurable stock mismatch reduction "
    "is not achieved, implementation support continues at no additional setup fee.\n\n"
    "No empty promises. Only measurable architecture.",
    "E8F8F5"
)

add_minor_header("Urgency (Why Now)")
add_bullet("Operational leakage compounds monthly")
add_bullet("Fixing control at ₹50L revenue is easier than fixing it at ₹1Cr")
add_bullet("Delay increases structural chaos")

add_minor_header("Scarcity (Why Not Everyone)")
add_body("Only 3 clothing retailers onboarded per month due to hands-on implementation depth.")
add_body("No bulk SaaS onboarding.")

add_minor_header("Pricing")
styled_table(
    ["Item", "Details"],
    [
        ["Total Perceived Value", "₹3–6L annually (leakage reduction + time savings)"],
        ["Setup Investment", "₹60K – ₹1L"],
        ["Monthly Subscription", "₹5K – ₹8K/month"],
        ["Payment Options", "Full pay OR Milestone split"],
    ],
    col_widths=[2.0, 4.0]
)


# Step 6
add_sub_header("STEP 6: MAKE IT PITCH-READY")

add_minor_header("30-Second Pitch")
add_highlight(
    '"This is a 30-day Retail Control Architecture™ installation.\n\n'
    "We measure your stock mismatch, restructure your inventory and supplier workflows, "
    "enforce a single source of truth, and validate measurable improvement.\n\n"
    'It\'s designed for established clothing retailers who want real control — not just billing software."',
    "FFF3E0"
)

add_minor_header("Call-to-Action")
add_body_bold("Book a Retail Control Audit Call.")
add_body("If your store doesn't have a real Control Gap, I'll tell you directly.")


# Step 7
add_sub_header("STEP 7: VALUE EQUATION CHECK (FINAL)")

styled_table(
    ["Variable", "Rating", "How"],
    [
        ["Dream Outcome", "HIGH", "Operational control + reduced leakage + independence"],
        ["Perceived Likelihood", "HIGH", "Done-with-you + enforcement + measurable baseline"],
        ["Time Delay", "SHORT", "30 days"],
        ["Effort Required", "MODERATE\nbut structured", "Owner approves. System enforces."],
    ],
    col_widths=[1.5, 1.0, 3.5]
)

add_body_bold("Result: Premium, outcome-driven, authority-positioned offer.")


# ── Final Positioning ──────────────────────────────────────────
add_spacer(12)

add_highlight(
    "FINAL POSITIONING LINE — USE EVERYWHERE\n\n"
    "I don't sell retail software.\n"
    "I install Retail Control Architecture™ inside clothing stores.",
    "1B2A4A"
)

# Fix the text color in that dark box - need to override
# Let's just do a clean centered final statement instead
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(12)
r1 = p.add_run("I don't sell retail software.\n")
r1.bold = True
r1.font.size = Pt(16)
r1.font.color.rgb = DEEP_BLUE
r2 = p.add_run("I install Retail Control Architecture™ inside clothing stores.")
r2.bold = True
r2.font.size = Pt(16)
r2.font.color.rgb = ACCENT_GOLD

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— Aakash Savant, Retail Control Architect")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = MED_GRAY


# ── Save ───────────────────────────────────────────────────────
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Week_2_Training_Grand_Slam_Offer.docx"
)
doc.save(output_path)
print(f"✅ Successfully created: {output_path}")
