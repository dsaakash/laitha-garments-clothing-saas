#!/usr/bin/env python3
"""
Generate a YouTube Script in Steve Jobs Style for 'Aakash Savant' channel.
Output: .docx
Tone: Authoritative, Visionary, Story-driven.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_script():
    doc = Document()

    # Document Styling
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # --- Header ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("YOUTUBE SCRIPT: THE RETAIL REVOLUTION\n")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0, 0, 0)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = subtitle.add_run("Channel: Aakash Savant | Style: Steve Jobs Keynote")
    r_sub.italic = True
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Act 1: The Hook ---
    doc.add_heading("ACT 1: THE SILENT LEAK", level=1)
    
    p = doc.add_paragraph()
    p.add_run("(Visual: Aakash stands in front of a pitch-black background. He looks directly into the lens. Silence for 3 seconds.)").italic = True
    
    p = doc.add_paragraph()
    p.add_run("AAKASH: ").bold = True
    p.add_run("You spend your whole life building a store. You select the best fabric. You hire the best staff. You work from 9 AM to 10 PM. ")
    
    p = doc.add_paragraph()
    p.add_run("But there’s something you don’t see. ")
    
    p = doc.add_paragraph()
    p.add_run("(Visual: A slide appears with a single, massive number: ₹2,739)").italic = True
    
    p = doc.add_paragraph()
    p.add_run("AAKASH: ").bold = True
    p.add_run("That is the amount of money leaving your pocket... every single day. Not because you aren’t selling. But because your shop doesn’t match your system.")

    # --- Act 2: The Villain ---
    doc.add_heading("ACT 2: THE SOFTWARE LIE", level=1)
    
    p = doc.add_paragraph()
    p.add_run("AAKASH: ").bold = True
    p.add_run("For twenty years, the industry told you that 'Software' would solve this. They gave you complex buttons, confusing dashboards, and data that nobody reads. ")
    
    p = doc.add_paragraph()
    p.add_run("What happened? You still look at a computer screen that says you have 5 shirts, while your hand is holding 3. ")
    
    p = doc.add_paragraph()
    p.add_run("Software didn't fix retail. It just gave you a digital version of chaos.")

    # --- Act 3: The Revelation ---
    doc.add_heading("ACT 3: INTRODUCING THE AUTHORITY", level=1)
    
    p = doc.add_paragraph()
    p.add_run("(Visual: Dramatic music builds. The logo 'RETAIL CONTROL AUTHORITY' fades in.)").italic = True
    
    p = doc.add_paragraph()
    p.add_run("AAKASH: ").bold = True
    p.add_run("Today, we’re changing that. We’re not giving you another software. We’re installing a Single Source of Truth.")
    
    p = doc.add_paragraph()
    p.add_run("We call it the Retail Control Authority System. It is the most advance way to run an Indian clothing store. Ever.")

    # --- Act 4: The 10-Minute Demonstation ---
    doc.add_heading("ACT 4: THE BEAUTY OF SIMPLICITY", level=1)
    
    p = doc.add_paragraph()
    p.add_run("AAKASH: ").bold = True
    p.add_run("We looked at the most successful stores in the world. They don't have smarter staff. They have stricter systems.")
    
    p = doc.add_paragraph()
    p.add_run("We've distilled everything down to three rules. Inward gatekeeper. Sales sync. Daily pulse. ")
    
    p = doc.add_paragraph()
    p.add_run("(Visual: Aakash pulls out his phone)").italic = True
    
    p = doc.add_paragraph()
    p.add_run("AAKASH: ").bold = True
    p.add_run("Our owners don't spend 2 hours reconciling stock anymore. They spend 10 minutes. Before they leave for home, they check their phone. If it's green, the store is safe. If it's not, they know exactly where the leakage happened... before it becomes a habit.")

    # --- Act 5: One More Thing ---
    doc.add_heading("ACT 5: THE CERTAINTY", level=1)
    
    p = doc.add_paragraph()
    p.add_run("AAKASH: ").bold = True
    p.add_run("Now... there’s one more thing.")
    
    p = doc.add_paragraph()
    p.add_run("(Visual: A slide with 'STOCK CERTAINTY GUARANTEE')").italic = True
    
    p = doc.add_paragraph()
    p.add_run("AAKASH: ").bold = True
    p.add_run("I’m not a software vendor. I’m a partner in your profit. If after 30 days of our installation, your system stock doesn't match your physical stock—we don't want your money. We work for free until it does.")

    # --- Act 6: The Call to Action ---
    doc.add_heading("ACT 6: THE ENCLOSURE", level=1)
    
    p = doc.add_paragraph()
    p.add_run("AAKASH: ").bold = True
    p.add_run("If you are tired of working for your store, and you want your store to start working for you... let’s talk. ")
    
    p = doc.add_paragraph()
    p.add_run("I only take on 3 new installations a month. Why? Because forensic accuracy can't be rushed. ")
    
    p = doc.add_paragraph()
    p.add_run("Click the link below. Book your Stock Audit. Let’s stop the leaking. ")
    
    p = doc.add_paragraph()
    p.add_run("I'm Aakash Savant. Let's build your authority.")

    # Save
    save_path = "/Users/aakash/Desktop/Week_3/YouTube_Script_Aakash_Savant_Steve_Jobs_Style.docx"
    doc.save(save_path)
    return save_path

if __name__ == "__main__":
    path = create_script()
    print(f"Script created: {path}")
