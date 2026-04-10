#!/usr/bin/env python3
"""
Generate 4 files for Week 1:
1. 6-7 Figure Blueprint Session DOCX — Niche (Retail Control Architect)
2. 6-7 Figure Blueprint Session DOCX — Template
3. 6-7 Figure Business Blueprint XLSX — Niche (Retail Control Architect)
4. 6-7 Figure Business Blueprint XLSX — Template
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import os

ORANGE = RGBColor(0xFF, 0x6B, 0x35)
GOLD = RGBColor(0xD4, 0xA5, 0x00)
GREEN = RGBColor(0x00, 0xA8, 0x55)
BLUE = RGBColor(0x3D, 0x8B, 0xE8)
RED = RGBColor(0xE8, 0x3D, 0x3D)
PURPLE = RGBColor(0x6C, 0x5C, 0xE7)
BLACK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, hex_color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), hex_color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)


def styled(doc, text, sz=12, c=BLACK, b=False, i=False, a=WD_ALIGN_PARAGRAPH.LEFT, sa=6, sb=0):
    p = doc.add_paragraph()
    p.alignment = a
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.color.rgb = c
    r.bold = b
    r.italic = i
    r.font.name = 'Calibri'
    return p


def section_header(doc, text, color=ORANGE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"━━━  {text}")
    r.font.size = Pt(20)
    r.font.color.rgb = color
    r.bold = True
    r.font.name = 'Calibri'


def add_pillar_section(doc, week_num, pillar_name, emoji, color, description, tasks, is_template=False):
    """Add a pillar/week section with tasks."""
    section_header(doc, f"WEEK {week_num} — PILLAR {week_num}: {pillar_name} {emoji}", color)
    styled(doc, description, sz=12, c=GRAY, i=True, sa=8)

    tbl = doc.add_table(rows=len(tasks) + 1, cols=4 if is_template else 3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    if is_template:
        headers = ["#", "TASK", "STATUS", "NOTES / LINK"]
    else:
        headers = ["#", "TASK", "STATUS"]

    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        shade(cell, '1A1A2E')
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.size = Pt(10)
                r.bold = True
                r.font.name = 'Calibri'

    for idx, task_info in enumerate(tasks):
        task = task_info[0]
        status = task_info[1] if len(task_info) > 1 else ""
        row_data = [str(idx + 1), task, status]
        if is_template:
            row_data.append("")

        for j, val in enumerate(row_data):
            cell = tbl.cell(idx + 1, j)
            cell.text = val
            # Color status
            bg = 'FFFFFF'
            if status == "Done" or status == "✅ Done":
                bg = 'E8F5E9'
            elif "process" in status.lower() or status == "⬜ In Progress":
                bg = 'FFF3E0'
            elif "not" in status.lower() or status == "⬜ Not Started":
                bg = 'FFEBEE'
            if j >= 2:
                shade(cell, bg)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Calibri'

    styled(doc, "", sz=4, sa=8)


def create_blueprint_docx(is_template=False):
    """Create the 6-7 Figure Blueprint Session DOCX."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(12)

    # ═══ COVER ═══
    for _ in range(3):
        doc.add_paragraph()

    if is_template:
        styled(doc, "IEC DIAMOND", sz=14, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER)
        styled(doc, "THE 6-7 FIGURE", sz=40, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        styled(doc, "BLUEPRINT SESSION", sz=36, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
        styled(doc, "TEMPLATE GUIDE", sz=20, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
        styled(doc, "💎", sz=30, a=WD_ALIGN_PARAGRAPH.CENTER, sa=16)
        styled(doc, "Fill in your own details and follow the 9 pillars step-by-step",
               sz=14, c=GRAY, i=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=20)
    else:
        styled(doc, "IEC DIAMOND — RETAIL CONTROL ARCHITECT", sz=14, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER)
        styled(doc, "THE 6-7 FIGURE", sz=40, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        styled(doc, "BLUEPRINT SESSION", sz=36, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
        styled(doc, "by Aakash Savant", sz=16, c=GRAY, i=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
        styled(doc, "💎", sz=30, a=WD_ALIGN_PARAGRAPH.CENTER, sa=20)

    styled(doc, "━" * 35, sz=12, c=ORANGE, a=WD_ALIGN_PARAGRAPH.CENTER, sa=12)

    # 9 Pillars Overview
    styled(doc, "THE 9 PILLARS TO 6-7 FIGURES", sz=18, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=10)

    pillars_overview = [
        ("Week 1", "Pillar 1", "Foundation (YOU)", "🧱"),
        ("Week 2", "Pillar 2", "Grand Slam Offer", "🎯"),
        ("Week 3", "Pillar 3", "Funnel", "🔗"),
        ("Week 4", "Pillar 4", "Marketing Systems", "📣"),
        ("Week 5", "Pillar 5", "Sales System", "💰"),
        ("Week 6", "Pillar 6", "One-to-Many Selling", "🎤"),
        ("Week 7", "Pillar 7", "Personal Branding", "⭐"),
        ("Week 8", "Pillar 8", "Automation", "⚙️"),
        ("Week 9", "Pillar 9", "Delivery Mechanism", "📦"),
    ]

    tbl = doc.add_table(rows=10, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["WEEK", "PILLAR", "FOCUS AREA", ""]):
        cell = tbl.cell(0, i)
        shade(cell, '1A1A2E')
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.font.size = Pt(10)
                r.bold = True
                r.font.name = 'Calibri'

    colors_list = ['FFF3E0', 'E3F2FD', 'F3E5F5', 'E8F5E9', 'FFF8E1', 'FCE4EC', 'E0F7FA', 'F5F5F5', 'FBE9E7']
    for i, (week, pillar, focus, emoji) in enumerate(pillars_overview):
        for j, val in enumerate([week, pillar, focus, emoji]):
            cell = tbl.cell(i + 1, j)
            shade(cell, colors_list[i])
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Calibri'

    doc.add_page_break()

    # ═══ PERSONAL DETAILS ═══
    section_header(doc, "YOUR PERSONAL DETAILS", RED)

    if is_template:
        fields = [
            ("✏️ Your Name:", "_" * 40),
            ("✏️ When Started:", "_" * 40),
            ("✏️ Your Goal:", "_" * 40),
            ("✏️ Current Revenue:", "_" * 40),
            ("✏️ Revenue Goal (6 months):", "_" * 40),
        ]
    else:
        fields = [
            ("Name:", "Aakash Savant"),
            ("When Started:", "February 2026"),
            ("Your Goal:", "Build a 7-figure retail consulting business"),
            ("Current Revenue:", "Starting phase — building case studies"),
            ("Revenue Goal (6 months):", "Rs 3 Lakh / month from Retail Control Architect"),
        ]

    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(fields):
        tbl.cell(i, 0).text = label
        tbl.cell(i, 1).text = value
        shade(tbl.cell(i, 0), 'FFF3E0')
        for j in range(2):
            for p in tbl.cell(i, j).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Calibri'
                    if j == 0:
                        r.bold = True

    doc.add_page_break()

    # ═══ PILLAR DETAILS ═══
    # Define all pillars with their tasks

    if is_template:
        # TEMPLATE VERSION — generic tasks with blank status
        all_pillars = [
            (1, "FOUNDATION (YOU)", "🧱", ORANGE,
             "Build your personal foundation — goals, mindset, habits, and vision.",
             [
                 ("Fill in your 6-7 Figure Business Blueprint sheet", "⬜ Not Started"),
                 ("Implement time management system", "⬜ Not Started"),
                 ("Complete Old You vs New You exercise", "⬜ Not Started"),
                 ("Take Myers-Briggs Personality Test", "⬜ Not Started"),
                 ("Create your Vision Map", "⬜ Not Started"),
                 ("Create your Vision Board", "⬜ Not Started"),
                 ("Set up your Tracking Sheet", "⬜ Not Started"),
                 ("Define Business Revenue & Profitability Goals", "⬜ Not Started"),
                 ("Optimize Body, Mind, Place & Work habits", "⬜ Not Started"),
                 ("Practice Inversion Technique (monthly)", "⬜ Not Started"),
                 ("Select a book to read this month", "⬜ Not Started"),
             ]),
            (2, "GRAND SLAM OFFER", "🎯", BLUE,
             "Create an irresistible offer that your market cannot refuse.",
             [
                 ("Select your Niche", "⬜ Not Started"),
                 ("Create your Marketing Message", "⬜ Not Started"),
                 ("Build your Grand Slam Offer", "⬜ Not Started"),
                 ("Prepare Proof of Concept (POC)", "⬜ Not Started"),
                 ("Set the Price of your Offer", "⬜ Not Started"),
                 ("Prepare your entire Sales Pitch", "⬜ Not Started"),
                 ("Complete Competitor Research Sheet", "⬜ Not Started"),
             ]),
            (3, "FUNNEL", "🔗", PURPLE,
             "Build your lead generation and sales funnel.",
             [
                 ("Study and understand the Funnel world", "⬜ Not Started"),
                 ("Select which funnel to follow (VSL/Webinar/Event/Direct)", "⬜ Not Started"),
                 ("Write headline & sub-headline for landing page", "⬜ Not Started"),
                 ("Write complete landing page copy", "⬜ Not Started"),
                 ("Build complete landing page", "⬜ Not Started"),
                 ("Create your VSL (Video Sales Letter)", "⬜ Not Started"),
                 ("Launch your funnel", "⬜ Not Started"),
                 ("Build your sales page", "⬜ Not Started"),
                 ("Build your official website", "⬜ Not Started"),
                 ("Complete competitor funnel hack research", "⬜ Not Started"),
             ]),
            (4, "MARKETING SYSTEMS", "📣", GREEN,
             "Set up systems to attract leads consistently.",
             [
                 ("Create your Lead Magnet", "⬜ Not Started"),
                 ("Implement Warm Reachout strategy", "⬜ Not Started"),
                 ("Create 3 months of content (GAP Framework)", "⬜ Not Started"),
                 ("Start posting content on all platforms", "⬜ Not Started"),
                 ("Prepare Content Distribution sheet", "⬜ Not Started"),
                 ("Set up Outreach Tracking sheet", "⬜ Not Started"),
                 ("Start daily outreach", "⬜ Not Started"),
                 ("Create 10 Facebook/Instagram ads (written + filmed)", "⬜ Not Started"),
                 ("Do marketing activity daily (min 3 hours/day)", "⬜ Not Started"),
             ]),
            (5, "SALES SYSTEM", "💰", GOLD,
             "Build your closing system to convert leads into clients.",
             [
                 ("Personalize the 8-Figure Closing Script for your industry", "⬜ Not Started"),
                 ("Prepare Objection Handling list + practice", "⬜ Not Started"),
                 ("Write 10 Follow-up Messages", "⬜ Not Started"),
                 ("Create Pre-framing Video", "⬜ Not Started"),
                 ("Build Follow-up & Pre-framing sheet", "⬜ Not Started"),
                 ("Watch Role Play Training", "⬜ Not Started"),
                 ("Do live Role Play practice", "⬜ Not Started"),
                 ("Write Reminder Messages for before calls", "⬜ Not Started"),
                 ("Prepare 1:1 Zoom call procedure + practice", "⬜ Not Started"),
                 ("Set up 1:1 Strategy Call tracking sheet", "⬜ Not Started"),
             ]),
            (6, "ONE-TO-MANY SELLING", "🎤", RED,
             "Learn to sell to many people at once through webinars and events.",
             [
                 ("Create Webinar PPT as per strategy", "⬜ Not Started"),
                 ("Set up Webinar Tracking sheet", "⬜ Not Started"),
                 ("Set up Webinar Sales tracking sheet", "⬜ Not Started"),
                 ("Write Webinar WhatsApp reminder messages", "⬜ Not Started"),
                 ("Practice your first webinar", "⬜ Not Started"),
                 ("Conduct your first webinar", "⬜ Not Started"),
                 ("Purchase Zoom & start using it", "⬜ Not Started"),
             ]),
            (7, "PERSONAL BRANDING", "⭐", PURPLE,
             "Build your personal brand to attract clients organically.",
             [
                 ("Complete Content Competitor Analysis", "⬜ Not Started"),
                 ("Create Content Distribution sheet", "⬜ Not Started"),
                 ("Set up Editor Work sheet", "⬜ Not Started"),
                 ("Start posting content on all platforms consistently", "⬜ Not Started"),
                 ("Optimize all social media profiles", "⬜ Not Started"),
             ]),
            (8, "AUTOMATION", "⚙️", BLUE,
             "Automate your business to save time and scale.",
             [
                 ("List all tools you will use (after Tools Training)", "⬜ Not Started"),
                 ("Start automation using Business Automation Mapping", "⬜ Not Started"),
             ]),
            (9, "CLIENT DELIVERY", "📦", GREEN,
             "Set up world-class delivery for your clients.",
             [
                 ("Create Proposal Template for your offer", "⬜ Not Started"),
                 ("Create Agreement/Contract template", "⬜ Not Started"),
                 ("Create Timeline & Delivery template for onboarding", "⬜ Not Started"),
                 ("Create Client Onboarding Process", "⬜ Not Started"),
             ]),
        ]
    else:
        # NICHE VERSION — Retail Control Architect specific
        all_pillars = [
            (1, "FOUNDATION (YOU)", "🧱", ORANGE,
             "Build your personal foundation as the Retail Control Architect.",
             [
                 ("Fill in 6-7 Figure Business Blueprint sheet", "✅ Done"),
                 ("Implement time management — 3-hour deep work blocks", "✅ Done"),
                 ("Old You vs New You — from confused to retail authority", "✅ Done"),
                 ("Myers-Briggs Personality Test completed", "✅ Done"),
                 ("Vision Map — Retail Control Architect in 500+ stores by 2028", "✅ Done"),
                 ("Vision Board — store visits, happy owners, team growth", "✅ Done"),
                 ("Daily + Weekly Tracking Sheet set up", "✅ Done"),
                 ("Revenue Goal: Rs 3L/month in 6 months from consulting", "✅ Done"),
                 ("Morning routine optimized: Exercise + Reading + Planning", "✅ Done"),
                 ("Inversion Technique: What could kill my business? Prevent it.", "✅ Done"),
                 ("Reading: $100M Offers by Alex Hormozi", "✅ Done"),
             ]),
            (2, "GRAND SLAM OFFER", "🎯", BLUE,
             "Create an irresistible offer for retail store owners.",
             [
                 ("Niche: Clothing & retail store owners with stock mismatch problems", "✅ Done"),
                 ("Marketing Message: Stop losing lakhs from stock mismatch in 30 days", "✅ Done"),
                 ("Grand Slam Offer: 5-Step Enforcement Framework + 30-day guarantee", "✅ Done"),
                 ("POC: Rs 3.2L recovered for client, 17% to 2% mismatch reduction", "✅ Done"),
                 ("Pricing: Free audit + paid implementation package", "✅ Done"),
                 ("15-Slide Sales Pitch Deck ready", "✅ Done"),
                 ("Competitor Sheet: Analyzed ERP sellers, Vyapar, Busy software", "✅ Done"),
             ]),
            (3, "FUNNEL", "🔗", PURPLE,
             "Build lead generation funnel for Retail Control Architect.",
             [
                 ("Studied VSL, Webinar, Direct booking funnels", "✅ Done"),
                 ("Selected: VSL funnel + Direct call booking", "✅ Done"),
                 ("Landing page headline: Stop Losing Money From Your Own Store", "✅ Done"),
                 ("Complete landing page copy written", "✅ Done"),
                 ("VSL landing page built and live", "✅ Done"),
                 ("138-slide VSL created + 15-slide pitch version", "✅ Done"),
                 ("Funnel launched and collecting leads", "✅ Done"),
                 ("Quiz-based lead capture system built", "✅ Done"),
                 ("Admin panel for lead management built", "✅ Done"),
                 ("Competitor funnel analysis completed", "✅ Done"),
             ]),
            (4, "MARKETING SYSTEMS", "📣", GREEN,
             "Attract store owners through content and outreach.",
             [
                 ("Lead Magnet: Free Stock Leakage Audit (30 min)", "✅ Done"),
                 ("Warm Reachout scripts ready (WhatsApp + Phone + In-person)", "✅ Done"),
                 ("3 months content ready using GAP Framework + ChatGPT", "⬜ In Progress"),
                 ("Started posting on Instagram + YouTube", "⬜ In Progress"),
                 ("Content Distribution sheet", "⬜ Not Started"),
                 ("Cold Call Qualification Questionnaire ready", "✅ Done"),
                 ("WhatsApp Cold DM scripts ready (7 templates)", "✅ Done"),
                 ("10 ad creatives for Instagram/Facebook", "⬜ Not Started"),
                 ("Daily marketing: 3 hours outreach + content", "✅ Done"),
             ]),
            (5, "SALES SYSTEM", "💰", GOLD,
             "Convert interested store owners into paying clients.",
             [
                 ("Closing Script customized for retail store owners", "✅ Done"),
                 ("Objection Handling: 'staff won't follow', 'too expensive', 'my store is small'", "✅ Done"),
                 ("10 Follow-up messages for WhatsApp", "✅ Done"),
                 ("Pre-framing Video: Leaking bucket story", "✅ Done"),
                 ("Follow-up & Pre-framing tracking sheet", "⬜ In Progress"),
                 ("Role Play Training watched", "⬜ Not Started"),
                 ("Role Play practice with partner", "⬜ Not Started"),
                 ("Reminder messages before store visit", "✅ Done"),
                 ("Store visit procedure prepared", "✅ Done"),
                 ("Strategy call tracking sheet", "⬜ Not Started"),
             ]),
            (6, "ONE-TO-MANY SELLING", "🎤", RED,
             "Sell to multiple store owners at once.",
             [
                 ("Webinar PPT: 'How To Stop Losing Money From Your Store'", "⬜ In Progress"),
                 ("Webinar tracking sheet", "⬜ Not Started"),
                 ("Webinar sales tracking sheet", "⬜ Not Started"),
                 ("WhatsApp reminder messages for webinar", "⬜ Not Started"),
                 ("Practice first webinar", "⬜ Not Started"),
                 ("Conduct first webinar for store owners", "⬜ Not Started"),
                 ("Zoom account set up", "✅ Done"),
             ]),
            (7, "PERSONAL BRANDING", "⭐", PURPLE,
             "Build Aakash Savant as the go-to retail stock control expert.",
             [
                 ("Content competitor analysis (retail consultants)", "⬜ Not Started"),
                 ("Content distribution: IG + YT + LinkedIn + WhatsApp", "⬜ In Progress"),
                 ("Editor workflow set up", "⬜ Not Started"),
                 ("Consistent posting started", "⬜ In Progress"),
                 ("All social profiles optimized for 'Retail Control Architect'", "⬜ In Progress"),
             ]),
            (8, "AUTOMATION", "⚙️", BLUE,
             "Automate lead flow and client management.",
             [
                 ("Tools: Next.js app + PostgreSQL + Admin Panel + WhatsApp API", "✅ Done"),
                 ("Lead capture automation from quiz to admin panel", "✅ Done"),
             ]),
            (9, "CLIENT DELIVERY", "📦", GREEN,
             "Deliver world-class results for store owner clients.",
             [
                 ("Proposal Template: Retail Control Architect engagement proposal", "⬜ Not Started"),
                 ("Agreement: 30-day stock control implementation contract", "⬜ Not Started"),
                 ("Timeline: Week-by-week 30-day implementation roadmap", "⬜ In Progress"),
                 ("Client Onboarding: Store audit → System setup → Training → Review", "⬜ In Progress"),
             ]),
        ]

    for week_num, name, emoji, color, desc, tasks in all_pillars:
        add_pillar_section(doc, week_num, name, emoji, color, desc, tasks, is_template=is_template)
        if week_num < 9:
            doc.add_page_break()

    # Final page
    doc.add_page_break()
    for _ in range(5):
        doc.add_paragraph()

    if is_template:
        styled(doc, "YOUR 6-7 FIGURE BLUEPRINT IS READY 💎", sz=24, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
        styled(doc, "Follow the 9 pillars. Complete each task.", sz=16, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
        styled(doc, "Track your progress in the Excel sheet.", sz=16, c=GRAY, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
        styled(doc, "One pillar at a time. One week at a time.", sz=16, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        styled(doc, "RETAIL CONTROL ARCHITECT", sz=24, c=ORANGE, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
        styled(doc, "6-7 FIGURE BLUEPRINT — IN PROGRESS 💎", sz=20, c=BLACK, b=True, a=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
        styled(doc, "Keep going, Aakash. One pillar at a time.", sz=16, c=GRAY, i=True, a=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


def create_blueprint_xlsx(is_template=False):
    """Create the 6-7 Figure Business Blueprint Excel tracker."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Blueprint Tracker"

    # Styles
    header_font = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1A1A2E', end_color='1A1A2E', fill_type='solid')
    pillar_font = Font(name='Calibri', size=12, bold=True, color='FFFFFF')
    normal_font = Font(name='Calibri', size=10)
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_wrap = Alignment(horizontal='left', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin', color='CCCCCC'),
        right=Side(style='thin', color='CCCCCC'),
        top=Side(style='thin', color='CCCCCC'),
        bottom=Side(style='thin', color='CCCCCC'),
    )

    done_fill = PatternFill(start_color='E8F5E9', end_color='E8F5E9', fill_type='solid')
    progress_fill = PatternFill(start_color='FFF3E0', end_color='FFF3E0', fill_type='solid')
    not_fill = PatternFill(start_color='FFEBEE', end_color='FFEBEE', fill_type='solid')

    pillar_colors = {
        1: 'FF6B35', 2: '3D8BE8', 3: '6C5CE7', 4: '00A855',
        5: 'D4A500', 6: 'E83D3D', 7: '6C5CE7', 8: '3D8BE8', 9: '00A855'
    }

    # Column widths
    ws.column_dimensions['A'].width = 35
    ws.column_dimensions['B'].width = 6
    ws.column_dimensions['C'].width = 55
    ws.column_dimensions['D'].width = 15
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 15
    ws.column_dimensions['G'].width = 30
    ws.column_dimensions['H'].width = 35

    # Headers
    headers = ['IEC Pillar', 'S.NO', 'Task to Complete', 'Status', 'When Started', 'When Completed', 'Link / Attachment', 'Notes']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = thin_border

    # Personal details
    ws.cell(row=3, column=2, value="Personal Details").font = Font(name='Calibri', size=11, bold=True, color='FF6B35')

    if is_template:
        personal = [
            ("Name", "[Your Name]"),
            ("When Started", "[Date]"),
            ("Your Goal", "[Your business goal]"),
            ("Current Revenue", "[Current monthly revenue]"),
            ("Revenue Goal (6 months)", "[Target monthly revenue]"),
        ]
    else:
        personal = [
            ("Name", "Aakash Savant"),
            ("When Started", "February 2026"),
            ("Your Goal", "Build a 7-figure retail consulting company with Retail Control Architect"),
            ("Current Revenue", "Building case studies — starting phase"),
            ("Revenue Goal (6 months)", "Rs 3 Lakh / month"),
        ]

    for i, (label, value) in enumerate(personal):
        row = 4 + i
        ws.cell(row=row, column=2, value=label).font = Font(name='Calibri', size=10, bold=True)
        ws.cell(row=row, column=3, value=value).font = normal_font
        ws.cell(row=row, column=2).border = thin_border
        ws.cell(row=row, column=3).border = thin_border

    # Pillar data
    if is_template:
        pillars_data = [
            (1, "Foundation (YOU)", [
                "Fill in 6-7 Figure Business Blueprint sheet",
                "Time management system implementation",
                "Old You vs New You exercise",
                "Myers-Briggs Personality Test",
                "Create Vision Map",
                "Create Vision Board",
                "Set up Tracking Sheet",
                "Business Revenue & Profitability Goals",
                "Optimize Body, Mind, Place & Work habits",
                "Practice Inversion Technique (monthly)",
                "Select a book to read",
            ]),
            (2, "Grand Slam Offer", [
                "Niche selection",
                "Marketing Message",
                "Grand Slam Offer creation",
                "Proof of Concept (POC) ready",
                "Price of the Offer",
                "Entire Sales Pitch ready",
                "Competitor Research Sheet",
            ]),
            (3, "Funnels", [
                "Understand the Funnel world",
                "Select which funnel to follow",
                "Headline & sub-headline for landing page",
                "Copy of entire landing page",
                "Complete Landing Page built",
                "VSL (Video Sales Letter) ready",
                "Launch your Funnel",
                "Sales page ready",
                "Official website ready",
                "Competitor funnel hack research",
            ]),
            (4, "Marketing Systems", [
                "Lead Magnet ready",
                "Warm reachout strategy implemented",
                "3 months Content ready (GAP Framework)",
                "Started posting content on all platforms",
                "Content Distribution sheet",
                "Outreach tracking ready",
                "Outreach started",
                "10 Facebook/Instagram ads ready",
                "Daily marketing activity (min 3 hours/day)",
            ]),
            (5, "Sales System", [
                "8-Figure Closing Script personalized",
                "Objection Handling list prepared & practice",
                "10 Follow-up Messages ready",
                "Pre-framing Video ready",
                "Follow-up & Pre-framing sheet",
                "Role Play Training watched",
                "Role Play practice live",
                "Reminder Messages before calls",
                "1:1 Zoom call procedure prepared",
                "1:1 Strategy call tracking sheet",
            ]),
            (6, "One-to-Many Selling", [
                "Webinar PPT ready",
                "Webinar Tracking sheet",
                "Webinar Sales tracking sheet",
                "Webinar WhatsApp reminder messages",
                "Practice first webinar",
                "Conduct first webinar",
                "Purchase Zoom & start using",
            ]),
            (7, "Personal Branding", [
                "Content Competitor Analysis",
                "Content Distribution sheet",
                "Editor Work sheet",
                "Start posting content consistently",
                "Optimize all social profiles",
            ]),
            (8, "Automation", [
                "List all tools for your business",
                "Start automation using Business Automation Mapping",
            ]),
            (9, "Client Delivery", [
                "Proposal Template ready",
                "Agreement / Contract ready",
                "Timeline & Delivery template for onboarding",
                "Client Onboarding Process ready",
            ]),
        ]
        default_status = ""
    else:
        pillars_data = [
            (1, "Foundation (YOU)", [
                ("6-7 Figure Business Blueprint sheet", "Done"),
                ("Time management — 3hr deep work blocks", "Done"),
                ("Old You vs New You — confused to authority", "Done"),
                ("Myers-Briggs Personality Test", "Done"),
                ("Vision Map — RCA in 500+ stores by 2028", "Done"),
                ("Vision Board created", "Done"),
                ("Daily + Weekly Tracking Sheet", "Done"),
                ("Revenue Goal: Rs 3L/month in 6 months", "Done"),
                ("Morning routine: Exercise+Reading+Planning", "Done"),
                ("Inversion Technique practiced", "Done"),
                ("Reading: $100M Offers by Hormozi", "Done"),
            ]),
            (2, "Grand Slam Offer", [
                ("Niche: Clothing/retail stores with stock mismatch", "Done"),
                ("Message: Stop losing lakhs from stock mismatch in 30 days", "Done"),
                ("Grand Slam: 5-Step Framework + 30-day guarantee", "Done"),
                ("POC: Rs 3.2L recovered, 17% to 2% mismatch", "Done"),
                ("Pricing: Free audit + paid implementation", "Done"),
                ("15-Slide Sales Pitch Deck", "Done"),
                ("Competitor Sheet: ERP sellers, Vyapar, Busy", "Done"),
            ]),
            (3, "Funnels", [
                ("Studied VSL, Webinar, Direct booking funnels", "Done"),
                ("Selected: VSL funnel + Direct call booking", "Done"),
                ("Headline: Stop Losing Money From Your Own Store", "Done"),
                ("Complete landing page copy written", "Done"),
                ("VSL landing page built and live", "Done"),
                ("138-slide VSL + 15-slide pitch version", "Done"),
                ("Funnel launched and collecting leads", "Done"),
                ("Quiz-based lead capture built", "Done"),
                ("Admin panel for lead management", "Done"),
                ("Competitor funnel analysis done", "Done"),
            ]),
            (4, "Marketing Systems", [
                ("Lead Magnet: Free 30-min Stock Leakage Audit", "Done"),
                ("Warm Reachout scripts (WA+Phone+In-person)", "Done"),
                ("3 months content (GAP Framework)", "In process"),
                ("Posting on Instagram + YouTube", "In process"),
                ("Content Distribution sheet", "Not Yet"),
                ("Cold Call Qualification Questionnaire", "Done"),
                ("WhatsApp Cold DM scripts (7 templates)", "Done"),
                ("10 ad creatives for IG/FB", "Not Yet"),
                ("Daily marketing: 3hrs outreach+content", "Done"),
            ]),
            (5, "Sales System", [
                ("Closing Script for retail store owners", "Done"),
                ("Objection Handling list", "Done"),
                ("10 Follow-up Messages", "Done"),
                ("Pre-framing Video: Leaking bucket", "Done"),
                ("Follow-up tracking sheet", "In process"),
                ("Role Play Training", "Not Yet"),
                ("Role Play practice", "Not Yet"),
                ("Reminder messages before store visit", "Done"),
                ("Store visit procedure", "Done"),
                ("Strategy call tracking", "Not Yet"),
            ]),
            (6, "One-to-Many Selling", [
                ("Webinar PPT for store owners", "In process"),
                ("Webinar tracking sheet", "Not Yet"),
                ("Webinar sales tracking", "Not Yet"),
                ("WA reminder messages", "Not Yet"),
                ("Practice first webinar", "Not Yet"),
                ("Conduct first webinar", "Not Yet"),
                ("Zoom account set up", "Done"),
            ]),
            (7, "Personal Branding", [
                ("Content competitor analysis", "Not Yet"),
                ("Content distribution plan", "In process"),
                ("Editor workflow", "Not Yet"),
                ("Consistent posting started", "In process"),
                ("Social profiles optimized for RCA", "In process"),
            ]),
            (8, "Automation", [
                ("Tools: Next.js + PostgreSQL + Admin Panel", "Done"),
                ("Lead capture automation from quiz", "Done"),
            ]),
            (9, "Client Delivery", [
                ("Proposal Template for RCA engagement", "Not Yet"),
                ("Agreement: 30-day implementation contract", "Not Yet"),
                ("Week-by-week 30-day implementation roadmap", "In process"),
                ("Client Onboarding: Audit>Setup>Train>Review", "In process"),
            ]),
        ]
        default_status = None

    current_row = 11
    for pillar_num, pillar_name, tasks in pillars_data:
        # Pillar header
        color_hex = pillar_colors.get(pillar_num, 'FF6B35')
        pillar_fill = PatternFill(start_color=color_hex, end_color=color_hex, fill_type='solid')

        ws.cell(row=current_row, column=1, value=f"Week {pillar_num} = Pillar {pillar_num} = ({pillar_name})")
        ws.cell(row=current_row, column=1).font = pillar_font
        ws.cell(row=current_row, column=1).fill = pillar_fill
        ws.cell(row=current_row, column=1).alignment = left_wrap
        for col in range(1, 9):
            ws.cell(row=current_row, column=col).border = thin_border
            if col > 1:
                ws.cell(row=current_row, column=col).fill = pillar_fill

        # Tasks
        for i, task_data in enumerate(tasks):
            row = current_row + i
            if row == current_row and i == 0:
                pass  # First task on same row as header? No, put on next row
            row = current_row + i

            if default_status is not None:
                task_name = task_data
                status = ""
            else:
                task_name = task_data[0]
                status = task_data[1]

            ws.cell(row=row, column=2, value=i + 1).font = normal_font
            ws.cell(row=row, column=2).alignment = center
            ws.cell(row=row, column=3, value=task_name).font = normal_font
            ws.cell(row=row, column=3).alignment = left_wrap
            ws.cell(row=row, column=4, value=status).font = normal_font
            ws.cell(row=row, column=4).alignment = center

            # Color the status
            if status == "Done":
                ws.cell(row=row, column=4).fill = done_fill
            elif "process" in status.lower():
                ws.cell(row=row, column=4).fill = progress_fill
            elif "not" in status.lower():
                ws.cell(row=row, column=4).fill = not_fill

            for col in range(1, 9):
                ws.cell(row=row, column=col).border = thin_border

        current_row += len(tasks) + 3  # gap between pillars

    # Freeze panes
    ws.freeze_panes = 'A2'

    return wb


# ═══ GENERATE ALL 4 FILES ═══
if __name__ == "__main__":
    out_dir = "/Users/aakash/Desktop/Week_3/Week1 "
    os.makedirs(out_dir, exist_ok=True)

    # 1. Niche DOCX
    doc1 = create_blueprint_docx(is_template=False)
    p1 = os.path.join(out_dir, "Retail_Control_6-7_Figure_Blueprint.docx")
    doc1.save(p1)
    print(f"Created: {p1}")

    # 2. Template DOCX
    doc2 = create_blueprint_docx(is_template=True)
    p2 = os.path.join(out_dir, "Template_6-7_Figure_Blueprint_Guide.docx")
    doc2.save(p2)
    print(f"Created: {p2}")

    # 3. Niche XLSX
    wb1 = create_blueprint_xlsx(is_template=False)
    p3 = os.path.join(out_dir, "Retail_Control_Blueprint_Tracker.xlsx")
    wb1.save(p3)
    print(f"Created: {p3}")

    # 4. Template XLSX
    wb2 = create_blueprint_xlsx(is_template=True)
    p4 = os.path.join(out_dir, "Template_Blueprint_Tracker.xlsx")
    wb2.save(p4)
    print(f"Created: {p4}")

    print("\nAll 4 files generated!")
