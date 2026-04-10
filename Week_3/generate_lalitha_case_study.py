#!/usr/bin/env python3
"""
Generate a Case Study/Testimonial document for Lalitha Garments.
Showcases the transformation using the Retail Control Authority SaaS.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_case_study():
    doc = Document()

    # Document Styling
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # --- Header ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header.add_run("SUCCESS CASE STUDY: LALITHA GARMENTS\n")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(27, 42, 74) # Dark Navy
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = subtitle.add_run("From Operational Chaos to 98% Stock Accuracy in 30 Days")
    r_sub.font.size = Pt(14)
    r_sub.font.color.rgb = RGBColor(212, 160, 30) # Gold

    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Section 1: The Profile ---
    doc.add_heading("Client Profile", level=2)
    p = doc.add_paragraph()
    p.add_run("Client: ").bold = True
    p.add_run("Lalitha Garments\n")
    p.add_run("Niche: ").bold = True
    p.add_run("Multi-Brand Clothing Retail\n")
    p.add_run("Challenge: ").bold = True
    p.add_run("Severe Stock Mismatch, Manual Register Overload, and Silent Leakage.")

    # --- Section 2: The Initial Audit ---
    doc.add_heading("The Forensic Audit (The Discovery)", level=2)
    p = doc.add_paragraph()
    p.add_run("When we first walked into Lalitha Garments, the owner felt 'something was wrong' but couldn't put a number on it. Our Forensic Audit revealed the following:")
    
    # Stats Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    def set_cell(row, label, value):
        table.rows[row].cells[0].text = label
        table.rows[row].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[row].cells[1].text = value

    set_cell(0, "System Stock (Software)", "₹32.4 Lakh")
    set_cell(1, "Physical Stock (Actual)", "₹28.1 Lakh")
    set_cell(2, "The Gap (Mismatch)", "₹4.3 Lakh")
    set_cell(3, "Annual Leakage Projection", "₹6.8 Lakh")

    # --- Section 3: The Problem ---
    doc.add_heading("The Problem: 'The Double Reality'", level=2)
    p = doc.add_paragraph()
    p.add_run("Lalitha Garments was suffering from a 'Double Reality.' They had expensive billing software, but the actual operations were happening in raw registers and manual slips. Stock was being sold faster than it was being entered, and returns were being 'adjusted' verbally. This created a massive mismatch that made the owner a prisoner of his own store.")

    # --- Section 4: The Solution (The SaaS Intervention) ---
    doc.add_heading("The Solution: Retail Control Authority SaaS", level=2)
    p = doc.add_paragraph()
    p.add_run("We didn't just give them software; we installed an Authority System. Using our proprietary SaaS Platform, we enforced:")
    
    bullets = [
        "Digital Gatekeeper: 100% of inward stock verified vs. invoice before hitting racks.",
        "Automatic Sync: Every barcode scan deducted stock in real-time—zero manual adjustments.",
        "Staff Control Rulebook: High-accountability protocols for every staff member.",
        "Owner Visibility: A 10-minute nightly routine for the owner to verify the 'Pulse' of the store."
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # --- Section 5: The Result ---
    doc.add_heading("The Result: Total Control", level=2)
    p = doc.add_paragraph()
    p.add_run("Within 30 days of the installation at Lalitha Garments:")
    
    results = [
        "Stock Accuracy soared to 98.4%.",
        "Reconciliation time dropped from 90 minutes to 10 minutes daily.",
        "Stopped hidden leakage worth ₹50,000+ per month.",
        "Owner freedom: The owner can now leave the store for 3 days without calling his manager every hour."
    ]
    for res in results:
        doc.add_paragraph(res, style='List Bullet')

    # --- Section 6: The Testimonial Quote ---
    doc.add_paragraph("\n")
    quote_box = doc.add_paragraph()
    quote_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = quote_box.add_run("\"Before Aakash and the Retail Control System, I was working FOR my store. I was a slave to my mismatch. Now, the system works for ME. I finally have the confidence to open my second branch because I know the foundation is solid. This isn't software—it's freedom.\"")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(34, 139, 34) # Green

    p_owner = doc.add_paragraph()
    p_owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_owner = p_owner.add_run("- Owner, Lalitha Garments")
    r_owner.bold = True

    # Save
    save_path = "/Users/aakash/Desktop/Week_3/Usecase_Lalitha_Garments_Success_Story.docx"
    doc.save(save_path)
    return save_path

if __name__ == "__main__":
    path = create_case_study()
    print(f"Case Study created: {path}")
