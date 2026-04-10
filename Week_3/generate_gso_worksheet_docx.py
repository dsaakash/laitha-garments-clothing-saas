#!/usr/bin/env python3
"""
Generate the filled-in GSO Building Worksheet .docx
For Aakash Savant — Retail Control Architect
All 7 steps filled with niche-specific content.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colours ─────────────────────────────────────────────────────
DEEP_BLUE   = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT_GOLD = RGBColor(0xD4, 0xA0, 0x1E)
ACCENT_TEAL = RGBColor(0x17, 0xA2, 0xB8)
WARM_RED    = RGBColor(0xC0, 0x39, 0x2B)
SUCCESS_GRN = RGBColor(0x27, 0xAE, 0x60)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
MED_GRAY    = RGBColor(0x66, 0x66, 0x66)
BODY        = RGBColor(0x2C, 0x2C, 0x2C)
ORANGE      = RGBColor(0xE6, 0x7E, 0x22)
PURPLE      = RGBColor(0x8E, 0x44, 0xAD)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11); style.font.color.rgb = BODY

# ── Helpers ─────────────────────────────────────────────────────
def shade(cell, hx):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hx}"/>'))

def spacer(pts=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)

def cover_title(txt, sz=26):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(4); p.space_after = Pt(4)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = DEEP_BLUE

def cover_sub(txt, sz=13, c=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    r = p.add_run(txt); r.italic = True; r.font.size = Pt(sz); r.font.color.rgb = c or ACCENT_GOLD

def sec_hdr(txt, emoji=""):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.space_before = Pt(16); p.space_after = Pt(3)
    full = f"{emoji}  {txt}" if emoji else txt
    r = p.add_run(full); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = DEEP_BLUE
    ln = doc.add_paragraph(); ln.space_after = Pt(6)
    rl = ln.add_run("━" * 75); rl.font.size = Pt(5); rl.font.color.rgb = ACCENT_GOLD

def sub_hdr(txt, c=None):
    p = doc.add_paragraph()
    p.space_before = Pt(12); p.space_after = Pt(3)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = c or ACCENT_TEAL

def minor_hdr(txt, c=None):
    p = doc.add_paragraph()
    p.space_before = Pt(8); p.space_after = Pt(2)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = c or DEEP_BLUE

def body(txt):
    p = doc.add_paragraph(txt)
    p.paragraph_format.space_after = Pt(5)
    return p

def body_b(txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(txt); r.bold = True; r.font.size = Pt(11)
    return p

def highlight(txt, bg="FFF8E1"):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]; c.text = ''
    p = c.paragraphs[0]
    r = p.add_run(txt); r.font.size = Pt(10.5); r.font.color.rgb = BODY
    shade(c, bg); c.width = Inches(6.0)
    spacer(6)

def numbered_fill(num_val, text):
    """A numbered item styled as filled-in worksheet entry."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r_num = p.add_run(f"{num_val}:- ")
    r_num.bold = True
    r_num.font.size = Pt(11)
    r_num.font.color.rgb = ACCENT_TEAL
    r_txt = p.add_run(text)
    r_txt.font.size = Pt(11)

def qa_row(question, answer):
    """Question-Answer pair for Step 4."""
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    # Question cell
    cq = t.rows[0].cells[0]; cq.text = ''
    pq = cq.paragraphs[0]
    rq_label = pq.add_run("Q: "); rq_label.bold = True; rq_label.font.size = Pt(10); rq_label.font.color.rgb = DEEP_BLUE
    rq = pq.add_run(question); rq.font.size = Pt(10)
    shade(cq, "EBF5FB")

    # Answer cell
    ca = t.rows[1].cells[0]; ca.text = ''
    pa = ca.paragraphs[0]
    ra_label = pa.add_run("A: "); ra_label.bold = True; ra_label.font.size = Pt(10); ra_label.font.color.rgb = SUCCESS_GRN
    ra = pa.add_run(answer); ra.font.size = Pt(10)
    shade(ca, "F0FFF0")

    cq.width = Inches(6.0); ca.width = Inches(6.0)
    spacer(6)

def problem_solution_row(problem, solution):
    """Problem → Solution mapping."""
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'

    cp = t.rows[0].cells[0]; cp.text = ''
    pp = cp.paragraphs[0]
    rp = pp.add_run(problem); rp.font.size = Pt(10)
    shade(cp, "FDEDED"); cp.width = Inches(2.5)

    ca = t.rows[0].cells[1]; ca.text = ''
    pa = ca.paragraphs[0]
    ra = pa.add_run("→"); ra.bold = True; ra.font.size = Pt(14); ra.font.color.rgb = ACCENT_GOLD
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ca.width = Inches(0.4)

    cs = t.rows[0].cells[2]; cs.text = ''
    ps = cs.paragraphs[0]
    rs = ps.add_run(solution); rs.font.size = Pt(10)
    shade(cs, "E8F8F5"); cs.width = Inches(3.1)

    spacer(3)

def stbl(headers, rows, cw=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
        shade(c, "1B2A4A")
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = table.rows[ri + 1].cells[ci]; c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(str(v)); r.font.size = Pt(10)
            if ri % 2 == 0: shade(c, "F7F9FC")
    if cw:
        for i, w in enumerate(cw):
            for row in table.rows: row.cells[i].width = Inches(w)
    spacer(6)
    return table


# ════════════════════════════════════════════════════════════════
#                        COVER PAGE
# ════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

cover_title("BUILDING A GRAND SLAM OFFER", 28)
cover_title("7-STEP WORKSHEET", 22)
spacer(8)
cover_sub("Filled & Completed for:", 12, MED_GRAY)
cover_title("The 30-Day Stock Certainty System™", 20)
spacer(4)
cover_sub("By Aakash Savant — Retail Control Architect", 12, MED_GRAY)
cover_sub("Niche: Retail Operational Control for Clothing Stores", 11, ACCENT_TEAL)

doc.add_paragraph()
highlight(
    "THE 7 STEPS TO A GRAND SLAM OFFER\n\n"
    "Step 1 → Identify the Dream Outcome\n"
    "Step 2 → List the Problems\n"
    "Step 3 → Turn Problems Into Solutions\n"
    "Step 4 → Create Delivery Vehicles (The How)\n"
    "Step 5 → Supercharge the Offer\n"
    "Step 6 → Trim & Stack\n"
    "Step 7 → Increase Perceived Value",
    "F0F4FF"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 1: DREAM OUTCOME
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 1: IDENTIFY THE DREAM OUTCOME", "🎯")
body("What does your ideal client want more than anything? Their \"heaven.\" Their transformation.")
spacer(4)

dream_outcomes = [
    "System stock always matches physical stock — every single day, without manual checking.",
    "Daily reconciliation drops from 1–2 hours to under 15 minutes.",
    "No surprise stock shortages when a customer asks for a specific size or color.",
    "Every purchase entry recorded the moment stock enters the store — zero delays.",
    "Every sale automatically reduces inventory in real time — no batch entries.",
    "Zero parallel systems — no manual register running alongside billing software.",
    "Clean, organized item structure with proper SKU names, sizes, and categories.",
    "Accurate profit visibility — knowing real margins without guesswork.",
    "Hidden leakage completely stopped — no silent money walking out the door.",
    "Dead stock identified early before capital gets permanently stuck.",
    "Supplier outstanding tracked clearly — no overbilling goes unnoticed.",
    "Owner can verify stock, sales, and supplier balances independently — without asking staff.",
    "Business runs smoothly even if the most trusted staff member leaves tomorrow.",
    "No fear of internal manipulation or silent theft.",
    "Ability to leave the store for 2–3 days without anxiety or constant phone calls.",
    "A structured daily control routine that takes 10 minutes and covers everything.",
    "Full confidence to open a second store — knowing the first runs on system, not memory.",
    "Better purchase decisions based on real data — not gut feeling or staff memory.",
    "Working capital freed from wrong inventory — cash flowing properly.",
    "A store that runs on SYSTEM — not on the OWNER.",
]

for i, dream in enumerate(dream_outcomes, 1):
    numbered_fill(i, dream)

spacer(6)
highlight(
    "EMOTIONAL CORE:\n"
    "Control • Clarity • Certainty • Independence\n\n"
    "At the deepest level — they want to feel like the REAL OWNER of their business.\n"
    "Not a prisoner of it.",
    "FFF8E1"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 2: LIST THE PROBLEMS
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 2: LIST THE PROBLEMS", "😰")
body("All the pain points, frustrations, and obstacles standing between your client and their dream outcome.")
spacer(4)

problems = [
    "System stock showing 47 pieces but physical count reveals 39 — nobody knows where the difference came from.",
    "Wrong sizes showing as available — customer comes in, item doesn't exist, sale lost.",
    "Duplicate item codes creating ghost inventory that looks real on screen but doesn't exist on shelves.",
    "Inward entries done at end of day — creating a window where stock is completely invisible.",
    "Cash flow tightening despite decent sales — capital stuck in wrong inventory and dead stock.",
    "Supplier invoices paid without verification — overbilling going unnoticed for months.",
    "1–2 hours spent every day on reconciliation that should take 15 minutes.",
    "Manual register running parallel to billing software — two different realities, neither accurate.",
    "One staff member holds all operational knowledge in their head — owner lives in fear of them leaving.",
    "Owner cannot independently verify ANY number without asking staff first — zero control.",
    "Bought billing software but mismatch continued because entry behavior didn't change.",
    "Switched to different billing software — same chaos, different interface.",
    "Tried hiring additional staff for data entry — more people, more errors, more cost.",
    "Believes \"Retail mein thoda mismatch toh hoga hi\" — normalized dysfunction.",
    "Believes stricter systems will slow down billing and frustrate customers.",
    "Watching competitors open second and third branches while they struggle to control one store.",
    "GST compliance pressure increasing — authorities demanding accurate records they don't have.",
    "Fear that if they confront the key employee about processes, that person will leave.",
    "Profit calculation always feels wrong — numbers don't add up confidently at tax time.",
    "Cannot scale, cannot delegate, cannot step away — trapped inside their own business.",
]

for i, prob in enumerate(problems, 1):
    numbered_fill(i, prob)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 3: TURN PROBLEMS INTO SOLUTIONS
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 3: TURN PROBLEMS INTO SOLUTIONS", "🔧")
body("Every problem from Step 2 has a direct, structured solution. Zero gaps.")
spacer(4)

ps_pairs = [
    ("Stock mismatch unknown. Owner doesn't know actual % or rupee cost of leakage.",
     "Control Gap Audit™ — Count selected SKUs physically, compare to system, calculate exact mismatch % and annual rupee cost."),

    ("Duplicate item codes creating ghost inventory. Wrong categorization. No SKU discipline.",
     "Inventory Foundation Reset — Clean every item, standardize SKU naming, map categories/sizes/colors, remove duplicates, correct opening stock."),

    ("Stock enters store without being entered in system immediately. No defined entry responsibility.",
     "Supplier Entry Lock™ — Non-negotiable rule: no stock physically enters without simultaneous system entry. Defined purchase workflow."),

    ("Billing happens without automatic stock deduction. Manual adjustments after the fact.",
     "Sales Deduction Lock™ — No sale processed outside system. Every billing transaction deducts at SKU level in real time. Zero exceptions."),

    ("Manual register running parallel to billing software. Excel sheets alongside system.",
     "Single Source of Truth Enforcement — Physically remove parallel registers. Delete duplicate tracking. One system. One number. One reality."),

    ("Owner cannot independently verify stock without asking staff. Business knowledge locked in one person.",
     "Owner Visibility Dashboard — Real-time stock view, supplier summary, daily sales movement. 10-Minute Daily Control Routine™ installed."),

    ("Rules set previously but never followed consistently. Staff reverts to old habits within days.",
     "30-Day Discipline Installation™ — Weekly structured reviews, daily compliance tracking, error correction, discipline reinforced until habit."),

    ("No measurable proof of improvement. Owner unsure if investment was worth it.",
     "Before & After Validation Report™ — Re-measure mismatch %, reconciliation time, supplier clarity, independence score. Proof in rupees."),

    ("Supplier outstanding unclear. Overbilling unnoticed for months.",
     "Supplier Outstanding Dashboard — Real-time tracking, verification step before any payment released, clear purchase history."),

    ("Dead stock invisible. Capital stuck in inventory that hasn't moved for 120+ days.",
     "Dead Stock Recovery Protocol — Tag slow-moving/dead inventory clearly, aging analysis activated, recovery plan created."),

    ("Staff entering data late / inconsistently. No enforcement mechanism.",
     "Staff Control Rulebook™ — Pre-written discipline rules for billing, inward, reconciliation. System enforces, not owner."),

    ("Owner spending 1–2 hours on daily reconciliation. Exhausting and unsustainable.",
     "10-Minute Daily Control Routine™ — Structured checklist replacing chaotic 2-hour process. Everything covered in 10 minutes."),
]

for prob, sol in ps_pairs:
    problem_solution_row(prob, sol)

spacer(6)
highlight(
    "COMPLETE COVERAGE CHECK:\n"
    "Stock mismatch → Audit + Structure + Lock\n"
    "Supplier confusion → Supplier Protocol + Dashboard\n"
    "Manual duplication → One-System Enforcement\n"
    "Staff dependency → Owner Dashboard + Rulebook\n"
    "Hidden leakage → Measured Validation\n"
    "Reconciliation stress → Daily Control Routine\n"
    "Scaling fear → Before-After Report + Expansion Scorecard\n\n"
    "NOTHING LEFT UNRESOLVED.",
    "E8F8F5"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 4: DELIVERY VEHICLES
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 4: CREATE DELIVERY VEHICLES (THE HOW)", "🚀")
body("Define exactly how you deliver the transformation.")
spacer(4)

qa_row(
    "What level of personal attention do I want to provide? One-on-one, small group, one-to-many?",
    "1:1 (One-on-One) Consulting-Led Implementation.\n\n"
    "Why: Each store has different chaos. Inventory setup differs. Staff discipline varies. "
    "You must diagnose and enforce individually. This is not scalable SaaS initially — "
    "this is controlled installation. Group/course models fail because retail is behavior-driven "
    "and discipline-sensitive."
)

qa_row(
    "What level of effort is expected from them? DIY, DWY, or DFY?",
    "Done-With-You (DWY) + Strict Enforcement.\n\n"
    "Why: Cannot fully DFY because they control physical inventory and staff behavior must be internalized. "
    "Cannot DIY because they won't implement properly — structure breaks without supervision.\n\n"
    "YOU: Design structure, set up system, define rules, train staff, monitor compliance.\n"
    "THEY: Follow rules, enter data properly, stop parallel systems.\n"
    "You enforce. They comply."
)

qa_row(
    "If doing something live, what environment or medium to deliver in?",
    "Hybrid: Physical + Remote.\n\n"
    "• Phase 1 (Audit): In-person visit to the store (physical stock count required)\n"
    "• Phase 2 (Setup): Can be done remotely via Zoom/AnyDesk + WhatsApp coordination\n"
    "• Phase 3 (Enforcement): Weekly Zoom review calls + daily WhatsApp compliance checks\n"
    "• Support: WhatsApp chat support during 30-day implementation\n\n"
    "Start local/regional for first 10–15 installs. Design SOPs so remote works later."
)

qa_row(
    "If doing a recording, how do I want them to consume it?",
    "Video (short Loom/WhatsApp videos) for:\n"
    "• Staff training on entry protocols (5–10 min each)\n"
    "• 10-Minute Daily Routine walkthrough\n"
    "• Dashboard navigation tutorial\n\n"
    "Written (PDF) for:\n"
    "• Staff Control Rulebook™\n"
    "• Supplier Entry Protocol checklist\n"
    "• Before-After Validation Report™"
)

qa_row(
    "How quickly do we want to reply? On what days? During what hours?",
    "During 30-Day Installation:\n"
    "• WhatsApp: Reply within 2 hours during business hours (10 AM – 7 PM, Mon–Sat)\n"
    "• Weekly Review Calls: Scheduled on fixed day (e.g., every Monday at 11 AM)\n"
    "• Urgent Issues: Same-day response via WhatsApp voice note\n\n"
    "Post-Installation (Monthly Maintenance):\n"
    "• Reply within 24 hours\n"
    "• Monthly review call"
)

qa_row(
    "0x to 10x Test: If customers paid ₹10 Lakh, what would I provide?",
    "For ₹10 Lakh, I would provide:\n\n"
    "✅ Full-time on-site presence for 30 days (physically in-store)\n"
    "✅ Complete inventory recount and correction (every single SKU)\n"
    "✅ Personal staff training — individual, not group\n"
    "✅ Custom-built dashboard with owner's specific KPIs\n"
    "✅ Daily compliance enforcement in-person\n"
    "✅ Supplier renegotiation support based on audit findings\n"
    "✅ 90-day on-call support post-installation\n"
    "✅ Second-store expansion blueprint with implementation\n"
    "✅ Quarterly on-site stability audits for 1 year\n"
    "✅ Video case study production for their marketing\n\n"
    "→ This reveals what's TRULY valuable. Pull elements down into your actual offer."
)

sub_hdr("Delivery Vehicle Map (Final)")
stbl(
    ["Phase", "Duration", "Format", "Effort Model", "Medium"],
    [
        ["Phase 1: Control Gap Audit", "Week 1", "1:1 Consulting", "Done-With-You", "In-Person + WhatsApp"],
        ["Phase 2: Structure & Setup", "Week 2–3", "1:1 Consulting", "Done-With-You", "Zoom/AnyDesk + WhatsApp"],
        ["Phase 3: Enforcement", "Week 3–4", "1:1 Monitoring", "Strict Enforcement", "Weekly Zoom + Daily WhatsApp"],
        ["Phase 4: Validation", "Day 30", "1:1 Review", "Done-For-You Report", "Zoom Call + PDF Report"],
        ["Post-Install", "Monthly", "Maintenance", "Light Oversight", "Monthly Call + WhatsApp"],
    ],
    cw=[1.5, 0.7, 1.0, 1.2, 1.6]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 5: SUPERCHARGE THE OFFER
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 5: SUPERCHARGE THE OFFER", "⚡")

# Urgency
sub_hdr("1. Urgency — Why They Must Act NOW", WARM_RED)
highlight(
    "THE LEAKAGE MULTIPLIER EFFECT:\n\n"
    "Every month you delay fixing stock control, the leakage continues silently.\n\n"
    "• Wrong purchase decisions compound on inaccurate data\n"
    "• Dead stock accumulates further\n"
    "• Staff habits harden — harder to change later\n"
    "• Capital stays frozen in wrong inventory\n\n"
    "A retailer with 10% mismatch on ₹40L inventory loses ~₹1L–1.5L in margin impact annually.\n"
    "Every month of delay costs ₹8,000–12,000 in SILENT leakage.\n\n"
    "IMPLEMENTATION RULE:\n"
    "New installations begin on the 1st of every month ONLY.\n"
    "No rolling start. Miss it → wait 30 days while leakage continues.",
    "FDEDED"
)

# Scarcity
sub_hdr("2. Scarcity — Why Not Everyone Gets In", ORANGE)
highlight(
    "Maximum 3 Clothing Stores Per Month.\n\n"
    "Not because of marketing. Because enforcement requires real attention.\n\n"
    "Script: \"I personally install this system in each store. That means I can only work with "
    "3 clothing stores per month. If all 3 slots are filled when you decide, the next available "
    "window is next month's cycle.\"\n\n"
    "This is REAL capacity-based scarcity. Not a fake countdown timer.",
    "FFF8E1"
)

# Guarantee
sub_hdr("3. Guarantee — The Stock Certainty Guarantee™", SUCCESS_GRN)
highlight(
    "THE STOCK CERTAINTY GUARANTEE™\n\n"
    "\"If after full 30-day compliance your stock mismatch does not reduce measurably — "
    "we continue working with you at zero additional setup fee until control is achieved.\"\n\n"
    "Close line: \"I don't win unless you win.\"\n\n"
    "This is effort-continuation guarantee. Not refund gimmick. Not unrealistic % claim.\n"
    "Builds trust without destroying margins.",
    "E8F8F5"
)

# Bonuses
sub_hdr("4. Bonuses — All Aligned to Control", ACCENT_TEAL)
stbl(
    ["#", "Bonus Name", "Purpose", "Value"],
    [
        ["🎁 1", "Hidden Leakage Exposure Report™", "Primary sales closing tool. Shows annual loss in ₹ — reframes your price as an investment.", "₹5,000"],
        ["🎁 2", "Staff Control Rulebook™", "Pre-written discipline rules for billing + inward + reconciliation. Permanent enforcement after engagement ends.", "₹3,000"],
        ["🎁 3", "10-Minute Daily Control Routine™", "Repeatable owner independence habit. Replaces 2-hour reconciliation chaos forever.", "₹2,000"],
        ["🎁 4", "90-Day Stability Audit™", "Re-check store after 90 days. Proves improvement is not temporary. Prevents system decay.", "₹8,000"],
        ["🎁 5", "Expansion Readiness Scorecard™", "If planning second store, this ensures structure scales. Activates growth dream.", "₹4,000"],
    ],
    cw=[0.3, 2.0, 2.5, 0.5]
)

# Naming
sub_hdr("5. Naming — Unique Mechanism", PURPLE)
stbl(
    ["Element", "Name"],
    [
        ["Offer Name", "The 30-Day Stock Certainty System™"],
        ["Tagline", "\"Make your system stock match your shop — permanently.\""],
        ["Identity", "Retail Control Architect™"],
        ["Category", "Retail Control Architecture™"],
        ["Method", "The Control Gap Method™ / Stock Certainty Method™"],
        ["One-Liner", "\"I fix stock mismatch in clothing stores in 30 days — permanently.\""],
    ],
    cw=[1.3, 4.7]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 6: TRIM & STACK
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 6: TRIM & STACK", "✂️")
body("Remove high-cost/low-value items. Keep only what drives the 4 value levers:")
body("1) Financial value  2) Likelihood of success  3) Less effort  4) Less time")

sub_hdr("❌ What Was REMOVED (Trimmed)")
stbl(
    ["Removed Item", "Why"],
    [
        ["Payroll / HR modules", "Not related to stock control. Distracts from core promise."],
        ["Advanced analytics & BI dashboards", "Over-engineering. Retailer needs simple control, not data science."],
        ["Full ERP implementation", "Scope creep. Increases time, effort, and confusion."],
        ["Multi-store management features", "Client is single-store. Add later for expansion clients."],
        ["CRM / customer management", "Not the problem being solved. Keep focus narrow."],
        ["Accounting / GST filing", "Adjacent problem. Refer to CA. Don't dilute positioning."],
    ],
    cw=[2.2, 3.8]
)

sub_hdr("✅ Final Stacked Offer — Main Offering")
stbl(
    ["#", "Component", "What It Does", "Value"],
    [
        ["1", "Control Gap Audit", "Measure exact mismatch % and annual rupee cost of leakage", "₹5,000"],
        ["2", "Inventory Foundation Reset", "Clean SKU logic, categories, sizes, remove duplicates", "₹15,000"],
        ["3", "Supplier Entry Lock™", "No stock enters without immediate system entry", "₹10,000"],
        ["4", "Sales Deduction Lock™", "Every sale reduces inventory automatically in real time", "₹10,000"],
        ["5", "Single System Enforcement", "Remove every parallel register and duplicate system", "₹8,000"],
        ["6", "30-Day Compliance Monitoring", "Weekly calls, WhatsApp support, daily compliance tracking", "₹20,000"],
        ["7", "Before & After Validation Report™", "Measurable proof of transformation in rupees and %", "₹7,000"],
    ],
    cw=[0.3, 2.0, 2.7, 0.7]
)

sub_hdr("🎁 Bonuses (Stacked)")
stbl(
    ["#", "Bonus", "Value"],
    [
        ["Bonus 1", "Hidden Leakage Exposure Report™", "₹5,000"],
        ["Bonus 2", "Staff Control Rulebook™", "₹3,000"],
        ["Bonus 3", "10-Minute Daily Control Routine™", "₹2,000"],
        ["Bonus 4", "90-Day Stability Audit™", "₹8,000"],
        ["Bonus 5", "Expansion Readiness Scorecard™", "₹4,000"],
    ],
    cw=[0.8, 3.5, 0.8]
)

# Total + Price
spacer(4)
highlight(
    "TOTAL VALUE:  ₹97,000\n\n"
    "PRICE (Your Investment):  ₹85,000  (One-time setup)\n\n"
    "INVESTMENT LOGIC:\n"
    "• 10% mismatch on ₹40L inventory = ₹4L uncontrolled stock annually\n"
    "• Even 50% correction in Year 1 recovers far more than ₹85,000\n"
    "• If Hidden Leakage Report shows < ₹85,000 in annual impact → we tell you honestly and walk away\n\n"
    "This is not an expense. It's leakage control.",
    "FFF3E0"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# STEP 7: INCREASE PERCEIVED VALUE
# ════════════════════════════════════════════════════════════════
sec_hdr("STEP 7: INCREASE PERCEIVED VALUE", "📈")
body("Apply Hormozi's Value Equation:")
highlight(
    "Value = (Dream Outcome × Perceived Likelihood) ÷ (Time Delay × Effort & Sacrifice)\n\n"
    "↑ Numerator (increase)     ↓ Denominator (decrease)     =     MASSIVE perceived value",
    "F0F4FF"
)

# 1. Reduce Effort
sub_hdr("No. 1: REDUCE Effort & Sacrifice", ACCENT_TEAL)
stbl(
    ["What Creates Effort", "How We Reduce It"],
    [
        ["Learning new software is scary", "We configure everything. No technical knowledge required."],
        ["Changing staff behavior feels risky", "Staff Control Rulebook™ makes YOU the system, not the bad guy."],
        ["Owner thinks THEY have to do all the work", "\"Your job is 3 rules. We handle everything else.\""],
        ["Daily control feels like more work", "10-Minute Daily Routine™ — replaces 2-hour chaos with 10-min check"],
        ["Fear of disrupting current operations", "We install alongside operations. No downtime. No disruption."],
    ],
    cw=[2.5, 3.5]
)
body_b("Owner effort = 3 rules + 10 minutes/day. That's it.")

# 2. Reduce Time
sub_hdr("No. 2: REDUCE Time Delay", ORANGE)
stbl(
    ["What Creates Delay", "How We Reduce It"],
    [
        ["\"30 days sounds long\"", "\"Most stores feel the difference within the first 14 days.\""],
        ["No visible milestones", "Weekly milestones: Week 1 = audit, Week 2 = structure, Week 3 = enforce, Week 4 = validate"],
        ["\"When does it start?\"", "Day 1 win: Hidden Leakage Report reveals exact rupee cost immediately."],
        ["\"What if it takes months?\"", "30-day fixed cycle. Not open-ended consulting."],
    ],
    cw=[2.5, 3.5]
)
body_b("Perceived timeline: 14 days (visible shift) → 30 days (full control).")

# 3. Increase Likelihood
sub_hdr("No. 3: INCREASE Perceived Likelihood of Success", SUCCESS_GRN)
stbl(
    ["What Creates Doubt", "How We Increase Certainty"],
    [
        ["\"I've tried software before\"", "\"This is not software. This is a Control Installation with enforcement.\""],
        ["\"Will this work for MY store?\"", "Diagnosis-first approach: we measure YOUR mismatch before pitching."],
        ["\"What if staff doesn't follow?\"", "30-Day Enforcement Period™ — we monitor and correct daily."],
        ["\"What if it doesn't work?\"", "Stock Certainty Guarantee™ — continue free until it works."],
        ["\"Where's the proof?\"", "Case studies + Before-After Validation Report™ + video testimonials."],
    ],
    cw=[2.5, 3.5]
)
body_b("Diagnosis + Process Detail + Enforcement + Guarantee + Proof = maximum certainty.")

# 4. Increase Dream
sub_hdr("No. 4: INCREASE Dream Outcome", PURPLE)
stbl(
    ["Current Perception", "Expanded Vision You Create"],
    [
        ["\"Fix my stock problem\"", "\"Run a store where system stock = reality, every day, forever.\""],
        ["\"Save some money\"", "\"Recover ₹1–4L annually in hidden leakage — this pays for itself in 44 days.\""],
        ["\"Get better software\"", "\"Get a system that makes you the REAL owner — not a prisoner of your own store.\""],
        ["\"One store is enough\"", "\"Build the foundation so Store #2 is a copy-paste success.\""],
        ["\"I just want less stress\"", "\"Walk in, check one screen in 10 minutes, know everything is under control, go home on time.\""],
    ],
    cw=[2.2, 3.8]
)
body_b("Identity upgrade: From 'shopkeeper in chaos' → 'professional retailer in control.'")

spacer(8)

# Final Value Equation Summary
sub_hdr("VALUE EQUATION: FINAL SCORE")
stbl(
    ["Variable", "BEFORE\n(Unoptimized)", "AFTER\n(Your Offer)", "How"],
    [
        ["Dream Outcome", "6/10", "9/10", "Identity upgrade + expansion anchor + visual specificity"],
        ["Perceived Likelihood", "5/10", "9/10", "Diagnosis first + process detail + guarantee + proof"],
        ["Time Delay", "6/10\n(high = bad)", "2/10\n(low = good)", "14-day visible shift + Day 1 win + weekly milestones"],
        ["Effort & Sacrifice", "5/10\n(high = bad)", "2/10\n(low = good)", "3 rules + 10 min daily + you handle everything"],
    ],
    cw=[1.2, 0.9, 0.9, 3.0]
)

highlight(
    "RESULT:\n"
    "(9 × 9) ÷ (2 × 2)?  =  81 ÷ 4  =  20.25×\n\n"
    "vs. Unoptimized: (6 × 5) ÷ (6 × 5)  =  30 ÷ 30  =  1×\n\n"
    "Your offer has 20× HIGHER perceived value than an average unoptimized offer in the same market.\n\n"
    "That's a Grand Slam Offer.",
    "E8F8F5"
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# FINAL SUMMARY PAGE
# ════════════════════════════════════════════════════════════════
sec_hdr("THE COMPLETE GRAND SLAM OFFER — ONE PAGE", "🏆")

cheat = doc.add_table(rows=14, cols=2)
cheat.alignment = WD_TABLE_ALIGNMENT.CENTER
cheat.style = 'Table Grid'

data = [
    ("OFFER NAME", "The 30-Day Stock Certainty System™"),
    ("TAGLINE", "Make your system stock match your shop — permanently."),
    ("CORE PROMISE", "In 30 days: system stock = physical stock, reconciliation under 15 min, one disciplined system."),
    ("TARGET CLIENT", "Owner-operated clothing stores, ₹30L–1.5Cr revenue, ₹20L+ inventory, stock mismatch."),
    ("FORMAT", "1:1 Consulting-Led Implementation, Done-With-You + Strict Enforcement."),
    ("DURATION", "30 Days Installation + Optional Monthly Maintenance."),
    ("CAPACITY", "3 stores per month maximum. Cycles begin on the 1st only."),
    ("INVESTMENT", "₹85,000 one-time setup fee."),
    ("TOTAL VALUE", "₹97,000 (7 components + 5 bonuses)."),
    ("GUARANTEE", "Stock Certainty Guarantee™ — continue free if mismatch doesn't reduce."),
    ("URGENCY", "Leakage costs ₹8K–12K/month. Delay has a rupee cost."),
    ("SCARCITY", "3 stores/month max — real capacity limit, not marketing."),
    ("BONUSES", "Hidden Leakage Report + Staff Rulebook + Daily Routine + 90-Day Audit + Expansion Scorecard."),
    ("ONE SENTENCE", "\"I fix stock mismatch in clothing stores in 30 days — permanently.\""),
]

for i, (label, value) in enumerate(data):
    cl = cheat.rows[i].cells[0]; cr = cheat.rows[i].cells[1]
    cl.text = ''; cr.text = ''
    pl = cl.paragraphs[0]; pr = cr.paragraphs[0]
    rl = pl.add_run(label); rl.bold = True; rl.font.size = Pt(9.5); rl.font.color.rgb = WHITE
    shade(cl, "1B2A4A"); cl.width = Inches(1.5)
    rr = pr.add_run(value); rr.font.size = Pt(9.5)
    cr.width = Inches(4.5)

spacer(12)

# Final positioning
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p.add_run("I don't sell retail software.\n")
r1.bold = True; r1.font.size = Pt(15); r1.font.color.rgb = DEEP_BLUE
r2 = p.add_run("I install Retail Control Architecture™ inside clothing stores.")
r2.bold = True; r2.font.size = Pt(15); r2.font.color.rgb = ACCENT_GOLD

spacer(8)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p2.add_run("— Aakash Savant, Retail Control Architect —")
r3.italic = True; r3.font.size = Pt(12); r3.font.color.rgb = MED_GRAY


# ── Save ───────────────────────────────────────────────────────
out = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "GSO_Building_Worksheet_Filled.docx"
)
doc.save(out)
print(f"✅ Successfully created: {out}")
